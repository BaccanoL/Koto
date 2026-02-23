#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
æ–‡æ¡£å·¥ä½œæµæ‰§è¡Œå™¨ - Document Workflow Executor
è‡ªåŠ¨è¯†åˆ«æ–‡æ¡£ä¸­çš„å·¥ä½œæµè§„åˆ’å¹¶æ‰§è¡Œ

åŠŸèƒ½ï¼š
1. è¯»å–Word/PDF/Markdownæ–‡æ¡£
2. æå–å·¥ä½œæµæ­¥éª¤ï¼ˆå®éªŒæ­¥éª¤ã€è®¡åˆ’ã€æµç¨‹ï¼‰
3. è‡ªåŠ¨åˆ†è§£ä¸ºå¯æ‰§è¡Œä»»åŠ¡
4. æŒ‰é¡ºåºæ‰§è¡Œå¹¶æ”¶é›†ç»“æœ
5. ç”Ÿæˆå®Œæ•´çš„æ‰§è¡ŒæŠ¥å‘Š
"""

import os
import json
import asyncio
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path


class WorkflowStep:
    """å·¥ä½œæµæ­¥éª¤"""
    def __init__(self, step_id: int, description: str, step_type: str, 
                 input_data: Any = None, expected_output: str = None):
        self.step_id = step_id
        self.description = description
        self.step_type = step_type  # VLM, SEARCH, CODE, FILE_GEN, etc.
        self.input_data = input_data
        self.expected_output = expected_output
        self.status = "pending"  # pending, running, completed, failed
        self.result = None
        self.error = None
        self.start_time = None
        self.end_time = None
    
    def to_dict(self):
        return {
            "step_id": self.step_id,
            "description": self.description,
            "step_type": self.step_type,
            "status": self.status,
            "result": self.result,
            "error": self.error,
            "duration": self._duration()
        }
    
    def _duration(self):
        if self.start_time and self.end_time:
            return (self.end_time - self.start_time).total_seconds()
        return None


class DocumentWorkflowExecutor:
    """æ–‡æ¡£å·¥ä½œæµæ‰§è¡Œå™¨"""
    
    # å…³é”®è¯æ˜ å°„åˆ°ä»»åŠ¡ç±»å‹
    KEYWORD_TO_TASK = {
        # VLMç›¸å…³
        "è¯†åˆ«": "VLM",
        "åˆ†æå›¾ç‰‡": "VLM",
        "å›¾åƒè¯†åˆ«": "VLM",
        "çœ‹å›¾": "VLM",
        "æè¿°å›¾ç‰‡": "VLM",
        "è§†è§‰": "VLM",
        
        # æœç´¢ç›¸å…³
        "æœç´¢": "WEB_SEARCH",
        "æŸ¥æ‰¾": "WEB_SEARCH",
        "æŸ¥è¯¢": "WEB_SEARCH",
        "æ£€ç´¢": "WEB_SEARCH",
        
        # ä»£ç æ‰§è¡Œ
        "è¿è¡Œä»£ç ": "CODE",
        "æ‰§è¡Œä»£ç ": "CODE",
        "è®¡ç®—": "CODE",
        "python": "CODE",
        
        # æ–‡ä»¶ç”Ÿæˆ
        "ç”Ÿæˆæ–‡æ¡£": "FILE_GEN",
        "åˆ›å»ºæ–‡ä»¶": "FILE_GEN",
        "å†™å…¥æ–‡ä»¶": "FILE_GEN",
        "ç”ŸæˆæŠ¥å‘Š": "FILE_GEN",
        "åšppt": "FILE_GEN",
        "åšè¡¨æ ¼": "FILE_GEN",
        
        # æ•°æ®å¤„ç†
        "å¤„ç†æ•°æ®": "DATA",
        "åˆ†ææ•°æ®": "DATA",
        "ç»Ÿè®¡": "DATA",
        
        # é€šç”¨
        "æ¯”è¾ƒ": "COMPARE",
        "å¯¹æ¯”": "COMPARE",
        "æ€»ç»“": "SUMMARY",
        "æ±‡æ€»": "SUMMARY",
    }
    
    def __init__(self, client, workspace_dir: str = "workspace"):
        self.client = client
        self.workspace_dir = workspace_dir
        self.steps: List[WorkflowStep] = []
        self.workflow_name = ""
        self.workflow_context = ""
    
    async def load_from_document(self, file_path: str) -> Dict[str, Any]:
        """
        ä»æ–‡æ¡£åŠ è½½å·¥ä½œæµ
        æ”¯æŒ .docx, .md, .txt, .json
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == ".docx":
                content = self._read_docx(file_path)
            elif file_ext == ".md":
                content = self._read_text(file_path)
            elif file_ext == ".txt":
                content = self._read_text(file_path)
            elif file_ext == ".json":
                return self._load_json_workflow(file_path)
            else:
                return {
                    "success": False,
                    "error": f"ä¸æ”¯æŒçš„æ–‡ä»¶ç±»å‹: {file_ext}"
                }
            
            # ä½¿ç”¨LLMæå–å·¥ä½œæµæ­¥éª¤
            workflow_data = await self._extract_workflow_with_llm(content, file_path)
            
            if workflow_data.get("success"):
                self.workflow_name = workflow_data.get("name", os.path.basename(file_path))
                self.workflow_context = workflow_data.get("context", content[:500])
                
                # åˆ›å»ºå·¥ä½œæµæ­¥éª¤
                for i, step_info in enumerate(workflow_data.get("steps", []), 1):
                    step = WorkflowStep(
                        step_id=i,
                        description=step_info.get("description"),
                        step_type=step_info.get("type", "GENERAL"),
                        input_data=step_info.get("input"),
                        expected_output=step_info.get("expected_output")
                    )
                    self.steps.append(step)
            
            return workflow_data
            
        except Exception as e:
            return {
                "success": False,
                "error": f"åŠ è½½æ–‡æ¡£å¤±è´¥: {str(e)}"
            }
    
    def _read_docx(self, file_path: str) -> str:
        """è¯»å–Wordæ–‡æ¡£"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # è¯»å–è¡¨æ ¼
            for table in doc.tables:
                content.append("\n[è¡¨æ ¼]")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    content.append(" | ".join(cells))
            
            return "\n".join(content)
        except ImportError:
            raise Exception("éœ€è¦å®‰è£… python-docx: pip install python-docx")
    
    def _read_text(self, file_path: str) -> str:
        """è¯»å–æ–‡æœ¬æ–‡ä»¶"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def _load_json_workflow(self, file_path: str) -> Dict[str, Any]:
        """ä»JSONåŠ è½½é¢„å®šä¹‰çš„å·¥ä½œæµ"""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        self.workflow_name = data.get("name", "æœªå‘½åå·¥ä½œæµ")
        self.workflow_context = data.get("context", "")
        
        for i, step_data in enumerate(data.get("steps", []), 1):
            step = WorkflowStep(
                step_id=i,
                description=step_data.get("description"),
                step_type=step_data.get("type"),
                input_data=step_data.get("input"),
                expected_output=step_data.get("expected_output")
            )
            self.steps.append(step)
        
        return {"success": True, "steps_count": len(self.steps)}
    
    async def _extract_workflow_with_llm(self, content: str, file_name: str) -> Dict[str, Any]:
        """ä½¿ç”¨LLMæå–å·¥ä½œæµæ­¥éª¤"""
        
        prompt = f"""
åˆ†æä»¥ä¸‹æ–‡æ¡£å†…å®¹ï¼Œæå–å…¶ä¸­æè¿°çš„å·¥ä½œæµã€å®éªŒæ­¥éª¤æˆ–æ‰§è¡Œè®¡åˆ’ã€‚

æ–‡æ¡£åç§°: {file_name}
æ–‡æ¡£å†…å®¹:
{content[:3000]}  # é™åˆ¶é•¿åº¦

è¯·è¯†åˆ«ï¼š
1. å·¥ä½œæµåç§°å’Œç›®æ ‡
2. æŒ‰é¡ºåºåˆ—å‡ºçš„æ­¥éª¤
3. æ¯ä¸ªæ­¥éª¤çš„ç±»å‹ï¼ˆVLMè§†è§‰åˆ†æã€WEB_SEARCHæœç´¢ã€CODEä»£ç æ‰§è¡Œã€FILE_GENæ–‡ä»¶ç”Ÿæˆã€DATAæ•°æ®å¤„ç†ã€COMPAREå¯¹æ¯”ã€SUMMARYæ€»ç»“ç­‰ï¼‰
4. æ¯ä¸ªæ­¥éª¤çš„è¾“å…¥å’Œé¢„æœŸè¾“å‡º

ä»¥JSONæ ¼å¼è¿”å›ï¼š
{{
    "name": "å·¥ä½œæµåç§°",
    "context": "å·¥ä½œæµèƒŒæ™¯å’Œç›®æ ‡ï¼ˆ50å­—å†…ï¼‰",
    "steps": [
        {{
            "description": "æ­¥éª¤æè¿°",
            "type": "ä»»åŠ¡ç±»å‹ï¼ˆVLM/WEB_SEARCH/CODE/FILE_GENç­‰ï¼‰",
            "input": "è¾“å…¥è¯´æ˜",
            "expected_output": "é¢„æœŸè¾“å‡º"
        }}
    ]
}}

å¦‚æœæ–‡æ¡£ä¸åŒ…å«æ˜ç¡®çš„å·¥ä½œæµï¼Œè¿”å›ï¼š
{{
    "success": false,
    "reason": "æœªå‘ç°æ˜ç¡®çš„å·¥ä½œæµæ­¥éª¤"
}}
"""
        
        try:
            response = self.client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.3,
                    response_mime_type="application/json"
                )
            )
            
            result = json.loads(response.text)
            result["success"] = True
            return result
            
        except Exception as e:
            return {
                "success": False,
                "error": f"LLMè§£æå¤±è´¥: {str(e)}"
            }
    
    async def execute_workflow(self, task_orchestrator=None) -> Dict[str, Any]:
        """
        æ‰§è¡Œæ•´ä¸ªå·¥ä½œæµ
        
        Args:
            task_orchestrator: TaskOrchestratorå®ä¾‹ï¼ˆä»app.pyä¼ å…¥ï¼‰
        """
        print(f"\n{'='*70}")
        print(f"ğŸš€ å¼€å§‹æ‰§è¡Œå·¥ä½œæµ: {self.workflow_name}")
        print(f"{'='*70}")
        print(f"ğŸ“‹ æ€»æ­¥éª¤æ•°: {len(self.steps)}")
        print(f"ğŸ“ èƒŒæ™¯: {self.workflow_context}\n")
        
        results = {
            "workflow_name": self.workflow_name,
            "start_time": datetime.now().isoformat(),
            "steps": [],
            "overall_status": "running"
        }
        
        for step in self.steps:
            print(f"\n[æ­¥éª¤ {step.step_id}/{len(self.steps)}] {step.description}")
            print(f"â””â”€ ç±»å‹: {step.step_type}")
            print(f"   â³ æ‰§è¡Œä¸­...")
            
            step.status = "running"
            step.start_time = datetime.now()
            
            try:
                # æ‰§è¡Œæ­¥éª¤
                if task_orchestrator:
                    step_result = await self._execute_step_with_orchestrator(
                        step, task_orchestrator
                    )
                else:
                    step_result = await self._execute_step_standalone(step)
                
                step.result = step_result
                step.status = "completed"
                
                print(f"   âœ… å®Œæˆ")
                if isinstance(step_result, dict):
                    if step_result.get("output"):
                        print(f"   ğŸ“„ è¾“å‡º: {str(step_result['output'])[:100]}...")
                
            except Exception as e:
                step.status = "failed"
                step.error = str(e)
                print(f"   âŒ å¤±è´¥: {e}")
                
                # å¯é€‰ï¼šå¤±è´¥åæ˜¯å¦ç»§ç»­
                if not self._should_continue_on_error():
                    results["overall_status"] = "failed"
                    break
            
            finally:
                step.end_time = datetime.now()
                results["steps"].append(step.to_dict())
        
        # å®Œæˆ
        results["end_time"] = datetime.now().isoformat()
        if results["overall_status"] == "running":
            results["overall_status"] = "completed"
        
        # ç”Ÿæˆæ€»ç»“æŠ¥å‘Š
        results["summary"] = self._generate_summary(results)
        
        print(f"\n{'='*70}")
        print(f"ğŸ“Š å·¥ä½œæµæ‰§è¡Œå®Œæˆ")
        print(f"{'='*70}")
        print(f"âœ… æˆåŠŸæ­¥éª¤: {sum(1 for s in results['steps'] if s['status']=='completed')}/{len(self.steps)}")
        print(f"âŒ å¤±è´¥æ­¥éª¤: {sum(1 for s in results['steps'] if s['status']=='failed')}/{len(self.steps)}")
        
        return results
    
    async def _execute_step_with_orchestrator(self, step: WorkflowStep, 
                                             orchestrator) -> Dict[str, Any]:
        """ä½¿ç”¨TaskOrchestratoræ‰§è¡Œæ­¥éª¤"""
        
        # æ„å»ºä»»åŠ¡è¯·æ±‚
        task_input = step.input_data or step.description
        
        # æ ¹æ®æ­¥éª¤ç±»å‹è°ƒç”¨ä¸åŒçš„å¤„ç†æ–¹æ³•
        if step.step_type == "VLM":
            # VLMä»»åŠ¡
            return await self._execute_vlm_step(step, orchestrator)
        
        elif step.step_type == "WEB_SEARCH":
            # æœç´¢ä»»åŠ¡
            result = await orchestrator.handle_search(task_input)
            return {"output": result}
        
        elif step.step_type == "CODE":
            # ä»£ç æ‰§è¡Œ
            result = await orchestrator.handle_code_execution(task_input)
            return {"output": result}
        
        elif step.step_type == "FILE_GEN":
            # æ–‡ä»¶ç”Ÿæˆ
            result = await orchestrator.handle_file_generation(task_input)
            return {"output": result}
        
        else:
            # é€šç”¨å¤„ç†
            return await self._execute_step_standalone(step)
    
    async def _execute_vlm_step(self, step: WorkflowStep, orchestrator) -> Dict[str, Any]:
        """æ‰§è¡ŒVLMæ­¥éª¤"""
        
        # æŸ¥æ‰¾å›¾ç‰‡æ–‡ä»¶
        image_path = None
        if step.input_data and isinstance(step.input_data, str):
            if os.path.exists(step.input_data):
                image_path = step.input_data
        
        # å¦‚æœæ²¡æœ‰æŒ‡å®šå›¾ç‰‡ï¼Œå°è¯•æŸ¥æ‰¾æœ€è¿‘çš„å›¾ç‰‡
        if not image_path:
            image_path = self._find_recent_image()
        
        if not image_path:
            return {
                "success": False,
                "error": "æœªæ‰¾åˆ°å›¾ç‰‡æ–‡ä»¶"
            }
        
        # è¯»å–å›¾ç‰‡
        with open(image_path, "rb") as f:
            image_data = f.read()
        
        # è°ƒç”¨VLM
        from google.genai import types
        
        response = self.client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=[
                {
                    "mime_type": f"image/{'jpeg' if image_path.endswith('.jpg') else 'png'}",
                    "data": image_data
                },
                step.description
            ],
            config=types.GenerateContentConfig(
                temperature=0.7,
                max_output_tokens=1000
            )
        )
        
        return {
            "success": True,
            "output": response.text,
            "image": image_path
        }
    
    async def _execute_step_standalone(self, step: WorkflowStep) -> Dict[str, Any]:
        """ç‹¬ç«‹æ‰§è¡Œæ­¥éª¤ï¼ˆä¸ä¾èµ–orchestratorï¼‰"""
        
        # ç®€åŒ–ç‰ˆæ‰§è¡Œ
        if step.step_type == "VLM":
            return {"output": f"[VLM] {step.description} - éœ€è¦å›¾ç‰‡è¾“å…¥"}
        
        elif step.step_type == "WEB_SEARCH":
            return {"output": f"[æœç´¢] {step.description} - éœ€è¦æœç´¢å¼•æ“"}
        
        else:
            return {"output": f"[{step.step_type}] {step.description} - å¾…å®ç°"}
    
    def _find_recent_image(self) -> Optional[str]:
        """æŸ¥æ‰¾æœ€è¿‘çš„å›¾ç‰‡æ–‡ä»¶"""
        image_dirs = [
            "workspace/images",
            "workspace/uploads",
            "uploads",
            "."
        ]
        
        for dir_path in image_dirs:
            if not os.path.exists(dir_path):
                continue
            
            images = []
            for ext in [".jpg", ".jpeg", ".png", ".gif"]:
                images.extend(Path(dir_path).glob(f"**/*{ext}"))
            
            if images:
                # è¿”å›æœ€æ–°çš„å›¾ç‰‡
                latest = max(images, key=lambda p: p.stat().st_mtime)
                return str(latest)
        
        return None
    
    def _should_continue_on_error(self) -> bool:
        """æ­¥éª¤å¤±è´¥åæ˜¯å¦ç»§ç»­"""
        # å¯é…ç½®ç­–ç•¥
        return True  # é»˜è®¤ç»§ç»­
    
    def _generate_summary(self, results: Dict[str, Any]) -> str:
        """ç”Ÿæˆæ‰§è¡Œæ€»ç»“"""
        
        total = len(results["steps"])
        completed = sum(1 for s in results["steps"] if s["status"] == "completed")
        failed = sum(1 for s in results["steps"] if s["status"] == "failed")
        
        summary = f"""
å·¥ä½œæµæ‰§è¡Œæ€»ç»“

åç§°: {results['workflow_name']}
çŠ¶æ€: {results['overall_status']}
æ€»æ­¥éª¤: {total}
æˆåŠŸ: {completed}
å¤±è´¥: {failed}
æˆåŠŸç‡: {completed/total*100:.1f}%

è¯¦ç»†ç»“æœ:
"""
        
        for step in results["steps"]:
            status_icon = "âœ…" if step["status"] == "completed" else "âŒ"
            summary += f"\n{status_icon} æ­¥éª¤{step['step_id']}: {step['description']}"
            if step.get("error"):
                summary += f"\n   é”™è¯¯: {step['error']}"
        
        return summary
    
    async def save_results(self, results: Dict[str, Any], 
                          output_dir: str = "workspace/workflows") -> str:
        """ä¿å­˜æ‰§è¡Œç»“æœ"""
        
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"workflow_{timestamp}.json"
        output_path = os.path.join(output_dir, filename)
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        
        print(f"\nğŸ“ ç»“æœå·²ä¿å­˜: {output_path}")
        
        # åŒæ—¶ç”Ÿæˆæ–‡æœ¬æŠ¥å‘Š
        report_path = output_path.replace(".json", "_report.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(results.get("summary", ""))
        
        print(f"ğŸ“„ æŠ¥å‘Šå·²ä¿å­˜: {report_path}")
        
        return output_path


# =============== è¾…åŠ©å‡½æ•° ===============

async def execute_document_workflow(file_path: str, client, 
                                    task_orchestrator=None) -> Dict[str, Any]:
    """
    å¿«é€Ÿæ‰§è¡Œæ–‡æ¡£å·¥ä½œæµçš„ä¾¿æ·å‡½æ•°
    
    Args:
        file_path: æ–‡æ¡£è·¯å¾„
        client: Geminiå®¢æˆ·ç«¯
        task_orchestrator: TaskOrchestratorå®ä¾‹ï¼ˆå¯é€‰ï¼‰
    
    Returns:
        æ‰§è¡Œç»“æœå­—å…¸
    """
    
    executor = DocumentWorkflowExecutor(client)
    
    # åŠ è½½å·¥ä½œæµ
    load_result = await executor.load_from_document(file_path)
    
    if not load_result.get("success"):
        return load_result
    
    print(f"\nâœ… å·¥ä½œæµåŠ è½½æˆåŠŸ")
    print(f"   åç§°: {executor.workflow_name}")
    print(f"   æ­¥éª¤: {len(executor.steps)}ä¸ª")
    
    # æ‰§è¡Œå·¥ä½œæµ
    results = await executor.execute_workflow(task_orchestrator)
    
    # ä¿å­˜ç»“æœ
    output_path = await executor.save_results(results)
    
    return {
        "success": True,
        "workflow_name": executor.workflow_name,
        "steps_count": len(executor.steps),
        "results": results,
        "output_path": output_path
    }


# =============== å‘½ä»¤è¡Œæµ‹è¯• ===============

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("ç”¨æ³•: python document_workflow_executor.py <document_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # ç®€å•æµ‹è¯•ï¼ˆä¸ä½¿ç”¨orchestratorï¼‰
    async def test():
        from google import genai
        
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        client = genai.Client(api_key=api_key)
        
        result = await execute_document_workflow(file_path, client)
        
        if result.get("success"):
            print(f"\nğŸ‰ å·¥ä½œæµæ‰§è¡Œå®Œæˆï¼")
            print(f"   ç»“æœæ–‡ä»¶: {result['output_path']}")
        else:
            print(f"\nâŒ æ‰§è¡Œå¤±è´¥: {result.get('error')}")
    
    asyncio.run(test())
