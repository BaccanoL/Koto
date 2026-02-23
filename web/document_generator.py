#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Document Generator - 高质量文档生成器
- 完整保留 Markdown 格式
- 支持粗体、斜体、链接、列表、代码块
- 智能标题提取
"""
import os
import re
from datetime import datetime
from typing import List, Tuple, Optional

DEFAULT_OUTPUT_DIR = None

try:
    # Try to derive workspace/documents from project layout
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
    DEFAULT_OUTPUT_DIR = os.path.join(PROJECT_ROOT, 'workspace', 'documents')
    os.makedirs(DEFAULT_OUTPUT_DIR, exist_ok=True)
except Exception:
    DEFAULT_OUTPUT_DIR = os.getcwd()


def _split_lines(text: str) -> List[str]:
    lines = []
    for line in text.splitlines():
        lines.append(line.rstrip())
    return lines


def _is_cjk_char(ch: str) -> bool:
    return bool(re.match(r"[\u4e00-\u9fff]", ch))


def _join_text_lines(prev: str, curr: str) -> str:
    if not prev:
        return curr
    if not curr:
        return prev
    prev_end = prev[-1]
    curr_start = curr[0]
    if _is_cjk_char(prev_end) and _is_cjk_char(curr_start):
        return prev + curr
    return f"{prev} {curr}".strip()


def _normalize_markdown_lines(text: str) -> List[str]:
    raw_lines = _split_lines(text)
    normalized = []
    buffer = ""
    in_code_block = False

    def flush_buffer():
        nonlocal buffer
        if buffer.strip():
            normalized.append(buffer.strip())
        buffer = ""

    for line in raw_lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_buffer()
            normalized.append(line)
            in_code_block = not in_code_block
            continue

        if in_code_block:
            normalized.append(line)
            continue

        if not stripped:
            flush_buffer()
            normalized.append("")
            continue

        is_heading = stripped.startswith("# ") or stripped.startswith("## ") or stripped.startswith("### ") or stripped.startswith("#### ")
        is_list = stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• ") or re.match(r"^\d+[\.)]\s", stripped)
        is_quote = stripped.startswith("> ")
        is_rule = stripped in ["---", "***", "___"]
        is_table = "|" in stripped and re.search(r"\|", stripped)

        if is_heading or is_list or is_quote or is_rule or is_table:
            flush_buffer()
            normalized.append(line)
            continue

        buffer = _join_text_lines(buffer, stripped)

    flush_buffer()

    # Collapse consecutive blank lines
    compact = []
    prev_blank = False
    for line in normalized:
        is_blank = not line.strip()
        if is_blank and prev_blank:
            continue
        compact.append(line)
        prev_blank = is_blank

    return compact


def _extract_title_from_content(text: str) -> Optional[str]:
    """从内容中智能提取标题"""
    lines = text.strip().split('\n')
    for line in lines[:10]:  # 只检查前10行
        line = line.strip()
        if line.startswith('# '):
            return line[2:].strip()[:60]
        if line.startswith('## '):
            return line[3:].strip()[:60]
    return None


def _sanitize_filename(name: str) -> str:
    """清理文件名，移除非法字符"""
    # 移除非法字符
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    # 替换空格
    name = name.replace(' ', '_')
    # 限制长度
    return name[:50] if len(name) > 50 else name


def save_docx(text: str, title: str = None, output_dir: str = None, filename: str = None, 
              add_toc: bool = False, add_page_numbers: bool = False) -> str:
    """
    Save structured Markdown text to a DOCX file with proper formatting.
    
    Features:
    - Headings (#, ##, ###)
    - Bold (**text**) and Italic (*text*)
    - Bullet lists (- item) and numbered lists (1. item)
    - Nested lists (indentation preserved)
    - Code blocks (```code```)
    - Inline code (`code`)
    - Links [text](url)
    - Tables (| header | header |)
    - Images (![alt](path))
    - Table of Contents (optional)
    - Page numbers (optional)
    
    Args:
        text: Markdown text content
        title: Document title
        output_dir: Output directory
        filename: Output filename
        add_toc: Add table of contents
        add_page_numbers: Add page numbers to footer
    """
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # 设置默认字体为中文友好字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5  # 行距 1.5
    style.paragraph_format.space_after = Pt(6)  # 段后间距
    
    # 添加页眉页脚
    if add_page_numbers or title:
        section = doc.sections[0]
        
        # 添加页眉（显示标题）
        if title:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = title
            header_para.style = doc.styles['Header']
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in header_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        # 添加页脚（页码）
        if add_page_numbers:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段
            run = footer_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            
            run.font.size = Pt(9)
    
    # 自动提取标题
    if title is None:
        title = _extract_title_from_content(text)
    
    # 添加文档标题（美观的格式）
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题字体和样式
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(0)
        heading_format.space_after = Pt(12)
        heading_format.line_spacing = 1.5
        
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(18)  # 标题更大
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x78)  # 专业蓝色
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 添加目录
    if add_toc:
        # 添加目录标题
        toc_heading = doc.add_paragraph()
        toc_heading.add_run("目录").bold = True
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.space_after = Pt(12)
        
        # 添加目录字段
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run._r.append(OxmlElement('w:t'))
        
        run._r.append(fldChar3)
        
        # 分页
        doc.add_page_break()
    
    def add_formatted_paragraph(doc, text_content: str, style_name: str = None, indent_level: int = 0):
        """添加带格式的段落，支持粗体、斜体、行内代码"""
        p = doc.add_paragraph(style=style_name)
        
        # 设置缩进
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
        
        # 解析 Markdown 格式
        # 顺序: 先处理粗斜体，再处理粗体，再处理斜体，最后处理行内代码
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***粗斜体***
            (r'\*\*(.+?)\*\*', 'bold'),              # **粗体**
            (r'\*(.+?)\*', 'italic'),                # *斜体*
            (r'`([^`]+)`', 'code'),                  # `行内代码`
            (r'\[([^\]]+)\]\(([^\)]+)\)', 'link'),   # [链接](url)
        ]
        
        # 简化处理：逐段解析
        remaining = text_content
        
        while remaining:
            # 找最早出现的格式标记
            earliest_match = None
            earliest_pos = len(remaining)
            earliest_type = None
            
            for pattern, fmt_type in patterns:
                match = re.search(pattern, remaining)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    earliest_type = fmt_type
            
            if earliest_match:
                # 添加匹配前的普通文本
                if earliest_pos > 0:
                    run = p.add_run(remaining[:earliest_pos])
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # 添加格式化文本
                if earliest_type == 'bold_italic':
                    run = p.add_run(earliest_match.group(1))
                    run.bold = True
                    run.italic = True
                elif earliest_type == 'bold':
                    run = p.add_run(earliest_match.group(1))
                    run.bold = True
                elif earliest_type == 'italic':
                    run = p.add_run(earliest_match.group(1))
                    run.italic = True
                elif earliest_type == 'code':
                    run = p.add_run(earliest_match.group(1))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
                elif earliest_type == 'link':
                    link_text = earliest_match.group(1)
                    link_url = earliest_match.group(2)
                    run = p.add_run(f"{link_text} ({link_url})")
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)
                    run.underline = True
                
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                remaining = remaining[earliest_match.end():]
            else:
                # 没有更多格式标记，添加剩余文本
                if remaining:
                    run = p.add_run(remaining)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                break
        
        return p

    lines = _normalize_markdown_lines(text)
    in_code_block = False
    code_block_content = []
    skip_first_title = title is not None  # 如果已添加标题，跳过内容中的第一个标题
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_block_content:
                    code_text = '\n'.join(code_block_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    # 添加背景色效果（通过段落底纹）
                    pPr = p._element.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F5F5F5')
                    pPr.append(shd)
                    # 添加左边框 (模拟代码块样式)
                    pBdr = OxmlElement('w:pBdr')
                    left_bdr = OxmlElement('w:left')
                    left_bdr.set(qn('w:val'), 'single')
                    left_bdr.set(qn('w:sz'), '12')
                    left_bdr.set(qn('w:space'), '4')
                    left_bdr.set(qn('w:color'), '4A90D9')
                    pBdr.append(left_bdr)
                    pPr.append(pBdr)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # 空行
        if not line.strip():
            doc.add_paragraph('')
            i += 1
            continue
        
        # 标题处理
        if line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(8)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x78)  # 蓝色标题
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(8)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x78)  # 蓝色标题
        elif line.startswith('# '):
            if skip_first_title:
                skip_first_title = False
            else:
                h = doc.add_heading(line[2:], level=1)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(10)
        
        # 列表处理 - 支持多级缩进和行内格式
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('• ') or line.lstrip().startswith('* '):
            # 计算缩进级别
            indent = len(line) - len(line.lstrip())
            indent_level = indent // 2  # 每2个空格一级
            
            # 提取文本
            stripped = line.lstrip()
            text_content = stripped[2:]
            
            # 根据嵌套层级选择不同的列表样式
            if indent_level == 0:
                p = doc.add_paragraph(style='List Bullet')
            else:
                try:
                    p = doc.add_paragraph(style='List Bullet 2')
                except KeyError:
                    p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            
            # 设置列表项的间距
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.3
            
            # 处理行内Markdown格式 (粗体、斜体、行内代码)
            remaining = text_content
            patterns = [
                (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),
                (r'\*\*(.+?)\*\*', 'bold'),
                (r'\*(.+?)\*', 'italic'),
                (r'`([^`]+)`', 'code'),
            ]
            while remaining:
                earliest_match = None
                earliest_pos = len(remaining)
                earliest_type = None
                for pattern, fmt_type in patterns:
                    match = re.search(pattern, remaining)
                    if match and match.start() < earliest_pos:
                        earliest_match = match
                        earliest_pos = match.start()
                        earliest_type = fmt_type
                if earliest_match:
                    if earliest_pos > 0:
                        run = p.add_run(remaining[:earliest_pos])
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run = p.add_run(earliest_match.group(1))
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    if earliest_type == 'bold_italic':
                        run.bold = True; run.italic = True
                    elif earliest_type == 'bold':
                        run.bold = True
                    elif earliest_type == 'italic':
                        run.italic = True
                    elif earliest_type == 'code':
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
                    remaining = remaining[earliest_match.end():]
                else:
                    if remaining:
                        run = p.add_run(remaining)
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    break
        
        # 数字列表
        elif re.match(r'^\s*\d+[\.\)]\s', line):
            match = re.match(r'^(\s*)(\d+[\.\)]\s*)(.*)', line)
            if match:
                indent = len(match.group(1))
                text_content = match.group(3)
                p = doc.add_paragraph(style='List Number')
                if indent > 0:
                    p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
                run = p.add_run(text_content)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)  # 灰色引用
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Markdown表格处理
        elif '|' in line and re.search(r'\|', line):
            # 收集表格行
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1
            
            # 解析表格
            if len(table_lines) >= 2:
                # 处理表头
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                
                # 跳过分隔线（第二行）
                data_rows = []
                for tline in table_lines[2:]:
                    cells = [cell.strip() for cell in tline.split('|') if cell.strip()]
                    if cells:
                        data_rows.append(cells)
                
                # 创建表格
                if headers and data_rows:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # 设置表头
                    hdr_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        hdr_cells[idx].text = header
                        # 表头粗体
                        for paragraph in hdr_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = 'Microsoft YaHei'
                    
                    # 填充数据
                    for row_idx, row_data in enumerate(data_rows, 1):
                        row_cells = table.rows[row_idx].cells
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = cell_data
                                for paragraph in row_cells[col_idx].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Microsoft YaHei'
                
                i = j - 1  # 跳到表格结束
        
        # 图片处理 ![alt](path)
        elif line.strip().startswith('!['):
            match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line.strip())
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                
                # 处理相对路径
                if not os.path.isabs(image_path):
                    # 尝试从多个位置查找图片
                    search_paths = [
                        os.path.join(PROJECT_ROOT, 'workspace', 'images', image_path),
                        os.path.join(PROJECT_ROOT, 'assets', image_path),
                        os.path.join(os.getcwd(), image_path),
                        image_path
                    ]
                    for search_path in search_paths:
                        if os.path.exists(search_path):
                            image_path = search_path
                            break
                
                # 插入图片
                if os.path.exists(image_path):
                    try:
                        # 添加图片（自动调整大小）
                        doc.add_picture(image_path, width=Inches(5))
                        
                        # 添加图片说明
                        if alt_text:
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = caption.add_run(f"图: {alt_text}")
                            run.font.size = Pt(9)
                            run.font.italic = True
                            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                    except Exception as e:
                        print(f"⚠️ 无法插入图片 {image_path}: {e}")
                        # 添加占位文本
                        p = doc.add_paragraph(f"[图片: {alt_text or image_path}]")
                        p.paragraph_format.left_indent = Inches(0.5)
                else:
                    # 图片不存在，添加占位
                    p = doc.add_paragraph(f"[图片未找到: {image_path}]")
                    p.paragraph_format.left_indent = Inches(0.5)
        
        # 分隔线
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        # 普通段落
        else:
            p = add_formatted_paragraph(doc, line)
            if p:
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.5
        
        i += 1

    # 生成文件名
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    if filename:
        final_filename = filename if filename.endswith('.docx') else f"{filename}.docx"
    elif title:
        safe_title = _sanitize_filename(title)
        final_filename = f"{safe_title}_{ts}.docx"
    else:
        final_filename = f"document_{ts}.docx"
    
    full_path = os.path.join(output_dir, final_filename)
    doc.save(full_path)
    print(f"[DocumentGenerator] ✅ DOCX saved: {full_path}")
    return full_path


def save_pdf(text: str, title: str = None, output_dir: str = None, filename: str = None) -> str:
    """
    Save structured Markdown text to a PDF file using reportlab.
    
    Features:
    - Headings with proper sizing
    - Bold and italic text
    - Lists (bullet and numbered)
    - Code blocks with monospace font
    - Chinese font support
    """
    if output_dir is None:
        output_dir = DEFAULT_OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, ListFlowable, ListItem, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register Chinese fonts with multiple fallbacks
    font_candidates = [
        ('MSYH', 'C:/Windows/Fonts/msyh.ttf'),
        ('MSYH', 'C:/Windows/Fonts/msyh.ttc'),
        ('SimHei', 'C:/Windows/Fonts/simhei.ttf'),
        ('SimSun', 'C:/Windows/Fonts/simsun.ttc'),
    ]
    base_font = 'Helvetica'
    
    for name, path in font_candidates:
        try:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont(name, path))
                base_font = name
                print(f"[DocumentGenerator] PDF using font: {name}")
                break
        except Exception:
            continue

    styles = getSampleStyleSheet()
    
    # 定义各种样式
    styles.add(ParagraphStyle(
        name='BodyCN',
        fontName=base_font,
        fontSize=11,
        leading=18,
        wordWrap='CJK',
        spaceBefore=6,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='TitleCN',
        fontName=base_font,
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        name='Heading1CN',
        fontName=base_font,
        fontSize=16,
        leading=22,
        spaceBefore=16,
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name='Heading2CN',
        fontName=base_font,
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name='Heading3CN',
        fontName=base_font,
        fontSize=12,
        leading=18,
        spaceBefore=10,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='CodeCN',
        fontName='Courier',
        fontSize=9,
        leading=12,
        leftIndent=20,
        backColor='#f5f5f5',
    ))
    styles.add(ParagraphStyle(
        name='BulletCN',
        fontName=base_font,
        fontSize=11,
        leading=16,
        leftIndent=20,
        bulletIndent=10,
    ))
    styles.add(ParagraphStyle(
        name='QuoteCN',
        fontName=base_font,
        fontSize=11,
        leading=16,
        leftIndent=30,
        textColor='#666666',
    ))

    # 自动提取标题
    if title is None:
        title = _extract_title_from_content(text)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    if filename:
        final_filename = filename if filename.endswith('.pdf') else f"{filename}.pdf"
    elif title:
        safe_title = _sanitize_filename(title)
        final_filename = f"{safe_title}_{ts}.pdf"
    else:
        final_filename = f"document_{ts}.pdf"
    
    full_path = os.path.join(output_dir, final_filename)
    doc = SimpleDocTemplate(full_path, pagesize=A4, 
                           leftMargin=0.75*inch, rightMargin=0.75*inch,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)

    story = []
    
    # 添加标题
    if title:
        story.append(Paragraph(title, styles['TitleCN']))
        story.append(Spacer(1, 12))

    def escape_xml(text):
        """转义 XML 特殊字符"""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    def process_inline_formatting(text):
        """处理行内格式（粗体、斜体、代码、可点击链接）"""
        # 粗体 **text**
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # 斜体 *text*
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # 行内代码 `code`
        text = re.sub(r'`([^`]+)`', r'<font name="Courier" size="9" color="#880000">\1</font>', text)
        # 链接 [text](url) → 可点击超链接
        text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2" color="blue"><u>\1</u></a>', text)
        return text

    in_code_block = False
    code_block_content = []
    skip_first_title = title is not None
    
    pdf_lines = _normalize_markdown_lines(text)
    i = 0
    while i < len(pdf_lines):
        line = pdf_lines[i]
        
        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                if code_block_content:
                    code_text = '\n'.join(code_block_content)
                    story.append(Preformatted(code_text, styles['CodeCN']))
                    story.append(Spacer(1, 8))
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # 空行
        if not line.strip():
            story.append(Spacer(1, 8))
            i += 1
            continue
        
        # 标题
        if line.startswith('#### '):
            story.append(Paragraph(escape_xml(line[5:]), styles['Heading3CN']))
        elif line.startswith('### '):
            story.append(Paragraph(escape_xml(line[4:]), styles['Heading3CN']))
        elif line.startswith('## '):
            story.append(Paragraph(escape_xml(line[3:]), styles['Heading2CN']))
        elif line.startswith('# '):
            if skip_first_title:
                skip_first_title = False
            else:
                story.append(Paragraph(escape_xml(line[2:]), styles['Heading1CN']))
        
        # 列表项
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('• ') or line.lstrip().startswith('* '):
            stripped = line.lstrip()
            text_content = stripped[2:]
            formatted = process_inline_formatting(escape_xml(text_content))
            story.append(Paragraph(f"• {formatted}", styles['BulletCN']))
        
        # 数字列表
        elif re.match(r'^\s*\d+[\.\)]\s', line):
            match = re.match(r'^(\s*)(\d+[\.\)]\s*)(.*)', line)
            if match:
                num = match.group(2).strip()
                text_content = match.group(3)
                formatted = process_inline_formatting(escape_xml(text_content))
                story.append(Paragraph(f"{num} {formatted}", styles['BulletCN']))
        
        # 引用
        elif line.startswith('> '):
            formatted = process_inline_formatting(escape_xml(line[2:]))
            story.append(Paragraph(formatted, styles['QuoteCN']))
        
        # Markdown 表格
        elif '|' in line and re.search(r'\|.*\|', line):
            table_lines = [line]
            j = i + 1
            while j < len(pdf_lines) and '|' in pdf_lines[j] and pdf_lines[j].strip():
                table_lines.append(pdf_lines[j])
                j += 1
            
            if len(table_lines) >= 2:
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                data_rows = []
                for tline in table_lines[2:]:  # skip separator row
                    cells = [cell.strip() for cell in tline.split('|') if cell.strip()]
                    if cells:
                        data_rows.append(cells)
                
                if headers:
                    # 构建reportlab Table
                    col_count = len(headers)
                    table_data = [headers]
                    for row in data_rows:
                        # 确保列数一致
                        padded = row + [''] * (col_count - len(row))
                        table_data.append(padded[:col_count])
                    
                    # 计算列宽 (平均分配可用宽度)
                    avail_width = A4[0] - 1.5 * inch
                    col_widths = [avail_width / col_count] * col_count
                    
                    # 将文本转为Paragraph以支持自动换行
                    cell_style = ParagraphStyle('TableCell', parent=styles['BodyCN'], fontSize=9, leading=13)
                    cell_header_style = ParagraphStyle('TableHeader', parent=cell_style, fontName=base_font, fontSize=9, leading=13)
                    
                    formatted_data = []
                    for r_idx, row in enumerate(table_data):
                        formatted_row = []
                        for cell in row:
                            st = cell_header_style if r_idx == 0 else cell_style
                            formatted_row.append(Paragraph(escape_xml(cell), st))
                        formatted_data.append(formatted_row)
                    
                    t = Table(formatted_data, colWidths=col_widths)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#4A90D9')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), white),
                        ('FONTNAME', (0, 0), (-1, 0), base_font),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('TOPPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), HexColor('#F8F9FA')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#FFFFFF'), HexColor('#F2F2F2')]),
                        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CCCCCC')),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(t)
                    story.append(Spacer(1, 6))
            
            i = j
            continue
        
        # 分隔线
        elif line.strip() in ['---', '***', '___']:
            story.append(Paragraph('─' * 60, styles['BodyCN']))
        
        # 普通段落
        else:
            formatted = process_inline_formatting(escape_xml(line))
            story.append(Paragraph(formatted, styles['BodyCN']))
        
        i += 1

    try:
        doc.build(story)
        print(f"[DocumentGenerator] ✅ PDF saved: {full_path}")
    except Exception as e:
        print(f"[DocumentGenerator] ⚠️ PDF generation error: {e}")
        # 尝试简化版本
        try:
            simple_story = []
            if title:
                simple_story.append(Paragraph(escape_xml(title), styles['TitleCN']))
            for line in text.split('\n'):
                if line.strip():
                    simple_story.append(Paragraph(escape_xml(line), styles['BodyCN']))
            doc.build(simple_story)
            print(f"[DocumentGenerator] ✅ PDF saved (simplified): {full_path}")
        except Exception as e2:
            print(f"[DocumentGenerator] ❌ PDF generation failed: {e2}")
            raise
    
    return full_path


def to_workspace_rel(full_path: str) -> str:
    """Return relative path under workspace for UI exposure."""
    try:
        workspace_dir = os.path.join(PROJECT_ROOT, 'workspace')
        rel = os.path.relpath(full_path, workspace_dir)
        return rel.replace('\\', '/')
    except Exception:
        return os.path.basename(full_path)
