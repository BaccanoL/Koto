#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word Track Changes 修订模式实现（改进版）
在文档中插入可以 accept/decline 的修改建议

改进点：
1. 保留原有格式（粗体、斜体、颜色等）
2. 支持同一段落内多处修改
3. 更精确的文本定位
4. 详细的成功/失败统计
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree
from datetime import datetime
from typing import List, Dict, Any
import zipfile
import shutil
import tempfile
import os


class TrackChangesEditor:
    """Word Track Changes 修订编辑器（改进版）"""
    
    def __init__(self, author: str = "Koto AI"):
        self.author = author
        self.change_id = 0
    
    def apply_comment_changes(
        self,
        file_path: str,
        annotations: List[Dict[str, str]],
        progress_callback=None
    ) -> Dict[str, Any]:
        """
        以右侧批注气泡的方式标注修改建议
        
        原文保持不变，修改建议以 Word 批注(Comment)显示在右侧边栏。
        用户在 Word 中可在「审阅」里逐条查看、接受或忽略。
        
        Args:
            file_path: Word 文档路径
            annotations: 标注列表，同 apply_tracked_changes
            progress_callback: 可选的进度回调函数 callback(current, total, status, detail)
        
        Returns:
            修改统计
        """
        try:
            doc = Document(file_path)
            
            applied_count = 0
            failed_count = 0
            
            print(f"[Comments] 💬 开始添加批注...")
            print(f"[Comments] 📊 共 {len(annotations)} 条修改建议")
            
            # 预处理标注
            normalized = []
            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                modified = anno.get("修改后文本",
                          anno.get("修改建议",
                          anno.get("改为", ""))).strip()
                reason = anno.get("修改原因", anno.get("原因", "")).strip()
                
                if original and modified and original != modified:
                    normalized.append({
                        "original": original,
                        "modified": modified,
                        "reason": reason
                    })
            
            print(f"[Comments] ✅ 有效修改: {len(normalized)} 条")
            
            if not normalized:
                doc.save(file_path)
                return {"success": True, "applied": 0, "failed": 0, "total": 0}
            
            # 获取或创建 comments part
            comments_el, comments_part_ref = self._get_or_create_comments_part(doc)
            
            # 通知开始应用
            if progress_callback:
                progress_callback(0, len(normalized), "start", f"开始添加 {len(normalized)} 条批注")
            
            for idx, anno in enumerate(normalized, 1):
                original = anno["original"]
                modified = anno["modified"]
                reason = anno["reason"]
                
                # 通知当前进度
                if progress_callback:
                    progress_callback(
                        idx, len(normalized), "processing",
                        f"正在处理: {original[:30]}..."
                    )
                
                found = False
                
                # 先在正文段落中查找
                for para in doc.paragraphs:
                    if original in para.text:
                        self.change_id += 1
                        cid = self.change_id
                        
                        # 1) 在 comments.xml 里添加批注内容
                        self._add_comment_element(comments_el, cid, modified, reason)
                        
                        # 2) 在段落中标记批注范围
                        success = self._add_comment_markers_to_paragraph(
                            para, original, cid
                        )
                        if success:
                            applied_count += 1
                            found = True
                            detail_msg = f"✅ #{idx}/{len(normalized)}: '{original[:25]}...'"
                            print(f"  💬 {detail_msg}")
                            if progress_callback:
                                progress_callback(idx, len(normalized), "success", detail_msg)
                            break
                
                # 再在表格中查找
                if not found:
                    for table in doc.tables:
                        if found:
                            break
                        for row in table.rows:
                            if found:
                                break
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if original in para.text:
                                        self.change_id += 1
                                        cid = self.change_id
                                        self._add_comment_element(comments_el, cid, modified, reason)
                                        success = self._add_comment_markers_to_paragraph(
                                            para, original, cid
                                        )
                                        if success:
                                            applied_count += 1
                                            found = True
                                            detail_msg = f"✅ (表格) #{idx}/{len(normalized)}: '{original[:20]}...'"
                                            print(f"  💬 {detail_msg}")
                                            if progress_callback:
                                                progress_callback(idx, len(normalized), "success", detail_msg)
                                            break
                
                if not found:
                    failed_count += 1
                    detail_msg = f"⚠️ #{idx} 未找到: '{original[:30]}...'"
                    print(f"  {detail_msg}")
                    if progress_callback:
                        progress_callback(idx, len(normalized), "failed", detail_msg)
            
            # 将 comments XML 写回 part
            comments_bytes = etree.tostring(
                comments_el,
                xml_declaration=True,
                encoding='UTF-8',
                standalone=True
            )
            
            if comments_part_ref is not None:
                # 已有 comments part，更新内容
                comments_part_ref._blob = comments_bytes
            else:
                # 新建 comments part
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                
                COMMENTS_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
                COMMENTS_RT = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
                
                new_part = Part(
                    PackURI('/word/comments.xml'),
                    COMMENTS_CT,
                    comments_bytes,
                    doc.part.package
                )
                doc.part.relate_to(new_part, COMMENTS_RT)
            
            # 保存文档
            doc.save(file_path)
            
            success_rate = (applied_count / len(normalized) * 100) if normalized else 0
            print(f"\n[Comments] 💾 文档已保存")
            print(f"[Comments] 📊 成功: {applied_count}, 失败: {failed_count}, 成功率: {success_rate:.1f}%")
            
            return {
                "success": True,
                "applied": applied_count,
                "failed": failed_count,
                "total": len(normalized)
            }
        
        except Exception as e:
            print(f"[Comments] ❌ 错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"success": False, "error": str(e)}
    
    def _get_or_create_comments_part(self, doc):
        """获取或创建文档的 comments 部分"""
        for rel in doc.part.rels.values():
            if 'comments' in rel.reltype:
                part = rel.target_part
                el = etree.fromstring(part.blob)
                return el, part
        
        # 新建空的 comments 元素
        el = etree.fromstring(
            b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            b' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        )
        return el, None
    
    def _add_comment_element(self, comments_el, comment_id, modified, reason=""):
        """在 comments XML 里添加一条批注"""
        WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        
        comment = etree.SubElement(comments_el, qn('w:comment'))
        comment.set(qn('w:id'), str(comment_id))
        comment.set(qn('w:author'), self.author)
        comment.set(qn('w:date'), datetime.now().isoformat() + 'Z')
        comment.set(qn('w:initials'), 'K')
        
        # 第1段：建议改为
        p1 = etree.SubElement(comment, qn('w:p'))
        r1 = etree.SubElement(p1, qn('w:r'))
        # 加粗 "建议改为："
        rpr1 = etree.SubElement(r1, qn('w:rPr'))
        etree.SubElement(rpr1, qn('w:b'))
        t1 = etree.SubElement(r1, qn('w:t'))
        t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t1.text = '建议改为：'
        
        # 修改内容（不加粗）
        r1b = etree.SubElement(p1, qn('w:r'))
        t1b = etree.SubElement(r1b, qn('w:t'))
        t1b.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t1b.text = modified
        
        # 第2段：原因
        if reason:
            p2 = etree.SubElement(comment, qn('w:p'))
            r2 = etree.SubElement(p2, qn('w:r'))
            rpr2 = etree.SubElement(r2, qn('w:rPr'))
            # 灰色小字
            color = etree.SubElement(rpr2, qn('w:color'))
            color.set(qn('w:val'), '888888')
            sz = etree.SubElement(rpr2, qn('w:sz'))
            sz.set(qn('w:val'), '18')  # 9pt
            t2 = etree.SubElement(r2, qn('w:t'))
            t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t2.text = f'原因：{reason}'
    
    def _add_comment_markers_to_paragraph(self, para, original, comment_id):
        """
        在段落中为原文片段添加批注标记
        
        添加 commentRangeStart / commentRangeEnd / commentReference
        """
        try:
            p = para._element
            runs = list(p.findall(qn('w:r')))
            
            if not runs:
                return False
            
            # 构建文本 → run 映射
            text_parts = []
            run_map = []
            for run in runs:
                run_text = self._get_run_text(run)
                start = len("".join(text_parts))
                text_parts.append(run_text)
                end = len("".join(text_parts))
                run_map.append((start, end, run))
            
            full_text = "".join(text_parts)
            pos = full_text.find(original)
            if pos == -1:
                return False
            
            target_end = pos + len(original)
            
            # 找到起始和结束 run
            start_run_idx = None
            end_run_idx = None
            
            for i, (s, e, run) in enumerate(run_map):
                if start_run_idx is None and s <= pos < e:
                    start_run_idx = i
                if s < target_end <= e:
                    end_run_idx = i
                    break
            
            if start_run_idx is None:
                return False
            if end_run_idx is None:
                end_run_idx = len(run_map) - 1
            
            cid = str(comment_id)
            WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # 在起始 run 前插入 commentRangeStart
            range_start = parse_xml(
                f'<w:commentRangeStart w:id="{cid}" xmlns:w="{WNS}"/>'
            )
            start_run_el = run_map[start_run_idx][2]
            idx = list(p).index(start_run_el)
            p.insert(idx, range_start)
            
            # 在结束 run 后插入 commentRangeEnd
            range_end = parse_xml(
                f'<w:commentRangeEnd w:id="{cid}" xmlns:w="{WNS}"/>'
            )
            end_run_el = run_map[end_run_idx][2]
            idx = list(p).index(end_run_el)
            p.insert(idx + 1, range_end)
            
            # 紧跟 commentRangeEnd 后插入 commentReference run
            ref_run = parse_xml(
                f'<w:r xmlns:w="{WNS}">'
                f'  <w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
                f'  <w:commentReference w:id="{cid}"/>'
                f'</w:r>'
            )
            idx = list(p).index(range_end)
            p.insert(idx + 1, ref_run)
            
            return True
            
        except Exception as e:
            print(f"[Comments] ⚠️ 添加批注标记失败: {str(e)}")
            return False
    
    def apply_tracked_changes(
        self,
        file_path: str,
        annotations: List[Dict[str, str]],
        progress_callback=None
    ) -> Dict[str, Any]:
        """
        应用 Track Changes 修订到文档
        
        Args:
            file_path: Word 文档路径
            annotations: 标注列表，每个包含：
                - 原文片段: 要修改的文本
                - 修改后文本/修改建议/改为: 修改后的文本
            progress_callback: 可选的进度回调函数
        
        Returns:
            修改统计
        """
        try:
            doc = Document(file_path)
            
            applied_count = 0
            failed_count = 0
            
            print(f"[TrackChanges] 📝 开始应用修订...")
            print(f"[TrackChanges] 📊 共 {len(annotations)} 条修改建议")
            
            # 预处理标注：标准化字段名
            normalized = []
            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                # 支持多种字段名
                modified = anno.get("修改后文本", 
                          anno.get("修改建议", 
                          anno.get("改为", ""))).strip()
                
                if original and modified and original != modified:
                    normalized.append({"original": original, "modified": modified})
            
            print(f"[TrackChanges] ✅ 有效修改: {len(normalized)} 条")
            
            if progress_callback:
                progress_callback(0, len(normalized), "start", f"开始应用 {len(normalized)} 条修订")
            
            for idx, anno in enumerate(normalized, 1):
                original = anno["original"]
                modified = anno["modified"]
                
                if progress_callback:
                    progress_callback(
                        idx, len(normalized), "processing",
                        f"正在处理: {original[:30]}..."
                    )
                
                # 查找文本位置并应用修订
                found = False
                
                # 先在正文段落中查找
                for para in doc.paragraphs:
                    if original in para.text:
                        success = self._apply_change_to_paragraph(
                            para, original, modified
                        )
                        if success:
                            applied_count += 1
                            found = True
                            detail_msg = f"✅ #{idx}/{len(normalized)}: '{original[:25]}...' → '{modified[:25]}...'"
                            print(f"  {detail_msg}")
                            if progress_callback:
                                progress_callback(idx, len(normalized), "success", detail_msg)
                            break
                
                # 再在表格中查找
                if not found:
                    for table in doc.tables:
                        if found:
                            break
                        for row in table.rows:
                            if found:
                                break
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if original in para.text:
                                        success = self._apply_change_to_paragraph(
                                            para, original, modified
                                        )
                                        if success:
                                            applied_count += 1
                                            found = True
                                            detail_msg = f"✅ (表格) #{idx}/{len(normalized)}: '{original[:20]}...'"
                                            print(f"  {detail_msg}")
                                            if progress_callback:
                                                progress_callback(idx, len(normalized), "success", detail_msg)
                                            break
                
                if not found:
                    failed_count += 1
                    detail_msg = f"⚠️ #{idx} 未找到: '{original[:30]}...'"
                    print(f"  {detail_msg}")
                    if progress_callback:
                        progress_callback(idx, len(normalized), "failed", detail_msg)
            
            # 保存文档
            doc.save(file_path)
            
            success_rate = (applied_count / len(normalized) * 100) if normalized else 0
            print(f"\n[TrackChanges] 💾 文档已保存")
            print(f"[TrackChanges] 📊 成功: {applied_count}, 失败: {failed_count}, 成功率: {success_rate:.1f}%")
            
            return {
                "success": True,
                "applied": applied_count,
                "failed": failed_count,
                "total": len(normalized)
            }
            
        except Exception as e:
            print(f"[TrackChanges] ❌ 错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e)
            }
    
    def _apply_change_to_paragraph(
        self,
        para,
        original: str,
        modified: str
    ) -> bool:
        """
        在段落中应用 Track Changes 修订
        
        策略：尽可能保留格式，精确定位文本
        """
        try:
            p = para._element
            runs = list(p.findall(qn('w:r')))
            
            if not runs:
                return False
            
            # 构建完整文本和 run 映射
            text_parts = []
            run_map = []  # [(start_pos, end_pos, run_element)]
            
            for run in runs:
                run_text = self._get_run_text(run)
                start = len("".join(text_parts))
                text_parts.append(run_text)
                end = len("".join(text_parts))
                run_map.append((start, end, run))
            
            full_text = "".join(text_parts)
            
            # 查找目标文本
            pos = full_text.find(original)
            if pos == -1:
                return False
            
            target_end = pos + len(original)
            
            # 找到涉及的 run
            start_run = None
            end_run = None
            start_offset = 0
            end_offset = 0
            
            for i, (s, e, run) in enumerate(run_map):
                if start_run is None and s <= pos < e:
                    start_run = i
                    start_offset = pos - s
                if s < target_end <= e:
                    end_run = i
                    end_offset = target_end - s
                    break
            
            if start_run is None:
                return False
            if end_run is None:
                end_run = len(run_map) - 1
                end_offset = len(self._get_run_text(run_map[end_run][2]))
            
            # 生成修订 ID
            self.change_id += 1
            del_id = str(self.change_id)
            self.change_id += 1
            ins_id = str(self.change_id)
            date_str = datetime.now().isoformat()
            
            # 单 run 内的修改（最常见情况）
            if start_run == end_run:
                run = run_map[start_run][2]
                run_text = self._get_run_text(run)
                
                before = run_text[:start_offset]
                target = run_text[start_offset:end_offset]
                after = run_text[end_offset:]
                
                # 获取格式
                rPr = run.find(qn('w:rPr'))
                rPr_xml = self._clone_rPr(rPr)
                
                # 构建新元素
                new_elements = []
                
                if before:
                    new_elements.append(self._make_run(before, rPr_xml))
                
                # 删除标记
                new_elements.append(parse_xml(
                    f'''<w:del w:id="{del_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r>{rPr_xml}<w:delText xml:space="preserve">{self._esc(target)}</w:delText></w:r>
                    </w:del>'''
                ))
                
                # 插入标记
                new_elements.append(parse_xml(
                    f'''<w:ins w:id="{ins_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r>{rPr_xml}<w:t xml:space="preserve">{self._esc(modified)}</w:t></w:r>
                    </w:ins>'''
                ))
                
                if after:
                    new_elements.append(self._make_run(after, rPr_xml))
                
                # 替换原 run
                idx = list(p).index(run)
                p.remove(run)
                for i, elem in enumerate(new_elements):
                    p.insert(idx + i, elem)
                
                return True
            
            else:
                # 跨多个 run：简化处理
                para_text = para.text
                parts = para_text.split(original, 1)
                
                # 清空段落
                for run in list(p.findall(qn('w:r'))):
                    p.remove(run)
                
                if parts[0]:
                    p.append(self._make_run(parts[0], ""))
                
                p.append(parse_xml(
                    f'''<w:del w:id="{del_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r><w:delText xml:space="preserve">{self._esc(original)}</w:delText></w:r>
                    </w:del>'''
                ))
                
                p.append(parse_xml(
                    f'''<w:ins w:id="{ins_id}" w:author="{self._esc(self.author)}" w:date="{date_str}"
                        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:r><w:t xml:space="preserve">{self._esc(modified)}</w:t></w:r>
                    </w:ins>'''
                ))
                
                if len(parts) > 1 and parts[1]:
                    p.append(self._make_run(parts[1], ""))
                
                return True
            
        except Exception as e:
            print(f"[TrackChanges] ⚠️ 段落修订失败: {str(e)}")
            return False
    
    def _get_run_text(self, run) -> str:
        """获取 run 中的文本"""
        parts = []
        for t in run.findall(qn('w:t')):
            if t.text:
                parts.append(t.text)
        return "".join(parts)
    
    def _clone_rPr(self, rPr) -> str:
        """克隆格式属性"""
        if rPr is None:
            return ""
        from lxml import etree
        return etree.tostring(rPr, encoding='unicode')
    
    def _make_run(self, text: str, rPr_xml: str):
        """创建 run 元素"""
        return parse_xml(
            f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                {rPr_xml}<w:t xml:space="preserve">{self._esc(text)}</w:t>
            </w:r>'''
        )
    
    @staticmethod
    def _esc(text: str) -> str:
        """转义 XML"""
        if not text:
            return ""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&apos;'))
    
    def apply_hybrid_changes(
        self,
        file_path: str,
        annotations: List[Dict[str, str]],
        progress_callback=None
    ) -> Dict[str, Any]:
        """
        混合应用两种标注方式：
        1. 精确的短文本修改 → Track Changes（修订标记）
        2. 大段落方向建议 → Comments（批注气泡）
        
        自动判断规则：
        - 原文 <= 30字且有精确替换文本 → Track Changes
        - 原文 > 30字或只有方向建议 → Comment
        
        Args:
            file_path: Word文档路径
            annotations: 标注列表
            progress_callback: 进度回调
        
        Returns:
            {"success": True, "tracked": 15, "commented": 8, "failed": 2}
        """
        try:
            doc = Document(file_path)
            
            # 分类标注
            track_changes_items = []  # 精确修改
            comment_items = []         # 方向建议
            
            for anno in annotations:
                original = anno.get("原文片段", anno.get("原文", "")).strip()
                modified = anno.get("修改后文本", anno.get("修改建议", anno.get("改为", ""))).strip()
                reason = anno.get("修改原因", anno.get("原因", "")).strip()
                
                if not original:
                    continue
                
                # 判断标注类型
                # 策略调整：只要提供了具体修改文本且未显式标记为"建议"，均视为修订(Track Changes)
                # 放宽长度限制，允许整句重写
                is_suggestion = (
                    modified.startswith("建议") or
                    modified.startswith("批注") or
                    "建议：" in modified or 
                    "原因：" in modified  # 某些情况下AI可能会把原因混入
                )
                
                is_precise = (
                    modified and             # 有替换文本
                    modified != original and # 不是重复
                    not is_suggestion and    # 不是建议
                    len(original) < 500      # 长度安全限制，防止整页替换
                )
                
                if is_precise:
                    # 精确修改 → Track Changes
                    track_changes_items.append({
                        "original": original,
                        "modified": modified,
                        "reason": reason
                    })
                else:
                    # 方向建议 → Comment
                    comment_items.append({
                        "original": original,
                        "modified": modified,
                        "reason": reason
                    })
            
            print(f"\n[Hybrid] 🎯 混合标注模式")
            print(f"[Hybrid] ✏️  精确修改: {len(track_changes_items)} 条（Track Changes）")
            print(f"[Hybrid] 💬 方向建议: {len(comment_items)} 条（Comments）")
            
            # 先应用 Track Changes
            tracked_success = 0
            tracked_failed = 0
            
            if track_changes_items:
                print(f"\n[Hybrid] 📝 第1步：应用精确修改...")
                
                for idx, item in enumerate(track_changes_items, 1):
                    if progress_callback:
                        progress_callback(
                            idx, 
                            len(track_changes_items) + len(comment_items),
                            "tracking",
                            f"修订标记: {item['original'][:20]}..."
                        )
                    
                    success = self._apply_single_track_change(
                        doc, 
                        item['original'],
                        item['modified'],
                        item['reason']
                    )
                    
                    if success:
                        tracked_success += 1
                    else:
                        tracked_failed += 1
                        print(f"[Hybrid] ⚠️  修订失败: {item['original'][:30]}...")
            
            # 再应用 Comments
            commented_success = 0
            commented_failed = 0
            comments_el = None
            comments_part_ref = None
            
            if comment_items:
                print(f"\n[Hybrid] 💬 第2步：添加批注建议...")
                
                # 获取 comments part
                comments_el, comments_part_ref = self._get_or_create_comments_part(doc)
                
                for idx, item in enumerate(comment_items, 1):
                    if progress_callback:
                        progress_callback(
                            len(track_changes_items) + idx,
                            len(track_changes_items) + len(comment_items),
                            "commenting",
                            f"批注建议: {item['original'][:20]}..."
                        )
                    
                    success = self._apply_single_comment(
                        doc,
                        comments_el,
                        item['original'],
                        item['modified'],
                        item['reason']
                    )
                    
                    if success:
                        commented_success += 1
                    else:
                        commented_failed += 1
                        print(f"[Hybrid] ⚠️  批注失败: {item['original'][:30]}...")
            
            # 将 comments XML 写入文档 Part（python-docx OPC 方式）
            if commented_success > 0 and comments_el is not None:
                comments_bytes = etree.tostring(
                    comments_el,
                    xml_declaration=True,
                    encoding='UTF-8',
                    standalone=True
                )
                
                if comments_part_ref is not None:
                    # 已有 comments part，更新内容
                    comments_part_ref._blob = comments_bytes
                else:
                    # 新建 comments part（与非混合模式相同的方式）
                    from docx.opc.part import Part
                    from docx.opc.packuri import PackURI
                    
                    COMMENTS_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
                    COMMENTS_RT = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
                    
                    new_part = Part(
                        PackURI('/word/comments.xml'),
                        COMMENTS_CT,
                        comments_bytes,
                        doc.part.package
                    )
                    doc.part.relate_to(new_part, COMMENTS_RT)
            
            # 保存文档
            doc.save(file_path)
            
            total_success = tracked_success + commented_success
            total_failed = tracked_failed + commented_failed
            
            print(f"\n[Hybrid] ✅ 完成！")
            print(f"[Hybrid] 📊 修订标记: {tracked_success}成功 / {tracked_failed}失败")
            print(f"[Hybrid] 📊 批注建议: {commented_success}成功 / {commented_failed}失败")
            print(f"[Hybrid] 📊 总计: {total_success}成功 / {total_failed}失败\n")
            
            return {
                "success": True,
                "tracked": tracked_success,
                "commented": commented_success,
                "failed": total_failed,
                "total": len(annotations),
                "applied": total_success
            }
            
        except Exception as e:
            print(f"[Hybrid] ❌ 混合标注失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "tracked": 0,
                "commented": 0,
                "failed": len(annotations),
                "total": len(annotations),
                "applied": 0
            }
    
    def _apply_single_track_change(
        self,
        doc: Document,
        original_text: str,
        modified_text: str,
        reason: str = ""
    ) -> bool:
        """应用单个修订标记（保留格式版）"""
        try:
            from copy import deepcopy
            
            # 收集所有段落（正文 + 表格）
            all_paragraphs = list(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_paragraphs.extend(cell.paragraphs)
            
            for para in all_paragraphs:
                if original_text not in para.text:
                    continue
                    
                # 找到匹配的段落
                # 1. 映射 run
                run_map = []
                current_pos = 0
                full_text_parts = []
                for i, run in enumerate(para.runs):
                    text = run.text
                    run_map.append({
                        "start": current_pos,
                        "end": current_pos + len(text),
                        "run": run,
                        "index": i
                    })
                    full_text_parts.append(text)
                    current_pos += len(text)
                
                full_text = "".join(full_text_parts)
                start_idx = full_text.find(original_text)
                if start_idx == -1:
                    continue
                end_idx = start_idx + len(original_text)
                
                # 2. 找到涉及的 runs
                target_runs = []
                start_run_info = None
                end_run_info = None
                
                for info in run_map:
                    # 如果 run 与 [start_idx, end_idx] 有交集
                    if max(start_idx, info["start"]) < min(end_idx, info["end"]):
                        target_runs.append(info)
                        if info["start"] <= start_idx < info["end"]:
                            start_run_info = info
                        if info["start"] < end_idx <= info["end"]:
                            end_run_info = info
                
                if not target_runs:
                    continue

                # 3. 准备修改
                runs_to_move = []
                parent = para._element
                
                # Helper: Split run content
                def split_run_element(run, keep_start, keep_end):
                    """返回 (elem_to_keep_before, elem_to_move, elem_to_keep_after)"""
                    # 这比较复杂，简单处理：
                    # 对于被完全包含的 run，直接移动
                    # 对于部分包含的 run，修改原 run 并克隆出移动部分
                    pass

                # 逻辑简化：
                # 我们不再尝试完美分割，因为太复杂。
                # 采用如下策略：
                # 1. 对涉及的 runs，如果是首尾 run，且只涉及部分：
                #    - 修改原 run text 为剩余部分
                #    - 克隆一个新 run 包含被删除部分（保持格式）
                # 2. 中间的 runs 直接整个移动
                
                # 处理 Start Run
                processed_start_elem = None
                
                s_info = target_runs[0] # start run
                s_run = s_info["run"]
                s_offset = start_idx - s_info["start"]
                
                # 处理 End Run (可能是同一个)
                e_info = target_runs[-1]
                e_run = e_info["run"]
                e_offset = end_idx - e_info["start"]
                
                current_time = datetime.now().isoformat()
                self.change_id += 1
                cid = str(self.change_id)

                # -------------------------------------------------
                # 情况 A: 单个 run 内修改
                # -------------------------------------------------
                if s_info["index"] == e_info["index"]:
                    original_run_text = s_run.text
                    prefix = original_run_text[:s_offset]
                    middle = original_run_text[s_offset:e_offset]
                    suffix = original_run_text[e_offset:]
                    
                    # 1. 修改原 run 为 prefix
                    s_run.text = prefix
                    
                    # 2. 如果 prefix 为空，原 run 变空（可能会被 Word 清理，但为了插入点保留它）
                    # 插入点：s_run 之后
                    insert_point = s_run._element
                    
                    # 3. 创建 middle run (将被删除的)
                    middle_run_elem = deepcopy(s_run._element)
                    t = middle_run_elem.find(qn('w:t'))
                    if t is None: t = etree.SubElement(middle_run_elem, qn('w:t'))
                    t.text = middle
                    
                    # 4. 创建 suffix run
                    if suffix:
                        suffix_run_elem = deepcopy(s_run._element)
                        t = suffix_run_elem.find(qn('w:t'))
                        if t is None: t = etree.SubElement(suffix_run_elem, qn('w:t'))
                        t.text = suffix
                    else:
                        suffix_run_elem = None
                        
                    runs_to_move.append(middle_run_elem)
                    
                    # 插入 suffix
                    if suffix_run_elem is not None:
                        parent.insert(parent.index(insert_point) + 1, suffix_run_elem)
                        
                # -------------------------------------------------
                # 情况 B: 跨越多个 runs
                # -------------------------------------------------
                else:
                    # --- Start Run ---
                    s_text = s_run.text
                    s_prefix = s_text[:s_offset]
                    s_del = s_text[s_offset:]
                    
                    s_run.text = s_prefix
                    insert_point = s_run._element
                    
                    s_del_elem = deepcopy(s_run._element)
                    t = s_del_elem.find(qn('w:t'))
                    if t is None: t = etree.SubElement(s_del_elem, qn('w:t'))
                    t.text = s_del
                    runs_to_move.append(s_del_elem)
                    
                    # --- Middle Runs ---
                    for info in target_runs[1:-1]:
                        # 移动这些 run (先从文档移除，后续加到 del)
                        r_elem = info["run"]._element
                        parent.remove(r_elem)
                        runs_to_move.append(r_elem)
                        
                    # --- End Run ---
                    e_text = e_run.text
                    e_del = e_text[:e_offset]
                    e_suffix = e_text[e_offset:]
                    
                    # 这里稍微特别：end_run 应该保留 e_suffix，而 e_del 被移走
                    # 我们修改 end_run 为 e_suffix
                    # 并克隆一个 e_del_elem
                    
                    e_del_elem = deepcopy(e_run._element)
                    t = e_del_elem.find(qn('w:t'))
                    if t is None: t = etree.SubElement(e_del_elem, qn('w:t'))
                    t.text = e_del
                    runs_to_move.append(e_del_elem)
                    
                    e_run.text = e_suffix
                
                # =================================================
                # 构建 <w:del> 和 <w:ins>
                # =================================================
                
                # 找到插入位置：insert_point 之后
                base_idx = parent.index(insert_point)
                
                # 1. 构建 <w:del>
                del_el = parse_xml(f'<w:del w:id="{cid}" w:author="{self.author}" w:date="{current_time}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                
                for r in runs_to_move:
                    # 转换 w:t -> w:delText
                    for t in r.findall(qn('w:t')):
                        t.tag = qn('w:delText')
                    del_el.append(r)
                
                parent.insert(base_idx + 1, del_el)
                
                # 2. 构建 <w:ins>
                self.change_id += 1
                ins_el = parse_xml(f'''
                <w:ins w:id="{str(self.change_id)}" w:author="{self.author}" w:date="{current_time}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:r>
                        <w:t>{self._esc(modified_text)}</w:t>
                    </w:r>
                </w:ins>
                ''')
                
                parent.insert(base_idx + 2, ins_el)
                
                return True
                
            return False
        except Exception as e:
            print(f"[TrackChange] Error: {e}")
            import traceback
            traceback.print_exc()
            return False

    
    def _inject_comments_to_docx(self, file_path: str, comments_el) -> bool:
        """
        将 comments_el (lxml Element) 注入到 docx zip 包中。
        python-docx 默认不会保存手动创建的 comments part，
        所以需要在 doc.save() 之后，操作 zip 文件来注入:
        1. word/comments.xml — 批注内容
        2. [Content_Types].xml — 添加 comments 的 content type
        3. word/_rels/document.xml.rels — 添加批注的关系
        """
        try:
            # 序列化 comments XML
            comments_xml = etree.tostring(comments_el, xml_declaration=True, 
                                          encoding='UTF-8', standalone=True)
            
            # 使用临时文件来安全修改 zip
            tmp_path = file_path + '.tmp'
            
            with zipfile.ZipFile(file_path, 'r') as zin:
                with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        
                        if item.filename == '[Content_Types].xml':
                            data = self._add_comments_content_type(data)
                        elif item.filename == 'word/_rels/document.xml.rels':
                            data = self._add_comments_relationship(data)
                        
                        zout.writestr(item, data)
                    
                    # 添加 word/comments.xml
                    zout.writestr('word/comments.xml', comments_xml)
            
            # 替换原文件
            shutil.move(tmp_path, file_path)
            print(f"[Hybrid] 💾 comments.xml 已注入 ({len(comments_xml)} bytes)")
            return True
            
        except Exception as e:
            print(f"[Hybrid] ❌ 注入 comments.xml 失败: {e}")
            tmp_path = file_path + '.tmp'
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            return False
    
    def _add_comments_content_type(self, content_types_data: bytes) -> bytes:
        """在 [Content_Types].xml 中添加 comments 的 Override"""
        try:
            root = etree.fromstring(content_types_data)
            ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
            
            for override in root.findall(f'{{{ns}}}Override'):
                if override.get('PartName') == '/word/comments.xml':
                    return content_types_data
            
            override = etree.SubElement(root, f'{{{ns}}}Override')
            override.set('PartName', '/word/comments.xml')
            override.set('ContentType', 
                         'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
            
            return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        except Exception as e:
            print(f"[Hybrid] ⚠️ Content_Types 修改失败: {e}")
            return content_types_data
    
    def _add_comments_relationship(self, rels_data: bytes) -> bytes:
        """在 document.xml.rels 中添加 comments 关系"""
        try:
            root = etree.fromstring(rels_data)
            ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
            
            for rel in root.findall(f'{{{ns}}}Relationship'):
                if 'comments' in rel.get('Type', '').lower():
                    return rels_data
            
            existing_ids = [rel.get('Id', '') for rel in root.findall(f'{{{ns}}}Relationship')]
            max_id = 0
            for rid in existing_ids:
                if rid.startswith('rId'):
                    try:
                        max_id = max(max_id, int(rid[3:]))
                    except ValueError:
                        pass
            new_id = f'rId{max_id + 1}'
            
            rel = etree.SubElement(root, f'{{{ns}}}Relationship')
            rel.set('Id', new_id)
            rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
            rel.set('Target', 'comments.xml')
            
            return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        except Exception as e:
            print(f"[Hybrid] ⚠️ Rels 修改失败: {e}")
            return rels_data
    
    def _apply_single_comment(
        self,
        doc: Document,
        comments_el,
        original_text: str,
        suggestion_text: str,
        reason: str = ""
    ) -> bool:
        """应用单个批注（内部方法）"""
        try:
            for para in doc.paragraphs:
                if original_text in para.text:
                    self.change_id += 1
                    comment_id = str(self.change_id)
                    
                    # 构建批注内容
                    comment_content = suggestion_text
                    if reason:
                        comment_content = f"{suggestion_text}\n\n原因：{reason}"
                    
                    # 创建批注元素
                    comment_xml = f'''
                    <w:comment w:id="{comment_id}" w:author="{self.author}" w:date="{datetime.now().isoformat()}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:p>
                            <w:pPr>
                                <w:pStyle w:val="CommentText"/>
                            </w:pPr>
                            <w:r>
                                <w:t>{self._esc(comment_content)}</w:t>
                            </w:r>
                        </w:p>
                    </w:comment>
                    '''
                    comments_el.append(parse_xml(comment_xml))
                    
                    # 在段落中标记批注范围
                    pos = para.text.index(original_text)
                    
                    # 分割 runs
                    accumulated = 0
                    start_run_idx = -1
                    start_offset = 0
                    end_run_idx = -1
                    end_offset = 0
                    
                    for idx, run in enumerate(para.runs):
                        run_len = len(run.text)
                        
                        if start_run_idx == -1 and accumulated + run_len > pos:
                            start_run_idx = idx
                            start_offset = pos - accumulated
                        
                        if accumulated + run_len >= pos + len(original_text):
                            end_run_idx = idx
                            end_offset = pos + len(original_text) - accumulated
                            break
                        
                        accumulated += run_len
                    
                    if start_run_idx >= 0:
                        # 在起始位置插入 commentRangeStart
                        start_marker = parse_xml(
                            f'<w:commentRangeStart w:id="{comment_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        )
                        para.runs[start_run_idx]._element.addprevious(start_marker)
                        
                        # 在结束位置插入 commentRangeEnd
                        if end_run_idx >= 0:
                            end_marker = parse_xml(
                                f'<w:commentRangeEnd w:id="{comment_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                            )
                            para.runs[end_run_idx]._element.addnext(end_marker)
                            
                            # 插入 comment reference
                            ref_run_xml = f'''
                            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:rPr>
                                    <w:rStyle w:val="CommentReference"/>
                                </w:rPr>
                                <w:commentReference w:id="{comment_id}"/>
                            </w:r>
                            '''
                            para._element.append(parse_xml(ref_run_xml))
                            
                            return True
            
            return False
            
        except Exception as e:
            print(f"[Comment] 单条批注失败: {e}")
            return False


if __name__ == "__main__":
    print("Track Changes Editor 已准备就绪")
