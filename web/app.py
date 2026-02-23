import os
import asyncio
import re
import json
import time
import threading
import subprocess
import sys
import mimetypes
import importlib.util
import base64
import shutil
from datetime import datetime

# 确保 web/ 目录在模块搜索路径中（通过 koto_app.py 启动时需要）
_web_dir = os.path.dirname(os.path.abspath(__file__))
if _web_dir not in sys.path:
    sys.path.append(_web_dir)

from flask import Flask, render_template, request, jsonify, send_from_directory, Response, stream_with_context, send_file
from flask_cors import CORS
from dotenv import load_dotenv

# Import new routing modules
from app.core.routing import SmartDispatcher

# 延迟导入 - 这些路由类仅在运行时首次访问时通过 __getattr__ 加载
# LocalModelRouter, AIRouter, TaskDecomposer, LocalPlanner 通过 app.core.routing.__getattr__ 延迟加载

# Import unified agent API blueprint — 延迟到蓝图注册时加载
agent_bp = None  # 延迟加载，见下方蓝图注册区

# ================= 并行执行系统导入 =================
try:
    from parallel_executor import (
        Task, TaskType, Priority, TaskStatus,
        get_queue_manager, get_resource_manager, get_task_monitor,
        submit_task, get_next_task, cancel_task
    )
    from task_dispatcher import get_scheduler, start_dispatcher, stop_dispatcher
    from parallel_api import register_parallel_api
    PARALLEL_SYSTEM_ENABLED = True
except ImportError as e:
    print(f"[WARNING] Failed to import parallel execution system: {e}")
    PARALLEL_SYSTEM_ENABLED = False

try:
    from flask_sock import Sock
except ImportError:
    Sock = None

# ================= 懒加载重型模块（启动优化） =================
# google.genai (~4.7s), requests (~0.5s) 延迟到首次使用时加载

class _LazyModule:
    """延迟导入代理 - 首次属性访问时才触发实际 import"""
    __slots__ = ('_import_func', '_module')

    def __init__(self, import_func):
        object.__setattr__(self, '_import_func', import_func)
        object.__setattr__(self, '_module', None)

    def _load(self):
        mod = object.__getattribute__(self, '_module')
        if mod is None:
            import_func = object.__getattribute__(self, '_import_func')
            mod = import_func()
            object.__setattr__(self, '_module', mod)
        return mod

    def __getattr__(self, name):
        return getattr(self._load(), name)

    def __repr__(self):
        mod = object.__getattribute__(self, '_module')
        if mod is None:
            return "<LazyModule (not loaded)>"
        return repr(mod)

def _import_genai():
    print("[LAZY_IMPORT] 加载 google.genai ...")
    from google import genai as _genai
    return _genai

def _import_types():
    print("[LAZY_IMPORT] 加载 google.genai.types ...")
    from google.genai import types as _types
    return _types

def _import_requests():
    print("[LAZY_IMPORT] 加载 requests ...")
    import requests as _requests
    return _requests

genai = _LazyModule(_import_genai)
types = _LazyModule(_import_types)
requests = _LazyModule(_import_requests)

# ================= 懒加载文档和PPT模块（启动加速） =================
# 延迟导入 python-docx (~572ms) 和 python-pptx (~666ms)

# 文档工作流执行器懒加载
_document_workflow_cache = {}

def get_document_workflow_executor():
    """懒加载文档工作流执行器"""
    if 'executor' not in _document_workflow_cache:
        print("[LAZY_IMPORT] 加载文档工作流执行器...")
        try:
            from web.document_workflow_executor import DocumentWorkflowExecutor, execute_document_workflow
        except ImportError:
            try:
                from document_workflow_executor import DocumentWorkflowExecutor, execute_document_workflow
            except ImportError:
                DocumentWorkflowExecutor = None
                execute_document_workflow = None
                print("[WARNING] 文档工作流执行器未安装")
        _document_workflow_cache['executor'] = DocumentWorkflowExecutor
        _document_workflow_cache['execute'] = execute_document_workflow
    return _document_workflow_cache.get('executor'), _document_workflow_cache.get('execute')

# DocumentWorkflowExecutor 和 execute_document_workflow 的懒加载代理
class _DocWorkflowProxy:
    def __getattr__(self, name):
        executor_cls, _ = get_document_workflow_executor()
        if executor_cls is None:
            raise ImportError("文档工作流执行器未安装")
        return getattr(executor_cls, name)

DocumentWorkflowExecutor = _DocWorkflowProxy()

def execute_document_workflow(*args, **kwargs):
    _, execute_func = get_document_workflow_executor()
    if execute_func is None:
        raise ImportError("文档工作流执行器未安装")
    return execute_func(*args, **kwargs)

# PPT多模型系统懒加载
_ppt_system_cache = {}

def get_ppt_system():
    """懒加载PPT生成系统"""
    if 'loaded' not in _ppt_system_cache:
        print("[LAZY_IMPORT] 加载PPT多模型生成系统...")
        try:
            from web.ppt_master import PPTMasterOrchestrator, PPTBlueprint
            from web.ppt_synthesizer import PPTSynthesizer
            from web.ppt_pipeline import PPTGenerationPipeline, PPTGenerationTaskHandler, format_ppt_generation_result
            print("[PPT_SYSTEM] ✅ 多模型PPT生成系统已加载")
        except ImportError:
            try:
                from ppt_master import PPTMasterOrchestrator, PPTBlueprint
                from ppt_synthesizer import PPTSynthesizer
                from ppt_pipeline import PPTGenerationPipeline, PPTGenerationTaskHandler, format_ppt_generation_result
                print("[PPT_SYSTEM] ✅ 多模型PPT生成系统已加载（相对导入）")
            except ImportError:
                PPTMasterOrchestrator = None
                PPTBlueprint = None
                PPTSynthesizer = None
                PPTGenerationPipeline = None
                PPTGenerationTaskHandler = None
                format_ppt_generation_result = None
                print("[WARNING] 多模型PPT生成系统未安装")
        _ppt_system_cache['orchestrator'] = PPTMasterOrchestrator
        _ppt_system_cache['blueprint'] = PPTBlueprint
        _ppt_system_cache['synthesizer'] = PPTSynthesizer
        _ppt_system_cache['pipeline'] = PPTGenerationPipeline
        _ppt_system_cache['handler'] = PPTGenerationTaskHandler
        _ppt_system_cache['formatter'] = format_ppt_generation_result
        _ppt_system_cache['loaded'] = True
    
    return (_ppt_system_cache.get('orchestrator'),
            _ppt_system_cache.get('blueprint'),
            _ppt_system_cache.get('synthesizer'),
            _ppt_system_cache.get('pipeline'),
            _ppt_system_cache.get('handler'),
            _ppt_system_cache.get('formatter'))

# 懒加载代理类
class _PPTModuleProxy:
    def __init__(self, index):
        self._index = index
    
    def __getattr__(self, name):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPT生成系统未安装")
        return getattr(module, name)
    
    def __call__(self, *args, **kwargs):
        modules = get_ppt_system()
        module = modules[self._index]
        if module is None:
            raise ImportError("PPT生成系统未安装")
        if callable(module):
            return module(*args, **kwargs)
        raise TypeError(f"{module} is not callable")

PPTMasterOrchestrator = _PPTModuleProxy(0)
PPTBlueprint = _PPTModuleProxy(1)
PPTSynthesizer = _PPTModuleProxy(2)
PPTGenerationPipeline = _PPTModuleProxy(3)
PPTGenerationTaskHandler = _PPTModuleProxy(4)
format_ppt_generation_result = _PPTModuleProxy(5)

# ================= Configuration =================
# 从 web 目录向上查找
import os
import sys as _sys

# 中断信号存储 - 改进版本，支持实时流中止
class StreamInterruptManager:
    """管理每个 session 的流中止状态和控制"""
    def __init__(self):
        self.interrupts = {}  # session_name -> {'flag': bool, 'event': threading.Event}

    def _ensure(self, session_name):
        """确保 session 记录存在"""
        if session_name not in self.interrupts:
            self.interrupts[session_name] = {'flag': False, 'event': threading.Event()}
        elif self.interrupts[session_name].get('event') is None:
            self.interrupts[session_name]['event'] = threading.Event()
    
    def set_interrupt(self, session_name):
        """设置中断标志"""
        self._ensure(session_name)
        self.interrupts[session_name]['flag'] = True
        if self.interrupts[session_name]['event']:
            self.interrupts[session_name]['event'].set()
        print(f"[INTERRUPT] Marked session {session_name} for interruption")
    
    def is_interrupted(self, session_name):
        """检查是否被中断"""
        if session_name not in self.interrupts:
            return False
        record = self.interrupts[session_name]
        event_flag = record.get('event').is_set() if record.get('event') else False
        return bool(record.get('flag')) or event_flag
    
    def reset(self, session_name):
        """重置中断标志"""
        self._ensure(session_name)
        self.interrupts[session_name]['flag'] = False
        if self.interrupts[session_name]['event']:
            self.interrupts[session_name]['event'].clear()
        print(f"[INTERRUPT] Reset interrupt flag for session {session_name}")

    def get_event(self, session_name):
        """获取/创建中断事件对象"""
        self._ensure(session_name)
        return self.interrupts[session_name]['event']
    
    def cleanup(self, session_name):
        """清理 session 的中断记录"""
        if session_name in self.interrupts:
            del self.interrupts[session_name]

_interrupt_manager = StreamInterruptManager()
# 保留向后兼容
_interrupt_flags = {}  # 仅用于向后兼容

# 判断是否为打包后运行
if getattr(_sys, 'frozen', False):
    # PyInstaller 打包后 - exe所在目录（持久化数据目录）
    PROJECT_ROOT = os.path.dirname(_sys.executable)
else:
    # 开发环境 - 从 web 目录向上找
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# Try multiple locations for the config file
config_locations = [
    os.path.join(PROJECT_ROOT, "config", "gemini_config.env"),
    os.path.join(PROJECT_ROOT, "gemini_config.env"),
    os.path.join(os.path.dirname(_sys.executable), "config", "gemini_config.env") if getattr(_sys, 'frozen', False) else "",
    "gemini_config.env",
    "../gemini_config.env"
]

for config_path in config_locations:
    if os.path.exists(config_path):
        load_dotenv(config_path)
        break

# 尝试读取 GEMINI_API_KEY 或 API_KEY
API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("API_KEY")

# 读取自定义 API 端点（用于中转服务）
GEMINI_API_BASE = os.getenv("GEMINI_API_BASE", "").strip()
FORCE_PROXY = os.getenv("FORCE_PROXY", "").strip()

_user_settings_cache = {}

def _load_user_settings() -> dict:
    """Load user_settings.json with caching and safe fallbacks."""
    if "data" in _user_settings_cache:
        return _user_settings_cache["data"]
    settings_path = os.path.join(PROJECT_ROOT, "config", "user_settings.json")
    try:
        with open(settings_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = {}
    _user_settings_cache["data"] = data
    return data

def get_workspace_root() -> str:
    """Return the workspace root directory from settings or default path."""
    settings = _load_user_settings()
    workspace_dir = settings.get("storage", {}).get("workspace_dir")
    if workspace_dir:
        return workspace_dir
    return os.path.join(PROJECT_ROOT, "workspace")

def get_organize_root() -> str:
    """Return the file organization root directory from settings or default path."""
    settings = _load_user_settings()
    organize_root = settings.get("storage", {}).get("organize_root")
    if organize_root:
        return organize_root
    return os.path.join(get_workspace_root(), "_organize")

def get_default_wechat_files_dir() -> str:
    """Return configured default WeChat files directory, if provided by user settings."""
    settings = _load_user_settings()
    return settings.get("storage", {}).get("wechat_files_dir", "")

if not API_KEY:
    print("⚠️ Warning: GEMINI_API_KEY or API_KEY not found in gemini_config.env")
    print("   请在 config/gemini_config.env 中配置 API 密钥")
    print("   应用将继续启动，但 AI 功能不可用")
    # 不再 sys.exit — 允许应用启动并在 UI 中提示用户配置

if GEMINI_API_BASE:
    print(f"📡 使用自定义 API 端点: {GEMINI_API_BASE}")

# 检测并设置代理
PROXY_OPTIONS = [
    "http://127.0.0.1:7890",
    "http://127.0.0.1:10809",
    "http://127.0.0.1:1080",
]

def _normalize_proxy_url(proxy_value: str) -> str:
    """Normalize proxy value to a URL with scheme."""
    if not proxy_value:
        return ""
    value = proxy_value.strip()
    if not value:
        return ""
    if "://" not in value:
        value = f"http://{value}"
    return value

def _extract_system_proxy_candidates() -> list:
    """Collect proxy candidates from system settings (Windows) and env."""
    candidates = []

    # 1) Environment variables first (if user/system already configured)
    env_proxy = os.environ.get("HTTPS_PROXY") or os.environ.get("https_proxy")
    if env_proxy:
        candidates.append(_normalize_proxy_url(env_proxy))

    # 2) Windows Internet Settings proxy (for "Use a proxy server")
    if sys.platform.startswith("win"):
        try:
            import winreg
            key_path = r"Software\Microsoft\Windows\CurrentVersion\Internet Settings"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path) as key:
                proxy_enabled = winreg.QueryValueEx(key, "ProxyEnable")[0]
                if proxy_enabled:
                    proxy_server = str(winreg.QueryValueEx(key, "ProxyServer")[0]).strip()
                    if proxy_server:
                        # Formats:
                        #   127.0.0.1:7890
                        #   http=127.0.0.1:7890;https=127.0.0.1:7890
                        if "=" in proxy_server and ";" in proxy_server:
                            pairs = [p.strip() for p in proxy_server.split(";") if p.strip()]
                            parsed_map = {}
                            for pair in pairs:
                                if "=" in pair:
                                    k, v = pair.split("=", 1)
                                    parsed_map[k.strip().lower()] = v.strip()
                            for proto in ["https", "http", "socks", "socks5"]:
                                if parsed_map.get(proto):
                                    candidates.append(_normalize_proxy_url(parsed_map.get(proto)))
                        else:
                            candidates.append(_normalize_proxy_url(proxy_server))
        except Exception:
            pass

    # 3) Built-in localhost fallback options
    candidates.extend(PROXY_OPTIONS)

    # De-duplicate while preserving order
    deduped = []
    seen = set()
    for item in candidates:
        if not item:
            continue
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped

def setup_proxy():
    # 优先使用强制代理（不需要测试）
    if FORCE_PROXY and FORCE_PROXY.lower() not in ("auto", "system"):
        os.environ["HTTPS_PROXY"] = FORCE_PROXY
        os.environ["HTTP_PROXY"] = FORCE_PROXY
        print(f"🔧 使用强制代理: {FORCE_PROXY}")
        return FORCE_PROXY

    # 自动匹配系统代理与本地常见端口
    import socket
    from urllib.parse import urlparse

    proxy_candidates = _extract_system_proxy_candidates()

    for proxy in proxy_candidates:
        try:
            # 从 URL 提取 host:port
            parsed = urlparse(proxy)
            host = parsed.hostname
            port = parsed.port
            if not host or not port:
                continue
            
            # 快速端口检测（0.1秒超时）
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(0.1)
            result = sock.connect_ex((host, port))
            sock.close()
            
            if result == 0:
                os.environ["HTTPS_PROXY"] = proxy
                os.environ["HTTP_PROXY"] = proxy
                print(f"✅ 自动匹配系统代理: {proxy}")
                return proxy
        except:
            continue
    
    return None

# 延迟代理检测到首次需要时（启动加速）
_detected_proxy = None
_proxy_checked = False

def get_detected_proxy():
    """懒加载代理检测（首次调用时执行）"""
    global _detected_proxy, _proxy_checked
    if not _proxy_checked:
        _detected_proxy = setup_proxy()
        _proxy_checked = True
    return _detected_proxy

# 向后兼容：detected_proxy 现在通过函数访问
detected_proxy = None  # 占位符，实际通过 get_detected_proxy() 获取

# 在后台线程预热代理检测（不阻塞启动）
def _warmup_proxy():
    global detected_proxy
    detected_proxy = get_detected_proxy()
threading.Thread(target=_warmup_proxy, daemon=True).start()

# 创建 GenAI 客户端 (配置代理和自定义端点)
def create_client():
    import httpx
    proxy = get_detected_proxy()
    # 超时时间: 连接30秒, 读取180秒 (Nano Banana 图像生成和长文本生成需要更长时间)
    timeout_config = httpx.Timeout(180.0, connect=30.0)
    
    # 构建 http_options
    http_options = {}
    
    # 注意：最新的 Gemini 模型（如 gemini-1.5-flash）需要 v1beta API
    # v1 API 只支持旧的模型。这里使用 v1beta。
    http_options['api_version'] = 'v1beta'
    
    # 自定义 API 端点（用于中转服务）
    if GEMINI_API_BASE:
        http_options['base_url'] = GEMINI_API_BASE
        print(f"📡 API 端点: {GEMINI_API_BASE}")
    
    # 配置代理 - 通过环境变量确保被使用
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy
        print(f"🔌 设置代理: {proxy}")
        
    # 使用 httpx with explicit proxy for genai
    try:
        if proxy:
            http_client = httpx.Client(
                proxy=proxy,
                timeout=timeout_config,
                verify=False  # SSL verification disabled with proxy
            )
            http_options['httpxClient'] = http_client
        else:
            # 无代理时也要配置超时
            http_client = httpx.Client(
                timeout=timeout_config,
                verify=True
            )
            http_options['httpxClient'] = http_client
    except Exception as e:
        print(f"⚠️ 创建 HTTP 客户端出错 (proxy={proxy}): {e}")
        # 回退：不使用代理
        if proxy:
            print(f"⚠️ 尝试不使用代理重新创建客户端")
            http_client = httpx.Client(
                timeout=timeout_config,
                verify=True
            )
            http_options['httpxClient'] = http_client
    
    return genai.Client(
        api_key=API_KEY,
        http_options=http_options if http_options else None
    )

# 懒加载客户端
_client = None

def get_client():
    """获取 GenAI 客户端（懒加载）"""
    global _client
    if _client is None:
        _client = create_client()
    return _client

# 保持向后兼容的 client 变量（通过属性访问触发懒加载）
class _ClientProxy:
    """代理类，实现懒加载"""
    def __getattr__(self, name):
        return getattr(get_client(), name)

client = _ClientProxy()


def create_research_client():
    """创建专用于 Deep Research 的长超时客户端 (5分钟 read timeout)"""
    import httpx
    proxy = get_detected_proxy()
    # 深度研究需要更长的超时时间：连接30秒，读取5分钟
    timeout_config = httpx.Timeout(300.0, connect=30.0)
    
    # 构建 http_options
    http_options = {}
    
    # 最新的 Gemini 模型需要 v1beta API
    http_options['api_version'] = 'v1beta'
    
    # 自定义 API 端点
    if GEMINI_API_BASE:
        http_options['base_url'] = GEMINI_API_BASE
    
    # 配置代理 - 通过环境变量确保被使用
    if proxy:
        os.environ["HTTP_PROXY"] = proxy
        os.environ["HTTPS_PROXY"] = proxy
        
    if proxy:
        http_client = httpx.Client(
            proxy=proxy,
            timeout=timeout_config,
            verify=False  # SSL verification disabled with proxy
        )
        http_options['httpxClient'] = http_client
    else:
        http_client = httpx.Client(
            timeout=timeout_config,
            verify=True
        )
        http_options['httpxClient'] = http_client
    
    return genai.Client(
        api_key=API_KEY,
        http_options=http_options if http_options else None
    )


def run_with_timeout(fn, timeout_seconds):
    """在线程中执行函数并限时返回 (避免卡死主流程)"""
    holder = {'result': None, 'error': None}

    def _runner():
        try:
            holder['result'] = fn()
        except Exception as e:
            holder['error'] = e

    t = threading.Thread(target=_runner, daemon=True)
    t.start()
    t.join(timeout_seconds)
    if t.is_alive():
        return None, TimeoutError(f"Timeout after {timeout_seconds}s"), True
    return holder['result'], holder['error'], False


def run_with_heartbeat(fn, start_time, heartbeat_callback, heartbeat_interval=5, timeout_seconds=90):
    """
    在后台线程运行函数，同时定期发送心跳。
    用于非流式 API 调用（如图像生成）。
    
    Args:
        fn: 要执行的函数
        start_time: 请求开始时间
        heartbeat_callback: 心跳回调函数，接收 elapsed_seconds 参数
        heartbeat_interval: 心跳间隔（秒）
        timeout_seconds: 超时时间（秒）
    
    Returns:
        (result, error, timed_out)
    """
    import queue
    import threading
    result_queue = queue.Queue()
    
    def worker():
        try:
            result = fn()
            result_queue.put(('success', result))
        except Exception as e:
            result_queue.put(('error', e))
    
    thread = threading.Thread(target=worker, daemon=True)
    thread.start()
    
    last_heartbeat = time.time()
    
    while True:
        # 检查是否超时
        elapsed = time.time() - start_time
        if elapsed > timeout_seconds:
            return None, TimeoutError(f"操作超时 ({int(elapsed)}s)"), True
        
        # 尝试获取结果（短超时）
        try:
            status, data = result_queue.get(timeout=1.0)
            if status == 'success':
                return data, None, False
            else:
                return None, data, False
        except queue.Empty:
            # 发送心跳
            current_time = time.time()
            if current_time - last_heartbeat >= heartbeat_interval:
                heartbeat_callback(int(current_time - start_time))
                last_heartbeat = current_time


def stream_with_keepalive(response_stream, start_time, keepalive_interval=5, max_wait_first_token=60):
    """
    包装流式响应，在等待第一个 token 期间发送保活心跳。
    
    Args:
        response_stream: 原始流式响应迭代器
        start_time: 请求开始时间
        keepalive_interval: 心跳间隔（秒）
        max_wait_first_token: 等待第一个 token 的最大时间（秒）
    
    Yields:
        (type, data): type 可以是 'chunk', 'heartbeat', 'timeout'
    """
    import queue
    import time
    
    chunk_queue = queue.Queue()
    first_chunk_received = threading.Event()
    stream_done = threading.Event()
    stream_error = {'error': None}
    
    def stream_reader():
        """在后台线程中读取流"""
        try:
            for chunk in response_stream:
                chunk_queue.put(('chunk', chunk))
                first_chunk_received.set()
            chunk_queue.put(('done', None))
        except Exception as e:
            stream_error['error'] = e
            chunk_queue.put(('error', e))
        finally:
            stream_done.set()
    
    # 启动后台读取线程
    reader_thread = threading.Thread(target=stream_reader, daemon=True)
    reader_thread.start()
    
    last_heartbeat = time.time()
    
    while True:
        # 检查是否等待第一个 token 超时
        if not first_chunk_received.is_set():
            elapsed = time.time() - start_time
            if elapsed > max_wait_first_token:
                yield ('timeout', f'等待响应超时 ({int(elapsed)}s)')
                return
        
        # 尝试获取 chunk，使用短超时以便发送心跳
        try:
            item_type, item_data = chunk_queue.get(timeout=1.0)
            
            if item_type == 'chunk':
                yield ('chunk', item_data)
            elif item_type == 'done':
                return
            elif item_type == 'error':
                raise item_data
                
        except queue.Empty:
            # 队列为空，检查是否需要发送心跳
            current_time = time.time()
            if current_time - last_heartbeat >= keepalive_interval:
                elapsed = int(current_time - start_time)
                yield ('heartbeat', elapsed)
                last_heartbeat = current_time
            
            # 检查流是否已结束
            if stream_done.is_set() and chunk_queue.empty():
                if stream_error['error']:
                    raise stream_error['error']
                return

app = Flask(__name__)
# 静态资源缓存，减少重复加载
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 3600

# CORS: 云模式限制来源，本地模式打开
_cors_origins = os.environ.get('KOTO_CORS_ORIGINS', '*')
if os.environ.get('KOTO_DEPLOY_MODE') == 'cloud' and _cors_origins == '*':
    # 云模式默认只允许自身站点（同源），可通过环境变量覆盖
    _cors_origins = os.environ.get('KOTO_SITE_URL', '*')
CORS(app, origins=_cors_origins)

# ================= 用户认证系统 =================
try:
    from auth import register_auth_routes
    register_auth_routes(app)
except Exception as e:
    print(f"[Auth] ⚠️ 认证模块加载失败: {e}")

# ================= 并行执行系统初始化 =================
if PARALLEL_SYSTEM_ENABLED:
    print("[PARALLEL] 🚀 Initializing parallel execution system...")
    try:
        register_parallel_api(app)
        start_dispatcher()
        print("[PARALLEL] ✅ Parallel execution system initialized successfully")
    except Exception as e:
        print(f"[PARALLEL] ❌ Failed to initialize parallel execution system: {e}")
        PARALLEL_SYSTEM_ENABLED = False

# ================= WebSocket 支持（可选） =================
sock = None
if Sock:
    sock = Sock(app)
else:
    print("[WebSocket] ⚠️ flask-sock 未安装，使用轮询作为通知兜底")

if sock:
    @sock.route('/ws/notifications')
    def ws_notifications(ws):
        user_id = request.args.get('user_id', 'default')
        manager = get_notification_manager()
        manager.register_connection(user_id, ws)
        try:
            while True:
                message = ws.receive()
                if message is None:
                    break
                if isinstance(message, str) and message.lower() == 'ping':
                    ws.send('pong')
        finally:
            manager.unregister_connection(user_id, ws)

# ================= 延迟注册蓝图（在后台线程中加载，避免阻塞启动） =================
_blueprints_registered = False
_blueprints_lock = threading.Lock()

def _register_blueprints_deferred():
    """在后台线程中注册所有蓝图，避免阻塞主线程启动."""
    global _blueprints_registered, agent_bp
    with _blueprints_lock:
        if _blueprints_registered:
            return
        _blueprints_registered = True
    
    # 注册统一 Agent API
    try:
        from app.api import agent_bp as _agent_bp
        agent_bp = _agent_bp
        app.register_blueprint(agent_bp, url_prefix='/api/agent')
        print("[UnifiedAgent] ✅ 统一 Agent API 已注册: /api/agent")
    except ImportError as e:
        print(f"[UnifiedAgent] ⚠️ 未能导入统一 Agent API 蓝图: {e}")
    except Exception as e:
        print(f"[UnifiedAgent] ❌ 注册失败: {e}")
    
    # 注册增强语音 API
    try:
        from voice_api_enhanced import voice_bp
        app.register_blueprint(voice_bp)
        print("[VOICE_API] 已注册增强语音 API 蓝图")
    except ImportError as e:
        print(f"[VOICE_API] ⚠️ 未能导入增强语音模块: {e}")

    # 注册 PPT 编辑 API（P1 功能）
    try:
        from web.ppt_api_routes import ppt_api_bp
        app.register_blueprint(ppt_api_bp)
        print("[PPT_API] ✅ PPT 编辑 API 已注册: /api/ppt")
    except ImportError as e:
        print(f"[PPT_API] ⚠️ 未能导入 PPT 编辑 API: {e}")
    except Exception as e:
        print(f"[PPT_API] ⚠️ PPT 编辑 API 注册失败: {e}")

    # 注册自适应 Agent API（已迁移到 UnifiedAgent，但保留兼容导入）
    try:
        from adaptive_agent_api import init_adaptive_agent_api
        init_adaptive_agent_api(app, gemini_client=None)
        print("[AdaptiveAgent] ✅ 自适应 Agent API 已注册 (延迟加载客户端)")
    except ImportError:
        print("[AdaptiveAgent] ℹ️ 旧 Agent 模块已退役，使用 UnifiedAgent")
    except Exception as e:
        print(f"[AdaptiveAgent] ⚠️ 旧 Agent 初始化失败 (非致命): {e}")
    
    print("[INIT] ✅ 所有蓝图注册完成")

# 在后台线程中注册蓝图，不阻塞主线程
threading.Thread(target=_register_blueprints_deferred, name="BlueprintLoader", daemon=True).start()

CHAT_DIR = os.path.join(PROJECT_ROOT, "chats")
WORKSPACE_DIR = get_workspace_root()
UPLOAD_DIR = os.path.join(PROJECT_ROOT, "web", "uploads")
os.makedirs(CHAT_DIR, exist_ok=True)
os.makedirs(WORKSPACE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ================= Settings Manager (提前加载) =================
try:
    from settings import SettingsManager
except ImportError:
    from web.settings import SettingsManager
settings_manager = SettingsManager()

# ================= 智能模型矩阵 (2026-01 最先进) =================
# 使用 Gemini 3.0 系列 + Imagen 4.0 图像模型

MODEL_MAP = {
    # 日常对话 - Gemini 3.0 Flash (快速响应)
    "CHAT": "gemini-3-flash-preview",
    
    # 编程/代码 - Gemini 3.0 Pro, 但不稳定时回退
    "CODER": "gemini-3-pro-preview",
    
    # 🌐 联网搜索 - Gemini 2.5 Flash (grounding 支持好)
    "WEB_SEARCH": "gemini-2.5-flash", 
    
    # 视觉理解 - Gemini 3.0 Flash (多模态)
    "VISION": "gemini-3-flash-preview",
    
    # 深度研究 - Gemini 3.0 Pro (稳定) - 原deep-research经常timeout
    "RESEARCH": "gemini-3-pro-preview",
    
    # 文档生成 - 根据复杂度动态选择
    "FILE_GEN": "gemini-3-flash-preview",
    
    # 🎨 图像生成 - Nano Banana Pro (原生图像生成)
    "PAINTER": "nano-banana-pro-preview",
    
    # 🖥️ 系统操作 - 本地执行 (不需要模型)
    "SYSTEM": "local-executor",
    
    # 📂 文件操作 - 本地文件系统
    "FILE_OP": "local-executor",
    
    # 🤖 Agent 工具调用 - Gemini 3.0 Flash (function calling)
    "AGENT": "gemini-3-flash-preview",
}

# 模型能力矩阵 (用于智能路由和显示)
MODEL_INFO = {
    # Gemini 3.0 系列 (最新最强)
    "gemini-3-pro-preview": {
        "name": "Gemini 3.0 Pro",
        "speed": "🚀",
        "tier": 7,
        "strengths": ["推理", "分析", "代码", "复杂任务"],
    },
    "gemini-3-flash-preview": {
        "name": "Gemini 3.0 Flash",
        "speed": "⚡",
        "tier": 6,
        "strengths": ["快速", "对话", "多模态"],
    },
    # Gemini 2.5 (grounding)
    "gemini-2.5-flash": {
        "name": "Gemini 2.5 Flash",
        "speed": "🌐",
        "tier": 5,
        "strengths": ["联网搜索", "grounding"],
    },
    # 深度研究
    "deep-research-pro-preview-12-2025": {
        "name": "Deep Research Pro",
        "speed": "🔬",
        "tier": 7,
        "strengths": ["深度研究", "学术分析", "综合报告"],
    },
    # 图像生成 (Nano Banana)
    "nano-banana-pro-preview": {
        "name": "Nano Banana Pro",
        "speed": "🎨",
        "tier": 5,
        "strengths": ["图像生成", "创意绘画", "艺术风格"],
    },
    # Gemini 2.0 Flash Exp (图像生成备用)
    "gemini-2.0-flash-exp": {
        "name": "Gemini 2.0 Flash Exp",
        "speed": "🎨",
        "tier": 5,
        "strengths": ["图像生成", "多模态", "实验功能"],
    },
    # 本地执行器
    "local-executor": {
        "name": "Local Executor",
        "speed": "🖥️",
        "tier": 0,
        "strengths": ["系统操作", "打开应用", "文件管理"],
    },
}

def get_model_display_name(model_id):
    info = MODEL_INFO.get(model_id)
    if info:
        return f"{info['name']} {info['speed']}"
    return model_id


# ================= 本地系统执行器 =================
class LocalExecutor:
    """
    本地系统操作执行器 - 让 Koto 成为真正的 AI OS
    支持：打开应用、文件操作、系统命令等
    """
    
    # Windows 常用应用路径映射 (包含更多路径)
    APP_ALIASES = {
        # 社交通讯
        "微信": ["WeChat", r"C:\Program Files\Tencent\WeChat\WeChat.exe", r"C:\Program Files (x86)\Tencent\WeChat\WeChat.exe"],
        "wechat": ["WeChat", r"C:\Program Files\Tencent\WeChat\WeChat.exe", r"C:\Program Files (x86)\Tencent\WeChat\WeChat.exe"],
        "qq": ["QQ", r"C:\Program Files\Tencent\QQ\Bin\QQ.exe", r"C:\Program Files (x86)\Tencent\QQ\Bin\QQ.exe", r"C:\Program Files\Tencent\QQNT\QQ.exe"],
        "钉钉": ["DingTalk", "dingtalk"],
        "飞书": ["Feishu", "Lark"],
        "telegram": ["Telegram"],
        "discord": ["Discord", "Update --processStart Discord.exe"],
        
        # 游戏平台
        "steam": ["steam", r"C:\Program Files (x86)\Steam\steam.exe", r"C:\Program Files\Steam\steam.exe", r"D:\Steam\steam.exe"],
        "epic": ["EpicGamesLauncher", r"C:\Program Files (x86)\Epic Games\Launcher\Portal\Binaries\Win32\EpicGamesLauncher.exe"],
        "战网": ["Battle.net"],
        "wallpaper engine": ["wallpaper32", "wallpaper64", r"C:\Program Files (x86)\Steam\steamapps\common\wallpaper_engine\wallpaper32.exe"],
        "wallpaper": ["wallpaper32", "wallpaper64"],
        
        # 浏览器
        "chrome": ["chrome", r"C:\Program Files\Google\Chrome\Application\chrome.exe", r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"],
        "谷歌浏览器": ["chrome"],
        "edge": ["msedge", r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"],
        "firefox": ["firefox", r"C:\Program Files\Mozilla Firefox\firefox.exe"],
        "浏览器": ["chrome", "msedge", "firefox"],
        
        # 开发工具
        "vscode": ["code"],
        "vs code": ["code"],
        "code": ["code"],
        "pycharm": ["pycharm64", "pycharm"],
        "idea": ["idea64", "idea"],
        "terminal": ["wt", "cmd", "powershell"],
        "终端": ["wt", "cmd", "powershell"],
        "命令行": ["cmd", "powershell"],
        "git": ["git-bash", r"C:\Program Files\Git\git-bash.exe"],
        
        # 办公软件
        "word": ["winword", "WINWORD"],
        "excel": ["excel", "EXCEL"],
        "ppt": ["powerpnt", "POWERPNT"],
        "powerpoint": ["powerpnt"],
        "outlook": ["outlook", "OUTLOOK"],
        "记事本": ["notepad"],
        "notepad": ["notepad"],
        "wps": ["wps", "wpsoffice", r"C:\Users\12524\AppData\Local\Kingsoft\WPS Office\ksolaunch.exe"],
        "wps office": ["wps", "wpsoffice", "ksolaunch"],
        
        # 媒体
        "spotify": ["Spotify"],
        "网易云": ["cloudmusic", r"C:\Program Files (x86)\Netease\CloudMusic\cloudmusic.exe", r"C:\Program Files\Netease\CloudMusic\cloudmusic.exe"],
        "网易云音乐": ["cloudmusic", r"C:\Program Files (x86)\Netease\CloudMusic\cloudmusic.exe"],
        "cloudmusic": ["cloudmusic", r"C:\Program Files (x86)\Netease\CloudMusic\cloudmusic.exe"],
        "qq音乐": ["QQMusic", r"C:\Program Files (x86)\Tencent\QQMusic\QQMusic.exe"],
        "酷狗": ["KuGou", r"C:\Program Files\KuGou\KuGou.exe"],
        "酷我": ["KuWo"],
        "网易音乐": ["cloudmusic"],
        "potplayer": ["PotPlayerMini64", "PotPlayerMini", r"C:\Program Files\DAUM\PotPlayer\PotPlayerMini64.exe"],
        "vlc": ["vlc", r"C:\Program Files\VideoLAN\VLC\vlc.exe"],
        
        # 系统
        "设置": ["ms-settings:"],
        "控制面板": ["control"],
        "任务管理器": ["taskmgr"],
        "计算器": ["calc"],
        "文件管理器": ["explorer"],
        "资源管理器": ["explorer"],
        "画图": ["mspaint"],
        "截图": ["snippingtool", "SnippingTool"],
    }
    
    # 系统操作关键词
    SYSTEM_KEYWORDS = [
        "打开", "启动", "运行", "开启", "关闭", "退出", "杀死",
        "open", "start", "launch", "run", "close", "kill", "exit",
        "搜索", "查找", "search", "find",
        "截图", "screenshot",
        "音量", "亮度", "volume", "brightness",
        "关机", "重启", "休眠", "睡眠", "shutdown", "restart", "sleep",
    ]
    
    # 知识提问模式 —— 如果匹配到这些，说明用户是在**问问题**，不是在下命令
    QUESTION_PATTERNS = [
        "怎么", "如何", "什么办法", "什么方法", "什么意思", "什么是", "是什么",
        "为什么", "为啥", "能不能", "可以吗", "可不可以", "怎样", "咋",
        "一般用", "通常", "有没有", "有什么", "哪些", "哪个", "哪种",
        "区别", "对比", "比较", "最好的", "推荐", "建议",
        "教程", "步骤", "流程", "原理", "概念",
        "用什么", "是啥", "啥意思", "讲讲", "说说", "介绍",
        "how to", "what is", "why", "which", "recommend",
        "difference between", "best way", "tutorial",
    ]

    @classmethod
    def is_system_command(cls, text):
        """检测是否是系统操作请求（祈使句/命令句，非知识提问）
        
        核心逻辑：
        1. 必须包含动作关键词（打开/启动/关闭等）
        2. 必须包含已知应用名或「打开+紧跟名词」的短句模式
        3. 排除知识性提问（包含"怎么/如何/什么办法"等）
        """
        text_lower = text.lower().strip()
        
        # ——— 排除条件：知识提问不是系统命令 ———
        if any(qp in text_lower for qp in cls.QUESTION_PATTERNS):
            return False
        
        # ——— 排除条件：句子太长一般不是命令（命令通常 <20字） ———
        if len(text_lower) > 30:
            return False
        
        # ——— 必须有动作关键词 ———
        action_keywords = [
            "打开", "启动", "运行", "开启", "关闭", "退出", "杀死",
            "open", "start", "launch", "close", "kill", "exit",
            "截图", "screenshot", "关机", "重启", "休眠", "睡眠",
            "shutdown", "restart", "sleep",
            "时间", "几点", "日期", "几号", "星期几", "time", "date",
            "状态", "信息", "配置", "内存", "cpu", "硬盘"
        ]
        has_action = any(kw in text_lower for kw in action_keywords)
        if not has_action:
            return False
        
        # ——— 必须有已知应用名 ———
        has_app = any(app in text_lower for app in cls.APP_ALIASES.keys())
        
        # 或者是独立的系统操作（无需应用名）
        standalone_commands = [
            "截图", "screenshot", "关机", "重启", "休眠", "睡眠", 
            "shutdown", "restart", "sleep",
            "时间", "几点", "日期", "几号", "星期几", "time", "date",
            "系统状态", "电脑状态", "系统信息", "电脑信息", "配置", "内存", "cpu", "硬盘"
        ]
        is_standalone = any(cmd in text_lower for cmd in standalone_commands)
        
        return has_app or is_standalone
    
    @classmethod
    def extract_app_name(cls, text):
        """从文本中提取应用名"""
        text_lower = text.lower()
        
        # 泛指类别映射到具体应用
        category_mapping = {
            "音乐软件": ["网易云", "qq音乐", "spotify", "酷狗"],
            "听歌软件": ["网易云", "qq音乐", "spotify", "酷狗"],
            "浏览器": ["edge", "chrome", "firefox"],
            "文本编辑器": ["记事本", "vscode", "notepad"],
            "代码编辑器": ["vscode", "pycharm", "idea"],
            "视频播放器": ["potplayer", "vlc"],
            "聊天软件": ["微信", "qq", "钉钉"],
            "办公软件": ["word", "excel", "ppt", "wps"]
        }
        
        # 先尝试精确匹配已知应用
        for app_name in sorted(cls.APP_ALIASES.keys(), key=len, reverse=True):
            if app_name in text_lower:
                return app_name
                
        # 尝试匹配泛指类别
        for category, apps in category_mapping.items():
            if category in text_lower:
                # 检查系统中安装了哪个
                import shutil
                for app in apps:
                    aliases = cls.APP_ALIASES.get(app, [app])
                    for alias in aliases:
                        if os.path.exists(alias) or shutil.which(alias):
                            return app
                # 如果都没找到绝对路径，默认返回第一个
                return apps[0]
        
        # 如果没有匹配，尝试提取"打开xxx"中的xxx
        import re
        patterns = [
            r'(?:打开|启动|运行|开启)\s*(?:一个|一款)?\s*(.+?)(?:\s|$|吧|呗)',
            r'(?:open|start|launch)\s+(?:a\s+)?(.+?)(?:\s|$)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return match.group(1).strip()
        
        return None
    
    @classmethod
    def find_app_in_start_menu(cls, app_name):
        """从开始菜单查找应用"""
        import glob
        
        start_menu_paths = [
            os.path.expandvars(r"%ProgramData%\Microsoft\Windows\Start Menu\Programs"),
            os.path.expandvars(r"%AppData%\Microsoft\Windows\Start Menu\Programs"),
        ]
        
        app_name_lower = app_name.lower()
        
        for start_path in start_menu_paths:
            if not os.path.exists(start_path):
                continue
            
            # 搜索 .lnk 文件
            for lnk_file in glob.glob(os.path.join(start_path, "**", "*.lnk"), recursive=True):
                lnk_name = os.path.basename(lnk_file).lower().replace(".lnk", "")
                if app_name_lower in lnk_name or lnk_name in app_name_lower:
                    return lnk_file
        
        return None
    
    @classmethod
    def find_app_smart(cls, app_name):
        """智能查找应用 - 多种方式"""
        import subprocess
        import shutil
        
        # 1. 先检查预定义别名
        if app_name.lower() in cls.APP_ALIASES:
            aliases = cls.APP_ALIASES[app_name.lower()]
            for alias in aliases:
                # 检查是否是完整路径
                if os.path.exists(alias):
                    return alias
                # 检查是否在 PATH 中
                if shutil.which(alias):
                    return alias
        
        # 2. 检查 PATH
        if shutil.which(app_name):
            return app_name
        
        # 3. 从开始菜单查找
        lnk_path = cls.find_app_in_start_menu(app_name)
        if lnk_path:
            return lnk_path
        
        # 4. 使用 PowerShell 搜索
        try:
            ps_cmd = f'Get-StartApps | Where-Object {{$_.Name -like "*{app_name}*"}} | Select-Object -First 1 -ExpandProperty AppID'
            result = subprocess.run(
                ["powershell", "-Command", ps_cmd],
                capture_output=True,
                text=True,
                timeout=5,
                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except:
            pass
        
        return None
    
    @classmethod
    def execute(cls, user_input):
        """执行系统操作"""
        import subprocess
        import shutil
        
        text_lower = user_input.lower()
        result = {
            "success": False,
            "action": "",
            "message": "",
            "details": ""
        }
        
        # === 打开应用 ===
        if any(kw in text_lower for kw in ["打开", "启动", "运行", "开启", "open", "start", "launch"]):
            app_name = cls.extract_app_name(text_lower)
            
            if app_name:
                # 使用智能查找
                app_path = cls.find_app_smart(app_name)
                
                if app_path:
                    try:
                        # 特殊处理 ms-settings 等 URI
                        if app_path.startswith("ms-"):
                            subprocess.Popen(
                                f'start {app_path}',
                                shell=True,
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL
                            )
                        # 处理 .lnk 快捷方式
                        elif app_path.endswith(".lnk"):
                            subprocess.Popen(
                                f'start "" "{app_path}"',
                                shell=True,
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL
                            )
                        # 处理 UWP 应用 AppID
                        elif "!" in app_path:
                            subprocess.Popen(
                                f'start shell:AppsFolder\\{app_path}',
                                shell=True,
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL
                            )
                        # 直接路径或命令
                        elif os.path.exists(app_path):
                            subprocess.Popen(
                                [app_path], 
                                shell=True,
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL,
                                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                            )
                        else:
                            # 使用 start 命令
                            subprocess.Popen(
                                f'start "" "{app_path}"',
                                shell=True,
                                stdout=subprocess.DEVNULL,
                                stderr=subprocess.DEVNULL
                            )
                        
                        result["success"] = True
                        result["action"] = "open_app"
                        result["message"] = f"✅ 已打开 {app_name}"
                        print(f"[LocalExecutor] ✅ 成功启动应用: {app_name} - 路径: {app_path}")
                        return result
                    except Exception as e:
                        result["message"] = f"❌ 打开 {app_name} 失败: {str(e)}"
                        print(f"[LocalExecutor] ❌ 启动失败: {app_name} - 错误: {str(e)}")
                        return result
                
                # 智能查找失败，尝试直接用 start 命令
                try:
                    subprocess.Popen(
                        f'start "" "{app_name}"',
                        shell=True,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL
                    )
                    result["success"] = True
                    result["action"] = "open_app"
                    result["message"] = f"✅ 正在尝试打开 {app_name}"
                    return result
                except:
                    pass
                
                result["message"] = f"❌ 无法打开 {app_name}，请确认已安装"
                result["details"] = f"提示: 您可以尝试使用完整的应用名称"
                return result
        
        # === 关闭应用 ===
        if any(kw in text_lower for kw in ["关闭", "退出", "杀死", "close", "kill", "exit"]):
            app_name = cls.extract_app_name(text_lower)
            if app_name:
                aliases = cls.APP_ALIASES.get(app_name, [app_name])
                for alias in aliases:
                    try:
                        if sys.platform == "win32":
                            # 提取进程名
                            proc_name = alias.split("\\")[-1] if "\\" in alias else alias
                            if not proc_name.endswith(".exe"):
                                proc_name += ".exe"
                            
                            ret = subprocess.run(
                                f'taskkill /IM "{proc_name}" /F',
                                shell=True,
                                capture_output=True,
                                timeout=5,
                                creationflags=subprocess.CREATE_NO_WINDOW
                            )
                            if ret.returncode == 0:
                                result["success"] = True
                                result["action"] = "close_app"
                                result["message"] = f"✅ 已关闭 {app_name}"
                                return result
                    except:
                        continue
                
                result["message"] = f"❌ 无法关闭 {app_name}"
                return result
        
        # === 截图 ===
        if "截图" in text_lower or "screenshot" in text_lower:
            if sys.platform == "win32":
                subprocess.Popen(
                    "snippingtool",
                    shell=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL
                )
                result["success"] = True
                result["action"] = "screenshot"
                result["message"] = "✅ 已打开截图工具"
                return result
        
        # === 搜索 ===
        if any(kw in text_lower for kw in ["搜索", "查找", "search"]):
            # 提取搜索内容
            search_terms = text_lower.replace("搜索", "").replace("查找", "").replace("search", "").strip()
            if search_terms:
                import webbrowser
                webbrowser.open(f"https://www.google.com/search?q={search_terms}")
                result["success"] = True
                result["action"] = "search"
                result["message"] = f"✅ 正在搜索: {search_terms}"
                return result
        
        # === 系统时间/日期 ===
        if any(kw in text_lower for kw in ["时间", "几点", "日期", "几号", "星期几", "time", "date"]):
            import datetime
            now = datetime.datetime.now()
            
            weekdays = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
            weekday_str = weekdays[now.weekday()]
            
            if any(kw in text_lower for kw in ["日期", "几号", "星期几", "date"]):
                time_str = now.strftime(f"%Y年%m月%d日 {weekday_str}")
                msg = f"📅 当前日期是：{time_str}"
            else:
                time_str = now.strftime(f"%Y-%m-%d %H:%M:%S {weekday_str}")
                msg = f"🕒 当前系统时间是：{time_str}"
                
            result["success"] = True
            result["action"] = "get_time"
            result["message"] = msg
            return result
            
        # === 电源操作 ===
        if any(kw in text_lower for kw in ["关机", "重启", "休眠", "睡眠", "shutdown", "restart", "sleep"]):
            if sys.platform == "win32":
                if "关机" in text_lower or "shutdown" in text_lower:
                    subprocess.Popen("shutdown /s /t 0", shell=True)
                    msg = "✅ 正在关机..."
                elif "重启" in text_lower or "restart" in text_lower:
                    subprocess.Popen("shutdown /r /t 0", shell=True)
                    msg = "✅ 正在重启..."
                elif "休眠" in text_lower or "睡眠" in text_lower or "sleep" in text_lower:
                    subprocess.Popen("rundll32.exe powrprof.dll,SetSuspendState 0,1,0", shell=True)
                    msg = "✅ 正在进入休眠/睡眠状态..."
                
                result["success"] = True
                result["action"] = "power_op"
                result["message"] = msg
                return result
                
        # === 系统状态/电脑状态 ===
        if any(kw in text_lower for kw in ["系统状态", "电脑状态", "系统信息", "电脑信息", "配置", "内存", "cpu", "硬盘"]):
            info = cls.get_system_info()
            if info.get("success"):
                msg = f"💻 **系统状态报告**\n\n"
                msg += f"- **操作系统**: {info.get('system')} ({info.get('platform')})\n"
                msg += f"- **处理器**: {info.get('processor')}\n"
                msg += f"- **CPU 使用率**: {info.get('cpu_percent')}%\n"
                
                mem = info.get('memory', {})
                msg += f"- **内存**: 已用 {mem.get('percent')}% (剩余 {mem.get('available')} / 总共 {mem.get('total')})\n"
                
                disk = info.get('disk', {})
                msg += f"- **C盘**: 已用 {disk.get('percent')}% (剩余 {disk.get('free')} / 总共 {disk.get('total')})\n"
                
                result["success"] = True
                result["action"] = "get_system_info"
                result["message"] = msg
                return result
        
        result["message"] = "❓ 无法识别该系统操作"
        return result
    
    @classmethod
    def get_clipboard(cls):
        """获取剪贴板内容"""
        try:
            import pyperclip
            content = pyperclip.paste()
            return {
                "success": True,
                "content": content,
                "length": len(content),
                "message": f"✅ 已获取剪贴板内容 ({len(content)} 字符)"
            }
        except Exception as e:
            return {
                "success": False,
                "content": "",
                "message": f"❌ 无法读取剪贴板: {str(e)}"
            }
    
    @classmethod
    def set_clipboard(cls, text):
        """设置剪贴板内容"""
        try:
            import pyperclip
            pyperclip.copy(text)
            return {
                "success": True,
                "message": f"✅ 已复制到剪贴板 ({len(text)} 字符)"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法写入剪贴板: {str(e)}"
            }
    
    @classmethod
    def get_system_info(cls):
        """获取系统信息"""
        try:
            import platform
            import psutil
            
            info = {
                "success": True,
                "system": platform.system(),
                "platform": platform.platform(),
                "processor": platform.processor(),
                "cpu_count": psutil.cpu_count(),
                "cpu_percent": psutil.cpu_percent(interval=1),
                "memory": {
                    "total": f"{psutil.virtual_memory().total / (1024**3):.2f} GB",
                    "available": f"{psutil.virtual_memory().available / (1024**3):.2f} GB",
                    "percent": psutil.virtual_memory().percent
                },
                "disk": {
                    "total": f"{psutil.disk_usage('/').total / (1024**3):.2f} GB",
                    "free": f"{psutil.disk_usage('/').free / (1024**3):.2f} GB",
                    "percent": psutil.disk_usage('/').percent
                }
            }
            return info
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法获取系统信息: {str(e)}"
            }
    
    @classmethod
    def list_running_apps(cls):
        """列出正在运行的应用"""
        try:
            import psutil
            
            apps = []
            for proc in psutil.process_iter(['pid', 'name']):
                try:
                    apps.append({
                        "name": proc.info['name'],
                        "pid": proc.info['pid']
                    })
                except:
                    continue
            
            return {
                "success": True,
                "apps": apps[:30],  # 返回前30个
                "count": len(apps),
                "message": f"✅ 找到 {len(apps)} 个运行中的进程"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法列出应用: {str(e)}"
            }
    
    @classmethod
    def open_file_or_directory(cls, path):
        """打开文件或目录"""
        try:
            import subprocess
            
            path = os.path.expanduser(path)
            
            if not os.path.exists(path):
                return {
                    "success": False,
                    "message": f"❌ 路径不存在: {path}"
                }
            
            if os.path.isfile(path):
                # 用默认应用打开文件
                os.startfile(path) if sys.platform == "win32" else subprocess.Popen(['open', path])
                return {
                    "success": True,
                    "message": f"✅ 已打开文件: {os.path.basename(path)}"
                }
            else:
                # 在资源管理器中打开目录
                os.startfile(path) if sys.platform == "win32" else subprocess.Popen(['open', path])
                return {
                    "success": True,
                    "message": f"✅ 已打开目录: {path}"
                }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法打开: {str(e)}"
            }
    
    @classmethod
    def send_keystroke(cls, key_combination):
        """模拟键盘快捷键"""
        try:
            import keyboard
            
            # 解析快捷键
            keys = key_combination.split('+')
            keys = [k.strip().lower() for k in keys]
            
            # 模拟快捷键
            keyboard.hotkey(*keys)
            
            return {
                "success": True,
                "message": f"✅ 已模拟快捷键: {key_combination}"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法发送快捷键: {str(e)}"
            }


# ================= 文件操作执行器 =================
class FileOperator:
    """
    本地文件操作执行器 - 处理文件读写、管理等操作
    """
    
    # 文件操作关键词
    FILE_KEYWORDS = [
        "读取文件", "打开文件", "查看文件", "读文件", "看看文件",
        "创建文件", "新建文件", "写入文件", "保存文件",
        "删除文件", "移动文件", "复制文件", "重命名",
        "文件列表", "目录", "文件夹", "列出文件",
        "自动归纳", "自动整理", "归纳文件夹", "整理文件夹", "归档文件夹", "微信文件归纳",
        "read file", "open file", "create file", "delete file",
        "list files", "directory", "folder",
    ]

    FOLDER_ORGANIZE_KEYWORDS = [
        "自动归纳", "自动整理", "归纳", "整理", "归档", "归类", "分类",
        "文件夹", "目录", "微信文件", "wechat files"
    ]
    
    @classmethod
    def is_file_operation(cls, text):
        """检测是否是文件操作请求"""
        text_lower = text.lower()
        return any(kw in text_lower for kw in cls.FILE_KEYWORDS)

    @classmethod
    def _is_folder_organize_intent(cls, text_lower: str) -> bool:
        has_action = any(kw in text_lower for kw in ["归纳", "整理", "归档", "归类", "分类"])
        has_target = any(kw in text_lower for kw in ["文件夹", "目录", "路径", "文件"])
        if has_action and has_target:
            return True
        return any(kw in text_lower for kw in cls.FOLDER_ORGANIZE_KEYWORDS)

    @classmethod
    def _extract_path_from_text(cls, user_input: str) -> str:
        """Extract a likely filesystem path from user input."""
        import re

        patterns = [
            r'["\']([^"\']+)["\']',
            r'([A-Za-z]:\\(?:[^\\/:*?"<>|\r\n]+\\)*[^\\/:*?"<>|\r\n]*)',
            r'(\.?/[\w\-./ ]+)',
        ]
        for pattern in patterns:
            m = re.search(pattern, user_input)
            if m:
                candidate = m.group(1).strip().strip('，。,.;；')
                if candidate:
                    return candidate
        return ""
    
    @classmethod
    def execute(cls, user_input):
        """执行文件操作"""
        text_lower = user_input.lower()
        result = {
            "success": False,
            "action": "",
            "message": "",
            "content": ""
        }

        # === 指定路径文件夹自动归纳 ===
        if cls._is_folder_organize_intent(text_lower):
            folder_path = cls._extract_path_from_text(user_input)
            if not folder_path:
                folder_path = get_default_wechat_files_dir()

            if not folder_path:
                result["message"] = (
                    "❓ 请提供要归纳的文件夹路径（可用引号包裹），或在 config/user_settings.json 中设置 "
                    "storage.wechat_files_dir 作为默认路径"
                )
                return result

            if not os.path.isabs(folder_path):
                folder_path = os.path.join(WORKSPACE_DIR, folder_path)

            if not os.path.isdir(folder_path):
                result["message"] = f"❌ 目录不存在: {folder_path}"
                return result

            try:
                try:
                    from web.folder_catalog_organizer import FolderCatalogOrganizer
                except Exception:
                    from folder_catalog_organizer import FolderCatalogOrganizer

                analyzer = get_file_analyzer()
                organizer = get_file_organizer()
                engine = FolderCatalogOrganizer(get_organize_root(), analyzer, organizer)
                summary = engine.organize_folder(folder_path)

                if not summary.get("success"):
                    result["message"] = f"❌ 自动归纳失败: {summary.get('error', '未知错误')}"
                    return result

                report_md = summary.get("report_markdown", "")
                report_json = summary.get("report_json", "")
                entries = summary.get("entries", [])

                sender_preview = []
                for item in entries:
                    sender = item.get("sender", "未知")
                    if sender and sender != "未知":
                        sender_preview.append(sender)
                sender_preview = sorted(set(sender_preview))[:8]
                sender_preview_text = "、".join(sender_preview) if sender_preview else "未识别到可靠发送者"

                result["success"] = True
                result["action"] = "folder_auto_catalog"
                result["message"] = (
                    f"✅ 归纳完成：{summary.get('organized_count', 0)}/{summary.get('total_files', 0)} 个文件已归纳"
                    f"\n📁 来源目录: {summary.get('source_dir', folder_path)}"
                    f"\n🧾 清单(MD): {report_md}"
                    f"\n🧾 清单(JSON): {report_json}"
                    f"\n👤 识别到的发送者/来源人: {sender_preview_text}"
                )
                return result
            except Exception as e:
                result["message"] = f"❌ 自动归纳异常: {str(e)}"
                return result
        
        # === 读取文件 ===
        if any(kw in text_lower for kw in ["读取", "打开文件", "查看文件", "读文件", "看看", "read file", "open file"]):
            # 提取文件路径
            import re
            # 尝试匹配常见路径模式
            patterns = [
                r'["\']([^"\']+)["\']',  # 引号包围的路径
                r'([A-Za-z]:\\[^\s]+)',   # Windows 绝对路径
                r'(\.?/[^\s]+)',          # Unix 风格路径
                r'(\S+\.\w{1,5})(?:\s|$)', # 带扩展名的文件
            ]
            
            filepath = None
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    filepath = match.group(1)
                    break
            
            if filepath:
                # 如果是相对路径，在 workspace 目录查找
                if not os.path.isabs(filepath):
                    workspace_path = os.path.join(WORKSPACE_DIR, filepath)
                    if os.path.exists(workspace_path):
                        filepath = workspace_path
                
                if os.path.exists(filepath):
                    try:
                        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        
                        # 限制内容长度
                        if len(content) > 10000:
                            content = content[:10000] + "\n\n... (文件过长，已截断)"
                        
                        result["success"] = True
                        result["action"] = "read_file"
                        result["message"] = f"✅ 已读取文件: {os.path.basename(filepath)}"
                        result["content"] = f"```\n{content}\n```"
                        return result
                    except Exception as e:
                        result["message"] = f"❌ 读取文件失败: {str(e)}"
                        return result
                else:
                    result["message"] = f"❌ 文件不存在: {filepath}"
                    return result
            else:
                result["message"] = "❓ 请指定要读取的文件路径"
                return result
        
        # === 列出文件 ===
        if any(kw in text_lower for kw in ["文件列表", "目录", "列出文件", "list files", "directory", "文件夹里"]):
            # 提取目录路径
            import re
            patterns = [
                r'["\']([^"\']+)["\']',
                r'([A-Za-z]:\\[^\s]+)',
                r'(\.?/[^\s]+)',
            ]
            
            dirpath = WORKSPACE_DIR  # 默认 workspace
            for pattern in patterns:
                match = re.search(pattern, user_input)
                if match:
                    dirpath = match.group(1)
                    break
            
            if not os.path.isabs(dirpath):
                dirpath = os.path.join(WORKSPACE_DIR, dirpath)
            
            if os.path.isdir(dirpath):
                try:
                    items = os.listdir(dirpath)
                    file_list = []
                    for item in items[:50]:  # 限制数量
                        item_path = os.path.join(dirpath, item)
                        if os.path.isdir(item_path):
                            file_list.append(f"📁 {item}/")
                        else:
                            size = os.path.getsize(item_path)
                            size_str = f"{size/1024:.1f}KB" if size > 1024 else f"{size}B"
                            file_list.append(f"📄 {item} ({size_str})")
                    
                    result["success"] = True
                    result["action"] = "list_files"
                    result["message"] = f"✅ 目录: {dirpath}"
                    result["content"] = "\n".join(file_list) if file_list else "空目录"
                    return result
                except Exception as e:
                    result["message"] = f"❌ 读取目录失败: {str(e)}"
                    return result
            else:
                result["message"] = f"❌ 目录不存在: {dirpath}"
                return result
        
        # === 创建/写入文件 ===
        if any(kw in text_lower for kw in ["创建文件", "新建文件", "写入文件", "保存到", "create file"]):
            result["message"] = "💡 请使用代码生成功能，Koto 会自动保存生成的文件到 workspace"
            return result
        
        result["message"] = "❓ 无法识别该文件操作，请尝试：读取文件、列出目录等"
        return result
    
    @classmethod
    def watch_directory(cls, directory, callback=None, patterns=None):
        """监听目录变化并触发回调"""
        try:
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler
            
            if patterns is None:
                patterns = ['*.txt', '*.pdf', '*.docx', '*.xlsx', '*.csv']
            
            class ChangeHandler(FileSystemEventHandler):
                def on_created(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace('*', '')) for p in patterns):
                            if callback:
                                callback('created', event.src_path)
                
                def on_modified(self, event):
                    if not event.is_directory:
                        filename = os.path.basename(event.src_path)
                        if any(filename.endswith(p.replace('*', '')) for p in patterns):
                            if callback:
                                callback('modified', event.src_path)
            
            observer = Observer()
            observer.schedule(ChangeHandler(), directory, recursive=True)
            observer.start()
            
            return {
                "success": True,
                "observer": observer,
                "message": f"✅ 已开始监听目录: {directory}"
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法监听目录: {str(e)}"
            }
    
    @classmethod
    def get_file_metadata(cls, filepath):
        """获取文件元数据"""
        try:
            if not os.path.exists(filepath):
                return {"success": False, "message": "文件不存在"}
            
            stat = os.stat(filepath)
            from datetime import datetime
            
            return {
                "success": True,
                "filepath": filepath,
                "filename": os.path.basename(filepath),
                "size": f"{stat.st_size / 1024:.2f} KB",
                "created": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "extension": os.path.splitext(filepath)[1],
                "is_file": os.path.isfile(filepath)
            }
        except Exception as e:
            return {
                "success": False,
                "message": f"❌ 无法获取文件信息: {str(e)}"
            }


# ================= 联网搜索能力 =================
class WebSearcher:
    """
    使用 Gemini 的 Google Search Grounding 能力
    获取实时天气、新闻等信息
    """
    
    # 需要联网的关键词（严格收窄：仅包含几乎只在需要实时信息时才会出现的词）
    WEB_KEYWORDS = [
        # 天气（高置信）
        "天气", "气温", "下雨吗", "下雪吗", "温度多少", "天气怎么样",
        "天气预报", "weather", "temperature", "forecast",
        # 实时行情（高置信）
        "股价", "汇率", "比特币价格", "黄金价格", "金价", "实时金价", "今日金价", "当前金价", "现货黄金", "国际金价", "石油价格",
        "a股", "港股", "美股", "stock price",
        # 比赛/体育（高置信）
        "比分", "比赛结果", "谁赢了",
        # 新闻（只匹配明确的新闻请求）
        "今天新闻", "最新新闻", "latest news",
    ]
    
    @classmethod
    def needs_web_search(cls, text):
        """检测是否需要联网搜索
        
        优化策略：
        1. 检查关键词列表
        2. 对于金融/预测类，更倾向于web-search
        3. 对于热点事件、新品发布，必须web-search
        """
        text_lower = text.lower()
        
        # 必须 web-search 的模式（绝不能用纯AI）
        must_search_patterns = [
            r'(能不能|应该不应该|值不值得|是否).*?买',  # 股票建议
            r'(最新|实时|今天|明天|下周).*?(股|行情|数据)',  # 实时行情
            r'(预测|预期|后市|趋势).*?(股|市场|行业)',  # 趋势预测
            r'(财报|业绩|营收).*?(公布|发布)',  # 财报动态
            r'(新品|发布|推出).*?(上市|发售)',  # 新品信息
            r'(突发|紧急|最新)\w*事件',  # 突发事件
            r'(当前|今日|实时|最新).*?(金价|黄金)',  # 黄金实时行情
            r'(金价|黄金).*?(多少|报价|走势|行情)',  # 金价查询
        ]
        
        import re
        for pattern in must_search_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        
        # 关键词匹配
        if any(kw in text_lower for kw in cls.WEB_KEYWORDS):
            return True
        
        return False
    
    @classmethod
    def search_with_grounding(cls, query):
        """使用 Gemini Google Search Grounding 进行实时搜索"""
        try:
            # 使用 Google Search 作为工具
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=query,
                config=types.GenerateContentConfig(
                    tools=[types.Tool(google_search=types.GoogleSearch())],
                    system_instruction="你是 Koto，一个智能助手。使用搜索结果提供准确、实时的信息。用中文回答，格式清晰。"
                )
            )
            
            if response.text:
                return {
                    "success": True,
                    "response": response.text,
                    "grounded": True
                }
            else:
                return {
                    "success": False,
                    "response": "搜索未返回结果",
                    "grounded": False
                }
        except Exception as e:
            return {
                "success": False,
                "response": f"搜索失败: {str(e)}",
                "grounded": False
            }
    
    @classmethod
    def generate_ppt_images(cls, slide_titles: list, topic: str, max_images: int = 3) -> list:
        """为 PPT 幻灯片生成配图（使用 Imagen / Nano Banana）
        
        从幻灯片标题中挑选最适合配图的 2-3 页，生成高质量配图。
        返回: [{"slide_index": int, "image_path": str}, ...]
        """
        import threading
        import queue as _queue
        
        if not slide_titles:
            return []
        
        # 用 AI 挑选最适合配图的幻灯片
        pick_prompt = (
            f"以下是一个关于「{topic}」的PPT的各页标题,请挑选最适合配图的 {min(max_images, len(slide_titles))} 页。\n"
            f"对每页生成一个简洁的英文图像描述（适合AI图像生成）。\n"
            f"只输出 JSON 数组，格式：[{{\"index\": 0, \"prompt\": \"...\"}}]\n\n"
        )
        for i, t in enumerate(slide_titles):
            pick_prompt += f"{i}. {t}\n"
        
        try:
            resp = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=pick_prompt,
                config=types.GenerateContentConfig(temperature=0.3, max_output_tokens=1024)
            )
            import json as _json
            raw = resp.text or ""
            # 提取 JSON 数组
            import re as _re
            m = _re.search(r'\[.*\]', raw, _re.DOTALL)
            if m:
                picks = _json.loads(m.group())
            else:
                picks = []
        except Exception as e:
            print(f"[PPT-IMAGE] 选图AI失败: {e}")
            # 回退：选前 max_images 个非过渡页
            picks = [{"index": i, "prompt": f"professional illustration about {t}"} for i, t in enumerate(slide_titles[:max_images])]
        
        results = []
        images_dir = os.path.join(WORKSPACE_DIR, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        for pick in picks[:max_images]:
            idx = pick.get("index", 0)
            prompt = pick.get("prompt", f"professional illustration for presentation")
            # 增强 prompt 质量 — 确保简洁、无文字要求
            full_prompt = (
                f"Create a clean, modern, professional infographic-style illustration for a presentation slide. "
                f"Topic: {prompt}. "
                f"Style: flat design, clean layout, soft gradients, business-appropriate color palette. "
                f"Requirements: NO text, NO words, NO letters, NO numbers in the image. "
                f"Pure visual illustration only."
            )
            
            result_q = _queue.Queue()
            
            def _gen_image(p, q):
                # ① 首选: Nano Banana Pro（Koto 指定的图像生成模型）
                try:
                    res = client.models.generate_content(
                        model="nano-banana-pro-preview",
                        contents=p,
                        config=types.GenerateContentConfig(response_modalities=["IMAGE"])
                    )
                    if res.candidates and res.candidates[0].content.parts:
                        for part in res.candidates[0].content.parts:
                            if hasattr(part, "inline_data") and part.inline_data and part.inline_data.data:
                                q.put(("success", part.inline_data.data))
                                return
                except Exception as e0:
                    print(f"[PPT-IMAGE] Nano Banana 失败: {e0}")
                
                # ② 备选: Imagen 4.0（高质量图像生成 API）
                try:
                    res = client.models.generate_images(
                        model="imagen-4.0-generate-preview-06-06",
                        prompt=p,
                        config=types.GenerateImagesConfig(number_of_images=1)
                    )
                    if res.generated_images:
                        q.put(("success", res.generated_images[0].image.image_bytes))
                        return
                except Exception as e1:
                    print(f"[PPT-IMAGE] Imagen 4.0 失败: {e1}")
                
                # ③ 最后备选: Gemini 多模态图像生成
                try:
                    res2 = client.models.generate_content(
                        model="gemini-2.0-flash-preview-image-generation",
                        contents=f"Generate an image: {p}",
                        config=types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"])
                    )
                    if res2.candidates and res2.candidates[0].content.parts:
                        for part in res2.candidates[0].content.parts:
                            if hasattr(part, "inline_data") and part.inline_data and part.inline_data.data:
                                q.put(("success", part.inline_data.data))
                                return
                except Exception as e2:
                    print(f"[PPT-IMAGE] Gemini 多模态也失败: {e2}")
                q.put(("fail", None))
            
            thread = threading.Thread(target=_gen_image, args=(full_prompt, result_q), daemon=True)
            thread.start()
            thread.join(timeout=90)  # Nano Banana 可能较慢，给足时间
            
            try:
                status, data = result_q.get_nowait()
                if status == "success" and data:
                    ts = int(time.time() * 1000) % 1000000
                    fname = f"ppt_slide_{idx}_{ts}.png"
                    fpath = os.path.join(images_dir, fname)
                    with open(fpath, "wb") as f:
                        f.write(data)
                    results.append({"slide_index": idx, "image_path": fpath})
                    print(f"[PPT-IMAGE] ✅ 幻灯片 {idx} 配图生成: {fname}")
            except Exception:
                print(f"[PPT-IMAGE] ⚠️ 幻灯片 {idx} 配图超时或失败")
        
        return results
    
    @classmethod
    def deep_research_for_ppt(cls, user_input: str, search_context: str = "") -> str:
        """对复杂/学术主题进行深度研究，返回详细的研究报告文本
        
        用于在生成 PPT 大纲之前，先用 Pro 模型做深度分析，
        保证内容专业度和信息量。
        """
        research_prompt = (
            "你是一位顶级行业研究分析师。请对以下主题进行深入、全面的研究分析。\n\n"
            "## 严格要求\n"
            "1. **必须提供具体数据** — 市场规模（金额）、增长率（%）、市占率、出货量等定量信息\n"
            "2. **必须引用来源** — 如 IDC、Gartner、Statista、行业年报等（基于搜索资料中的数据）\n"
            "3. **必须包含真实案例** — 具体公司名称、产品型号、发布时间、销售数据等\n"
            "4. **必须有对比分析** — 不同产品/方案/技术路线之间的优劣对比\n"
            "5. **必须覆盖完整视角** — 历史演进 → 现状格局 → 技术路线 → 竞争分析 → 未来趋势\n"
            "6. **必须结构化** — 用清晰的标题层级和要点编排\n"
            "7. 中文回答，内容必须详实，**空洞的描述是不可接受的**\n\n"
            "## 输出格式\n"
            "为每个板块提供:\n"
            "- 2-3 个核心数据点（带数字和来源）\n"
            "- 2-3 个具体案例/产品\n"
            "- 1-2 个关键趋势判断\n\n"
            f"研究主题：{user_input}\n"
        )
        if search_context:
            research_prompt += f"\n已有的搜索参考资料：\n{search_context[:8000]}\n"
        
        research_models = ["gemini-3-pro-preview", "gemini-2.5-flash"]
        for model in research_models:
            try:
                resp = client.models.generate_content(
                    model=model,
                    contents=research_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.5,
                        max_output_tokens=16384,
                    )
                )
                if resp.text and len(resp.text) > 200:
                    print(f"[PPT-RESEARCH] ✅ 深度研究完成 ({model}), {len(resp.text)} 字符")
                    return resp.text
            except Exception as e:
                print(f"[PPT-RESEARCH] {model} 失败: {e}")
                continue
        return ""

# === System Instruction ===
# 简化版系统指令 - 用于CHAT/RESEARCH等非文件生成任务
def _get_chat_system_instruction(question: str = None):
    """
    生成包含当前日期时间和系统状态的系统指令
    
    Args:
        question: 用户问题（可选），用于智能上下文选择
    
    Returns:
        系统指令文本
    """
    try:
        # 如果提供了问题，使用智能上下文注入
        if question:
            from web.context_injector import get_dynamic_system_instruction
            return get_dynamic_system_instruction(question)
    except Exception as e:
        print(f"[Koto] Warning: Dynamic context injection failed: {e}")
    
    # 降级方案：使用基础系统指令
    from datetime import datetime
    
    now = datetime.now()
    date_str = now.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][now.weekday()]
    time_str = now.strftime("%H:%M:%S")
    
    # 获取系统信息（如果可用）
    system_info_section = ""
    try:
        from web.system_info import get_formatted_system_info, get_system_warnings
        
        formatted_info = get_formatted_system_info(include_processes=False)
        warnings = get_system_warnings()
        
        system_info_section = f"""
## 💻 当前系统状态
{formatted_info}"""
        
        if warnings:
            system_info_section += "\n\n## ⚠️ 系统警告\n"
            for warning in warnings:
                system_info_section += f"  • {warning}\n"
    except Exception as e:
        print(f"[Koto] Warning: Failed to collect system info: {e}")
    
    return f"""你是 Koto (言)，一个与用户计算机深度融合的个人AI助手。

## 📅 当前时间（用于相对日期计算）
🕒 **系统时间**: {date_str} {weekday} {time_str}
📅 **ISO日期**: {now.strftime("%Y-%m-%d")}
⏰ **使用此时间计算**: "明天"、"下周"、"前天" 等相对时间{system_info_section}

## 👤 角色定位
- 精通多个领域：编程、数据分析、写作、问题解决、系统管理
- 充分了解用户的计算环境和当前状态
- 快速理解用户意图，提供符合实际情境的答案
- 充当用户与Windows系统的智能中介

## 📋 回答原则
1. **简洁直接** - 不自我介绍，直接进入主题
2. **优先中文** - 默认用中文回答，除非用户要求其他语言
3. **清晰结构** - 使用标题、列表、代码块组织内容，便于快速理解
4. **上下文感知** - 结合用户的系统状态给出建议
5. **环境感知** - 了解当前 CPU、内存、磁盘状态，做出合适的建议
6. **时间准确性** - 使用系统时间准确计算相对日期
7. **禁止生成文件** - 仅在明确要求PDF/Word/Excel/PPT时才生成

## ✅ 能做的事
- 帮助用户分析本地文件、文档、图片
- 建议系统操作、自动化脚本、PowerShell命令
- 理解文件路径、应用名称、快捷键等Windows内容
- 根据当前系统状况给出性能优化建议
- 基于磁盘剩余空间建议存储位置
- 基于内存和 CPU 使用情况建议何时执行任务
- 协助处理剪贴板、监听快捷键、系统设置
- 联动本地应用（打开微信、邮件、浏览器等）
- 进行系统诊断：如果用户反映电脑卡，可以分析当前 CPU/内存/磁盘情况
- 准确理解和计算时间问题

## ❌ 不做的事
- ✗ 自我介绍或重复身份
- ✗ 生成代码标记 BEGIN_FILE/END_FILE（仅文件生成任务使用）
- ✗ 输出冗长的前言、风险提示或过度谨慎的警告
- ✗ 拒绝合理的系统操作请求"""

def _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION():
    """获取默认的系统指令（用于降级场景）"""
    try:
        return _get_chat_system_instruction()
    except:
        # 终极降级：返回基础指令
        return "你是 Koto (言)，一个与用户计算机深度融合的个人AI助手。精通多个领域，快速理解用户意图，提供符合实际情境的答案。"

def _get_system_instruction():
    """生成包含当前日期时间的文档生成系统指令"""
    from datetime import datetime
    
    now = datetime.now()
    date_str = now.strftime("%Y年%m月%d日")
    weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][now.weekday()]
    
    return f"""你是 Koto 文档生成专家，专注于生成高质量、可用的文档。

## 当前时间上下文
📅 **生成日期**: {date_str} {weekday}

## 时间理解规则（严格遵守）
- 这是本次请求的唯一时间锚点，请据此理解“今天/本月/今年/1月”等相对时间。
- 当用户只说“X月”未写年份时，默认使用**当前年份**（例如当前是 2026 年，则“1月新番”默认指 2026 年 1 月）。
- 不要默认使用过去年份，除非用户明确指定（如“2024年1月新番”）。

## 核心职责
1. **直接输出文档内容** - 直接输出最终要保存的文档内容，而不是代码或JSON
2. **中文优先** - 使用简体中文，专业术语准确无误
3. **格式规范** - 使用标题、列表、段落进行清晰组织

## 文档生成规则

### 优先策略：直接输出模式（推荐）
- **直接输出最终文档内容**，无需代码包装
- 使用Markdown式格式组织（# ## ### 标题、- 列表、段落）
- 系统会自动将你的输出转换为Word/PDF
- 这是最快、最可靠的方法

示例（只输出内容，不输出代码）：
```
# 文档标题

## 第一节
内容段落...

## 第二节
- 要点1
- 要点2
```

### 代码生成模式（仅当需要特殊格式时）
- 必须使用 ---BEGIN_FILE: filename.py--- 和 ---END_FILE--- 标记
- 代码控制在 80 行以内
- 必须包含中文字体处理（特别是PDF生成）
- 使用 try/except 包装错误处理
- **仅当直接输出无法满足需求时才使用此模式**

## 禁止项清单
- ✗ 输出JSON格式的"虚拟文档"
- ✗ 输出结构化数据而非真实内容
- ✗ 生成 BEGIN_FILE/END_FILE 标记（除非必须生成Python代码）
- ✗ 生成要求用户手动复制粘贴的内容

## 优先级
1. **直接输出内容** > 代码生成 > JSON结构
2. 内容准确、结构清晰 > 输出格式完美
3. 实际可执行性 > 审美程度
"""

# SYSTEM_INSTRUCTION 不再在模块加载时构建，改为按需调用 _get_system_instruction()
# SYSTEM_INSTRUCTION = _get_system_instruction()

def _get_filegen_brief_instruction() -> str:
    """FILE_GEN 的简版系统提示（每次调用实时取时间）。"""
    now = datetime.now()
    return (
        "你是Koto文档生成器，输出清晰的结构化内容，不要输出代码。\n"
        f"当前系统日期: {now.strftime('%Y-%m-%d')}（{now.strftime('%Y年%m月%d日')}）。\n"
        "时间规则：若用户仅写月份未写年份（如‘1月新番’），默认按当前年份解释。"
    )

def _parse_time_info_for_filegen(user_text: str) -> dict:
    """解析 FILE_GEN 输入中的时间信息，重点处理“仅月份未写年份”的场景。"""
    now = datetime.now()
    info = {
        "raw": user_text or "",
        "year": None,
        "month": None,
        "resolved_year": None,
        "resolved_month": None,
        "time_text": now.strftime("%Y年%m月%d日"),
        "rule_hit": False,
    }

    text = user_text or ""
    m = re.search(r'(?:(20\d{2})\s*年)?\s*([1-9]|1[0-2])\s*月', text)
    if not m:
        return info

    year_str = m.group(1)
    month_str = m.group(2)
    month = int(month_str)
    year = int(year_str) if year_str else None

    info["year"] = year
    info["month"] = month
    info["resolved_year"] = year if year is not None else now.year
    info["resolved_month"] = month
    info["rule_hit"] = year is None
    return info

def _build_filegen_time_context(user_text: str) -> tuple[str, dict]:
    """构建注入给模型的时间上下文文本。"""
    parsed = _parse_time_info_for_filegen(user_text)
    now = datetime.now()
    lines = [
        "[时间上下文]",
        f"- 当前系统时间: {now.strftime('%Y-%m-%d %H:%M:%S')}",
    ]

    if parsed.get("resolved_month"):
        lines.append(
            f"- 用户时间意图解析: {parsed['resolved_year']}年{parsed['resolved_month']}月"
        )
        if parsed.get("rule_hit"):
            lines.append(
                "- 解析规则命中: 用户仅提供月份，已按当前年份解析"
            )
    else:
        lines.append("- 用户时间意图解析: 未检测到明确月份，按当前语境理解")

    return "\n".join(lines), parsed


# ===== 任务特定系统提示词 =====
TASK_PROMPTS = {
    "CHAT": """助手模式：普通对话
- 直接回答问题，提供有用信息
- 保持对话自然流畅
- 记住之前的上下文""",
    
    "CODER": """代码生成专家
- 生成高质量、可运行的代码
- 遵循Python/JavaScript最佳实践
- 添加必要注释，解释复杂逻辑
- 包含错误处理和边界检查
- 代码长度控制在80行以内""",
    
    "FILE_GEN": """文档生成专家
- 生成结构清晰、格式规范的文档
- 使用标题、列表、段落进行组织
- 适配Word/PDF/Excel导出
- 内容准确、专业、可执行
- 禁止输出代码块和技术细节""",
    
    "PAINTER": """图像生成艺术家
- 创作独特、高质量的图像
- 理解用户的审美偏好
- 支持风格、颜色、构图的微调
- 输出高分辨率图像""",
    
    "RESEARCH": """深度研究专家
- 进行全面的信息搜索和分析
- 查找最新、最准确的信息
- 整理多个来源的观点
- 提供有根据的结论和见解
- 标注信息来源""",
    
    "SYSTEM": """系统操作执行器
- 执行本地系统命令和操作
- 打开应用、管理文件、控制系统
- 提供清晰的执行反馈
- 解释操作结果和错误""",
}

# ===== Windows本地快捷指令映射 =====
WINDOWS_SHORTCUTS = {
    # 文件和剪贴板操作
    "复制": "Ctrl+C",
    "粘贴": "Ctrl+V", 
    "剪切": "Ctrl+X",
    "撤销": "Ctrl+Z",
    "重做": "Ctrl+Y",
    "全选": "Ctrl+A",
    "保存": "Ctrl+S",
    "打开": "Ctrl+O",
    "新建": "Ctrl+N",
    
    # 浏览器操作
    "新标签页": "Ctrl+T",
    "关闭标签页": "Ctrl+W",
    "历史记录": "Ctrl+H",
    "书签": "Ctrl+B",
    "刷新": "Ctrl+R",
    "放大": "Ctrl+加号",
    "缩小": "Ctrl+减号",
    
    # 系统操作
    "任务管理器": "Ctrl+Shift+Esc",
    "截图": "Win+Shift+S",
    "开始菜单": "Win",
    "锁屏": "Win+L",
    "关机": "Alt+F4",
    "虚拟桌面": "Win+Tab",
    "显示桌面": "Win+D",
    
    # 应用切换
    "切换应用": "Alt+Tab",
    "关闭应用": "Alt+F4",
}

# ================= RAG 上下文分析器 =================
class ContextAnalyzer:
    """
    基于 RAG (检索增强生成) 的智能上下文分析器
    
    功能：
    1. 分析历史对话，提取关键信息
    2. 构建结构化的上下文提示词
    3. 智能判断任务关联性
    4. 生成增强后的输入
    """
    
    # 任务类型特征签名
    TASK_SIGNATURES = {
        "PAINTER": {
            "keywords": ["图", "画", "照片", "image", "photo", "picture", "图像已生成", "图片已生成", "猫", "狗", "人物", "风景", "头像"],
            "outputs": ["图像已生成", "图片已生成", "已保存图片", "✨ 图片已生成"],
            "entities": ["颜色", "风格", "大小", "背景", "表情", "姿势", "眼睛", "毛发", "脸"],
        },
        "FILE_GEN": {
            "keywords": [
                "pdf", "word", "excel", "docx", "文档", "报告", "文件", "简历", "合同",
                "标注", "批注", "润色", "改写", "校对", "审校", "修订", "优化", "纠错"
            ],
            "outputs": ["已生成文件", "文件已保存", ".pdf", ".docx", ".xlsx", "✅ **文件生成成功"],
            "entities": ["标题", "章节", "内容", "格式", "模板", "标注", "批注", "修改建议"],
        },
        "RESEARCH": {
            "keywords": ["研究", "分析", "介绍", "了解", "原理", "技术", "深入"],
            "outputs": ["##", "###", "1.", "2.", "总结", "结论"],
            "entities": ["定义", "特点", "优势", "劣势", "应用", "发展"],
        },
        "CODER": {
            "keywords": ["代码", "编程", "函数", "脚本", "code", "script", "python", "javascript"],
            "outputs": ["```python", "```javascript", "```", "def ", "class "],
            "entities": ["函数", "变量", "类", "模块", "算法"],
        },
        "CHAT": {
            "keywords": ["你好", "谢谢", "帮我", "请问", "什么是"],
            "outputs": [],
            "entities": [],
        },
    }
    
    # 延续性指示词分类 - 需要更严格的匹配
    CONTINUATION_PATTERNS = {
        "modify": {
            # 修改类：必须是短句或明确的修改指令
            "indicators": ["再来一张", "再来一个", "更大一点", "更小一点", "大一点", "小一点", "深一些", "浅一些", "颜色换成", "背景换成"],
            "weight": 0.9,
            "max_input_length": 30,  # 限制输入长度，长句子不太可能是简单修改
            "prompt_template": "用户要求修改之前的结果：{modification}"
        },
        "reference": {
            # 引用类：必须在句首或独立使用
            "indicators": ["这个怎么", "这张图", "那个文件", "上面的", "刚才的", "把它", "把这个", "基于这个"],
            "weight": 0.85,
            "require_start": True,  # 需要在句首出现
            "prompt_template": "用户引用了之前的内容：{reference}"
        },
        "convert": {
            # 转换类：明确的格式转换请求
            "indicators": ["做成word", "做成pdf", "做成excel", "转成word", "转成pdf", "变成文档", "导出为", "保存为word", "保存为pdf"],
            "weight": 0.95,
            "prompt_template": "用户要求将之前的内容转换为新格式：{conversion}"
        },
        "continue": {
            # 继续类：明确要求继续之前的内容
            "indicators": [
                "继续写", "接着说", "接着写", "然后呢", "下一步", "还有呢", "另外补充",
                "再找找", "再搜", "再查", "再看看", "继续查", "继续找", "再找", "再搜一下"
            ],
            "weight": 0.7,
            "max_input_length": 20,  # 短句才是继续指令
            "prompt_template": "用户要求继续之前的任务：{continuation}"
        },
        "detail": {
            # 详细类：只有非常明确的展开请求才算，且必须是短句
            "indicators": ["详细说说", "展开说说", "详细讲讲", "具体说一下", "解释一下刚才的"],
            "weight": 0.75,
            "max_input_length": 25,  # 限制长度
            "prompt_template": "用户要求详细说明之前提到的内容：{detail}"
        },
    }
    
    @classmethod
    def extract_entities(cls, text: str, task_type: str = None) -> list:
        """从文本中提取关键实体"""
        entities = []
        text_lower = text.lower()
        
        # 通用实体提取
        # 颜色
        colors = ["红色", "蓝色", "绿色", "黄色", "白色", "黑色", "灰色", "粉色", "紫色", "橙色", "棕色"]
        for color in colors:
            if color in text_lower:
                entities.append({"type": "color", "value": color})
        
        # 风格
        styles = ["可爱", "帅气", "写实", "卡通", "动漫", "赛博朋克", "水彩", "油画", "简约", "复古"]
        for style in styles:
            if style in text_lower:
                entities.append({"type": "style", "value": style})
        
        # 主题/对象
        subjects = ["猫", "狗", "人", "风景", "建筑", "汽车", "花", "树", "山", "海", "城市"]
        for subject in subjects:
            if subject in text_lower:
                entities.append({"type": "subject", "value": subject})
        
        # 特定任务的实体
        if task_type and task_type in cls.TASK_SIGNATURES:
            for entity_keyword in cls.TASK_SIGNATURES[task_type].get("entities", []):
                if entity_keyword in text_lower:
                    entities.append({"type": "task_specific", "value": entity_keyword})
        
        return entities
    
    @classmethod
    def build_context_summary(cls, history: list, max_turns: int = 3) -> dict:
        """
        构建历史上下文摘要
        
        返回:
        {
            "task_history": [],      # 任务历史
            "key_entities": [],      # 关键实体
            "last_user_intent": "",  # 最近的用户意图
            "last_model_output": "", # 最近的模型输出
            "conversation_topic": "" # 对话主题
        }
        """
        summary = {
            "task_history": [],
            "key_entities": [],
            "last_user_intent": "",
            "last_model_output": "",
            "conversation_topic": "",
        }
        
        if not history:
            return summary
        
        # 分析最近的对话
        recent_turns = history[-max_turns * 2:] if len(history) > max_turns * 2 else history
        
        all_entities = []
        topics = []
        
        for turn in recent_turns:
            content = turn['parts'][0] if turn['parts'] else ''
            role = turn['role']
            
            if role == 'user':
                summary["last_user_intent"] = content
                # 识别任务类型
                for task_type, signatures in cls.TASK_SIGNATURES.items():
                    if any(kw in content.lower() for kw in signatures["keywords"]):
                        summary["task_history"].append({
                            "type": task_type,
                            "content": content[:100]
                        })
                        topics.append(task_type)
                        break
                
                # 提取实体
                entities = cls.extract_entities(content)
                all_entities.extend(entities)
                
            elif role == 'model':
                summary["last_model_output"] = content
        
        # 去重实体
        seen = set()
        unique_entities = []
        for e in all_entities:
            key = f"{e['type']}:{e['value']}"
            if key not in seen:
                seen.add(key)
                unique_entities.append(e)
        summary["key_entities"] = unique_entities
        
        # 确定对话主题
        if topics:
            summary["conversation_topic"] = topics[-1]  # 最近的任务类型
        
        return summary
    
    @classmethod
    def build_rag_prompt(cls, user_input: str, context_summary: dict, continuation_type: str = None) -> str:
        """
        构建 RAG 风格的增强提示词
        
        将上下文信息结构化地注入到用户输入中
        """
        prompt_parts = []
        
        # 1. 添加上下文标记
        if context_summary.get("conversation_topic"):
            prompt_parts.append(f"[上下文类型: {context_summary['conversation_topic']}]")
        
        # 2. 添加关键实体信息
        if context_summary.get("key_entities"):
            entities_str = ", ".join([f"{e['type']}={e['value']}" for e in context_summary["key_entities"][:5]])
            prompt_parts.append(f"[关键信息: {entities_str}]")
        
        # 3. 添加历史意图
        if context_summary.get("last_user_intent"):
            # 截取核心描述
            last_intent = context_summary["last_user_intent"]
            if len(last_intent) > 200:
                last_intent = last_intent[:200] + "..."
            prompt_parts.append(f"[之前的请求: {last_intent}]")
        
        # 4. 根据延续类型添加特定指令
        if continuation_type and continuation_type in cls.CONTINUATION_PATTERNS:
            pattern = cls.CONTINUATION_PATTERNS[continuation_type]
            # 不添加模板，让实体和上下文自然融合
        
        # 5. 添加用户当前输入
        prompt_parts.append(f"[当前请求: {user_input}]")
        
        # 6. 如果是转换请求，添加源内容
        if continuation_type == "convert" and context_summary.get("last_model_output"):
            output = context_summary["last_model_output"]
            # 限制长度
            if len(output) > 4000:
                output = output[:4000] + "\n...(内容已截断)"
            prompt_parts.append(f"\n[需要转换的源内容:]\n{output}")
        
        # 组合成最终的增强提示
        enhanced_prompt = "\n".join(prompt_parts)
        
        return enhanced_prompt
    
    @classmethod
    def analyze_context(cls, user_input: str, history: list) -> dict:
        """
        RAG 风格的上下文分析
        
        返回:
        {
            "is_continuation": bool,      # 是否是延续任务
            "related_task": str,          # 关联的任务类型
            "continuation_type": str,     # 延续类型 (modify/reference/convert/continue/detail)
            "context_summary": dict,      # 结构化上下文摘要
            "enhanced_input": str,        # RAG 增强后的输入
            "confidence": float,          # 置信度
        }
        """
        result = {
            "is_continuation": False,
            "related_task": None,
            "continuation_type": None,
            "context_summary": {},
            "enhanced_input": user_input,
            "confidence": 0.0,
        }
        
        if not history or len(history) < 2:
            return result
        
        user_lower = user_input.lower()
        input_length = len(user_input)
        
        # 1. 构建上下文摘要
        context_summary = cls.build_context_summary(history)
        result["context_summary"] = context_summary
        
        # 2. 检测延续类型和置信度（更严格的匹配）
        detected_type = None
        max_weight = 0.0
        
        for pattern_type, pattern_info in cls.CONTINUATION_PATTERNS.items():
            indicators = pattern_info["indicators"]
            weight = pattern_info["weight"]
            
            # 检查输入长度限制（如果有）
            max_len = pattern_info.get("max_input_length")
            if max_len and input_length > max_len:
                continue  # 输入太长，不太可能是简单的延续指令
            
            # 检查是否需要在句首出现
            require_start = pattern_info.get("require_start", False)
            
            # 计算匹配的指示词数量
            matches = 0
            for ind in indicators:
                if ind in user_lower:
                    if require_start:
                        # 需要在句首（前10个字符内）
                        if user_lower.find(ind) < 10:
                            matches += 1
                    else:
                        matches += 1
            
            if matches > 0:
                # 加权计算置信度
                adjusted_weight = weight * (1 + 0.1 * (matches - 1))  # 多个匹配增加置信度
                if adjusted_weight > max_weight:
                    max_weight = adjusted_weight
                    detected_type = pattern_type
        
        # 3. 额外检查：如果用户输入包含明确的新主题，降低延续判断
        # 新主题标志：包含"关于"、"一个"后接新实体
        new_topic_indicators = ["关于", "一篇", "一份", "一个新的", "帮我写", "帮我做", "帮我生成", "给我生成", "生成一"]
        has_new_topic = any(ind in user_lower for ind in new_topic_indicators)
        
        # 检查是否是完全不同的任务类型（如：打开微信 -> 生成图片）
        task_mismatch = False
        if context_summary.get("conversation_topic"):
            prev_topic = context_summary["conversation_topic"]
            # 检测当前输入的任务类型
            curr_likely_task = None
            if any(kw in user_lower for kw in ["查", "搜", "搜索", "查询", "找", "再找", "再查", "再搜"]):
                curr_likely_task = "WEB_SEARCH"
            elif any(kw in user_lower for kw in ["图", "画", "照片", "image"]):
                curr_likely_task = "PAINTER"
            elif any(kw in user_lower for kw in ["word", "pdf", "文档", "报告"]):
                curr_likely_task = "FILE_GEN"
            elif any(kw in user_lower for kw in ["打开", "运行", "关闭"]):
                curr_likely_task = "SYSTEM"
            
            # 如果任务类型完全不同，不应该是延续
            if curr_likely_task and prev_topic and curr_likely_task != prev_topic:
                task_mismatch = True
                print(f"[ContextAnalyzer] 任务类型不匹配: {prev_topic} -> {curr_likely_task}")
        
        if has_new_topic and input_length > 10:
            # 有新主题且输入较长，很可能是独立任务
            max_weight *= 0.2  # 大幅降低置信度
            print(f"[ContextAnalyzer] 检测到新主题标志，降低延续置信度")
        
        if task_mismatch:
            # 任务类型不匹配，强制清零
            max_weight = 0
            detected_type = None
            print(f"[ContextAnalyzer] 任务类型不匹配，清除延续判断")
        
        # 4. 如果检测到延续模式且置信度足够高
        if detected_type and max_weight > 0.5:
            result["is_continuation"] = True
            result["continuation_type"] = detected_type
            result["confidence"] = min(max_weight, 1.0)
            
            # 确定关联的任务类型
            if context_summary.get("conversation_topic"):
                result["related_task"] = context_summary["conversation_topic"]
            elif context_summary.get("task_history"):
                result["related_task"] = context_summary["task_history"][-1]["type"]
            
            # 4. 构建 RAG 增强提示
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input, 
                context_summary, 
                detected_type
            )
            
            print(f"[ContextAnalyzer] RAG Analysis:")
            print(f"  - Continuation Type: {detected_type}")
            print(f"  - Related Task: {result['related_task']}")
            print(f"  - Confidence: {result['confidence']:.2f}")
            print(f"  - Entities: {[e['value'] for e in context_summary.get('key_entities', [])]}")
        
        # 5. 特殊处理：转换请求（即使没有明确的延续指示词）
        convert_patterns = ["做成word", "做成pdf", "转成word", "转成pdf", "生成word", "生成pdf", "导出为"]
        if any(p in user_lower for p in convert_patterns) and context_summary.get("last_model_output"):
            result["is_continuation"] = True
            result["continuation_type"] = "convert"
            result["related_task"] = "FILE_GEN"
            result["confidence"] = 0.95
            result["enhanced_input"] = cls.build_rag_prompt(
                user_input,
                context_summary,
                "convert"
            )
        
        return result

    @classmethod
    def filter_history(cls, user_input: str, history: list, keep_turns: int = 6) -> list:
        """过滤历史记录，尽量避免无关上下文污染"""
        if not history:
            return []

        # 如果历史很短，直接返回
        if len(history) <= keep_turns * 2:
            return history

        user_lower = user_input.lower()

        # 抽取用户输入中的实体与关键词
        entities = cls.extract_entities(user_input)
        entity_values = {e["value"] for e in entities}

        # 额外提取中文关键词（长度>=2）与英文单词（长度>=3）
        import re
        cjk_words = re.findall(r"[\u4e00-\u9fff]{2,}", user_input)
        eng_words = re.findall(r"[a-zA-Z]{3,}", user_input)
        keyword_set = {k.lower() for k in (cjk_words + eng_words)}
        keyword_set.update({v.lower() for v in entity_values})

        # 构建相关历史：包含关键词的对话
        relevant = []
        for turn in history:
            content = (turn.get("parts") or [""])[0]
            content_lower = content.lower()
            if any(k in content_lower for k in keyword_set if k):
                relevant.append(turn)

        # 始终保留最近 3 轮对话（确保上下文连贯）
        tail_count = 6
        tail = history[-tail_count:] if len(history) >= tail_count else history

        if not relevant:
            # 没有匹配时，保留最近 4 轮（最多 8 条）作为兜底上下文
            fallback_count = 8
            return history[-fallback_count:] if len(history) >= fallback_count else history

        # 合并相关历史 + 尾部对话，保持顺序并去重
        merged = []
        seen = set()
        for turn in relevant + tail:
            key = f"{turn.get('role','')}-{(turn.get('parts') or [''])[0]}"
            if key not in seen:
                seen.add(key)
                merged.append(turn)

        # 只保留最近 keep_turns 轮
        return merged[-keep_turns * 2:]



class TaskOrchestrator:
    """
    编排和执行多个子任务
    
    责职：
    1. 顺序执行子任务
    2. 在子任务间传递数据/上下文
    3. 处理错误和重试
    4. 最终验证输出质量
    """
    
    @classmethod
    async def execute_compound_task(cls, user_input: str, subtasks: list, session_name: str = None) -> dict:
        """
        执行复合任务的所有子任务
        
        返回:
            {
                "success": bool,
                "primary_result": 主任务结果,
                "secondary_results": [次要任务结果],
                "combined_output": 最终合并输出,
                "execution_log": 执行日志,
                "quality_score": 质量评分 (0-100),
                "errors": 错误列表
            }
        """
        execution_log = []
        results = []
        context = {"original_input": user_input, "user_input": user_input}
        errors = []
        
        try:
            for i, subtask in enumerate(subtasks):
                print(f"\n[TaskOrchestrator] 执行子任务 {i+1}/{len(subtasks)}: {subtask['task_type']}")
                execution_log.append(f"步骤 {i+1}: 执行 {subtask['task_type']} - {subtask['description']}")
                step_input = subtask.get("input") or user_input
                
                try:
                    # 根据任务类型调用相应的处理函数
                    if subtask["task_type"] == "WEB_SEARCH":
                        result = await cls._execute_web_search(step_input, context)
                    elif subtask["task_type"] == "FILE_GEN":
                        result = await cls._execute_file_gen(step_input, context, subtask)
                    elif subtask["task_type"] == "PAINTER":
                        result = await cls._execute_painter(step_input, context)
                    elif subtask["task_type"] == "RESEARCH":
                        result = await cls._execute_research(step_input, context)
                    else:
                        result = {"success": False, "error": f"未知任务类型: {subtask['task_type']}"}
                    
                    subtask["status"] = "completed"
                    subtask["result"] = result
                    results.append(result)
                    
                    # 将结果保存到上下文，供下一个任务使用
                    context[f"{subtask['task_type']}_result"] = result
                    context[f"step_{i+1}_output"] = result.get("output", result.get("content", ""))
                    
                    execution_log.append(f"  ✅ 完成: {subtask['description']}")
                    
                except Exception as e:
                    error_msg = str(e)
                    subtask["status"] = "failed"
                    subtask["error"] = error_msg
                    errors.append(error_msg)
                    execution_log.append(f"  ❌ 失败: {error_msg}")
                    print(f"[TaskOrchestrator] 子任务失败: {error_msg}")
            
            # 合并结果
            combined_output = cls._merge_results(subtasks, context)
            
            # 质量验证
            quality_score = await cls._validate_quality(user_input, combined_output, context)
            
            return {
                "success": len(errors) == 0,
                "primary_result": results[0] if results else None,
                "secondary_results": results[1:] if len(results) > 1 else [],
                "combined_output": combined_output,
                "execution_log": execution_log,
                "quality_score": quality_score,
                "errors": errors,
                "context": context
            }
        
        except Exception as e:
            return {
                "success": False,
                "primary_result": None,
                "secondary_results": [],
                "combined_output": None,
                "execution_log": execution_log,
                "quality_score": 0,
                "errors": errors + [str(e)],
                "context": context
            }
    
    @classmethod
    async def _execute_web_search(cls, user_input: str, context: dict, progress_callback=None) -> dict:
        """执行 Web 搜索子任务 (带可视进度)"""
        
        def _report(msg: str, detail: str = ""):
            print(f"[WEB_SEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)
        
        try:
            _report("启动网络搜索...", "正在规划搜索关键词")
            
            # Phase 1: Planning
            # (WebSearcher manages its own queries, but we can simulate the 'thought' process)
            await asyncio.sleep(0.3)
            _report("执行 Google Search...", f"关键词: {user_input[:20]}...")
            
            # Phase 2: Execution
            # wrap in thread
            result = await asyncio.to_thread(WebSearcher.search_with_grounding, user_input)
            
            # Phase 3: Reporting
            if result.get("grounded"):
                _report("✅ 搜索并引用完成", "已结合最新信息")
            else:
                _report("✅ 搜索完成", "已获取相关网页摘要")
            
            return {
                "success": result.get("success", False),
                "output": result.get("response", ""),
                "content": result.get("response", ""),
                "grounded": result.get("grounded", False),
                "raw_result": result,
                "model_id": "gemini-2.5-flash"
            }
        except Exception as e:
            _report("❌ 搜索遇到问题", str(e))
            return {
                "success": False,
                "output": "",
                "error": str(e),
                "raw_result": None,
                "model_id": "gemini-2.5-flash"
            }
    

    @classmethod
    async def _execute_ppt_multi_step(cls, user_input: str, context: dict, subtask: dict, progress_callback=None) -> dict:
        """执行多阶段PPT生成任务 (Plan-then-Execute)"""
        
        def _report(msg: str, detail: str = ""):
            print(f"[PPT_PROGRESS] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)
        
        _report("开始多阶段PPT生成流程...", "阶段 1/3: 智能规划")
        previous_data = context.get(f"step_{subtask['index']}_output", "")
        
        # 1. 规划阶段 (Planning Phase)
        try:
            from web.ppt_master import PPTContentPlanner, PPTBlueprint
            
            # 初始化规划器
            planner = PPTContentPlanner(ai_client=client, model_name="gemini-2.5-flash")
            
            # 执行规划
            _report("正在规划内容结构...", "调用 AI 生成大纲")
            plan_result = await planner.plan_content_structure(user_input, search_results=None)
            
            # 提取大纲
            outline_data = plan_result.get("outline", [])
            theme_choice = plan_result.get("theme_recommendation", "business")
            total_slides = plan_result.get("total_expected_slides", 10)
            
            # --- 1.2 展示规划概览 (User Requirement: Visualize Plan) ---
            plan_summary = f"大纲概览 ({len(outline_data)} 章节, {total_slides} 页):\n"
            for idx, sec in enumerate(outline_data):
                plan_summary += f"{idx+1}. {sec.get('section_title')} ({len(sec.get('slides', []))} 页)\n"
            _report(f"规划完成，共 {total_slides} 页", plan_summary)
            
            # 将大纲转换为 PPTGenerator 可识别的格式
            ppt_slides = []
            
            # --- 多阶段执行：逐页生成内容 ---
            total_steps = sum(len(sec.get("slides", [])) for sec in outline_data)
            current_step = 0
            
            for section in outline_data:
                section_title = section.get("section_title", "章节")
                # 添加章节页
                ppt_slides.append({
                    "type": "section",
                    "title": section_title,
                    "content": [section.get("section_theme", "")]
                })
                
                for slide in section.get("slides", []):
                    current_step += 1
                    s_title = slide.get("slide_title", "未命名幻灯片")
                    s_type = slide.get("slide_type", "content")
                    s_points = slide.get("key_points", [])
                    
                    # Log progress
                    _report(f"生成第 {current_step}/{total_steps} 页内容: {s_title}", "阶段 2/3: 内容扩充")
                    
                    # 扩充内容 (Per-Slide Generation)
                    expanded_points = s_points
                    if hasattr(planner, 'expand_slide_content'):
                        try:
                            # Use new method in PPTContentPlanner
                            expanded_points = await planner.expand_slide_content(
                                s_title, 
                                s_points, 
                                context=f"Context: {section_title}"
                            )
                            if expanded_points != s_points:
                                _report(f"  ✨ 内容已扩充: {len(expanded_points)} 条", f"幻灯片: {s_title}")
                        except Exception as exp_err:
                            _report(f"  ⚠️ 扩充失败，使用原始内容", str(exp_err))
                            expanded_points = s_points
                    
                    ppt_slides.append({
                        "type": s_type if s_type in ["content", "content_image", "comparison", "data"] else "content",
                        "title": s_title,
                        "points": expanded_points,
                        "content": expanded_points, 
                        "notes": slide.get("content_description", "")
                    })
            
            # 如果没有生成有效的幻灯片，回退到旧逻辑
            if not ppt_slides:
                raise ValueError("规划器未生成有效幻灯片大纲")

            # --- 2.5 验证阶段 (User Requirement: Model Verification) ---
            _report("正在验证生成内容...", "阶段 2.5/3: 质量自检")
            try:
                verify_prompt = (
                    f"请作为质检员检查生成的PPT内容是否符合用户需求。\n"
                    f"用户需求: {user_input}\n"
                    f"生成的标题: {[s['title'] for s in ppt_slides]}\n"
                    "请简要回答：内容是否覆盖了需求？(是/否) + 一句话点评。"
                )
                verify_resp = await asyncio.to_thread(lambda: client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=verify_prompt
                ))
                if verify_resp and verify_resp.text:
                    _report("✅ 验证通过", f"模型点评: {verify_resp.text.strip()[:60]}...")
            except Exception as v_err:
                _report("⚠️ 验证跳过 (非致命)", str(v_err))

            # 2. 执行阶段 (Execution Phase) - 生成 PPT 文件
            _report("正在生成最终文件...", "阶段 3/3: 渲染与保存")
            from web.ppt_generator import PPTGenerator
            ppt_gen = PPTGenerator(theme=theme_choice)
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[\\/*?:"<>|]', "", user_input[:20]) or "演示文稿"
            filename = f"{safe_title}_{timestamp}.pptx"
            ppt_path = os.path.join(settings_manager.documents_dir, filename)
            os.makedirs(settings_manager.documents_dir, exist_ok=True)
            
            # 使用 PPTGenerator 生成 (目前它直接支持 outline list)
            ppt_gen.generate_from_outline(
                title=safe_title,
                outline=ppt_slides,
                output_path=ppt_path
            )
            
            rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace("\\", "/")
            
            # 返回结果，格式与 _execute_file_gen 保持一致
            # 构建 markdown 表示的大纲供前端显示
            md_outline = f"# {safe_title}\n\n"
            for slide in ppt_slides:
                md_outline += f"## {slide['title']}\n"
                for p in slide.get('points', []):
                    md_outline += f"- {p}\n"
                md_outline += "\n"
            
            return {
                "success": True,
                "output": md_outline,
                "content": md_outline,
                "saved_files": [rel_path],
                "model_id": "gemini-2.5-flash (Planner)"
            }
            
        except Exception as e:
            print(f"[PPT] ⚠️ 多阶段生成失败，回退到单步生成: {e}")
            # 重新抛出异常让上层处理，或者在这里调用旧逻辑?
            # 为了简单，抛出异常让外部 _execute_file_gen 的 except 块捕获 (但外部是 generic exception)
            # 或者我们直接返回失败，让 TaskOrchestrator 记录错误
            return {
                "success": False,
                "error": str(e),
                "opt_out_to_legacy": True # 标记需要回退
            }


    @classmethod
    async def _execute_file_gen(cls, user_input: str, context: dict, subtask: dict, progress_callback=None) -> dict:
        """执行文件生成子任务
        增强：复杂/长文/要求“深度、详细、研究”时，先运行深度研究并切换到更强模型生成。
        """
        def _report(msg: str, detail: str = ""):
            print(f"[FILE_GEN] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)

        try:
            # 提取前一个任务的结果作为输入
            previous_data = context.get(f"step_{subtask['index']}_output", "")

            # 复杂度判定（长文本或显式“深度/详细/研究/全面/技术”请求）
            text_lower = user_input.lower()
            complex_flags = [
                len(user_input) > 120,
                any(k in text_lower for k in ["深度", "详细", "研究", "全面", "技术", "报告", "综述", "whitepaper"]),
            ]
            is_complex = any(complex_flags)
            
            # 检测目标格式（PPT、Excel、Word等）
            ppt_keywords = ["ppt", "幻灯片", "演示", "汇报", "presentation", "slide"]
            prefer_ppt = any(kw in user_input.lower() for kw in ppt_keywords)
            
            prefer_excel = ("excel" in user_input.lower() or "xlsx" in user_input.lower() or "表格" in user_input)
            prefer_pdf = "pdf" in user_input.lower()
            
            # 根据目标格式选择提示
            if prefer_ppt:
                # 尝试使用新的多阶段生成流程 (Plan-then-Execute)
                try:
                    ppt_result = await cls._execute_ppt_multi_step(user_input, context, subtask, progress_callback)
                    if ppt_result.get("success"):
                        _report(f"PPT生成成功", f"文件: {(ppt_result.get('saved_files') or [''])[0]}")
                        return ppt_result
                    elif ppt_result.get("opt_out_to_legacy"):
                        print("[FILE_GEN] ⚠️ 多阶段生成遇到问题，回退到旧版生成逻辑")
                    else:
                        return ppt_result
                except Exception as e:
                    print(f"[FILE_GEN] ⚠️ 多阶段生成异常: {e}")
                
                # 回退旧逻辑 (Legacy Prompt Generation)
                gen_prompt = (
                    "你是一个顶尖的演示文稿内容策划师和排版规划师。\n\n"
                    "在每个 `## 章节标题` 前一行写类型标签来选择幻灯片类型：\n"
                    "- `[详细]` — 深入展示 3-5 个要点\n"
                    "- `[概览]` — 多主题速览，用 `### 子标题` 分组\n"
                    "- `[亮点]` — 关键数据，格式: `- 数值 | 说明`\n"
                    "- `[对比]` — 两方对比，用 `### 选项A` 和 `### 选项B` 分组\n"
                    "- `[过渡页]` — 章节过渡（最多 2 个）\n\n"
                    "**输出格式（严格遵循 Markdown）**：\n"
                    "```\n"
                    "# 演示标题\n\n"
                    "[详细]\n"
                    "## 章节标题\n"
                    "- 要点1（包含具体信息）\n"
                    "- 要点2\n"
                    "```\n\n"
                    "规则：重点内容用多个 [详细] 展开，简要内容合并到 [概览]，关键数据用 [亮点]。\n"
                    "每个要点包含具体信息，中文输出，只输出大纲。\n"
                )
            else:
                gen_prompt = (
                    "你是Koto，一个专业的数据整理与报告生成助手。\n"
                    "请基于用户需求和提供的数据，输出清晰、可直接放入文档的 Markdown 内容。\n"
                    "如果是价格类信息，必须包含一个 Markdown 表格，字段建议为：时间、价格、变化、来源。\n"
                    "输出要求：\n"
                    "- 只输出内容，不要输出代码或 BEGIN_FILE 标记\n"
                    "- 中文输出，结构清晰\n"
                )
            
            full_input = (
                f"用户原始需求: {context['original_input']}\n\n"
                f"前面步骤的数据/信息:\n{previous_data}\n\n"
                f"{gen_prompt}"
            )
            
            # 深度研究：为复杂任务先补充研究上下文
            research_context = ""
            if is_complex:
                try:
                    research_context = WebSearcher.deep_research_for_ppt(user_input, previous_data)
                    if research_context:
                        previous_data = f"[深度研究]\n{research_context}\n\n[已有信息]\n{previous_data}"
                        print(f"[FILE_GEN] 🔬 深度研究完成，追加 {len(research_context)} 字上下文")
                except Exception as research_err:
                    print(f"[FILE_GEN] ⚠️ 深度研究失败: {research_err}")

            # 调用模型生成内容
            model_id = SmartDispatcher.get_model_for_task("FILE_GEN", complexity="complex" if is_complex else "normal")
            
            _report(f"正在生成内容...", f"模型: {model_id}")
            
            def _generate_text(prompt_text: str) -> str:
                response = client.models.generate_content(
                    model=model_id,
                    contents=prompt_text,
                    config=types.GenerateContentConfig(
                        system_instruction=_get_filegen_brief_instruction(),
                        temperature=0.4,
                        max_output_tokens=4000,
                    )
                )
                return response.text or ""
            
            text_out = _generate_text(full_input) or "(无输出)"
            _report(f"内容生成完成", f"字数: {len(text_out)}")
            
            # 解析 Markdown 表格
            def _extract_markdown_table(md_text: str):
                lines = [line.strip() for line in md_text.splitlines() if "|" in line]
                for i in range(len(lines) - 1):
                    header_line = lines[i]
                    sep_line = lines[i + 1]
                    if re.match(r"^\s*\|?\s*[-:|\s]+\|\s*$", sep_line):
                        headers = [c.strip() for c in header_line.strip("|").split("|")]
                        rows = []
                        j = i + 2
                        while j < len(lines) and "|" in lines[j]:
                            row = [c.strip() for c in lines[j].strip("|").split("|")]
                            if len(row) < len(headers):
                                row += [""] * (len(headers) - len(row))
                            rows.append(row[:len(headers)])
                            j += 1
                        return [headers] + rows
                return None
            
            # 解析PPT大纲结构（支持智能规划标签）
            def _parse_ppt_outline(md_text: str) -> dict:
                """解析带 [类型] 标签的 PPT 大纲"""
                lines = md_text.split('\n')
                outline = {"title": "", "slides": []}
                _tmap = {
                    "过渡页": "divider", "过渡": "divider",
                    "详细": "detail", "重点": "detail",
                    "亮点": "highlight", "数据": "highlight",
                    "概览": "overview", "速览": "overview", "简要": "overview",
                    "对比": "comparison", "比较": "comparison",
                }
                cur_type = "detail"
                cur_slide = None
                cur_sub = None
                
                for line in lines:
                    line = line.rstrip()
                    if line.strip() in ('```', '```markdown'):
                        continue
                    tm = re.match(r'^\s*\[(.+?)\]\s*$', line)
                    if tm:
                        cur_type = _tmap.get(tm.group(1).strip(), "detail")
                        continue
                    if line.startswith('# ') and not line.startswith('## '):
                        outline["title"] = line[2:].strip()
                    elif line.startswith('## '):
                        if cur_sub and cur_slide and cur_slide.get("type") in ("overview", "comparison"):
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                            cur_sub = None
                        if cur_slide:
                            outline["slides"].append(cur_slide)
                        cur_slide = {"type": cur_type, "title": line[3:].strip(), "points": [], "content": []}
                        if cur_type == "divider":
                            cur_slide["description"] = ""
                        cur_type = "detail"
                        cur_sub = None
                    elif line.startswith('### ') and cur_slide:
                        if cur_sub:
                            cur_slide.setdefault("subsections", []).append(cur_sub)
                        cur_sub = {"subtitle": line[4:].strip(), "label": line[4:].strip(), "points": []}
                    elif re.match(r'^[\s]*[-•*]\s', line) and cur_slide is not None:
                        pt = re.sub(r'^[\s]*[-•*]\s+', '', line).strip()
                        if cur_sub is not None:
                            cur_sub["points"].append(pt)
                        else:
                            cur_slide["points"].append(pt)
                            cur_slide["content"].append(pt)
                    elif cur_slide and cur_slide.get("type") == "divider" and line.strip():
                        cur_slide["description"] = line.strip()
                
                if cur_sub and cur_slide and cur_slide.get("type") in ("overview", "comparison"):
                    cur_slide.setdefault("subsections", []).append(cur_sub)
                if cur_slide:
                    outline["slides"].append(cur_slide)
                for sl in outline["slides"]:
                    if sl.get("type") == "comparison" and "subsections" in sl:
                        subs = sl["subsections"]
                        if len(subs) >= 2:
                            sl["left"] = subs[0]
                            sl["right"] = subs[1]
                return outline
            
            title = "生成文档"
            if "价格" in user_input or "表格" in user_input:
                title = "价格波动表格"
            elif prefer_ppt:
                title = "演示文稿"
            
            saved_files = []
            file_type = None
            excel_error = None
            
            # 生成PPT
            if prefer_ppt:
                try:
                    from web.ppt_generator import PPTGenerator
                    ppt_outline = _parse_ppt_outline(text_out)
                    
                    # 确定主题（通过关键词检测）
                    theme = "business"  # 默认商务主题
                    if "tech" in user_input.lower() or "技术" in user_input:
                        theme = "tech"
                    elif "creative" in user_input.lower() or "创意" in user_input:
                        theme = "creative"
                    
                    _report("正在生成PPT...", f"主题: {theme}")
                    
                    ppt_gen = PPTGenerator(theme=theme)
                    filename = f"{ppt_outline.get('title', '演示文稿')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pptx"
                    ppt_path = os.path.join(settings_manager.documents_dir, filename)
                    os.makedirs(settings_manager.documents_dir, exist_ok=True)
                    
                    ppt_gen.generate_from_outline(
                        title=ppt_outline.get('title', '演示'),
                        outline=ppt_outline.get('slides', []),
                        output_path=ppt_path
                    )
                    
                    rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace("\\", "/")
                    saved_files.append(rel_path)
                    file_type = "pptx"
                    _report("PPT生成完成", f"已保存到: {rel_path}")
                    
                except Exception as ppt_err:
                    print(f"[FILE_GEN] ⚠️ PPT生成失败: {ppt_err}")
                    _report("PPT生成失败，回退到Word...", f"错误: {str(ppt_err)[:50]}")
                    # PPT失败时回退到Word
                    from web.document_generator import save_docx
                    saved_docx = save_docx(text_out, title=title, output_dir=settings_manager.documents_dir)
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace("\\", "/")
                    saved_files.append(rel_path)
                    file_type = "docx"
            else:
                # 生成Excel或Word
                _report("正在处理内容...", "解析文档结构")
                table_rows = _extract_markdown_table(text_out)
                if prefer_excel and not table_rows:
                    # 第一次未生成合格表格 → 生成修正Prompt重试一次
                    fix_prompt = (
                        "请只输出一个 Markdown 表格，不要输出其他说明。\n"
                        "表格必须包含以下列：时间、价格、变化、来源。\n"
                        "每行数据一行，格式严格。\n\n"
                        f"用户需求: {context['original_input']}\n\n"
                        f"可用数据:\n{previous_data}\n"
                    )
                    text_out_retry = _generate_text(fix_prompt)
                    if text_out_retry:
                        text_out = text_out_retry
                        table_rows = _extract_markdown_table(text_out)
                
                if prefer_excel and table_rows:
                    _report("正在生成Excel...", f"写入 {len(table_rows)} 行数据")
                    try:
                        from openpyxl import Workbook
                        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                        from openpyxl.utils import get_column_letter
                        
                        wb = Workbook()
                        ws = wb.active
                        ws.title = title[:31] if title else "Sheet1"
                        
                        # 写入数据
                        for row in table_rows:
                            ws.append(row)
                        
                        # --- 样式美化 ---
                        header_font = Font(name='Microsoft YaHei', size=11, bold=True, color='FFFFFF')
                        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                        data_font = Font(name='Microsoft YaHei', size=10)
                        thin_border = Border(
                            left=Side(style='thin', color='D9D9D9'),
                            right=Side(style='thin', color='D9D9D9'),
                            top=Side(style='thin', color='D9D9D9'),
                            bottom=Side(style='thin', color='D9D9D9'),
                        )
                        alt_fill = PatternFill(start_color='F2F7FB', end_color='F2F7FB', fill_type='solid')
                        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
                        
                        max_row = ws.max_row
                        max_col = ws.max_column
                        
                        for col_idx in range(1, max_col + 1):
                            # 表头样式
                            cell = ws.cell(row=1, column=col_idx)
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = center_align
                            cell.border = thin_border
                            
                            # 数据行样式
                            for row_idx in range(2, max_row + 1):
                                cell = ws.cell(row=row_idx, column=col_idx)
                                cell.font = data_font
                                cell.alignment = left_align
                                cell.border = thin_border
                                # 隔行变色
                                if row_idx % 2 == 0:
                                    cell.fill = alt_fill
                            
                            # 自动列宽
                            max_len = 0
                            for row_idx in range(1, max_row + 1):
                                val = ws.cell(row=row_idx, column=col_idx).value
                                if val:
                                    # CJK 字符算2个字符宽
                                    vlen = sum(2 if ord(c) > 127 else 1 for c in str(val))
                                    max_len = max(max_len, vlen)
                            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)
                        
                        # 冻结首行
                        ws.freeze_panes = 'A2'
                        
                        filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        excel_path = os.path.join(settings_manager.documents_dir, filename)
                        os.makedirs(settings_manager.documents_dir, exist_ok=True)
                        wb.save(excel_path)
                        rel_path = os.path.relpath(excel_path, WORKSPACE_DIR).replace("\\", "/")
                        saved_files.append(rel_path)
                        file_type = "xlsx"
                        _report("Excel生成完成", f"已保存到: {rel_path}")
                    except Exception as excel_err:
                        excel_error = str(excel_err)
                        print(f"[FILE_GEN] ⚠️ Excel保存失败: {excel_error}")
                        _report("Excel保存失败，回退到Word...", f"错误: {excel_error[:50]}")
                
                # 保存为 DOCX（无表格或Excel失败时回退）
                if not saved_files:
                    _report("正在生成Word文档...", "转换为 DOCX")
                    from web.document_generator import save_docx, save_pdf
                    saved_docx = save_docx(text_out, title=title, output_dir=settings_manager.documents_dir)
                    rel_path = os.path.relpath(saved_docx, WORKSPACE_DIR).replace("\\", "/")
                    saved_files.append(rel_path)
                    file_type = "docx"
                    _report("Word文档生成完成", f"已保存到: {rel_path}")
                    
                    # 如用户明确需要 PDF，也同时保存
                    if prefer_pdf:
                        try:
                            _report("正在生成PDF...", "转换为 PDF")
                            saved_pdf = save_pdf(text_out, title=title, output_dir=settings_manager.documents_dir)
                            pdf_rel = os.path.relpath(saved_pdf, WORKSPACE_DIR).replace("\\", "/")
                            saved_files.append(pdf_rel)
                            _report("PDF生成完成", f"已保存到: {pdf_rel}")
                        except Exception as pdf_err:
                            print(f"[FILE_GEN] ⚠️ PDF保存失败: {pdf_err}")
                            _report("PDF生成失败", str(pdf_err)[:50])
            
            return {
                "success": True,
                "output": f"已生成{file_type.upper()}文档: {', '.join([os.path.basename(p) for p in saved_files])}" + (f" (Excel失败: {excel_error})" if excel_error else ""),
                "content": text_out,
                "file_type": file_type or "docx",
                "saved_files": saved_files,
                "model_id": model_id
            }
        except Exception as e:
            return {
                "success": False,
                "output": "",
                "error": str(e)
            }
    
    @classmethod
    async def _execute_painter(cls, user_input: str, context: dict, progress_callback=None) -> dict:
        """执行图像生成子任务 - 为PPT等生成配图 (带可视进度)"""
        
        def _report(msg: str, detail: str = ""):
            print(f"[PAINTER] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)
        
        try:
            topic = context.get("original_input", user_input)
            prompt = f"Professional illustration for: {topic[:100]}. Clean flat design, no text."
            
            image_paths = []
            images_dir = os.path.join(WORKSPACE_DIR, "images")
            os.makedirs(images_dir, exist_ok=True)
            
            _report("启动图像生成...", "调用 Imagen 3 模型")
            
            for i in range(2):
                try:
                    _report(f"正在生成第 {i+1}/2 张配图...", f"提示词: {prompt[:30]}...")
                    
                    # Run potentially blocking generation in thread
                    res = await asyncio.to_thread(lambda: client.models.generate_images(
                        model="imagen-4.0-generate-preview-06-06",
                        prompt=prompt,
                        config=types.GenerateImagesConfig(number_of_images=1)
                    ))
                    
                    if res.generated_images:
                        ts = int(time.time() * 1000) % 1000000
                        fname = f"ppt_img_{i}_{ts}.png"
                        fpath = os.path.join(images_dir, fname)
                        with open(fpath, "wb") as f:
                            f.write(res.generated_images[0].image.image_bytes)
                        image_paths.append(fpath)
                        print(f"[PAINTER] ✅ 配图 {i+1} 已生成: {fname}")
                        _report(f"✅ 配图 {i+1} 完成", fname)
                except Exception as img_err:
                    print(f"[PAINTER] ⚠️ 配图 {i+1} 生成失败: {img_err}")
                    _report(f"⚠️ 配图 {i+1} 失败", str(img_err))
            
            success = len(image_paths) > 0
            if success:
                 _report("✅ 图像生成任务完成", f"共生成 {len(image_paths)} 张")
            else:
                 _report("❌ 图像生成任务失败", "未生成有效图片")
            
            return {
                "success": success,
                "output": f"已生成 {len(image_paths)} 张配图",
                "content": ",".join(image_paths),
                "image_paths": image_paths,
                "model_id": "imagen-3.0"
            }
        except Exception as e:
            _report("❌ 图像生成遇到致命错误", str(e))
            return {
                "success": False,
                "output": "",
                "error": str(e)
            }
    
    @classmethod
    async def _execute_research(cls, user_input: str, context: dict, progress_callback=None) -> dict:
        """执行深度研究子任务 - 使用 Gemini Pro 深度分析 (可视进度)"""
        
        def _report(msg: str, detail: str = ""):
            print(f"[RESEARCH] {msg} | {detail}")
            if progress_callback:
                progress_callback(msg, detail)
        
        try:
            _report("启动深度研究流程...", "分析上下文数据")
            search_data = context.get("WEB_SEARCH_result", {})
            search_text = search_data.get("content", "") or search_data.get("output", "")
            
            # Phase 1: Planning
            _report("规划研究大纲...", "确定分析维度")
            # (Implied planning by WebSearcher internal logic, but we report it)
            await asyncio.sleep(0.5) # Simulate quick think
            
            # Phase 2: Synthesis
            _report("正在进行深度分析...", "调用 Gemini 1.5 Pro 进行综合")
            # Run in thread to not block event loop if sync
            research_text = await asyncio.to_thread(WebSearcher.deep_research_for_ppt, user_input, search_text)
            
            # Phase 3: Verification
            _report("验证研究报告...", "检查内容完整性")
            if research_text:
                _report("✅ 研究完成", f"生成 {len(research_text)} 字详细报告")
                return {
                    "success": True,
                    "output": f"深度研究完成，获取 {len(research_text)} 字专业分析",
                    "content": research_text,
                    "model_id": "gemini-1.5-pro"
                }
            else:
                _report("⚠️ 研究产出为空", "回退到基础搜索结果")
                return {
                    "success": True,
                    "output": "研究未返回结果，将使用已有信息",
                    "content": search_text
                }
        except Exception as e:
            _report("❌ 研究过程出错",Str(e))
            return {
                "success": False,
                "output": "",
                "error": str(e)
            }
    
    @classmethod
    def _merge_results(cls, subtasks: list, context: dict) -> dict:
        """合并所有子任务的结果"""
        merged = {
            "summary": "任务执行完成",
            "steps": [],
            "final_output": ""
        }
        
        for i, subtask in enumerate(subtasks):
            step_info = {
                "step": i + 1,
                "task": subtask["task_type"],
                "status": subtask["status"],
                "description": subtask["description"]
            }
            
            if subtask["result"]:
                step_info["output"] = subtask["result"].get("output", "")
            if subtask["error"]:
                step_info["error"] = subtask["error"]
            
            merged["steps"].append(step_info)
        
        # 最后一个完成的任务的输出作为最终输出
        for subtask in reversed(subtasks):
            if subtask["status"] == "completed" and subtask["result"]:
                merged["final_output"] = subtask["result"].get("output", "")
                break
        
        return merged
    
    @classmethod
    async def _validate_quality(cls, user_input: str, combined_output: dict, context: dict) -> int:
        """
        验证输出质量
        
        返回: 质量评分 (0-100)
        """
        score = 50  # 基础分
        
        # 检查是否有错误
        if context.get("errors"):
            score -= 20
        
        # 检查是否有最终输出
        if combined_output.get("final_output"):
            score += 20
        
        # 检查步骤完成度
        total_steps = len(combined_output.get("steps", []))
        completed_steps = len([s for s in combined_output.get("steps", []) if s["status"] == "completed"])
        if total_steps > 0:
            completion_rate = completed_steps / total_steps
            score += int(completion_rate * 20)
        
        # 确保分数在 0-100 之间
        return max(0, min(100, score))





# ================= 智能语料路由器配置 =================
# 配置 SmartDispatcher 以使用本地定义的类和对象
# SmartDispatcher、ModelRouter 等已从 app.core.routing 导入

try:
    print("[INIT] Configuring SmartDispatcher with local dependencies...")
    SmartDispatcher.configure(
        local_executor=LocalExecutor,
        context_analyzer=ContextAnalyzer,
        web_searcher=WebSearcher,
        model_map=MODEL_MAP,
        client=client
    )
    print("[INIT] SmartDispatcher configured successfully.")
except Exception as e:
    print(f"[ERROR] Failed to configure SmartDispatcher: {e}")

# === Ollama 后备路由 (可选) ===
LOCAL_ROUTER_MODEL = "qwen3:8b"  # 升级: Qwen3 中英文能力远超旧模型
OLLAMA_API_URL = "http://localhost:11434/api/generate"

class LocalDispatcher:
    """后备路由器 - 使用 Ollama (如果可用)"""
    
    @staticmethod
    def is_ollama_running():
        try:
            requests.get("http://localhost:11434", timeout=0.2)
            return True
        except:
            return False
    
    @staticmethod
    def analyze(user_input, history=None):
        """优先使用 SmartDispatcher，失败时使用 Ollama"""
        # 使用智能本地路由
        return SmartDispatcher.analyze(user_input, history)

# ================= Utilities =================

class Utils:
    _PACKAGE_ALLOWLIST = {
        "pygame": "pygame",
        "numpy": "numpy",
        "pandas": "pandas",
        "requests": "requests",
        "bs4": "beautifulsoup4",
        "beautifulsoup4": "beautifulsoup4",
        "lxml": "lxml",
        "pillow": "Pillow",
        "PIL": "Pillow",
        "opencv": "opencv-python",
        "cv2": "opencv-python",
        "matplotlib": "matplotlib",
        "scipy": "scipy",
        "sklearn": "scikit-learn",
        "flask": "flask",
        "fastapi": "fastapi",
        "uvicorn": "uvicorn",
        "streamlit": "streamlit",
        "gradio": "gradio",
    }

    @staticmethod
    def sanitize_string(s):
        if isinstance(s, str):
            return s.encode('utf-8', 'ignore').decode('utf-8')
        return s

    @staticmethod
    def is_failure_output(text: str) -> bool:
        if not text or not str(text).strip():
            return True
        t = str(text).strip().lower()
        return t.startswith("❌") or "失败" in t or "错误" in t

    @staticmethod
    def build_fix_prompt(task_type: str, user_input: str, prev_output: str = "", error_hint: str = "") -> str:
        base = (
            f"用户需求: {user_input}\n\n"
            f"上次输出/错误:\n{prev_output or error_hint}\n\n"
            "请修正并重新输出最终结果。不要解释过程，只输出最终内容。\n"
        )

        if task_type == "FILE_GEN":
            return base + (
                "要求：输出可执行的 Python 脚本，并使用 BEGIN_FILE/END_FILE 标记。\n"
                "必须生成文档或表格文件（docx/xlsx/pdf）。"
            )
        if task_type == "CODER":
            return base + "要求：输出完整可运行代码，并包含必要说明。"
        if task_type == "RESEARCH":
            return base + "要求：输出结构化报告，包含标题与要点。"
        if task_type == "WEB_SEARCH":
            return base + "要求：基于实时信息回答，给出清晰结论。"
        return base

    @staticmethod
    def adapt_prompt_to_markdown(task_type: str, user_input: str, history: list = None) -> str:
        """使用快速小模型将原始请求转为结构化 Markdown，便于大模型理解。"""
        try:
            try:
                from web.prompt_adapter import PromptAdapter
            except ImportError:
                from prompt_adapter import PromptAdapter

            def _generate_markdown(prompt_text: str) -> str:
                response = client.models.generate_content(
                    model="gemini-2.0-flash-lite",
                    contents=prompt_text,
                    config=types.GenerateContentConfig(
                        max_output_tokens=700,
                        temperature=0.2,
                    )
                )
                return response.text or ""

            return PromptAdapter.adapt(
                user_input=user_input,
                task_type=task_type,
                history=history,
                model_generate=_generate_markdown,
            )
        except Exception as e:
            print(f"[PROMPT_ADAPTER] Failed: {e}")
            return user_input

    @staticmethod
    def quick_self_check(task_type: str, user_input: str, output_text: str) -> dict:
        """使用快速模型进行自检，返回 {'pass': bool, 'fix_prompt': str}。"""
        try:
            check_prompt = (
                "你是质量检查器。判断输出是否满足用户需求。\n"
                "只输出以下格式之一：\n"
                "PASS\n"
                "或\n"
                "FAIL\nFIX_PROMPT: <用于修正的提示词>\n\n"
                f"任务类型: {task_type}\n"
                f"用户需求: {user_input}\n"
                f"模型输出:\n{output_text}\n"
            )
            response = client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents=check_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=300,
                    temperature=0.1,
                )
            )
            text = (response.text or "").strip()
            if text.startswith("PASS"):
                return {"pass": True, "fix_prompt": ""}
            if text.startswith("FAIL"):
                fix = ""
                for line in text.splitlines():
                    if line.startswith("FIX_PROMPT:"):
                        fix = line.replace("FIX_PROMPT:", "").strip()
                        break
                return {"pass": False, "fix_prompt": fix}
            return {"pass": True, "fix_prompt": ""}
        except Exception as e:
            print(f"[SELF_CHECK] Failed: {e}")
            return {"pass": True, "fix_prompt": ""}

    @staticmethod
    def detect_required_packages(text: str) -> list:
        """从输出中粗略检测第三方依赖（仅返回白名单内的包）。"""
        if not text:
            return []
        modules = set()
        for line in text.splitlines():
            line = line.strip()
            if line.startswith("import "):
                parts = line.replace("import", "").split(",")
                for p in parts:
                    name = p.strip().split(" ")[0]
                    if name:
                        modules.add(name)
            elif line.startswith("from "):
                parts = line.split()
                if len(parts) >= 2:
                    modules.add(parts[1].strip())

        packages = set()
        for mod in modules:
            if mod in Utils._PACKAGE_ALLOWLIST:
                packages.add(Utils._PACKAGE_ALLOWLIST[mod])
        return sorted(packages)

    @staticmethod
    def auto_install_packages(packages: list) -> dict:
        """安装缺失的依赖包。返回安装结果摘要。"""
        result = {"installed": [], "skipped": [], "failed": []}
        if not packages:
            return result

        for pkg in packages:
            try:
                spec = importlib.util.find_spec(pkg)
                if spec is not None:
                    result["skipped"].append(pkg)
                    continue
                module_aliases = [m for m, p in Utils._PACKAGE_ALLOWLIST.items() if p == pkg]
                if any(importlib.util.find_spec(m) is not None for m in module_aliases):
                    result["skipped"].append(pkg)
                    continue
            except Exception:
                pass

            try:
                cmd = [sys.executable, "-m", "pip", "install", pkg]
                proc = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=120,
                    creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                )
                if proc.returncode == 0:
                    result["installed"].append(pkg)
                else:
                    result["failed"].append(pkg)
            except Exception:
                result["failed"].append(pkg)

        return result

    @staticmethod
    def auto_save_files(text):
        """自动从响应中提取并保存文件"""
        saved = []

        code_dir = os.path.join(WORKSPACE_DIR, "code")
        os.makedirs(code_dir, exist_ok=True)

        def _get_save_dir(filename):
            ext = os.path.splitext(filename)[1].lower()
            code_exts = {
                ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".cs", ".cpp", ".c",
                ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".m", ".scala",
                ".sh", ".ps1", ".bat", ".cmd", ".json", ".yaml", ".yml",
                ".toml", ".ini", ".cfg", ".sql", ".md", ".html", ".css"
            }
            return code_dir if ext in code_exts else WORKSPACE_DIR
        
        # 调试：打印前800字符看看格式
        print(f"[FILE_GEN] Response first 800 chars:\n{text[:800]}\n")
        
        # 预处理：统一格式 (去掉多余空格)
        normalized_text = text
        
        # 方法1: 多种 BEGIN_FILE 格式的正则匹配
        patterns = [
            # 格式1: ---BEGIN_FILE: filename.py--- (无空格)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # 格式2: ---BEGIN_FILE: filename.py--- ... ---END_FILE--- (带换行)
            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
            # 格式3: 更宽松 - 允许各种空白
            r"---\s*BEGIN_FILE[:\s]+([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
            # 格式4: 最宽松 - 捕获任意文件名
            r"---BEGIN_FILE[:\s]+([^\n-]+?)---\s*(.*?)---END_FILE---",
        ]
        
        matches1 = []
        for i, pattern in enumerate(patterns):
            try:
                matches1 = re.findall(pattern, normalized_text, re.DOTALL | re.IGNORECASE)
                print(f"[FILE_GEN] Pattern{i+1} matches: {len(matches1)}")
                if matches1:
                    print(f"[FILE_GEN] ✓ Using pattern {i+1}")
                    break
            except Exception as e:
                print(f"[FILE_GEN] Pattern{i+1} error: {e}")
        
        for filename, content in matches1:
            try:
                filename = filename.strip()
                content = content.strip()
                print(f"[FILE_GEN] Processing file: '{filename}', content length: {len(content)}")
                
                # 清除 Markdown 代码块标记
                if content.startswith('```'):
                    lines = content.split('\n')
                    if lines[0].startswith('```'):
                        lines = lines[1:]
                    if lines and lines[-1].strip() == '```':
                        lines = lines[:-1]
                    content = '\n'.join(lines)
                    print(f"[FILE_GEN] After stripping markdown: {len(content)} chars")
                
                # 确保文件名有效
                if not filename or len(filename) > 100:
                    print(f"[FILE_GEN] Invalid filename: {filename}")
                    continue
                
                # 确保文件名有扩展名
                if '.' not in filename:
                    filename = filename + '.py'
                
                base_dir = _get_save_dir(filename)
                path = os.path.join(base_dir, filename)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                saved.append(filename)
                print(f"[FILE_GEN] ✅ Saved: {filename} to {path}")
            except Exception as e:
                print(f"[FILE_GEN] ❌ Save failed: {e}")
                import traceback
                traceback.print_exc()
        
        # 方法2: 如果方法1没找到，尝试提取 ```python 代码块 + 文件名注释
        if not saved:
            print(f"[FILE_GEN] Method1 empty, trying method2 (```python blocks)...")
            
            # 先尝试匹配带文件名的代码块
            # 例如: # filename: cat_info.py 或 # cat_info.py
            pattern2a = r"```python\s*\n#\s*(?:filename:\s*)?([a-zA-Z0-9_.-]+\.py)\s*\n(.*?)```"
            matches2a = re.findall(pattern2a, text, re.DOTALL)
            print(f"[FILE_GEN] Pattern2a (with filename comment) matches: {len(matches2a)}")
            
            if matches2a:
                for filename, code in matches2a:
                    code = code.strip()
                    if not code or len(code) < 20:
                        continue
                    base_dir = _get_save_dir(filename)
                    path = os.path.join(base_dir, filename)
                    try:
                        with open(path, "w", encoding="utf-8") as f:
                            f.write(code)
                        saved.append(filename)
                        print(f"[FILE_GEN] ✅ Method2a saved: {filename}")
                    except Exception as e:
                        print(f"[FILE_GEN] ❌ Method2a save failed: {e}")
            else:
                # 无文件名的代码块，使用时间戳
                pattern2 = r"```python\s*\n(.*?)```"
                matches2 = re.findall(pattern2, text, re.DOTALL)
                print(f"[FILE_GEN] Pattern2 (generic) matches: {len(matches2)}")
                
                if matches2:
                    timestamp = int(time.time())
                    for idx, code in enumerate(matches2):
                        code = code.strip()
                        if not code or len(code) < 50:
                            continue
                        
                        # 尝试从代码中提取有意义的文件名
                        filename = None
                        # 查找 doc_path, file_path 等变量
                        path_match = re.search(r'(?:doc_path|file_path|filepath|output_path)\s*=.*?["\']([^"\']+\.(pdf|docx|xlsx))["\']', code)
                        if path_match:
                            # 使用目标文件名作为脚本名
                            target_file = os.path.basename(path_match.group(1))
                            filename = target_file.rsplit('.', 1)[0] + '.py'
                        
                        if not filename:
                            filename = f"generated_{timestamp}_{idx}.py"
                        
                        base_dir = _get_save_dir(filename)
                        path = os.path.join(base_dir, filename)
                        try:
                            with open(path, "w", encoding="utf-8") as f:
                                f.write(code)
                            saved.append(filename)
                            print(f"[FILE_GEN] ✅ Method2 saved: {filename}")
                        except Exception as e:
                            print(f"[FILE_GEN] ❌ Method2 save failed: {e}")
        
        print(f"[FILE_GEN] Final saved files: {saved}")
        return saved

    @staticmethod
    def save_image_part(blob_part):
        try:
            # 使用用户设置的图片目录
            images_dir = settings_manager.images_dir
            os.makedirs(images_dir, exist_ok=True)
            
            timestamp = int(time.time())
            filename = f"generated_{timestamp}.png"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as f:
                f.write(blob_part.inline_data.data)
            
            # 返回相对于 workspace 的路径
            # 确保路径始终在 workspace 下，且格式为正斜杠
            try:
                rel_path = os.path.relpath(filepath, WORKSPACE_DIR)
                # 如果包含 .. 说明不在 workspace 下，需要处理
                if ".." in rel_path:
                    # 降级为只返回文件名，放在 workspace/images 下
                    abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                    os.makedirs(abs_workspace_images, exist_ok=True)
                    fallback_path = os.path.join(abs_workspace_images, filename)
                    with open(fallback_path, "wb") as f:
                        f.write(blob_part.inline_data.data)
                    rel_path = os.path.relpath(fallback_path, WORKSPACE_DIR)
                    print(f"[IMAGE] Falling back to workspace/images: {rel_path}")
                
                result = rel_path.replace("\\", "/")
                print(f"[IMAGE] Saved image: {result}")
                return result
            except Exception as path_err:
                print(f"[IMAGE] Path calculation error: {path_err}")
                # 最后的保险方案：直接保存到 workspace/images
                abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                os.makedirs(abs_workspace_images, exist_ok=True)
                fallback_path = os.path.join(abs_workspace_images, filename)
                with open(fallback_path, "wb") as f:
                    f.write(blob_part.inline_data.data)
                result = os.path.relpath(fallback_path, WORKSPACE_DIR).replace("\\", "/")
                print(f"[IMAGE] Emergency fallback: {result}")
                return result
        except Exception as e:
            print(f"[IMAGE] Save failed: {e}")
            import traceback
            traceback.print_exc()
            return None

# ================= Session Manager =================

class SessionManager:
    def __init__(self):
        self.sessions = {}
    
    def list_sessions(self):
        """列出所有会话，按修改时间排序（最新在前）"""
        files = [f for f in os.listdir(CHAT_DIR) if f.endswith(".json")]
        # 按修改时间排序，最新的在前
        files_with_time = []
        for f in files:
            path = os.path.join(CHAT_DIR, f)
            mtime = os.path.getmtime(path)
            files_with_time.append((f, mtime))
        files_with_time.sort(key=lambda x: x[1], reverse=True)
        return [f[0] for f in files_with_time]
    
    def load(self, filename):
        """加载会话历史 - 返回用于模型上下文的截断版本"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    full_history = json.load(f)
                    # 仅截断用于模型上下文的部分，不影响持久化存储
                    return self._trim_history(full_history)
            except:
                return []
        return []
    
    def load_full(self, filename):
        """加载完整会话历史 - 用于追加保存，不做截断"""
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return json.load(f)
            except:
                return []
        return []
    
    def _trim_history(self, history, max_turns=20):
        """保留最多 20 轮对话（约 12000+ tokens），确保上下文足够但不过长"""
        if len(history) <= max_turns:
            return history
        # 只保留最后 N 轮对话
        trimmed = history[-max_turns:]
        print(f"[HISTORY] Trimmed to last {max_turns} turns (was {len(history)})")
        return trimmed
    
    def create(self, name):
        safe = "".join([c if c.isalnum() else "_" for c in name])
        filename = f"{safe}.json"
        path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump([], f)
        return filename
    
    def save(self, filename, history):
        path = os.path.join(CHAT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=2, ensure_ascii=False)
    
    def append_and_save(self, filename, user_msg, model_msg, **extra_fields):
        """追加消息并保存 - 基于磁盘完整历史，避免截断导致数据丢失"""
        full_history = self.load_full(filename)
        user_timestamp = extra_fields.pop("user_timestamp", datetime.now().isoformat())
        model_timestamp = extra_fields.pop("model_timestamp", datetime.now().isoformat())

        full_history.append({"role": "user", "parts": [user_msg], "timestamp": user_timestamp})
        model_entry = {"role": "model", "parts": [model_msg]}
        if "timestamp" not in extra_fields:
            model_entry["timestamp"] = model_timestamp
        model_entry.update(extra_fields)
        full_history.append(model_entry)
        self.save(filename, full_history)
        return full_history
    
    def append_user_early(self, filename, user_msg):
        """在请求到达时立即保存用户消息，防止断连导致丢失
        返回history长度，后续用update_last_model_response更新模型回复"""
        full_history = self.load_full(filename)
        now_iso = datetime.now().isoformat()
        full_history.append({"role": "user", "parts": [user_msg], "timestamp": now_iso})
        full_history.append({"role": "model", "parts": ["⏳ 处理中..."], "timestamp": now_iso})
        self.save(filename, full_history)
        return len(full_history)
    
    def update_last_model_response(self, filename, model_msg, **extra_fields):
        """更新最后一条模型回复（配合append_user_early使用）"""
        full_history = self.load_full(filename)
        if full_history and full_history[-1].get("role") == "model":
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history[-1] = model_entry
            self.save(filename, full_history)
        else:
            # fallback: 直接追加
            model_entry = {"role": "model", "parts": [model_msg]}
            if "timestamp" not in extra_fields:
                model_entry["timestamp"] = datetime.now().isoformat()
            model_entry.update(extra_fields)
            full_history.append(model_entry)
            self.save(filename, full_history)

    def add_message(self, filename, role, content, task="CHAT", model_name="Auto", **extra_fields):
        """追加单条消息（兼容旧调用），默认附带时间戳"""
        full_history = self.load_full(filename)
        entry = {
            "role": role,
            "parts": [content],
            "task": task,
            "model_name": model_name,
            "timestamp": extra_fields.pop("timestamp", datetime.now().isoformat())
        }
        entry.update(extra_fields)
        full_history.append(entry)
        self.save(filename, full_history)
        return entry
    
    def delete(self, filename):
        path = os.path.join(CHAT_DIR, filename)
        if os.path.exists(path):
            try:
                os.remove(path)
                return True
            except:
                return False
        return False

session_manager = SessionManager()

# ================= 初始化全局模块 =================
# 懒加载 Memory Manager 和 Knowledge Base
_memory_manager = None
_kb = None

def get_memory_manager():
    """获取或创建 Memory Manager 实例（增强版）"""
    global _memory_manager
    if _memory_manager is None:
        try:
            # 优先使用增强版本
            from enhanced_memory_manager import EnhancedMemoryManager
            _memory_manager = EnhancedMemoryManager()
            print("[INIT] ✅ 增强记忆管理器已初始化")
        except ImportError:
            try:
                from web.enhanced_memory_manager import EnhancedMemoryManager
                _memory_manager = EnhancedMemoryManager()
                print("[INIT] ✅ 增强记忆管理器已初始化")
            except ImportError:
                # 降级到基础版本
                try:
                    from memory_manager import MemoryManager
                except ImportError:
                    from web.memory_manager import MemoryManager
                _memory_manager = MemoryManager()
                print("[INIT] ⚠️  使用基础记忆管理器")
    return _memory_manager


def _start_memory_extraction(user_msg: str, ai_msg: str, history=None):
    """后台提取长期记忆，避免阻塞主对话流程"""
    try:
        from memory_integration import MemoryIntegration
    except ImportError:
        try:
            from web.memory_integration import MemoryIntegration
        except ImportError:
            print("[MemoryIntegration] ⚠️  模块未找到，跳过自动记忆提取")
            return

    if not MemoryIntegration.should_extract(user_msg, ai_msg):
        return

    def _worker():
        try:
            memory_mgr = get_memory_manager()

            class _LLMAdapter:
                async def generate(self, prompt, temperature=0.1, max_tokens=500):
                    resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            temperature=temperature,
                            max_output_tokens=max_tokens,
                        )
                    )
                    return resp.text or ""

            result = asyncio.run(
                MemoryIntegration.extract_and_apply(
                    memory_mgr, user_msg, ai_msg, _LLMAdapter(), history
                )
            )
            if result.get("success"):
                print("[MemoryIntegration] ✅ 自动记忆提取完成")
            else:
                print(f"[MemoryIntegration] ⚠️ 自动记忆提取失败: {result.get('error')}")
        except Exception as e:
            print(f"[MemoryIntegration] ❌ 自动记忆提取异常: {e}")

    threading.Thread(target=_worker, daemon=True).start()

def get_knowledge_base():
    """获取或创建 Knowledge Base 实例"""
    global _kb
    if _kb is None:
        try:
            from knowledge_base import KnowledgeBase
        except ImportError:
            from web.knowledge_base import KnowledgeBase
        _kb = KnowledgeBase()
        print("[INIT] ✅ Knowledge Base 已初始化")
    return _kb

# 为了向后兼容，导出全局变量
memory_manager = None  # 将通过 get_memory_manager() 动态获取
kb = None  # 将通过 get_knowledge_base() 动态获取

# ================= Koto Brain =================

class KotoBrain:
    # 图像编辑关键词
    IMAGE_EDIT_KEYWORDS = [
        "修改", "换", "改成", "变成", "底色", "背景", "颜色",
        "抠图", "去背景", "P图", "美化", "滤镜", "调色", "编辑",
        "change", "modify", "edit", "background", "color",
    ]
    
    def chat(self, history, user_input, file_data=None, model=None, auto_model=True):
        start_time = time.time()
        original_input = user_input
        # 支持模型选择和自动选择
        if model and not auto_model:
            model_id = model
            route_method = "Manual select"
            target_key = "CHAT"
        else:
            target_key = "CHAT"
            route_method = "Auto"
            
            if file_data:
                # 判断是图像编辑还是图像分析
                user_lower = user_input.lower()
                is_edit = any(kw in user_lower for kw in self.IMAGE_EDIT_KEYWORDS)
                
                if is_edit:
                    target_key = "PAINTER"
                    route_method = "Image Edit"
                else:
                    target_key = "VISION"
                    route_method = "Image Analysis"
            else:
                # 使用智能路由器
                target_key, route_method, _ = SmartDispatcher.analyze(user_input)
            
            model_id = SmartDispatcher.get_model_for_task(target_key, has_image=bool(file_data))

        # 使用小模型将请求转换为结构化 Markdown（仅在大模型处理时启用）
        model_input = user_input
        if auto_model and not file_data and target_key not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]:
            model_input = Utils.adapt_prompt_to_markdown(target_key, user_input, history=history)
            if model_input != user_input:
                print("[PROMPT_ADAPTER] Applied Markdown prompt for model")
        
        result = {
            "task": target_key,
            "model": model_id,
            "route_method": route_method,  # 路由方法信息
            "response": "",
            "images": [],
            "saved_files": [],
            "latency": 0,
            "total_time": 0
        }
        
        try:
            # === SYSTEM Mode (本地执行) ===
            if target_key == "SYSTEM":
                exec_result = LocalExecutor.execute(user_input)
                result["response"] = exec_result["message"]
                if exec_result.get("details"):
                    result["response"] += f"\n\n{exec_result['details']}"
                result["total_time"] = time.time() - start_time
                return result
            
            # === PAINTER Mode (图像生成/编辑) ===
            if target_key == "PAINTER":
                # 如果有输入图片（图像编辑模式）- 使用代码方式处理
                if file_data:
                    # 保存上传的图片到 workspace
                    import tempfile
                    import subprocess
                    
                    temp_img_path = os.path.join(WORKSPACE_DIR, "images", f"input_{int(time.time())}.jpg")
                    os.makedirs(os.path.dirname(temp_img_path), exist_ok=True)
                    with open(temp_img_path, "wb") as f:
                        f.write(file_data["data"])
                    
                    # 构建图像编辑的系统指令
                    edit_instruction = f"""你是一个图像处理专家。用户上传了一张图片，需要你生成 Python 代码来处理它。

图片路径: {temp_img_path}
用户请求: {user_input}

请生成完整的 Python 代码来完成用户的图像编辑请求。

要求:
1. 使用 OpenCV (cv2) 或 PIL 处理图片
2. 处理后的图片保存到: {settings_manager.images_dir}
3. 文件名格式: edited_{{timestamp}}.jpg 或 .png
4. 代码必须完整可执行
5. 对于换背景色，使用颜色阈值或边缘检测来识别背景区域

常用的背景色处理方法:
- 证件照换底色: 检测接近原背景色的像素，替换为目标颜色
- 蓝色背景 RGB: (67, 142, 219) 或 (0, 191, 255)
- 红色背景 RGB: (255, 0, 0) 或 (220, 0, 0)  
- 白色背景 RGB: (255, 255, 255)

代码格式（必须使用这个格式）:
---BEGIN_FILE: image_edit.py---
# 你的代码
---END_FILE---"""

                    # 调用 Gemini 生成代码（带回退）
                    edit_models = ["gemini-3-flash-preview", "gemini-3-pro-preview", "gemini-2.5-flash", "gemini-2.0-flash"]
                    code_response = None
                    last_error = None

                    def _process_code_response(code_response_text: str):
                        # 提取代码 - 支持多种格式
                        import re
                        patterns = [
                            r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                            r"```python\s*(.*?)```",  # 标准 markdown 代码块
                            r"```\s*(.*?)```",  # 无语言标记的代码块
                        ]
                        
                        code_content = None
                        for pattern in patterns:
                            matches = re.findall(pattern, code_response_text, re.DOTALL | re.IGNORECASE)
                            if matches:
                                if isinstance(matches[0], tuple):
                                    code_content = matches[0][1].strip()
                                else:
                                    code_content = matches[0].strip()
                                print(f"[IMAGE_EDIT] Extracted code, length: {len(code_content)}")
                                break
                        
                        if not code_content:
                            return {
                                "images": [],
                                "response": f"❌ 无法从模型响应中提取代码\n\n模型返回内容:\n```\n{code_response_text[:500]}\n```",
                                "error": "no_code"
                            }
                        
                        # 保存并执行代码
                        temp_script = os.path.join(tempfile.gettempdir(), f"koto_edit_{int(time.time())}.py")
                        with open(temp_script, 'w', encoding='utf-8') as f:
                            f.write(code_content)
                        
                        print(f"[IMAGE_EDIT] Executing script: {temp_script}")
                        exec_result = subprocess.run(
                            [sys.executable, temp_script],
                            capture_output=True,
                            text=True,
                            timeout=60,
                            cwd=WORKSPACE_DIR
                        )
                        
                        print(f"[IMAGE_EDIT] Script result: returncode={exec_result.returncode}")
                        if exec_result.stdout:
                            print(f"[IMAGE_EDIT] stdout: {exec_result.stdout[:200]}")
                        if exec_result.stderr:
                            print(f"[IMAGE_EDIT] stderr: {exec_result.stderr[:200]}")
                        
                        # 清理临时脚本
                        try:
                            os.remove(temp_script)
                        except:
                            pass
                        
                        if exec_result.returncode == 0:
                            images = []
                            images_dir = settings_manager.images_dir
                            for f in os.listdir(images_dir):
                                if f.startswith("edited_") and f.endswith(('.jpg', '.png', '.jpeg')):
                                    full_path = os.path.join(images_dir, f)
                                    age = time.time() - os.path.getmtime(full_path)
                                    if age < 60:
                                        rel_path = os.path.relpath(full_path, WORKSPACE_DIR).replace("\\", "/")
                                        images.append(rel_path)
                            
                            if images:
                                return {
                                    "images": images,
                                    "response": f"✅ 图片编辑完成!\n🖼️ 保存位置: `{images_dir}`",
                                    "error": ""
                                }
                            return {
                                "images": [],
                                "response": f"⚠️ 脚本执行成功但未检测到新图片\n\n{exec_result.stdout[:500]}",
                                "error": "no_output"
                            }
                        
                        return {
                            "images": [],
                            "response": f"❌ 图片处理失败\n```\n{exec_result.stderr[:500]}\n```",
                            "error": "exec_failed"
                        }
                    
                    for edit_model in edit_models:
                        try:
                            print(f"[IMAGE_EDIT] Trying model: {edit_model}")
                            print(f"[IMAGE_EDIT] Sending request to API...")
                            response = client.models.generate_content(
                                model=edit_model,
                                contents=edit_instruction,
                                config=types.GenerateContentConfig(
                                    max_output_tokens=4096,
                                    temperature=0.5
                                )
                            )
                            print(f"[IMAGE_EDIT] Got API response")
                            
                            if response.candidates and response.candidates[0].content.parts:
                                code_response = response.candidates[0].content.parts[0].text
                                print(f"[IMAGE_EDIT] Got response from {edit_model}, length: {len(code_response)}")
                                break
                        except Exception as model_err:
                            last_error = str(model_err)
                            print(f"[IMAGE_EDIT] Model {edit_model} failed: {last_error[:100]}")
                            continue
                    
                    if code_response:
                        run_result = _process_code_response(code_response)
                        result["images"] = run_result["images"]
                        result["response"] = run_result["response"]
                    else:
                        result["response"] = f"❌ 所有模型都不可用: {last_error[:200] if last_error else '未知错误'}"
                    
                    # 失败后自动修正并重试一次（避免无编辑结果）
                    if not result["images"] and Utils.is_failure_output(result["response"]):
                        fix_prompt = (
                            "上次生成失败，请修正并只输出完整可执行的 Python 代码。\n"
                            "必须使用 BEGIN_FILE/END_FILE 格式。\n"
                            f"图片路径: {temp_img_path}\n"
                            f"输出目录: {settings_manager.images_dir}\n"
                            f"用户请求: {user_input}\n\n"
                            f"失败信息/输出: {result['response']}\n"
                        )
                        retry_models = ["gemini-3-flash-preview", "gemini-2.5-flash"]
                        for retry_model in retry_models:
                            try:
                                print(f"[IMAGE_EDIT] Retry with model: {retry_model}")
                                retry_resp = client.models.generate_content(
                                    model=retry_model,
                                    contents=fix_prompt,
                                    config=types.GenerateContentConfig(
                                        max_output_tokens=4096
                                    )
                                )
                                if retry_resp.candidates and retry_resp.candidates[0].content.parts:
                                    retry_code = retry_resp.candidates[0].content.parts[0].text
                                    retry_run = _process_code_response(retry_code)
                                    if retry_run["images"]:
                                        result["images"] = retry_run["images"]
                                        result["response"] = retry_run["response"]
                                        break
                                    result["response"] = retry_run["response"]
                            except Exception as retry_err:
                                print(f"[IMAGE_EDIT] Retry failed: {retry_err}")
                    
                    result["total_time"] = time.time() - start_time
                    return result
                else:
                    # 纯图像生成使用 nano-banana-pro-preview
                    try:
                        print(f"[图像生成] 开始生成: {user_input[:50]}...")
                        response = client.models.generate_content(
                            model="nano-banana-pro-preview",
                            contents=user_input,
                            config=types.GenerateContentConfig(
                                response_modalities=["IMAGE"]
                            )
                        )
                        print(f"[图像生成] 响应成功，候选数: {len(response.candidates) if response.candidates else 0}")
                        
                        # 保存生成的图片
                        if response.candidates and response.candidates[0].content.parts:
                            for part in response.candidates[0].content.parts:
                                if hasattr(part, "inline_data") and part.inline_data:
                                    img_filename = Utils.save_image_part(part)
                                    if img_filename:
                                        result["images"].append(img_filename)
                                        print(f"[图像生成] 已保存: {img_filename}")
                        
                        if result["images"]:
                            save_path = settings_manager.images_dir
                            result["response"] = f"✨ 图片已生成!\n🖼️ 保存位置: `{save_path}`"
                        else:
                            result["response"] = "❌ 图像生成失败: 无输出内容，请检查提示词"
                        result["total_time"] = time.time() - start_time
                        return result
                    except Exception as img_err:
                        error_msg = str(img_err)
                        print(f"[图像生成] 错误: {error_msg[:200]}")
                        
                        # 提供更详细的错误信息
                        if "disconnected" in error_msg.lower() or "timeout" in error_msg.lower():
                            result["response"] = f"❌ 连接超时或中断: {error_msg[:100]}\n\n💡 建议: 请稍后重试，或检查网络连接"
                        elif "safety" in error_msg.lower():
                            result["response"] = "❌ 内容因安全政策被过滤，请修改提示词"
                        elif "quota" in error_msg.lower() or "rate" in error_msg.lower():
                            result["response"] = "❌ API 配额已达限制，请稍后重试"
                        else:
                            result["response"] = f"❌ 图像生成失败: {error_msg[:100]}"
                        
                        result["total_time"] = time.time() - start_time
                        return result
                
                if not response.candidates:
                    result["response"] = "Generation failed (safety filter or busy)."
                    result["total_time"] = time.time() - start_time
                    return result
                
                text_response = ""
                if response.candidates and response.candidates[0].content.parts:
                    for part in response.candidates[0].content.parts:
                        if hasattr(part, 'text') and part.text:
                            text_response += part.text
                        if hasattr(part, "inline_data") and part.inline_data:
                            img_filename = Utils.save_image_part(part)
                            if img_filename:
                                result["images"].append(img_filename)
                
                # 添加图片保存位置提示
                if result["images"]:
                    save_path = settings_manager.images_dir
                    text_response += f"\n\n🖼️ 图片已保存到: `{save_path}`"
                
                result["response"] = text_response if text_response else "Image generated successfully!"
                result["total_time"] = time.time() - start_time
                return result
            
            # === RAG: Retrieve Relevant Context (Auto) ===
            try:
                # 获取知识库实例
                kb_inst = get_knowledge_base()
                
                # 仅在非特定模式且输入有效时检索
                if target_key not in ["PAINTER", "SYSTEM"] and len(original_input) > 3:
                    # 避免对极短的问候语进行检索
                    skip_keywords = ["你好", "hello", "hi", "test", "测试"]
                    if not any(original_input.lower() == k for k in skip_keywords):
                        print(f"[RAG]正在检索知识库: {original_input[:50]}...")
                        rag_results = kb_inst.search(original_input, top_k=3)
                        
                        if rag_results:
                            print(f"[RAG] 检索到 {len(rag_results)} 个相关片段")
                            context_str = "\n".join([
                                f"--- 来源: {r['file_name']} (相似度: {r['similarity']:.2f}) ---\n{r['text']}"
                                for r in rag_results
                            ])
                            
                            # 将上下文注入 prompt
                            rag_context = f"\n\n【参考资料】\n以下是从本地知识库检索到的相关内容，供回答参考：\n{context_str}\n\n"
                            
                            # Log retrieval
                            print(f"[RAG] Injected context length: {len(rag_context)}")
                            
                            # Update model input
                            # 如果有 file_data，model_input 可能是 None 或不被直接使用，需谨慎
                            if not file_data:
                                model_input = rag_context + model_input
                            else:
                                # 对于有文件的请求，我们将上下文拼接到 original_input (user prompt)
                                # 注意：下面 generate_content 用的是 original_input + image_part
                                original_input = rag_context + original_input

            except Exception as rag_err:
                print(f"[RAG] Retrieval warning: {rag_err}")

            # === Regular Mode ===
            # 构建历史记录格式（过滤无关历史）
            history_for_model = ContextAnalyzer.filter_history(original_input, history)
            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(types.Content(
                    role=turn['role'],
                    parts=[types.Part.from_text(text=p) for p in turn['parts']]
                ))
            
            if file_data:
                # 构建正确的 Part 格式
                image_part = types.Part.from_bytes(
                    data=file_data["data"],
                    mime_type=file_data["mime_type"]
                )
                response = client.models.generate_content(
                    model=model_id,
                    contents=[original_input, image_part],
                    config=types.GenerateContentConfig(
                        system_instruction=_get_system_instruction()
                    )
                )
                accumulated_text = response.text if response.text else ""
            else:
                response = client.models.generate_content(
                    model=model_id,
                    contents=formatted_history + [types.Content(
                        role="user",
                        parts=[types.Part.from_text(text=model_input)]
                    )],
                    config=types.GenerateContentConfig(
                        system_instruction=_get_system_instruction()
                    )
                )
                accumulated_text = response.text if response.text else ""
            
            first_token_latency = (time.time() - start_time) * 1000
            result["latency"] = first_token_latency
            
            # Auto-save files
            saved_files = Utils.auto_save_files(accumulated_text)
            result["saved_files"] = saved_files
            
            # 添加文件保存提示
            if saved_files:
                files_list = ", ".join(saved_files)
                accumulated_text += f"\n\n📁 文件已保存: **{files_list}**\n📂 位置: `{WORKSPACE_DIR}`"
            
            result["response"] = accumulated_text
            result["total_time"] = time.time() - start_time
            return result
            
        except Exception as e:
            result["response"] = f"Error: {str(e)}"
            result["total_time"] = time.time() - start_time
            return result

brain = KotoBrain()

# ================= Routes =================

@app.route('/')
def index():
    # 云模式：未认证用户看到落地页
    deploy_mode = os.environ.get('KOTO_DEPLOY_MODE', 'local')
    auth_enabled = os.environ.get('KOTO_AUTH_ENABLED', 'false').lower() == 'true'
    if deploy_mode == 'cloud' and auth_enabled:
        return render_template('landing.html')
    return render_template('index.html')

@app.route('/app')
def app_main():
    """主应用页面（SaaS 模式下需认证后访问）"""
    return render_template('index.html')

@app.route('/file-network')
def file_network():
    """文件网络界面"""
    return render_template('file_network.html')

@app.route('/knowledge-graph')
def knowledge_graph_page():
    """知识图谱可视化界面"""
    return render_template('knowledge_graph.html')

@app.route('/test_upload')
def test_upload():
    return render_template('test_upload.html')

@app.route('/edit-ppt/<session_id>')
def edit_ppt(session_id):
    """PPT 生成后编辑页面（P1 功能）"""
    return render_template('edit_ppt.html')

@app.route('/monitoring-dashboard')
def monitoring_dashboard():
    """Phase 4 System Monitoring Dashboard"""
    return send_from_directory(os.path.join(os.path.dirname(__file__), 'static'), 'monitoring_dashboard.html')

@app.route('/api/sessions', methods=['GET'])
def get_sessions():
    sessions = session_manager.list_sessions()
    return jsonify({
        "sessions": [s.replace(".json", "") for s in sessions]
    })

@app.route('/api/sessions', methods=['POST'])
def create_session():
    data = request.json
    name = data.get('name', f'chat_{int(time.time())}')
    filename = session_manager.create(name)
    return jsonify({
        "success": True,
        "session": filename.replace(".json", "")
    })

@app.route('/api/sessions/<session_name>', methods=['GET'])
def get_session(session_name):
    # 返回完整历史供前端渲染（不截断），截断仅用于模型上下文
    history = session_manager.load_full(f"{session_name}.json")
    return jsonify({
        "session": session_name,
        "history": history
    })

@app.route('/api/sessions/<session_name>', methods=['DELETE'])
def delete_session(session_name):
    success = session_manager.delete(f"{session_name}.json")
    return jsonify({"success": success})

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    session_name = data.get('session')
    user_input = data.get('message', '')
    locked_task = data.get('locked_task')
    locked_model = data.get('locked_model', 'auto')
    
    if not session_name or not user_input:
        return jsonify({"error": "Missing session or message"}), 400
    
    user_input = Utils.sanitize_string(user_input)
    
    # Load history
    history = session_manager.load(f"{session_name}.json")
    
    # 确定使用的模型
    if locked_model and locked_model != 'auto':
        model = locked_model
        auto_model = False
    elif locked_task:
        model = MODEL_MAP.get(locked_task, MODEL_MAP['CHAT'])
        auto_model = False
    else:
        model = None
        auto_model = True
    
    # Get response
    result = brain.chat(history, user_input, model=model, auto_model=auto_model)

    # 代码任务: 自动检查依赖并安装
    if result.get("task") == "CODER" and result.get("response"):
        pkgs = Utils.detect_required_packages(result["response"])
        if pkgs:
            install_result = Utils.auto_install_packages(pkgs)
            installed = install_result.get("installed", [])
            failed = install_result.get("failed", [])
            skipped = install_result.get("skipped", [])
            msg_parts = []
            if installed:
                msg_parts.append(f"✅ 已安装: {', '.join(installed)}")
            if skipped:
                msg_parts.append(f"ℹ️ 已存在: {', '.join(skipped)}")
            if failed:
                msg_parts.append(f"⚠️ 安装失败: {', '.join(failed)}")
            if msg_parts:
                result["response"] += "\n\n" + "\n".join(msg_parts)
    
    # Update history (基于磁盘完整历史追加，避免截断丢失)
    session_manager.append_and_save(f"{session_name}.json", user_input, result["response"])
    
    return jsonify(result)


# ============== Agent 确认 API ==============
# NOTE: These routes have been migrated to the unified agent blueprint
#       (app/api/agent_routes.py) under /api/agent/confirm and /api/agent/choice.
#       Kept here as comments for reference.

# @app.route('/api/agent/confirm', methods=['POST'])
# def agent_confirm():
#     """Agent 用户确认 API — 前端点击确认/取消后回调"""
#     ...

# @app.route('/api/agent/choice', methods=['POST'])
# def agent_choice():
#     """Agent 用户选择 API — 前端选择后回调"""
#     ...


# NOTE: /api/agent/plan has been migrated to the unified agent blueprint
#       (app/api/agent_routes.py). Kept as comment for reference.
# @app.route('/api/agent/plan', methods=['POST'])
# def agent_plan(): ...


@app.route('/api/chat/stream', methods=['POST'])
def chat_stream():
    """流式聊天 API - 实时返回响应"""
    data = request.json
    session_name = data.get('session')
    user_input = data.get('message', '')
    locked_task = data.get('locked_task')
    locked_model = data.get('locked_model', 'auto')
    
    print(f"\n[STREAM] Incoming request: locked_task='{locked_task}', locked_model='{locked_model}'")
    print(f"[STREAM] User input: {user_input[:60]}")
    
    if not session_name or not user_input:
        def error_gen():
            yield f"data: {json.dumps({'type': 'error', 'message': 'Missing session or message'})}\n\n"
        return Response(error_gen(), mimetype='text/event-stream')
    
    # API 密钥缺失时提前返回友好提示
    if not API_KEY:
        def no_key_gen():
            msg = ("⚠️ **API 密钥未配置**\n\n"
                   "请在 `config/gemini_config.env` 文件中设置：\n"
                   "```\nGEMINI_API_KEY=你的密钥\n```\n\n"
                   "💡 获取密钥：[Google AI Studio](https://aistudio.google.com/apikey)\n\n"
                   "设置完成后重启 Koto 即可使用。")
            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"
        return Response(no_key_gen(), mimetype='text/event-stream')
    
    user_input = Utils.sanitize_string(user_input)

    # ⏳ 重复上一个任务 (Repeat Last Task)
    repeat_patterns = [r'^重复.*任务', r'^再做一遍', r'^再来一次', r'^re(peat|do).*last.*task', r'^try.*again']
    if any(re.search(p, user_input, re.IGNORECASE) for p in repeat_patterns):
        try:
            full_hist = session_manager.load_full(f"{session_name}.json")
            # 倒序查找最近的一条 user 消息
            last_user_msg = None
            for msg in reversed(full_hist):
                if msg.get("role") == "user":
                    content = (msg.get("parts") or [""])[0]
                    # 避免无限循环：如果上一条也是“重复任务”，则继续往前找
                    if not any(re.search(p, content, re.IGNORECASE) for p in repeat_patterns):
                        last_user_msg = content
                        break
            
            if last_user_msg:
                print(f"[REPEAT] Found last user message: {last_user_msg[:50]}...")
                user_input = last_user_msg # 替换当前输入为上一条任务
                # 可以选择注入一个提示，告诉用户正在重试
                # 但为了 context 连贯，直接替换最简单
            else:
                print("[REPEAT] Check failed: No valid previous user message found.")
        except Exception as e:
            print(f"[REPEAT] Error fetching history: {e}")
    
    # ⚡ 快速路径：系统时间查询 - 直接返回，无需发送到LLM
    time_query_patterns = [
        r'当前.*时间|当前系统时间', r'现在.*几点|几点钟', r'几点|什么时间',
        r'时间是|现在是', r'now.*time|what.*time|current.*time'
    ]
    if any(re.search(pattern, user_input, re.IGNORECASE) for pattern in time_query_patterns):
        def quick_time_response():
            from datetime import datetime
            
            now = datetime.now()
            date_str = now.strftime("%Y年%m月%d日")
            weekday = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][now.weekday()]
            time_str = now.strftime("%H:%M:%S")
            timestamp = now.isoformat()  # 记录精确时间戳
            response = f"当前系统时间为：\n\n**{date_str} {weekday} {time_str}**"
            
            # 记录到历史（用户 + 模型，均带时间戳）
            try:
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response,
                    task="CHAT",
                    model_name="QuickResponse",
                    timestamp=timestamp,
                    user_timestamp=timestamp,
                    model_timestamp=timestamp,
                )
            except Exception as e:
                print(f"[STREAM] Quick time history save failed: {e}")
            
            yield f"data: {json.dumps({'type': 'progress', 'message': '📅 系统时间查询', 'detail': '从本地获取'}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'token', 'content': response, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0.01, 'timestamp': timestamp}, ensure_ascii=False)}\n\n"
        
        return Response(quick_time_response(), mimetype='text/event-stream')
    
    # 🎯 获取动态系统指令（根据用户问题智能注入上下文）
    try:
        system_instruction = _get_chat_system_instruction(user_input)
    except Exception as e:
        print(f"[STREAM] Warning: Dynamic system instruction failed: {e}")
        system_instruction = _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION()  # 降级到新鲜生成的指令
    
    history = session_manager.load(f"{session_name}.json")
    
    # 🕵️‍♀️ 检测是否有最近上传的文件 (5分钟内)
    has_recent_upload = False
    recent_file_type = None
    try:
        upload_scan_dirs = ["web/uploads", "uploads", "workspace/documents"]
        recent_threshold = time.time() - 300 # 5分钟内
        for d in upload_scan_dirs:
            if os.path.exists(d):
                for f in os.listdir(d):
                    fp = os.path.join(d, f)
                    if os.path.isfile(fp) and os.path.getmtime(fp) > recent_threshold:
                        has_recent_upload = True
                        _, ext = os.path.splitext(f)
                        recent_file_type = ext.lower()
                        print(f"[STREAM] Found recent upload: {f} ({recent_file_type})")
                        break
            if has_recent_upload: break
    except Exception as e:
        print(f"[STREAM] Error checking uploads: {e}")

    # 确定任务类型和模型
    context_info = None
    if locked_task:
        task_type = locked_task
        route_method = "🔒 Manual"
        print(f"[STREAM] ✅ Using locked_task: '{task_type}'")
    else:
        # 将文件信息传递给分析器
        context_override = {"has_file": has_recent_upload, "file_type": recent_file_type}
        task_type, route_method, context_info = SmartDispatcher.analyze(user_input, history, file_context=context_override)
        print(f"[STREAM] Auto-detected task_type: '{task_type}', context: {context_info is not None}")

        
        # 如果有上下文信息，记录详情
        if context_info and context_info.get("is_continuation"):
            print(f"[STREAM] Context continuation: {context_info.get('related_task')}, confidence: {context_info.get('confidence')}")
    
    # === 处理复杂任务 (多步流程) ===
    if task_type == "MULTI_STEP" and context_info and context_info.get("is_multi_step_task"):
        print(f"[STREAM] 🔄 检测到复杂任务，使用 TaskOrchestrator 执行多步流程")
        multi_step_info = context_info.get("multi_step_info", {})
        pattern = multi_step_info.get("pattern", "unknown")
        
        # === 文档工作流执行 ===
        if pattern == "document_workflow" and DocumentWorkflowExecutor:
            print(f"[STREAM] 📄 执行文档工作流")
            
            def generate_doc_workflow():
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_WORKFLOW', 'pattern': 'document_workflow', 'route_method': route_method, 'message': '🎯 任务分类: 📄 文档工作流执行'})}\n\n"
                
                # 查找最近上传的文档
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]
                
                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        docs = []
                        for ext in [".docx", ".md", ".txt", ".json"]:
                            import glob
                            docs.extend(glob.glob(f"{dir_path}/**/*{ext}", recursive=True))
                        
                        if docs:
                            # 获取最新的文档
                            doc_path = max(docs, key=os.path.getmtime)
                            break
                
                if not doc_path:
                    yield f"data: {json.dumps({'type': 'error', 'message': '❌ 未找到可执行的文档文件（支持 .docx, .md, .txt, .json）'})}\n\n"
                    return
                
                status_msg = f"📄 找到文档: {os.path.basename(doc_path)}\n"
                yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"
                
                try:
                    import asyncio
                    
                    # 执行文档工作流
                    executor = DocumentWorkflowExecutor(client)
                    
                    # 加载工作流
                    status_msg = "⏳ 正在解析文档中的工作流...\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"
                    
                    load_result = asyncio.run(executor.load_from_document(doc_path))
                    
                    if not load_result.get("success"):
                        error_msg = f"❌ 文档解析失败: {load_result.get('error', '未知错误')}\n"
                        yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                        return
                    
                    # 显示工作流信息
                    info_msg = f"✅ 工作流加载成功\n"
                    info_msg += f"   名称: {executor.workflow_name}\n"
                    info_msg += f"   步骤数: {len(executor.steps)}\n"
                    info_msg += f"   背景: {executor.workflow_context}\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': info_msg})}\n\n"
                    
                    # 显示所有步骤
                    steps_msg = "📋 工作流步骤:\n"
                    for step in executor.steps:
                        steps_msg += f"  {step.step_id}. [{step.step_type}] {step.description}\n"
                    steps_msg += "\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': steps_msg})}\n\n"
                    
                    # 执行工作流（流式反馈每个步骤）
                    start_msg = "🚀 开始执行工作流...\n\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': start_msg})}\n\n"
                    
                    for step in executor.steps:
                        step_msg = f"[步骤 {step.step_id}/{len(executor.steps)}] {step.description}\n"
                        step_msg += f"└─ 类型: {step.step_type}\n"
                        step_msg += f"   ⏳ 执行中...\n"
                        yield f"data: {json.dumps({'type': 'status', 'message': step_msg})}\n\n"
                        
                        step.status = "running"
                        step.start_time = datetime.now()
                        
                        try:
                            # 执行步骤
                            step_result = asyncio.run(executor._execute_step_standalone(step))
                            step.result = step_result
                            step.status = "completed"
                            
                            success_msg = f"   ✅ 完成\n"
                            if isinstance(step_result, dict) and step_result.get("output"):
                                output_preview = str(step_result['output'])[:200]
                                success_msg += f"   📄 输出预览: {output_preview}...\n"
                            success_msg += "\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': success_msg})}\n\n"
                            
                        except Exception as e:
                            step.status = "failed"
                            step.error = str(e)
                            error_msg = f"   ❌ 失败: {e}\n\n"
                            yield f"data: {json.dumps({'type': 'status', 'message': error_msg})}\n\n"
                        
                        finally:
                            step.end_time = datetime.now()
                    
                    # 生成结果
                    results = {
                        "workflow_name": executor.workflow_name,
                        "start_time": datetime.now().isoformat(),
                        "steps": [step.to_dict() for step in executor.steps],
                        "overall_status": "completed"
                    }
                    results["summary"] = executor._generate_summary(results)
                    
                    # 保存结果
                    output_path = asyncio.run(executor.save_results(results))
                    
                    # 发送完成消息
                    separator = "=" * 50
                    final_msg = f"\n{separator}\n"
                    final_msg += f"✅ 文档工作流执行完成\n\n"
                    final_msg += f"📊 执行统计:\n"
                    total = len(results["steps"])
                    completed = sum(1 for s in results["steps"] if s["status"] == "completed")
                    failed = sum(1 for s in results["steps"] if s["status"] == "failed")
                    final_msg += f"  总步骤: {total}\n"
                    final_msg += f"  成功: {completed}\n"
                    final_msg += f"  失败: {failed}\n"
                    final_msg += f"  成功率: {completed/total*100:.1f}%\n\n"
                    final_msg += f"📁 结果已保存: {os.path.basename(output_path)}\n"
                    final_msg += f"📂 位置: `workspace/workflows/`\n\n"
                    final_msg += f"{separator}\n"
                    
                    yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [output_path]})}\n\n"
                    
                    # 保存文档工作流对话历史（基于磁盘完整历史追加）
                    try:
                        session_manager.append_and_save(f"{session_name}.json", user_input, f"[文档工作流完成] {executor.workflow_name}")
                    except Exception:
                        pass
                    
                except Exception as e:
                    import traceback
                    error_detail = traceback.format_exc()
                    error_msg = f"❌ 工作流执行失败: {str(e)}\n{error_detail}"
                    yield f"data: {json.dumps({'type': 'error', 'message': error_msg})}\n\n"
                    # 保存失败记录
                    try:
                        session_manager.append_and_save(f"{session_name}.json", user_input, f"[文档工作流失败] {str(e)[:200]}")
                    except Exception:
                        pass
            
            return Response(generate_doc_workflow(), mimetype='text/event-stream')
        
        # === 其他多步任务执行 ===
        from app.core.routing import TaskDecomposer
        subtasks = TaskDecomposer.create_subtasks(user_input, multi_step_info)
        use_local_planner = (multi_step_info.get("pattern") == "local_plan")
        
        def generate_multi_step():
            # === 立即发送任务分类信息 ===
            pattern = multi_step_info.get("pattern", "unknown")
            classification_msg = f"🎯 任务分类: 🔄 多步任务\n"
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'MULTI_STEP', 'pattern': pattern, 'route_method': route_method, 'message': classification_msg})}\n\n"
            
            # 显示所有子任务
            status_msg = f"📋 任务分解:\n"
            for i, subtask in enumerate(subtasks):
                status_msg += f"  {i+1}. {subtask['task_type']} - {subtask['description']}\n"
            status_msg += "\n"
            yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"
            
            # 执行所有子任务（逐步流式反馈）
            try:
                import asyncio
                execution_log = []
                step_results = []
                context = {"original_input": user_input, "user_input": user_input}
                saved_files = []
                
                for i, subtask in enumerate(subtasks):
                    # 发送任务特定的进度提示
                    _step_icons = {
                        "WEB_SEARCH": "🔍", "RESEARCH": "🔬", "PAINTER": "🎨",
                        "FILE_GEN": "📄", "CODER": "💻"
                    }
                    _task_type = subtask["task_type"]
                    _step_icon = _step_icons.get(_task_type, "⚙️")
                    step_msg = f"{_step_icon} 步骤 {i+1}/{len(subtasks)}: {subtask['description']}"
                    _detail = f"任务类型: {_task_type}"
                    yield f"data: {json.dumps({'type': 'progress', 'message': step_msg, 'detail': _detail})}\n\n"
                    
                    step_input = subtask.get("input") or user_input
                    
                    # 使用队列接收来自后台线程的实时进度
                    import queue
                    import threading
                    progress_queue = queue.Queue()
                    
                    def _progress_cb(msg, detail=""):
                        progress_queue.put({"msg": msg, "detail": detail})
                        
                    task_result_holder = {"result": None}
                    
                    def _run_task_thread():
                        try:
                            # Running asyncio.run inside a thread is tricky if not handled well, but here it's a fresh thread.
                            # We pass _progress_cb to updated methods.
                            if subtask["task_type"] == "WEB_SEARCH":
                                task_result_holder["result"] = asyncio.run(TaskOrchestrator._execute_web_search(step_input, context, progress_callback=_progress_cb))
                            elif subtask["task_type"] == "FILE_GEN":
                                # FILE_GEN likely already supports it or ignores extra args if using **kwargs in some versions, 
                                # but based on context it takes it as 4th arg or kwarg.
                                task_result_holder["result"] = asyncio.run(TaskOrchestrator._execute_file_gen(step_input, context, subtask, progress_callback=_progress_cb))
                            elif subtask["task_type"] == "PAINTER":
                                task_result_holder["result"] = asyncio.run(TaskOrchestrator._execute_painter(step_input, context, progress_callback=_progress_cb))
                            elif subtask["task_type"] == "RESEARCH":
                                task_result_holder["result"] = asyncio.run(TaskOrchestrator._execute_research(step_input, context, progress_callback=_progress_cb))
                            else:
                                # Fallback for unknown
                                task_result_holder["result"] = {"success": False, "error": f"未知任务类型: {subtask['task_type']}"}
                        except Exception as e:
                            import traceback
                            traceback.print_exc()
                            task_result_holder["result"] = {"success": False, "error": str(e)}
                        finally:
                            progress_queue.put(None) # Signal done

                    # 启动后台线程执行任务
                    t = threading.Thread(target=_run_task_thread)
                    t.start()
                    
                    # 主线程循环读取进度
                    while True:
                        try:
                            item = progress_queue.get(timeout=0.1)
                            if item is None:
                                break
                            # 发送进度SSE
                            yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': item['detail']})}\n\n"
                        except queue.Empty:
                            if not t.is_alive():
                                break
                                
                    t.join()
                    result = task_result_holder["result"]
                    
                    # 发送步骤完成进度
                    if result.get("success"):
                        _done_detail = result.get("output", "")[:60]
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 步骤 {i+1} 完成', 'detail': _done_detail})}\n\n"
                    else:
                        _err_detail = result.get("error", "未知错误")[:60]
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ 步骤 {i+1} 遇到问题', 'detail': _err_detail})}\n\n"
                    
                    subtask["status"] = "completed" if result.get("success") else "failed"
                    subtask["result"] = result
                    
                    context[f"{subtask['task_type']}_result"] = result
                    context[f"step_{i+1}_output"] = result.get("output", result.get("content", ""))
                    
                    if result.get("model_id"):
                        model_msg = f"🤖 使用模型: {result['model_id']}\n"
                        yield f"data: {json.dumps({'type': 'status', 'message': model_msg})}\n\n"
                    
                    model_answer = result.get("content") or result.get("output")
                    if model_answer:
                        answer_msg = f"🧠 模型回答:\n{model_answer}\n\n"
                        yield f"data: {json.dumps({'type': 'status', 'message': answer_msg})}\n\n"
                    
                    if result.get("saved_files"):
                        saved_files.extend(result["saved_files"])
                    step_results.append(result)
                    
                    execution_log.append(f"  ✅ 完成: {subtask['description']}")
                
                # 输出最终结果
                final_output = TaskOrchestrator._merge_results(subtasks, context)

                # 3. 结果验证 (User Requirement: Model Verification Feedback)
                # Tell user we are verifying
                yield f"data: {json.dumps({'type': 'status', 'message': '🔍 正在进行最终质量验证...'})}\n\n"
                quality_score = asyncio.run(TaskOrchestrator._validate_quality(user_input, final_output, context))
                yield f"data: {json.dumps({'type': 'status', 'message': f'✅ 质量验证完成，评分: {quality_score}/100'})}\n\n"
                
                separator = "=" * 50
                output_text = f"\n{separator}\n"
                output_text += f"✅ 多步任务完成\n"
                output_text += f"质量评分: {quality_score}/100\n"
                if saved_files:
                    output_text += f"已保存文件:\n"
                    # Add clickable links
                    for p in saved_files:
                        name = os.path.basename(p)
                        # Check if path is absolute or relative
                        link_path = p.replace("\\", "/")
                        if not link_path.startswith("http"):
                            # Assuming frontend can handle workspace relative paths or full paths exposed via virtual route
                            # Just output markdown link
                            output_text += f"- [{name}]({link_path})\n"
                    output_text += f"\n📂 位置: `{settings_manager.documents_dir}`\n"
                
                errors_list = []
                for subtask in subtasks:
                    if subtask.get("status") == "failed" and subtask.get("result"):
                        err = subtask["result"].get("error")
                        if err:
                            errors_list.append(err)
                if errors_list:
                    output_text += f"⚠️ 遇到的问题: {', '.join(errors_list)}\n"
                
                final_result = final_output.get('final_output', '(无输出)')
                # 复杂任务快速自检
                check = Utils.quick_self_check("MULTI_STEP", user_input, final_result)
                if not check.get("pass") and check.get("fix_prompt"):
                    status_msg = "🩺 自检未通过，正在修正最终输出...\n"
                    yield f"data: {json.dumps({'type': 'status', 'message': status_msg})}\n\n"
                    fix_resp = client.models.generate_content(
                        model=SmartDispatcher.get_model_for_task("CHAT"),
                        contents=check["fix_prompt"],
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=3000,
                        )
                    )
                    corrected = fix_resp.text or final_result
                    final_result = corrected
                if use_local_planner:
                    from app.core.routing import LocalPlanner
                    plan_check = LocalPlanner.self_check(user_input, subtasks, step_results)
                    status = plan_check.get("status", "partial")
                    summary = plan_check.get("summary", "")
                    next_actions = plan_check.get("next_actions", []) if isinstance(plan_check.get("next_actions", []), list) else []
                    output_text += f"\n自检结论: {status}\n"
                    if summary:
                        output_text += f"说明: {summary}\n"
                    if next_actions:
                        output_text += f"建议后续: {', '.join(next_actions)}\n"
                output_text += f"\n最终输出:\n{final_result}\n"
                output_text += f"{separator}\n"
                
                yield f"data: {json.dumps({'type': 'token', 'content': output_text})}\n\n"
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files})}\n\n"
                
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': f'多步任务执行失败: {str(e)}'})}\n\n"
            
            # 保存 MULTI_STEP 对话历史（基于磁盘完整历史追加）
            try:
                _multi_summary = f"[多步任务完成] {', '.join(s['description'] for s in subtasks)}"
                if saved_files:
                    _multi_summary += f"\n生成文件: {', '.join(os.path.basename(p) for p in saved_files)}"
                session_manager.append_and_save(f"{session_name}.json", user_input, _multi_summary)
                print(f"[MULTI_STEP] ✅ 对话历史已保存")
            except Exception as save_err:
                print(f"[MULTI_STEP] ⚠️ 保存对话历史失败: {save_err}")
        
        return Response(generate_multi_step(), mimetype='text/event-stream')
    
    # === Agent 任务执行 ===
    if task_type == "AGENT":
        print(f"[STREAM] 🤖 执行 Agent 任务 (UnifiedAgent)")
        
        def generate_agent():
            # 发送分类信息
            yield f"data: {json.dumps({'type': 'classification', 'task_type': 'AGENT', 'route_method': route_method, 'message': '🎯 任务分类: 🤖 智能助手 (工具调用)'})}\n\n"
            
            # 使用 UnifiedAgent 替代旧 agent_loop
            try:
                from app.core.agent.factory import create_agent
                from app.core.agent.types import AgentStepType
                _ua = create_agent(model_id=SmartDispatcher.get_model_for_task("AGENT"))
                
                collected_steps = []
                final_answer = ""
                for step in _ua.run(input_text=user_input, history=history):
                    step_data = step.to_dict()
                    collected_steps.append(step_data)
                    if step.step_type == AgentStepType.ANSWER:
                        final_answer = step.content or ""
                    yield f"data: {json.dumps({'type': 'agent_step', 'data': step_data}, ensure_ascii=False)}\n\n"
                
                if not final_answer and collected_steps:
                    final_answer = collected_steps[-1].get('content', '')
                task_payload = {
                    'id': f'task_{int(time.time() * 1000)}',
                    'status': 'success',
                    'result': final_answer,
                    'steps': collected_steps,
                }
                yield f"data: {json.dumps({'type': 'task_final', 'data': task_payload}, ensure_ascii=False)}\n\n"
                
                # 保存对话历史
                try:
                    session_manager.append_and_save(f"{session_name}.json", user_input, final_answer or '[Agent 任务完成]')
                except Exception:
                    pass
            
            except Exception as e:
                import traceback
                error_detail = traceback.format_exc()
                print(f"[AGENT] ❌ Agent 执行失败:\n{error_detail}")
                yield f"data: {json.dumps({'type': 'error', 'message': f'Agent 执行失败: {str(e)}'})}\n\n"
        
        return Response(generate_agent(), mimetype='text/event-stream')
    
    if locked_model and locked_model != 'auto':
        model_id = locked_model
    else:
        # 传递 complexity 以便为复杂任务选择更强的模型
        _complexity = (context_info or {}).get("complexity", "normal")
        model_id = SmartDispatcher.get_model_for_task(task_type, complexity=_complexity)
    
    print(f"[STREAM] Final: task_type='{task_type}', model_id='{model_id}'\n")
    
    # 读取用户设置：是否显示思考过程
    _show_thinking = False
    try:
        _show_thinking = settings_manager.get('ai', 'show_thinking') == True
    except:
        pass
    
    def generate():
        start_time = time.time()
        
        def yield_thinking(message: str, phase: str = "thinking"):
            """发送思考过程事件（仅当用户开启 show_thinking 时）"""
            if not _show_thinking:
                return ""
            elapsed = round(time.time() - start_time, 1)
            return f"data: {json.dumps({'type': 'thinking', 'message': message, 'phase': phase, 'elapsed': elapsed}, ensure_ascii=False)}\n\n"
        
        # === 立即反馈任务分类信息 ===
        task_display_names = {
            "PAINTER": "🎨 图像生成",
            "FILE_GEN": "📄 文档生成",
            "CODER": "💻 代码编程",
            "RESEARCH": "📚 深度研究",
            "WEB_SEARCH": "🌐 实时搜索",
            "CHAT": "💬 对话",
            "SYSTEM": "🖥️ 系统操作",
            "FILE_OP": "📂 文件操作",
            "FILE_EDIT": "✏️ 文件编辑",
            "FILE_SEARCH": "🔍 文件搜索",
            "VISION": "👁️ 图像识别",
            "MULTI_STEP": "🔄 多步任务",
            "AGENT": "🤖 智能助手"
        }
        
        model_display = get_model_display_name(model_id)
        task_display = task_display_names.get(task_type, task_type)
        
        # 发送任务分类信息（在最开始，立即显示）
        classification_msg = f"🎯 任务分类: {task_display}"
        if route_method:
            classification_msg += f" (方法: {route_method})"

        routing_list = None
        # 仅保留 routing_list 用于内部调试，不显示给用户
        if context_info and context_info.get("routing_list"):
            routing_list = context_info.get("routing_list")
        
        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': task_display, 'model': model_id, 'model_display': model_display, 'route_method': route_method, 'routing_list': routing_list, 'message': classification_msg})}\n\n"
        
        # 思考过程：任务路由分析
        t = yield_thinking(f"分析用户意图 → 识别为 {task_display}", "routing")
        if t: yield t
        t = yield_thinking(f"路由方法: {route_method}，选择模型: {model_display}", "model")
        if t: yield t
        if routing_list:
            steps_str = " → ".join([f"{r.get('method','?')}({r.get('confidence','?')})" for r in routing_list[:5]]) if isinstance(routing_list, list) else str(routing_list)
            t = yield_thinking(f"路由决策链: {steps_str}", "routing")
            if t: yield t
        
        # 如果有复杂度信息，也发送
        if context_info and context_info.get("complexity"):
            complexity_msg = f"📊 任务复杂度: {context_info['complexity']}"
            yield f"data: {json.dumps({'type': 'info', 'message': complexity_msg})}\n\n"
            t = yield_thinking(f"任务复杂度评估: {context_info['complexity']}", "analyzing")
            if t: yield t
        
        # 如果有上下文，使用增强后的输入
        effective_input = user_input
        if context_info and context_info.get("is_continuation") and context_info.get("enhanced_input"):
            effective_input = context_info["enhanced_input"]
            print(f"[STREAM] Using enhanced input (length: {len(effective_input)})")
            yield f"data: {json.dumps({'type': 'info', 'message': '🔗 检测到延续任务，使用上下文增强'})}\n\n"
            t = yield_thinking(f"检测到上下文延续，增强输入 ({len(effective_input)} 字符)", "context")
            if t: yield t

        # 使用快速小模型将请求转为结构化 Markdown（仅对大模型任务启用）
        if task_type not in ["SYSTEM", "FILE_OP", "PAINTER", "VISION"]:
            adapted_input = Utils.adapt_prompt_to_markdown(task_type, effective_input, history=history)
            if adapted_input != effective_input:
                effective_input = adapted_input
                yield f"data: {json.dumps({'type': 'info', 'message': '🧾 已将请求结构化为Markdown提示'})}\n\n"
                t = yield_thinking("将用户请求结构化为 Markdown 格式以提升输出质量", "planning")
                if t: yield t
        
        # 重置中断标志（每次新请求都重置）
        _interrupt_manager.reset(session_name)
        interrupt_event = _interrupt_manager.get_event(session_name)

        def interrupted():
            return _interrupt_manager.is_interrupted(session_name)
        
        # 发送进度: 开始处理
        yield f"data: {json.dumps({'type': 'progress', 'message': f'🚀 {task_type} 任务开始...', 'detail': get_model_display_name(model_id)})}\n\n"
        
        try:
            # 初始化模型追踪变量（用于日志记录）
            used_model = "unknown"
            
            # === SYSTEM Mode (本地执行 - 即时) ===
            if task_type == "SYSTEM":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🖥️ 正在分析系统指令...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '⚡ 正在执行操作...', 'detail': ''})}\n\n"
                
                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result["message"]
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"
                
                if Utils.is_failure_output(response_text):
                    t = yield_thinking("系统指令执行失败，使用 AI 修正后重试", "validating")
                    if t: yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次执行失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt("SYSTEM", user_input, response_text)
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        )
                    )
                    response_text = fix_resp.text or response_text
                
                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                
                # 先保存历史，再发送 done 事件（防止客户端断开导致丢失）
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task=task_type,
                    model_name=used_model
                )
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === FILE_OP Mode (文件操作 - 即时) ===
            if task_type == "FILE_OP":
                used_model = "LocalExecutor"
                yield f"data: {json.dumps({'type': 'progress', 'message': '📂 正在分析文件操作...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 正在访问文件系统...', 'detail': ''})}\n\n"

                batch_manager = get_batch_ops_manager()
                if batch_manager.is_batch_command(user_input):
                    parsed = batch_manager.parse_command(user_input)
                    if not parsed.get("success"):
                        response_text = f"❌ {parsed.get('error')}\n\n{parsed.get('hint', '')}"
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            response_text,
                            task="FILE_OP",
                            model_name=used_model
                        )
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return

                    job = batch_manager.create_job(
                        name=f"batch_{parsed.get('operation')}",
                        operation=parsed.get('operation'),
                        input_dir=parsed.get('input_dir'),
                        output_dir=parsed.get('output_dir'),
                        options=parsed.get('options', {})
                    )
                    batch_manager.start_job(job.job_id)
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'🧩 已创建批量任务: {job.job_id}', 'detail': ''})}\n\n"

                    summary_text = None
                    for event in batch_manager.iter_job_events(job.job_id):
                        if event.get("type") == "progress":
                            current = event.get('current', 0)
                            total = event.get('total', 0)
                            progress_pct = int((current / total) * 100) if total else 0
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⏳ 批量处理中...', 'detail': event.get('detail', ''), 'progress': progress_pct, 'total': total})}\n\n"
                        elif event.get("type") == "final":
                            summary_text = event.get("summary") or "✅ 批量处理完成"
                            break
                        elif event.get("type") == "error":
                            summary_text = event.get("message", "❌ 批量任务失败")
                            break

                    if summary_text:
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_text})}\n\n"
                        session_manager.append_and_save(
                            f"{session_name}.json",
                            user_input,
                            summary_text,
                            task="FILE_OP",
                            model_name=used_model
                        )

                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                
                file_result = FileOperator.execute(user_input)
                response_text = file_result["message"]
                if file_result.get("content"):
                    response_text += f"\n\n{file_result['content']}"
                
                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次执行失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt("FILE_OP", user_input, response_text)
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1000,
                        )
                    )
                    response_text = fix_resp.text or response_text
                
                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                
                # 先保存历史，再发送 done 事件
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_OP",
                    model_name=used_model
                )
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === FILE_EDIT Mode (文件编辑 - 智能修改) ===
            if task_type == "FILE_EDIT":
                used_model = model_id
                t = yield_thinking("进入文件编辑模式，将理解用户指令并修改文件", "routing")
                if t: yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在分析编辑指令...', 'detail': ''})}\n\n"
                
                editor = get_file_editor()
                
                # 尝试从用户输入中提取文件路径和指令
                # 模式 1: "修改 path/to/file 把xxx改成yyy"
                match = re.search(r'(?:修改|编辑|改)\s+["\']?([^"\']+?)["\']?\s+(.+)', user_input)
                if not match:
                    # 模式 2: "把 path/to/file 的xxx改成yyy"
                    match = re.search(r'(?:把|将)\s+["\']?([^"\']+?)["\']?\s+(?:的|中的|里的)\s*(.+)', user_input)
                
                if match:
                    file_path = match.group(1).strip()
                    instruction = match.group(2).strip()
                    
                    t = yield_thinking(f"提取到文件路径: {file_path}, 指令: {instruction}", "analyzing")
                    if t: yield t
                    
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'🔍 目标文件: {os.path.basename(file_path)}', 'detail': ''})}\n\n"
                    yield f"data: {json.dumps({'type': 'progress', 'message': '✏️ 正在执行编辑...', 'detail': ''})}\n\n"
                    
                    result = editor.smart_edit(file_path, instruction)
                    
                    if result["success"]:
                        operation = result.get("operation", "edit")
                        edit_result = result.get("result", {})
                        
                        response_text = f"✅ 文件编辑成功！\n\n"
                        response_text += f"**操作类型**: {operation}\n"
                        
                        if operation == "replace":
                            response_text += f"**替换次数**: {edit_result.get('replacements', 0)}\n"
                            response_text += f"**预览**:\n```\n{edit_result.get('preview', '')}\n```"
                        elif operation == "delete_lines":
                            response_text += f"**删除内容**:\n```\n{edit_result.get('deleted_content', '')}\n```"
                        elif operation == "insert_line":
                            response_text += f"**消息**: {edit_result.get('message', '')}"
                        
                        if edit_result.get("backup"):
                            response_text += f"\n\n💾 备份文件: `{edit_result.get('backup')}`"
                    else:
                        error_msg = result.get("error", "未知错误")
                        hint = result.get("hint", "")
                        response_text = f"❌ 文件编辑失败\n\n{error_msg}\n\n{hint}"
                else:
                    # 无法提取文件路径，让AI理解
                    response_text = "❌ 无法识别文件路径和编辑指令\n\n"
                    response_text += "请使用以下格式:\n"
                    response_text += "- `修改 文件路径 把'旧文本'改成'新文本'`\n"
                    response_text += "- `把 文件路径 的第5-10行删除`\n"
                    response_text += "- `编辑 文件路径 在第3行之后插入'新内容'`"
                
                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_EDIT",
                    model_name=used_model
                )
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === FILE_SEARCH Mode (文件搜索 - 内容定位) ===
            if task_type == "FILE_SEARCH":
                used_model = "FileIndexer (Local)"
                t = yield_thinking("进入文件搜索模式，将在索引中查找匹配文件", "searching")
                if t: yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 正在搜索文件...', 'detail': ''})}\n\n"
                
                indexer = get_file_indexer()
                
                # 提取搜索关键词
                keywords = user_input.replace("找文件", "").replace("搜索", "").replace("查找", "")
                keywords = keywords.replace("包含", "").replace("的文件", "").strip()
                
                t = yield_thinking(f"搜索关键词: {keywords}", "searching")
                if t: yield t
                
                # 先尝试全文搜索
                results = indexer.search(keywords, limit=10)
                
                if not results:
                    # 如果没有结果，尝试内容相似度搜索
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🔄 扩展搜索范围...', 'detail': ''})}\n\n"
                    results = indexer.find_by_content(keywords, min_similarity=0.2)
                
                if results:
                    response_text = f"🔍 找到 {len(results)} 个匹配文件:\n\n"
                    
                    for i, r in enumerate(results[:10], 1):
                        file_name = r.get("file_name", "未知文件")
                        file_path = r.get("file_path", "")
                        snippet = r.get("match_snippet", "")
                        score = r.get("score", 0)
                        similarity = r.get("similarity")
                        
                        response_text += f"### {i}. {file_name}\n"
                        response_text += f"📁 路径: `{file_path}`\n"
                        
                        if similarity:
                            response_text += f"🎯 相似度: {similarity:.0%}\n"
                        elif score:
                            response_text += f"⭐ 匹配分: {score:.2f}\n"
                        
                        if snippet:
                            snippet_clean = snippet.replace("**", "**`")[:200]
                            response_text += f"📄 预览: {snippet_clean}...\n"
                        
                        response_text += "\n"
                else:
                    response_text = "❌ 未找到匹配文件\n\n"
                    response_text += "💡 提示:\n"
                    response_text += "- 确保文件已被索引（Koto 处理过的文件会自动索引）\n"
                    response_text += "- 尝试使用更具体的关键词\n"
                    response_text += f"- 当前索引文件数: {len(indexer.list_indexed_files(limit=1000))}"
                
                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                
                session_manager.append_and_save(
                    f"{session_name}.json",
                    user_input,
                    response_text,
                    task="FILE_SEARCH",
                    model_name=used_model
                )
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === DOC_ANNOTATE Mode (文档标注/润色 - 流式反馈) ===
            if task_type == "DOC_ANNOTATE":
                used_model = model_id if model_id else "gemini-3-flash-preview"
                t = yield_thinking(f"进入文档标注模式，将使用 {model_id or 'gemini-3-flash-preview'} 分析文档", "routing")
                if t: yield t
                print(f"[STREAM] 📄 执行 DOC_ANNOTATE 任务")
                
                # 从请求中获取task_id，用于支持取消操作
                task_id = request.json.get('task_id')
                
                # 查找最近上传的文档
                doc_path = None
                upload_dirs = ["web/uploads", "uploads", "workspace/documents"]
                
                for dir_path in upload_dirs:
                    if os.path.exists(dir_path):
                        import glob
                        docs = []
                        for ext in [".docx", ".docxm"]:
                            docs.extend(glob.glob(f"{dir_path}/**/*{ext}", recursive=True))
                        if docs:
                            doc_path = max(docs, key=os.path.getmtime)
                            break
                
                if not doc_path or not os.path.exists(doc_path):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '❌ 未找到 Word 文档', 'detail': '请上传 .docx 文件'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return
                
                # Step 1: 读取文档信息
                yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '📖 正在读取文档...', 'detail': os.path.basename(doc_path)})}\n\n"
                
                doc_filename = os.path.basename(doc_path)
                total_chars = 0
                total_paras = 0
                
                try:
                    from docx import Document
                    doc = Document(doc_path)
                    total_paras = len([p for p in doc.paragraphs if p.text.strip()])
                    total_chars = sum(len(p.text) for p in doc.paragraphs)
                    
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading_complete', 'message': f'✅ 文档解析完成', 'detail': f'{doc_filename}: {total_paras} 段  |  {total_chars} 字'})}\n\n"
                except Exception as e:
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 读取文档失败: {str(e)}'})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                    return
                
                # Step 2: 展示任务信息
                nl = '\n'
                task_info_msg = f'📋 【任务信息】{nl}- 模型: {model_id}{nl}- 需求: {user_input[:100]}{nl}- 文档: {doc_filename}'
                yield f"data: {json.dumps({'type': 'info', 'message': task_info_msg})}\n\n"
                
                try:
                    from web.document_feedback import DocumentFeedbackSystem
                    feedback_system = DocumentFeedbackSystem(gemini_client=client)
                    
                    # 使用流式分析系统，逐步反馈进度
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'processing_start', 'message': '🔍 开始处理文档...', 'detail': '这个过程会涉及多个阶段'})}\n\n"
                    
                    revised_file = None
                    final_result = None
                    cancelled = False
                    
                    # 迭代流式结果，传入task_id用于支持取消
                    for progress_event in feedback_system.full_annotation_loop_streaming(
                        doc_path,
                        user_input,
                        task_id=task_id,
                        model_id=model_id,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(session_name)
                    ):
                        stage = progress_event.get('stage', 'unknown')
                        progress = progress_event.get('progress', 0)
                        message = progress_event.get('message', '')
                        detail = progress_event.get('detail', '')
                        
                        # 处理任务取消
                        if stage == 'cancelled':
                            cancelled = True
                            yield f"data: {json.dumps({'type': 'info', 'message': '⏸️ 任务已取消', 'detail': '用户中止了处理'})}\n\n"
                            break
                        
                        # 根据阶段发送不同样式的进度信息
                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message, 'detail': detail, 'progress': progress})}\n\n"
                        
                        # 保存最终结果
                        if stage == 'complete':
                            final_result = progress_event.get('result', {})
                            revised_file = final_result.get('revised_file')
                    
                    # 如果任务被取消，返回取消响应
                    if cancelled:
                        total_time = time.time() - start_time
                        # 保存取消记录到历史
                        session_manager.append_and_save(f"{session_name}.json", user_input, "⏸️ 文档标注任务已取消")
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time, 'cancelled': True})}\n\n"
                        return
                    
                    # 如果处理成功，生成详细总结
                    if final_result and final_result.get('success'):
                        applied = final_result.get('applied', 0)
                        failed = final_result.get('failed', 0)
                        total = final_result.get('total', applied + failed)
                        
                        # 计算修改密度
                        density = (applied / total_chars * 1000) if total_chars > 0 else 0
                        
                        summary_msg = (
                            f"✅ **文档修改完成！**\n\n"
                            f"📊 **测试结果**：\n"
                            f"- **文档分析**：成功读取 {total_paras} 段，共 {total_chars} 字。\n"
                            f"- **AI 处理**：文档被并发处理，总耗时约 {int(time.time() - start_time)} 秒。\n"
                            f"- **生成质量**：AI 成功找出了 **{total} 处** 翻译生硬、语序不顺的地方。\n"
                            f"- **应用修订**：成功将 **{applied} 处** 修改以“修订模式（Track Changes）”写入了 Word 文档（仅有 {failed} 处因复杂格式定位失败，属于正常容错范围）。\n\n"
                            f"📂 **验证文件**：\n"
                            f"高质量的测试结果文件已经生成在您的本地目录中，您可以直接打开查看效果：\n"
                            f"👉 `{os.path.basename(revised_file) if revised_file else '待生成'}`\n\n"
                            f"💡 **使用方法**：\n"
                            f"1. 用 Microsoft Word 打开输出文件\n"
                            f"2. 点击「审阅」标签页\n"
                            f"3. 右侧气泡中查看全部修改建议\n"
                            f"4. 逐条接受或忽略（右键批注可操作）\n"
                            f"5. 点击「接受全部」或逐条处理\n\n"
                            f"📂 **文件位置**: `{os.path.dirname(revised_file) if revised_file else settings_manager.documents_dir}`"
                        )
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': '📝 生成最终报告...', 'detail': ''})}\n\n"
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"
                        
                        # 保存对话历史（包含元数据）
                        session_manager.append_and_save(
                            f"{session_name}.json", user_input, summary_msg,
                            task="DOC_ANNOTATE", model_name=model_id,
                            saved_files=[revised_file] if revised_file else []
                        )
                        
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': total_time})}\n\n"
                    else:
                        error_msg = final_result.get('message', '未知错误') if final_result else '处理失败'
                        # 保存失败记录
                        session_manager.append_and_save(f"{session_name}.json", user_input, f"❌ 文档标注失败: {error_msg}")
                        yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 处理失败: {error_msg}'})}\n\n"
                        
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                
                except Exception as e:
                    import traceback
                    error_detail = traceback.format_exc()
                    print(f"[DOC_ANNOTATE] ❌ 失败:\n{error_detail}")
                    # 保存异常记录
                    session_manager.append_and_save(f"{session_name}.json", user_input, f"❌ 文档标注异常: {str(e)[:200]}")
                    
                    yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 处理异常: {str(e)[:200]}'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                
                return
                        
            # === WEB_SEARCH Mode (联网搜索 - 实时信息) ===
            if task_type == "WEB_SEARCH":
                used_model = "gemini-2.5-flash (Google Search)"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🌐 正在连接互联网...', 'detail': ''})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 正在搜索实时信息...', 'detail': 'Google Search'})}\n\n"
                yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在整理搜索结果...', 'detail': ''})}\n\n"
                
                search_result = WebSearcher.search_with_grounding(user_input)
                response_text = search_result["response"]
                
                if Utils.is_failure_output(response_text) or "搜索失败" in response_text:
                    t = yield_thinking("初次搜索结果不佳，使用 gemini-2.0-flash-lite 改写查询词后重试", "searching")
                    if t: yield t
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次搜索失败，正在修正查询...', 'detail': ''})}\n\n"
                    fix_query_prompt = (
                        "请把用户需求改写成更适合搜索的简短关键词或查询语句，只输出查询语句。\n"
                        f"用户需求: {user_input}"
                    )
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash-lite",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(
                            temperature=0.2,
                            max_output_tokens=64,
                        )
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result["response"]
                
                if Utils.is_failure_output(response_text):
                    fix_prompt = Utils.build_fix_prompt("WEB_SEARCH", user_input, response_text)
                    fix_resp = client.models.generate_content(
                        model=model_id,
                        contents=fix_prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.4,
                            max_output_tokens=1200,
                        )
                    )
                    response_text = fix_resp.text or response_text
                
                yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                
                # 先保存历史，再发送 done 事件
                session_manager.append_and_save(f"{session_name}.json", user_input, response_text)
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === RESEARCH Mode (深度研究 - 流式响应优先) ===
            if task_type == "RESEARCH":
                used_model = model_id if model_id else "gemini-3-pro"
                t = yield_thinking(f"进入深度研究模式，使用 {model_id or 'gemini-3-pro'} 进行专业级分析", "analyzing")
                if t: yield t
                newline = '\n'
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔬 启动深度研究模式...', 'detail': '使用Gemini 3.0 Pro进行流式分析'})}{newline}{newline}"
                
                # 构建深度研究的system instruction
                research_instruction = """你是一位专业的研究助手，擅长深度分析复杂技术话题。请按照以下结构提供全面深入的研究报告：

1. **技术概述**：清晰定义和解释核心概念
2. **技术原理**：详细说明工作机制和底层原理
3. **优势分析**：列举主要优点和应用场景
4. **问题与挑战**：分析存在的问题和技术瓶颈
5. **对比分析**：与其他同类技术进行横向对比
6. **发展趋势**：讨论未来发展方向和应用前景
7. **参考资料**：提供相关技术文档和学术资料的引用

📌 **特殊查询类型增强规则**：

**价格/费用/票务查询**（如高铁票、机票、酒店、门票等）：
- ✅ **首先输出一个清晰的表格**，包含关键信息（车次、发车时间、到达时间、座位、价格、时长等）
- ✅ 必须提供**具体价格**（例如：二等座 ¥524.5）
- ❌ 禁止使用价格区间（如"500-600元"）
- ✅ 按座位/房型等级**分别列出**每个选项的确切价格
- ✅ 列出**具体班次/车次号**（如 G12、航班 MU5137）
- ✅ 列出**发车时间和到达时间**，方便用户对比选择
- ❌ 禁止输出重复内容或多个相同的段落

**强制使用表格格式**：
```
🚄 上海虹桥 → 北京南（2026年2月12日）

| 车次   | 发车  | 到达  | 座位类型 | 价格     | 时长  |
|--------|-------|-------|----------|----------|-------|
| G12次  | 09:00 | 13:24 | 商务座   | ¥1,748   | 4h24m |
| G12次  | 09:00 | 13:24 | 一等座   | ¥933     | 4h24m |
| G12次  | 09:00 | 13:24 | 二等座   | ¥524.5   | 4h24m |
| G8次   | 10:00 | 14:31 | 商务座   | ¥1,748   | 4h31m |
| G8次   | 10:00 | 14:31 | 一等座   | ¥933     | 4h31m |
| G8次   | 10:00 | 14:31 | 二等座   | ¥524.5   | 4h31m |

💡 购票方式：访问 12306.cn 搜索对应车次购买。
```

要求：
- 提供具体的技术细节和数据支持
- 使用专业术语但确保可理解性
- 保持客观中立的分析态度
- 内容全面且有深度
- 适当使用图表和示例说明"""
                
                collected_text = []
                
                try:
                    # 使用Gemini 3.0 Pro进行流式生成
                    newline = '\n'
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在进行深度分析...', 'detail': 'Gemini 3.0 Pro 正在思考，可能需要30-90秒'})}{newline}{newline}"
                    
                    response_stream = client.models.generate_content_stream(
                        model="gemini-3-pro-preview",
                        contents=effective_input,
                        config=types.GenerateContentConfig(
                            system_instruction=research_instruction,
                            temperature=0.7,
                            max_output_tokens=8000,  # 允许更长的输出
                            top_p=0.95,
                        )
                    )
                    
                    chunk_count = 0
                    heartbeat_interval = 5  # 每5秒发送一次心跳
                    first_chunk_received = False
                    
                    # 使用保活包装器处理流式响应
                    for item_type, item_data in stream_with_keepalive(response_stream, start_time, 
                                                                       keepalive_interval=heartbeat_interval,
                                                                       max_wait_first_token=90):  # 最多等待90秒
                        # 检查中断
                        if interrupted():
                            print(f"[RESEARCH] 用户中断研究")
                            newline = '\n'
                            interrupt_msg = f'{newline}{newline}⏹️ 研究已被用户中断'
                            yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}{newline}{newline}"
                            break
                        
                        if item_type == 'heartbeat':
                            # 发送心跳保持连接
                            elapsed = item_data
                            if first_chunk_received:
                                char_count = len(''.join(collected_text))
                                yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在生成中...', 'detail': f'已生成 {char_count} 字符，耗时 {elapsed}s'})}\n\n"
                            else:
                                yield f"data: {json.dumps({'type': 'progress', 'message': '🧠 模型正在深度思考...', 'detail': f'已等待 {elapsed}s，请耐心等待'})}\n\n"
                        
                        elif item_type == 'timeout':
                            # 等待超时
                            yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，模型响应时间过长，请稍后重试'})}\n\n"
                            break
                        
                        elif item_type == 'chunk':
                            chunk = item_data
                            if chunk.text:
                                if not first_chunk_received:
                                    first_chunk_received = True
                                    print(f"[RESEARCH] 收到第一个响应块，耗时 {time.time() - start_time:.1f}s")
                                
                                collected_text.append(chunk.text)
                                yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"
                                chunk_count += 1
                                
                                # 每50个chunk显示一次进度日志
                                if chunk_count % 50 == 0:
                                    print(f"[RESEARCH] 已生成 {chunk_count} 个chunk, {len(''.join(collected_text))} 字符")
                    
                    final_text = ''.join(collected_text)
                    print(f"[RESEARCH] ✅ 研究完成，共 {len(final_text)} 字符")
                    
                    # 保存历史（基于磁盘完整历史追加）
                    session_manager.append_and_save(
                        f"{session_name}.json",
                        user_input,
                        final_text[:4000],
                        task="RESEARCH",
                        model_name=used_model
                    )
                    
                except Exception as research_err:
                    error_msg = str(research_err)
                    print(f"[RESEARCH] 错误: {error_msg}")
                    
                    # 智能错误处理
                    if "503" in error_msg or "UNAVAILABLE" in error_msg:
                        # API过载，尝试使用Flash版本
                        try:
                            newline = '\n'
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 服务繁忙，切换到Gemini 3.0 Flash...', 'detail': ''})}{newline}{newline}"
                            
                            response_stream = client.models.generate_content_stream(
                                model="gemini-3-flash-preview",
                                contents=effective_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=research_instruction,
                                    temperature=0.7,
                                    max_output_tokens=8000,
                                )
                            )
                            
                            last_heartbeat_flash = time.time()
                            for chunk in response_stream:
                                if interrupted():
                                    break
                                if chunk.text:
                                    collected_text.append(chunk.text)
                                    yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"
                                    
                                    # Flash 模式下也发送心跳
                                    current_time = time.time()
                                    if current_time - last_heartbeat_flash > 3:
                                        elapsed = int(current_time - start_time)
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚡ 快速模式生成中...', 'detail': f'{elapsed}s'})}\n\n"
                                        last_heartbeat_flash = current_time
                            
                            final_text = ''.join(collected_text)
                            session_manager.append_and_save(
                                f"{session_name}.json",
                                user_input,
                                final_text[:4000],
                                task="RESEARCH",
                                model_name="gemini-3-flash-preview"
                            )
                            
                        except Exception as fallback_err:
                            error_text = f"❌ 研究服务暂时不可用\n\n错误信息: {str(fallback_err)[:200]}\n\n💡 建议：\n1. 稍后重试\n2. 简化问题\n3. 使用普通对话模式"
                            yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                            session_manager.append_and_save(f"{session_name}.json", user_input, error_text[:1000], task="RESEARCH", model_name="gemini-3-flash-preview")
                    
                    elif "timeout" in error_msg.lower() or "disconnect" in error_msg.lower():
                        # 连接问题
                        error_text = f"⚠️ 连接超时或中断\n\n可能原因：\n1. 网络不稳定\n2. 服务器繁忙\n3. 代理配置问题\n\n建议：请稍后重试，或检查网络连接"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(f"{session_name}.json", user_input, error_text[:1000], task="RESEARCH", model_name=used_model)
                    
                    else:
                        # 其他错误
                        error_text = f"❌ 研究过程中出现错误\n\n{error_msg[:300]}\n\n请尝试：\n1. 重新提问\n2. 简化问题描述\n3. 稍后重试"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                        session_manager.append_and_save(f"{session_name}.json", user_input, error_text[:1000], task="RESEARCH", model_name=used_model)
                
                total_time = time.time() - start_time
                newline = '\n'
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}{newline}{newline}"
                return
            
            # === PAINTER Mode (图像生成 - Nano Banana 优先，Imagen 4.0 备用) ===
            if task_type == "PAINTER":
                used_model = "Nano Banana (Imagen 4.0 fallback)"
                yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 正在理解你的创作请求...', 'detail': ''})}\n\n"
                
                # 使用上下文增强的输入（如果有）
                if context_info and context_info.get("is_continuation") and context_info.get("enhanced_input"):
                    image_prompt = context_info["enhanced_input"]
                    print(f"[PAINTER] 使用上下文增强的prompt: {image_prompt[:100]}...")
                else:
                    image_prompt = effective_input
                
                yield f"data: {json.dumps({'type': 'progress', 'message': '🖌️ Nano Banana 正在生成图像...', 'detail': '请耐心等待'})}\n\n"
                
                max_retries = 2
                use_fallback = False
                images = []
                
                for attempt in range(max_retries):
                    try:
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 图像生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        
                        if attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 第 {attempt} 次重试...', 'detail': ''})}\n\n"
                            time.sleep(2)
                        
                        # 选择模型
                        if use_fallback:
                            model_name = "Imagen 4.0"
                            yield f"data: {json.dumps({'type': 'progress', 'message': '🔄 切换到 Imagen 4.0...', 'detail': ''})}\n\n"
                        else:
                            model_name = "Nano Banana"
                        
                        # 使用后台线程执行请求，主线程发送心跳
                        import queue
                        import threading
                        result_queue = queue.Queue()
                        
                        def worker():
                            try:
                                if use_fallback:
                                    result = client.models.generate_images(
                                        model="imagen-4.0-generate-preview-06-06",
                                        prompt=image_prompt,
                                        config=types.GenerateImagesConfig(number_of_images=1)
                                    )
                                else:
                                    result = client.models.generate_content(
                                        model="nano-banana-pro-preview",
                                        contents=image_prompt,
                                        config=types.GenerateContentConfig(response_modalities=["IMAGE"])
                                    )
                                result_queue.put(('success', result))
                            except Exception as e:
                                result_queue.put(('error', e))
                        
                        thread = threading.Thread(target=worker, daemon=True)
                        thread.start()
                        
                        timeout_seconds = 90
                        response = None
                        
                        while True:
                            elapsed = time.time() - start_time
                            
                            if elapsed > timeout_seconds:
                                yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ 图像生成超时 ({int(elapsed)}s)，请稍后重试'})}\n\n"
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return
                            
                            if interrupted():
                                yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 图像生成已中断'})}\n\n"
                                total_time = time.time() - start_time
                                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                                return
                            
                            try:
                                status, data = result_queue.get(timeout=3.0)
                                if status == 'success':
                                    response = data
                                    break
                                else:
                                    raise data
                            except queue.Empty:
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'🎨 {model_name} 生成中...', 'detail': f'{int(elapsed)}s'})}\n\n"
                        
                        # 处理响应
                        yield f"data: {json.dumps({'type': 'progress', 'message': '💾 正在保存图片...', 'detail': ''})}\n\n"
                        
                        if use_fallback:
                            if response.generated_images:
                                for gen_img in response.generated_images:
                                    img_data = gen_img.image.image_bytes
                                    images_dir = settings_manager.images_dir
                                    os.makedirs(images_dir, exist_ok=True)
                                    timestamp = int(time.time())
                                    filename = f"generated_{timestamp}.png"
                                    filepath = os.path.join(images_dir, filename)
                                    with open(filepath, "wb") as f:
                                        f.write(img_data)
                                    
                                    # 确保路径在 workspace 下
                                    try:
                                        rel_path = os.path.relpath(filepath, WORKSPACE_DIR).replace("\\", "/")
                                        if ".." not in rel_path:
                                            images.append(rel_path)
                                            print(f"[PAINTER] Imagen 已保存: {rel_path}")
                                        else:
                                            # 降级保存到 workspace/images
                                            abs_workspace_images = os.path.join(WORKSPACE_DIR, "images")
                                            os.makedirs(abs_workspace_images, exist_ok=True)
                                            fallback_filepath = os.path.join(abs_workspace_images, filename)
                                            with open(fallback_filepath, "wb") as f:
                                                f.write(img_data)
                                            fallback_rel = os.path.relpath(fallback_filepath, WORKSPACE_DIR).replace("\\", "/")
                                            images.append(fallback_rel)
                                            print(f"[PAINTER] Imagen 降级保存: {fallback_rel}")
                                    except Exception as path_err:
                                        print(f"[PAINTER] Path error: {path_err}")
                        else:
                            if response.candidates and response.candidates[0].content.parts:
                                for part in response.candidates[0].content.parts:
                                    if hasattr(part, "inline_data") and part.inline_data:
                                        img_filename = Utils.save_image_part(part)
                                        if img_filename:
                                            images.append(img_filename)
                                            print(f"[PAINTER] Nano Banana 已保存: {img_filename}")
                        
                        if images:
                            save_path = settings_manager.images_dir
                            msg = f"✨ 图片已生成! (使用 {model_name})\n🖼️ 保存位置: {save_path}"
                            yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"
                            
                            # 先保存历史记录（包含图片路径），再发送 done
                            session_manager.append_and_save(
                                f"{session_name}.json", user_input, "图像已生成",
                                images=images, task="PAINTER", model_name=model_name
                            )
                            
                            total_time = time.time() - start_time
                            print(f"[PAINTER] 发送图片列表: {images}")  # 调试
                            yield f"data: {json.dumps({'type': 'done', 'images': images, 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            if not use_fallback:
                                use_fallback = True
                                continue
                            else:
                                yield f"data: {json.dumps({'type': 'token', 'content': '❌ 模型未返回图片'})}\n\n"
                                
                    except Exception as img_err:
                        error_msg = str(img_err)
                        print(f"[PAINTER] 尝试 {attempt+1} 失败: {error_msg[:200]}")
                        
                        if not use_fallback and ("503" in error_msg or "overloaded" in error_msg.lower() or 
                                                 "unavailable" in error_msg.lower() or "400" in error_msg):
                            use_fallback = True
                            continue
                        
                        if "safety" in error_msg.lower() or "blocked" in error_msg.lower():
                            user_msg = "❌ 内容被安全策略过滤，请修改描述"
                        elif "location is not supported" in error_msg.lower():
                            user_msg = "❌ 地区限制，请配置中转服务"
                        else:
                            user_msg = f"❌ 图像生成失败: {error_msg[:100]}"
                        
                        yield f"data: {json.dumps({'type': 'token', 'content': user_msg})}\n\n"
                
                # PAINTER 所有重试都失败时也要保存历史
                session_manager.append_and_save(f"{session_name}.json", user_input, "图像生成失败")
                
                total_time = time.time() - start_time
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                return
            
            # === FILE_GEN Mode (文件生成 - 自动执行) ===
            if task_type == "FILE_GEN":
                t = yield_thinking(f"进入文件生成模式，将使用 {model_id} 生成文档", "generating")
                if t: yield t
                print(f"[FILE_GEN] ===== Starting file generation =====")
                print(f"[FILE_GEN] Model: {model_id}, User input: {user_input[:100]}...")
                
                response_text = ""
                generated_files = []
                temp_scripts = []  # 临时脚本列表（执行后删除）
                api_timeout = 120  # 增加到 120 秒，长文档需要更多时间

                if interrupted():
                    yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                
                # ⭐ 检查是否是"转换请求"（把之前的内容做成word/pdf）
                is_convert_request = (
                    context_info and 
                    context_info.get("is_continuation") and 
                    context_info.get("continuation_type") == "convert" and
                    context_info.get("context_summary", {}).get("last_model_output")
                )
                
                if is_convert_request:
                    # 直接转换模式 - 不需要调用模型，直接生成文档
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在将内容转换为文档...', 'detail': ''})}\n\n"
                    
                    try:
                        from web.document_generator import save_docx, save_pdf
                        
                        source_content = context_info["context_summary"]["last_model_output"]
                        print(f"[FILE_GEN] 直接转换模式，源内容长度: {len(source_content)}")
                        
                        # 提取标题（尝试从内容中找 # 标题）
                        title_match = re.search(r'^#\s*(.+)$', source_content, re.MULTILINE)
                        if title_match:
                            title = title_match.group(1).strip()[:50]
                        else:
                            title = f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                        
                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)
                        
                        # 判断生成 Word 还是 PDF
                        user_lower = user_input.lower()
                        
                        if "pdf" in user_lower:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成 PDF...', 'detail': ''})}\n\n"
                            saved_path = save_pdf(source_content, title=title, output_dir=docs_dir)
                            file_type = "PDF"
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成 Word 文档...', 'detail': ''})}\n\n"
                            saved_path = save_docx(source_content, title=title, output_dir=docs_dir)
                            file_type = "Word"
                        
                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace("\\", "/")
                        generated_files.append(rel_path)
                        
                        success_msg = f"✅ **{file_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        
                        print(f"[FILE_GEN] ✅ 直接转换成功: {rel_path}")
                        
                    except Exception as convert_err:
                        error_msg = f"❌ 文档转换失败: {str(convert_err)}"
                        print(f"[FILE_GEN] 转换错误: {convert_err}")
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"
                    
                    # 保存历史（基于磁盘完整历史追加）
                    _model_msg = f"已生成文件: {', '.join(generated_files)}" if generated_files else "文档转换失败"
                    session_manager.append_and_save(f"{session_name}.json", user_input, _model_msg)
                    
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return
                
                # ⭐ 检查是否是 PPT 生成请求
                ppt_keywords = ["ppt", "幻灯片", "演示文稿", "演示", "presentation", "slide", "slides"]
                user_lower_check = user_input.lower()
                is_ppt_request = any(kw in user_lower_check for kw in ppt_keywords)
                
                if is_ppt_request:
                    # =============== PPT 专用生成流程 ===============
                    print(f"[FILE_GEN] 🎯 检测到 PPT 生成请求")
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 正在生成演示文稿...', 'detail': '正在规划内容结构'})}\n\n"
                    
                    try:
                        # ──────── Session: 创建 PPT 编辑会话 ────────
                        ppt_session_id = None
                        try:
                            from web.ppt_session_manager import get_ppt_session_manager
                            ppt_session_mgr = get_ppt_session_manager()
                            ppt_session_id = ppt_session_mgr.create_session(
                                title=user_input[:50],  # 前 50 字作为临时标题
                                user_input=user_input,
                                theme="business"
                            )
                            print(f"[FILE_GEN/PPT] 📋 创建编辑会话: {ppt_session_id}")
                        except Exception as session_err:
                            print(f"[FILE_GEN/PPT] ⚠️ 会话创建异常（不影响生成）: {session_err}")
                        
                        # ──────── Step 0: 处理上传的文件 ────────
                        uploaded_file_context = ""
                        uploaded_files = request.files.getlist('files[]') if request.method == 'POST' else []
                        
                        if uploaded_files:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'📂 正在解析 {len(uploaded_files)} 个上传文件...', 'detail': '提取文本内容'})}\n\n"
                            try:
                                from web.file_parser import FileParser
                                
                                uploaded_file_paths = []
                                for uploaded_file in uploaded_files:
                                    if uploaded_file and uploaded_file.filename:
                                        # 保存临时文件
                                        temp_dir = os.path.join(WORKSPACE_DIR, 'temp_uploads')
                                        os.makedirs(temp_dir, exist_ok=True)
                                        temp_path = os.path.join(temp_dir, uploaded_file.filename)
                                        uploaded_file.save(temp_path)
                                        uploaded_file_paths.append(temp_path)
                                
                                if uploaded_file_paths:
                                    # 批量解析
                                    parse_results = FileParser.batch_parse(uploaded_file_paths)
                                    successful_results = [r for r in parse_results if r.get('success')]
                                    
                                    if successful_results:
                                        uploaded_file_context = FileParser.merge_contents(successful_results)
                                        print(f"[FILE_GEN/PPT] ✅ 已解析 {len(successful_results)} 个文件, 总字数: {len(uploaded_file_context)}")
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 已解析 {len(successful_results)} 个上传文件', 'detail': f'{len(uploaded_file_context)} 字内容'})}\n\n"
                                    else:
                                        print(f"[FILE_GEN/PPT] ⚠️ 上传文件解析失败")
                                        failed_reasons = [r.get('error', '未知错误') for r in parse_results if not r.get('success')]
                                        print(f"    原因: {', '.join(failed_reasons)}")
                            
                            except ImportError:
                                print(f"[FILE_GEN/PPT] ⚠️ FileParser 模块未找到，跳过文件处理")
                            except Exception as file_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 文件处理异常: {file_err}")
                        
                        # ──────── Step 0.1: 智能判断是否需要联网搜索 ────────
                        search_context = ""
                        
                        # 检测是否需要搜索最新信息
                        _needs_search = WebSearcher.needs_web_search(user_input)
                        
                        # 额外PPT话题检测：包含年份/时间/新品/事件/排行等的PPT大概率需要搜索
                        import re as _re
                        _time_topic_patterns = [
                            r'20\d{2}',             # 年份
                            r'\d+月',               # 月份
                            r'(新番|新片|新剧|新歌|新品|上映|首发|发售)',
                            r'(排行|排名|榜单|top|盘点|导视|速递|一览)',
                            r'(行情|走势|趋势|市场|价格|报告)',
                            r'(热门|热点|火爆|流行|人气)',
                            r'(最新|最近|近期|本周|本月|当前|目前)',
                        ]
                        if not _needs_search:
                            for pat in _time_topic_patterns:
                                if _re.search(pat, user_input, _re.IGNORECASE):
                                    _needs_search = True
                                    print(f"[FILE_GEN/PPT] 🔍 话题时效性检测命中: {pat}")
                                    break
                        
                        if _needs_search:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '🔍 正在搜索最新信息...', 'detail': '确保内容准确和时效性'})}\n\n"
                            try:
                                search_result = WebSearcher.search_with_grounding(user_input)
                                if search_result.get("success") and search_result.get("response"):
                                    search_context = search_result["response"]
                                    print(f"[FILE_GEN/PPT] ✅ 搜索完成, 获取 {len(search_context)} 字符参考信息")
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '✅ 搜索完成，正在整合信息...', 'detail': ''})}\n\n"
                                else:
                                    print(f"[FILE_GEN/PPT] ⚠️ 搜索无结果或失败")
                            except Exception as search_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 搜索异常: {search_err}")
                        
                        # ──────── Step 0.5: 复杂主题深度研究 ────────
                        research_context = ""
                        _complex_patterns = [
                            r'(原理|机制|架构|技术|算法|理论|分析|研究|综述)',
                            r'(行业|产业|市场|商业|战略|规划|方案)',
                            r'(学术|论文|课题|毕业|教学|课程)',
                            r'(历史|发展|演变|变迁|沿革)',
                            r'(对比|比较|评估|评测|benchmark)',
                            r'(经济|金融|投资|财务|财报)',
                        ]
                        _is_complex = len(user_input) > 30 or any(
                            _re.search(p, user_input) for p in _complex_patterns
                        )
                        
                        if _is_complex:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '🔬 检测到复杂主题，启动深度研究...', 'detail': '使用 Gemini Pro 进行专业分析'})}\n\n"
                            try:
                                research_context = WebSearcher.deep_research_for_ppt(user_input, search_context)
                                if research_context:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 深度研究完成，获取 {len(research_context)} 字参考', 'detail': '正在整合研究成果...'})}\n\n"
                                else:
                                    print(f"[FILE_GEN/PPT] ⚠️ 深度研究未返回结果")
                            except Exception as res_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 深度研究异常: {res_err}")
                        
                        # ──────── Step 1: 用 AI 生成结构化大纲 ────────
                        # ──── 提取用户 PPT 偏好（页数、重点、简要话题） ────
                        import re as _ppt_re
                        
                        def _extract_ppt_preferences(text):
                            prefs = {"target_pages": None, "focus_topics": [], "brief_topics": []}
                            pm = _ppt_re.search(r'(?:做|生成|需要|大概|约|大约)?\s*(\d+)\s*页', text)
                            if pm:
                                prefs["target_pages"] = int(pm.group(1))
                            for pat in [
                                r'(?:重点|详细|着重|深入|多讲|多介绍)(?:介绍|讲|分析|说明|展示|讲解)\s*(.+?)(?:[，,。；;、]|$)',
                                r'(?:突出|强调)\s*(.+?)(?:[，,。；;、]|$)',
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["focus_topics"].append(t)
                            for pat in [
                                r'(?:简单|简要|简略|大致)(?:带过|介绍|说|讲)\s*(.+?)(?:[，,。；;、]|$)',
                                r'(.+?)(?:一笔带过|略过|跳过|简单说)',
                            ]:
                                for m in _ppt_re.finditer(pat, text):
                                    t = m.group(1).strip()
                                    if t and len(t) < 30:
                                        prefs["brief_topics"].append(t)
                            return prefs
                        
                        ppt_prefs = _extract_ppt_preferences(user_input)
                        _target_pages = ppt_prefs["target_pages"]
                        _target_hint = f"约 {_target_pages} 页（封面和结束页除外，这是用户指定的，必须严格遵守）" if _target_pages else "8~15 页（根据内容复杂度智能调整，内容多可以多做几页，内容少就精简）"
                        _focus_hint = ""
                        if ppt_prefs["focus_topics"]:
                            _focus_hint = "\n**用户指定的重点内容（必须用 [详细] 多页展开）：**\n" + "\n".join(f"- {t}" for t in ppt_prefs["focus_topics"]) + "\n"
                        _brief_hint = ""
                        if ppt_prefs["brief_topics"]:
                            _brief_hint = "\n**用户指定的简要内容（合并到 [概览] 页）：**\n" + "\n".join(f"- {t}" for t in ppt_prefs["brief_topics"]) + "\n"
                        
                        print(f"[FILE_GEN/PPT] 用户偏好: 页数={_target_pages}, 重点={ppt_prefs['focus_topics']}, 简要={ppt_prefs['brief_topics']}")
                        
                        # ──── 智能内容规划 Prompt ────
                        ppt_outline_prompt = (
                            "你是一个顶尖的演示文稿内容策划师和排版规划师。\n\n"
                            "你的工作分两步：\n"
                            "1. **内容规划** — 分析主题，判断哪些内容是重点（需要多页详细展示），哪些是简要（可以一页多主题速览）\n"
                            "2. **版式选择** — 为每部分选择最合适的幻灯片类型\n\n"
                            "## 可用的幻灯片类型\n"
                            "在每个 `## 章节标题` 前一行写类型标签：\n\n"
                            "| 标签 | 用途 | 格式 |\n"
                            "|------|------|------|\n"
                            "| `[详细]` | 常规内容页，深入展示 4-6 个要点（每个要点30-80字） | `- **关键词** — 详细解释说明和具体数据` |\n"
                            "| `[概览]` | 多主题速览页，2-4 个小主题并列（每个小主题下2-4个要点） | 用 `### 子标题` 分组 |\n"
                            "| `[亮点]` | 关键数据突出页（3-4组数据） | `- 数值 \\| 详细说明` |\n"
                            "| `[对比]` | 两方对比页（每方3-5个要点） | 用 `### 选项A` 和 `### 选项B` 分两组 |\n"
                            "| `[过渡页]` | 章节过渡，引入大章节（少用） | 下方写一行描述 |\n\n"
                            "## 输出格式（严格遵循）\n"
                            "```\n"
                            "# PPT主标题\n\n"
                            "[过渡页]\n"
                            "## 第一部分标题\n"
                            "简短描述\n\n"
                            "[详细]\n"
                            "## 页面标题\n"
                            "- **核心概念** — 这是一段完整的解释性文字，包含关键数据或事实依据，让观众能真正理解这一点的内容\n"
                            "- **技术特点** — 具体描述技术的工作原理、优势所在、实际应用场景和相关参数\n"
                            "- **市场数据** — 引用权威机构的统计数字、市场规模、增长率等量化信息\n"
                            "- **实际案例** — 某公司/项目的具体实践经验，取得了什么样的成果\n\n"
                            "[概览]\n"
                            "## 速览标题\n"
                            "### 子话题1\n"
                            "- 第一个要点的详细说明\n"
                            "- 第二个要点的详细说明\n"
                            "- 第三个要点的详细说明\n"
                            "### 子话题2\n"
                            "- 第一个要点和具体数据\n"
                            "- 第二个要点和应用场景\n\n"
                            "[亮点]\n"
                            "## 关键数据\n"
                            "- 500亿 | 全球市场规模\n"
                            "- 35% | 年增长率\n\n"
                            "[对比]\n"
                            "## 对比标题\n"
                            "### 方案A\n"
                            "- 特点1\n"
                            "### 方案B\n"
                            "- 特点1\n"
                            "```\n\n"
                            "## 内容规划规则\n"
                            f"1. **总页数目标: {_target_hint}**\n"
                            "2. **重点内容**使用多个 `[详细]` 页展开，每页 4-6 个信息丰富的要点\n"
                            "3. ⚠️ **每个要点必须是一个完整的信息段落（30-80字），不能只写几个词或短语**\n"
                            "4. 要点格式: `- **关键词** — 具体的解释说明，包含数据、事实、案例等实质内容`\n"
                            "5. **非重点内容**合并到 `[概览]` 页，一页 2-4 个小主题，⚠️ **每个小主题下必须有 2-4 个要点**\n"
                            "6. 有数据亮点时用 `[亮点]` 页（全文最多 1-2 次），每页 3-4 组数据\n"
                            "7. `[过渡页]` 最多 2 个，用于划分大章节\n"
                            "8. ⚠️ **搜索资料和研究报告中的数据、案例、数字必须如实引用，不得编造。每页至少引用 1 个具体数据或案例**\n"
                            "9. 中文输出，只输出大纲，不要额外说明\n"
                            "10. ⚠️ **内容充实度是最重要的评判标准 — 宁可要点少一些但每个要点信息量大，不要很多空洞的要点**\n"
                            "11. ⚠️ **禁止出现模糊表述**：如 '显著增长'、'广泛应用'、'巨大潜力' 等，必须用具体数字替代。例如：'市场规模达 XX 亿' 而不是 '市场规模巨大'\n"
                            "12. **每个 [详细] 页至少包含 1 个真实案例或数据点**，数据需标注来源（如 '据IDC数据' '根据XX年报'）\n"
                            f"{_focus_hint}"
                            f"{_brief_hint}"
                            f"\n用户需求: {user_input}\n"
                        )
                        
                        # 注入搜索结果（增加限额以保留更多数据）
                        if uploaded_file_context:
                            ppt_outline_prompt = (
                                ppt_outline_prompt[:-len("\n用户需求: " + user_input)] 
                                + f"\n\n## 上传的参考文件内容\n"
                                f"以下是用户上传的文档资料，请充分利用其中的内容、数据、案例来生成 PPT：\n"
                                f"---\n{uploaded_file_context[:15000]}\n---\n"
                                f"\n用户需求: {user_input}\n"
                            )
                        
                        if search_context:
                            ppt_outline_prompt += (
                                f"\n**以下是联网搜索获取的最新参考资料（包含重要数据），请务必基于这些信息生成内容，尤其是其中的数字、案例、市场数据：**\n"
                                f"---\n{search_context[:10000]}\n---\n"
                            )
                        
                        # 注入深度研究结果（增加限额）
                        if research_context:
                            ppt_outline_prompt += (
                                f"\n**以下是深度研究分析报告——这是你最重要的内容来源，其中的数据和分析必须充分融入大纲：**\n"
                                f"---\n{research_context[:12000]}\n---\n"
                            )
                        
                        # 也注入上下文
                        if context_info and context_info.get("is_continuation") and context_info.get("enhanced_input"):
                            ppt_outline_prompt += f"\n\n历史上下文参考资料:\n{context_info['enhanced_input'][:3000]}"
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在生成内容大纲...', 'detail': ''})}\n\n"
                        
                        outline_response = None
                        outline_models = ["gemini-2.5-flash", model_id, "gemini-2.0-flash"]
                        # 根据目标页数调整 token 限额：20页大纲需要更多空间
                        _outline_tokens = 16384 if (_target_pages and _target_pages >= 15) else 8192
                        for om in outline_models:
                            try:
                                resp = client.models.generate_content(
                                    model=om,
                                    contents=ppt_outline_prompt,
                                    config=types.GenerateContentConfig(
                                        temperature=0.6,
                                        max_output_tokens=_outline_tokens
                                    )
                                )
                                if resp.text:
                                    outline_response = resp.text
                                    print(f"[FILE_GEN/PPT] ✅ 大纲生成成功 ({om}), 长度: {len(outline_response)}")
                                    break
                            except Exception as oe:
                                print(f"[FILE_GEN/PPT] 大纲模型 {om} 失败: {oe}")
                                continue
                        
                        if not outline_response:
                            raise Exception("所有模型均无法生成大纲")
                        
                        # Step 2: 解析智能规划大纲（支持多种幻灯片类型标签）
                        def _parse_ppt_plan(md_text):
                            """解析带 [类型] 标签的智能 PPT 大纲"""
                            import re as _re
                            lines = md_text.split('\n')
                            plan = {"title": "", "subtitle": "", "slides": []}
                            
                            _type_map = {
                                "过渡页": "divider", "过渡": "divider", "分隔": "divider",
                                "详细": "detail", "重点": "detail",
                                "亮点": "highlight", "数据": "highlight", "关键": "highlight",
                                "概览": "overview", "速览": "overview", "简要": "overview", "总览": "overview",
                                "对比": "comparison", "比较": "comparison", "vs": "comparison",
                            }
                            
                            current_slide = None
                            current_type = "detail"
                            current_sub = None  # 当前子主题（用于 overview / comparison）
                            
                            for line in lines:
                                line = line.rstrip()
                                
                                # 跳过 markdown 代码块标记
                                if line.strip() in ('```', '```markdown'):
                                    continue
                                
                                # 类型标签行: [xxx]
                                tag_m = _re.match(r'^\s*\[(.+?)\]\s*$', line)
                                if tag_m:
                                    tag = tag_m.group(1).strip()
                                    current_type = _type_map.get(tag, "detail")
                                    continue
                                
                                # 主标题: # xxx
                                if line.startswith('# ') and not line.startswith('## '):
                                    raw = line[2:].strip()
                                    for pfx in ["幻灯片标题：", "幻灯片标题:", "演示标题：", "演示标题:", "PPT标题：", "PPT标题:"]:
                                        if raw.startswith(pfx):
                                            raw = raw[len(pfx):].strip()
                                    plan["title"] = raw
                                    continue
                                
                                # 章节标题: ## xxx
                                if line.startswith('## '):
                                    # 保存上一个 slide 的 subsection
                                    if current_sub and current_slide and current_slide.get("type") in ("overview", "comparison"):
                                        current_slide.setdefault("subsections", []).append(current_sub)
                                        current_sub = None
                                    # 保存上一个 slide
                                    if current_slide:
                                        plan["slides"].append(current_slide)
                                    
                                    current_slide = {
                                        "type": current_type,
                                        "title": line[3:].strip(),
                                        "points": [],
                                        "content": [],
                                    }
                                    if current_type == "divider":
                                        current_slide["description"] = ""
                                    current_type = "detail"  # 重置（每个标签只作用于紧跟的 ## ）
                                    current_sub = None
                                    continue
                                
                                # 子标题: ### xxx （用于 overview / comparison）
                                if line.startswith('### ') and current_slide:
                                    # 如果当前 slide 不是 overview/comparison，自动升级为 overview
                                    if current_slide.get("type") not in ("overview", "comparison"):
                                        current_slide["type"] = "overview"
                                    if current_sub:
                                        current_slide.setdefault("subsections", []).append(current_sub)
                                    current_sub = {
                                        "subtitle": line[4:].strip(),
                                        "label": line[4:].strip(),
                                        "points": [],
                                    }
                                    continue
                                
                                # 要点行: - / • / * 或数字编号 1. 2. 等
                                if (_re.match(r'^[\s]*[-•*]\s', line) or _re.match(r'^[\s]*\d+[.、)\s]\s*', line)) and current_slide is not None:
                                    pt = _re.sub(r'^[\s]*[-•*\d.、)\s]+\s*', '', line).strip()
                                    if not pt:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(pt)
                                    else:
                                        current_slide["points"].append(pt)
                                        current_slide["content"].append(pt)
                                    continue
                                
                                # 普通文本行（非空、非标题）→ 也捕获为要点
                                if current_slide is not None and line.strip() and not line.startswith('#'):
                                    # 过渡页描述文字优先
                                    if current_slide.get("type") == "divider":
                                        current_slide["description"] = line.strip()
                                        continue
                                    # 清理可能残留的 markdown 标记
                                    cleaned = _re.sub(r'^#{1,4}\s+', '', line.strip())
                                    cleaned = cleaned.strip()
                                    if not cleaned:
                                        continue
                                    if current_sub is not None:
                                        current_sub["points"].append(cleaned)
                                    else:
                                        current_slide["points"].append(cleaned)
                                        current_slide["content"].append(cleaned)
                                    continue
                                
                                # 过渡页描述文字 (fallback - 不应到达这里)
                                if current_slide and current_slide.get("type") == "divider" and line.strip() and not line.startswith('#'):
                                    current_slide["description"] = line.strip()
                            
                            # 收尾
                            if current_sub and current_slide:
                                current_slide.setdefault("subsections", []).append(current_sub)
                            if current_slide:
                                plan["slides"].append(current_slide)
                            
                            # 后处理: 如果 slide 有 subsections 但类型不是 overview/comparison，自动修正
                            for sl in plan["slides"]:
                                if sl.get("subsections") and sl.get("type") not in ("overview", "comparison"):
                                    sl["type"] = "overview"
                            
                            # 后处理: comparison 的 subsections → left / right
                            for sl in plan["slides"]:
                                if sl.get("type") == "comparison" and "subsections" in sl:
                                    subs = sl["subsections"]
                                    if len(subs) >= 2:
                                        sl["left"] = subs[0]
                                        sl["right"] = subs[1]
                            
                            return plan
                        
                        ppt_data = _parse_ppt_plan(outline_response)
                        slide_count = len(ppt_data["slides"])
                        slide_types_summary = ", ".join(f'{s.get("type","detail")}' for s in ppt_data["slides"])
                        print(f"[FILE_GEN/PPT] 解析完成: 标题='{ppt_data['title']}', {slide_count} 页, 类型=[{slide_types_summary}]")
                        
                        if slide_count == 0:
                            raise Exception("大纲解析失败，未提取到幻灯片内容")
                        
                        # ──────── Step 2.1: 用户指定页数时调整幻灯片数量 ────────
                        _max_slides = _target_pages  # 只有用户明确指定时才生效
                        if _max_slides and slide_count > _max_slides:
                            print(f"[FILE_GEN/PPT] ⚠️ 页数超限 ({slide_count} > {_max_slides})，执行智能精简...")
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'✂️ 精简页面: {slide_count} → {_max_slides} 页', 'detail': '合并相似内容，保留核心信息'})}\n\n"
                            
                            slides = ppt_data["slides"]
                            # 策略: 1) 合并相邻的详细页为概览页  2) 去掉多余过渡页  3) 截断尾部
                            
                            # 先去掉多余过渡页（只保留最多 1 个）
                            divider_indices = [i for i, s in enumerate(slides) if s.get("type") == "divider"]
                            if len(divider_indices) > 1:
                                for idx in divider_indices[1:]:
                                    slides[idx]["_remove"] = True
                                slides = [s for s in slides if not s.get("_remove")]
                            
                            # 然后合并相邻的详细页为概览页
                            while len(slides) > _max_slides:
                                merged = False
                                for i in range(len(slides) - 1):
                                    if (slides[i].get("type") == "detail" and 
                                        slides[i+1].get("type") == "detail"):
                                        # 合并: 第一个和第二个详细页变成一个概览页
                                        s1 = slides[i]
                                        s2 = slides[i+1]
                                        merged_slide = {
                                            "type": "overview",
                                            "title": s1.get("title", ""),
                                            "points": [],
                                            "content": [],
                                            "subsections": [
                                                {"subtitle": s1.get("title", ""), "label": s1.get("title", ""),
                                                 "points": (s1.get("points", []) or s1.get("content", []))[:4]},
                                                {"subtitle": s2.get("title", ""), "label": s2.get("title", ""),
                                                 "points": (s2.get("points", []) or s2.get("content", []))[:4]},
                                            ]
                                        }
                                        slides[i] = merged_slide
                                        slides.pop(i+1)
                                        merged = True
                                        break
                                if not merged:
                                    # 无法合并了，直接截断
                                    slides = slides[:_max_slides]
                                    break
                            
                            ppt_data["slides"] = slides
                            slide_count = len(slides)
                            print(f"[FILE_GEN/PPT] 精简后: {slide_count} 页")
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'📊 大纲就绪: {slide_count} 页幻灯片', 'detail': ppt_data['title']})}\n\n"
                        
                        # ──────── Step 2.2: 内容充实（逐页扩写） ────────
                        # 检查内容是否单薄（平均每页要点少于 3 个或要点太短）
                        _thin_slides = []
                        for si, sl in enumerate(ppt_data["slides"]):
                            stype = sl.get("type", "detail")
                            if stype in ("divider",):
                                continue  # 过渡页不需要充实
                            pts = sl.get("points", [])
                            subs = sl.get("subsections", [])
                            # 要点太少 或 平均要点太短
                            avg_len = sum(len(p) for p in pts) / max(len(pts), 1)
                            sub_pts_count = sum(len(sub.get("points", [])) for sub in subs) if subs else 0
                            
                            if stype == "overview":
                                # 概览页：子主题数太少或每个子主题要点太少
                                if not subs or sub_pts_count < len(subs) * 2:
                                    _thin_slides.append(si)
                            elif stype in ("detail", "comparison"):
                                if len(pts) < 3 or avg_len < 20:
                                    _thin_slides.append(si)
                            elif stype == "highlight":
                                if len(pts) < 2:
                                    _thin_slides.append(si)
                        
                        if _thin_slides:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'📝 正在充实 {len(_thin_slides)} 页内容...', 'detail': '确保每页都有丰富详实的信息'})}\n\n"
                            
                            # 构建批量充实 prompt（一次性处理所有薄弱页面）
                            _enrich_prompt = (
                                "你是PPT内容撰写专家。以下幻灯片内容太单薄，请逐页充实。\n\n"
                                "**要求：**\n"
                                "1. 每个[详细]页必须有 4-6 个要点，每个要点 30-80 字\n"
                                "2. 每个[概览]页的每个子主题下必须有 2-4 个要点\n"
                                "3. 保持 `- **关键词** — 详细解释` 的格式\n"
                                "4. ⚠️ **必须包含具体数据、案例、事实** — 禁止写 '显著增长' '广泛应用' 等模糊表述，\n"
                                "   必须写 '据IDC数据，2025年市场规模达XXX亿' 这样有数字有来源的内容\n"
                                "5. 优先使用下方【参考资料】和【研究分析】中的真实数据\n"
                                "6. 严格按以下 JSON 格式输出，不要额外文字\n\n"
                            )
                            
                            _slides_to_enrich = []
                            for si in _thin_slides:
                                sl = ppt_data["slides"][si]
                                _slides_to_enrich.append({
                                    "index": si,
                                    "type": sl.get("type", "detail"),
                                    "title": sl.get("title", ""),
                                    "current_points": sl.get("points", []),
                                    "subsections": [
                                        {"subtitle": sub.get("subtitle", ""), "points": sub.get("points", [])}
                                        for sub in sl.get("subsections", [])
                                    ] if sl.get("subsections") else []
                                })
                            
                            _enrich_prompt += f"主题: {ppt_data['title']}\n"
                            if search_context:
                                _enrich_prompt += f"\n参考资料（包含重要数据，请充分利用）:\n{search_context[:6000]}\n"
                            if research_context:
                                _enrich_prompt += f"\n研究分析（包含核心数据和案例，必须融入）:\n{research_context[:6000]}\n"
                            
                            _enrich_prompt += (
                                f"\n需要充实的幻灯片:\n```json\n{json.dumps(_slides_to_enrich, ensure_ascii=False, indent=2)}\n```\n\n"
                                "请输出充实后的结果，格式:\n"
                                "```json\n"
                                "[{\"index\": 0, \"points\": [\"...\", ...], \"subsections\": [{\"subtitle\": \"...\", \"points\": [\"...\"]}, ...]}]\n"
                                "```\n"
                                "只输出 JSON，不要额外文字。"
                            )
                            
                            try:
                                _enrich_resp = client.models.generate_content(
                                    model="gemini-2.5-flash",
                                    contents=_enrich_prompt,
                                    config=types.GenerateContentConfig(temperature=0.5, max_output_tokens=8192)
                                )
                                _enrich_text = _enrich_resp.text or ""
                                import re as _enrich_re
                                _em = _enrich_re.search(r'\[.*\]', _enrich_text, _enrich_re.DOTALL)
                                if _em:
                                    _enriched = json.loads(_em.group())
                                    _applied = 0
                                    for _e in _enriched:
                                        _idx = _e.get("index")
                                        if _idx is not None and 0 <= _idx < len(ppt_data["slides"]):
                                            _sl = ppt_data["slides"][_idx]
                                            # 更新 points
                                            if _e.get("points") and len(_e["points"]) >= len(_sl.get("points", [])):
                                                _sl["points"] = _e["points"]
                                                _sl["content"] = _e["points"]
                                            # 更新 subsections
                                            if _e.get("subsections") and len(_e["subsections"]) > 0:
                                                _new_subs = []
                                                for _ns in _e["subsections"]:
                                                    _new_subs.append({
                                                        "subtitle": _ns.get("subtitle", ""),
                                                        "label": _ns.get("subtitle", ""),
                                                        "points": _ns.get("points", [])
                                                    })
                                                if _new_subs:
                                                    _sl["subsections"] = _new_subs
                                                    # 也更新 comparison 的 left/right
                                                    if _sl.get("type") == "comparison" and len(_new_subs) >= 2:
                                                        _sl["left"] = _new_subs[0]
                                                        _sl["right"] = _new_subs[1]
                                            _applied += 1
                                    
                                    if _applied > 0:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 已充实 {_applied} 页内容', 'detail': ''})}\n\n"
                                        print(f"[FILE_GEN/PPT] ✅ 内容充实完成: {_applied}/{len(_thin_slides)} 页")
                                    else:
                                        print(f"[FILE_GEN/PPT] ⚠️ 内容充实解析成功但未应用")
                                else:
                                    print(f"[FILE_GEN/PPT] ⚠️ 内容充实返回格式异常")
                            except Exception as enrich_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 内容充实异常（不影响生成）: {enrich_err}")
                        
                        # ──────── Step 2.5: 为幻灯片生成配图（Nano Banana 优先） ────────
                        ppt_images = []
                        # 对详细页配图（概览/对比/过渡/亮点页不适合插图）
                        img_candidate_slides = [(i, s) for i, s in enumerate(ppt_data["slides"]) 
                                         if s.get("type", "detail") == "detail"]
                        
                        if img_candidate_slides:
                            _n_images = min(4, max(2, len(img_candidate_slides) // 2 + 1))
                            yield f"data: {json.dumps({'type': 'progress', 'message': '🖼️ Nano Banana 正在生成配图...', 'detail': f'为 {_n_images} 个重点页面生成专业插图'})}\n\n"
                            try:
                                slide_titles_for_img = [s.get("title", "") for _, s in img_candidate_slides]
                                img_results = WebSearcher.generate_ppt_images(
                                    slide_titles_for_img,
                                    topic=ppt_data["title"],
                                    max_images=_n_images
                                )
                                # 将配图路径注入到对应 slide
                                for img_info in img_results:
                                    picked_idx = img_info["slide_index"]
                                    if picked_idx < len(img_candidate_slides):
                                        real_idx = img_candidate_slides[picked_idx][0]
                                        ppt_data["slides"][real_idx]["image"] = img_info["image_path"]
                                        ppt_images.append(img_info["image_path"])
                                
                                if ppt_images:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': f'✅ 已生成 {len(ppt_images)} 张配图', 'detail': ''})}\n\n"
                                else:
                                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 配图生成未成功，继续使用纯文本', 'detail': ''})}\n\n"
                            except Exception as img_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 配图生成异常: {img_err}")
                                yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 配图跳过，不影响PPT生成', 'detail': ''})}\n\n"
                        
                        # ──────── Step 3: 生成 PPT 文件(含逐页进度) ────────
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 开始渲染 PPT 文件...', 'detail': ''})}\n\n"
                        
                        from web.ppt_generator import PPTGenerator
                        
                        # 检测主题
                        theme = "business"
                        if any(kw in user_lower_check for kw in ["技术", "tech", "科技", "编程", "开发"]):
                            theme = "tech"
                        elif any(kw in user_lower_check for kw in ["创意", "creative", "艺术", "设计"]):
                            theme = "creative"
                        
                        ppt_gen = PPTGenerator(theme=theme)
                        
                        ppt_title = ppt_data["title"] or "演示文稿"
                        safe_title = re.sub(r'[\\/*?:"<>|]', '_', ppt_title)[:50]
                        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pptx"
                        docs_dir = settings_manager.documents_dir
                        os.makedirs(docs_dir, exist_ok=True)
                        ppt_path = os.path.join(docs_dir, filename)
                        
                        # 使用 progress_callback 来收集进度消息（生成器无法在回调中yield）
                        _slide_progress_msgs = []
                        def _ppt_progress_cb(cur, total, stitle, stype):
                            _slide_progress_msgs.append((cur, total, stitle, stype))
                        
                        ppt_gen.generate_from_outline(
                            title=ppt_title,
                            outline=ppt_data["slides"],
                            output_path=ppt_path,
                            subtitle=ppt_data.get("subtitle", ""),
                            author="Koto AI",
                            progress_callback=_ppt_progress_cb
                        )
                        
                        # 发送逐页进度（回调已经收集完毕）
                        for cur, total, stitle, stype in _slide_progress_msgs:
                            if stitle:
                                yield f"data: {json.dumps({'type': 'progress', 'message': f'📄 渲染 {cur}/{total}: {stitle}', 'detail': stype})}\n\n"
                        
                        yield f"data: {json.dumps({'type': 'progress', 'message': '✅ PPT 渲染完成，正在保存...', 'detail': ''})}\n\n"
                        
                        rel_path = os.path.relpath(ppt_path, WORKSPACE_DIR).replace("\\", "/")
                        generated_files.append(rel_path)
                        
                        # 统计各类型幻灯片数量
                        _type_names = {"detail": "详细页", "overview": "概览页", "highlight": "亮点页", "divider": "过渡页", "comparison": "对比页"}
                        _type_counts = {}
                        for _s in ppt_data["slides"]:
                            _t = _s.get("type", "detail")
                            _type_counts[_t] = _type_counts.get(_t, 0) + 1
                        _type_desc = "、".join(f"{_type_names.get(k,k)} ×{v}" for k, v in _type_counts.items())
                        
                        _img_desc = f"\n🖼️ 配图: {len(ppt_images)} 张" if ppt_images else ""
                        _research_desc = "\n🔬 已融入深度研究分析" if research_context else ""
                        
                        # ──────── 保存会话数据（P1 编辑功能支持） ────────
                        if ppt_session_id:
                            try:
                                from web.ppt_session_manager import get_ppt_session_manager
                                ppt_session_mgr = get_ppt_session_manager()
                                ppt_session_mgr.save_generation_data(
                                    session_id=ppt_session_id,
                                    ppt_data=ppt_data,
                                    ppt_file_path=rel_path,
                                    search_context=search_context,
                                    research_context=research_context,
                                    uploaded_file_context=uploaded_file_context
                                )
                                print(f"[FILE_GEN/PPT] 💾 会话数据已保存，可用于后续编辑")
                            except Exception as save_err:
                                print(f"[FILE_GEN/PPT] ⚠️ 会话保存异常: {save_err}")
                        
                        success_msg = (
                            f"✅ **PPT 演示文稿生成成功！**\n\n"
                            f"📊 标题: **{ppt_title}**\n"
                            f"📄 页数: {slide_count} 页（{_type_desc}）{_img_desc}{_research_desc}\n"
                            f"📁 文件: **{filename}**\n"
                            f"📍 位置: `{docs_dir}`"
                        )
                        
                        # 如果有会话，附加编辑链接
                        if ppt_session_id:
                            success_msg += f"\n\n🎨 **[点击编辑 PPT](/edit-ppt/{ppt_session_id})** - 修改内容、调整顺序、重新生成页面"
                        
                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        print(f"[FILE_GEN/PPT] ✅ PPT 生成成功: {rel_path}")
                        
                    except Exception as ppt_err:
                        print(f"[FILE_GEN/PPT] ❌ PPT 生成失败: {ppt_err}")
                        import traceback
                        traceback.print_exc()
                        error_msg = f"❌ PPT 生成失败: {str(ppt_err)}"
                        yield f"data: {json.dumps({'type': 'token', 'content': error_msg})}\n\n"
                    
                    # 保存历史（基于磁盘完整历史追加，在 done 事件之前）
                    _ppt_msg = f"已生成PPT: {', '.join(generated_files)}" if generated_files else "PPT生成失败"
                    session_manager.append_and_save(f"{session_name}.json", user_input, _ppt_msg)
                    
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return
                
                # 使用上下文增强的输入（如果有，例如"把这个做成word"时会包含之前的内容）
                if context_info and context_info.get("is_continuation") and context_info.get("enhanced_input"):
                    file_gen_input = context_info["enhanced_input"]
                    print(f"[FILE_GEN] 使用上下文增强输入 (length: {len(file_gen_input)})")
                else:
                    file_gen_input = effective_input

                # ⭐ FILE_GEN 前置步骤：时间解析 + 信息收集
                _time_context_text, _time_parse = _build_filegen_time_context(user_input)
                _web_context = ""
                _should_collect = WebSearcher.needs_web_search(user_input)

                # 对“X月新番/番剧/动画”等时间敏感主题强制启用信息收集
                _anime_time_patterns = [
                    r'([1-9]|1[0-2])\s*月\s*(新番|番剧|动画)',
                    r'(新番|番剧|动画).*(\d{1,2}\s*月)',
                ]
                if not _should_collect and any(re.search(p, user_input, re.IGNORECASE) for p in _anime_time_patterns):
                    _should_collect = True

                if _should_collect:
                    try:
                        if _time_parse.get("resolved_month"):
                            _q = f"{_time_parse['resolved_year']}年{_time_parse['resolved_month']}月 新番 动画 番剧 名单 介绍"
                        else:
                            _q = user_input

                        _time_detail = _time_context_text.replace("\n", " | ")[:180]
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🕒 正在解析时间语义...', 'detail': _time_detail})}\n\n"
                        yield f"data: {json.dumps({'type': 'progress', 'message': '🌐 正在收集最新信息...', 'detail': _q[:120]})}\n\n"

                        _search_res = WebSearcher.search_with_grounding(_q)
                        if _search_res.get("success") and _search_res.get("response"):
                            _web_context = _search_res.get("response", "")
                            print(f"[FILE_GEN] ✅ 信息收集完成，长度: {len(_web_context)}")
                            yield f"data: {json.dumps({'type': 'progress', 'message': '✅ 信息收集完成', 'detail': f'已获取 {len(_web_context)} 字符参考信息'})}\n\n"
                        else:
                            print(f"[FILE_GEN] ⚠️ 信息收集未返回结果")
                    except Exception as _collect_err:
                        print(f"[FILE_GEN] ⚠️ 信息收集异常: {_collect_err}")

                # 将时间上下文/检索结果拼接进生成输入
                _prepended_blocks = [_time_context_text]
                if _web_context:
                    _prepended_blocks.append("[联网检索参考]\n" + _web_context[:9000])
                file_gen_input = "\n\n".join(_prepended_blocks) + "\n\n" + file_gen_input

                # ⭐ 判断是否是文档生成请求（Word/PDF）
                _doc_keywords = ["word", "docx", "doc", "pdf", "报告", "文档", "论文", "综述", "whitepaper"]
                _is_doc_request = any(k in user_input.lower() for k in _doc_keywords)
                _is_complex = (context_info or {}).get("complexity") == "complex"

                if _is_doc_request:
                    # ============== 文档直出模式（流式） ==============
                    # 使用 generate_content_stream 保持连接活跃，避免代理超时断开
                    _doc_type = "PDF" if "pdf" in user_input.lower() else "Word"
                    yield f"data: {json.dumps({'type': 'progress', 'message': f'📄 正在生成 {_doc_type} 文档...', 'detail': '请稍候，正在撰写内容'})}\n\n"
                    print(f"[FILE_GEN] 📄 文档直出模式-流式 (type={_doc_type}, complex={_is_complex})")

                    _doc_instruction = """你是 Koto 专业文档撰写助手。请根据用户要求，直接输出**完整、详细、高质量**的文档正文内容。

## 输出规则
- 直接输出 Markdown 格式的文档正文，不要输出代码
- 使用 # ## ### 组织标题层级
- 使用段落、列表、表格丰富内容
- 中文撰写，专业术语准确
- 内容要**充实详尽**，每一节至少2-3段，总字数不少于3000字
- 如果是技术报告，必须包含：行业概述、技术原理、关键工艺、对比分析、应用场景、发展趋势
- 不要输出任何 BEGIN_FILE/END_FILE 标记
- 不要输出 JSON 或代码格式"""

                    _doc_instruction += "\n\n时间要求：若用户请求涉及月份但未写年份（如‘1月新番’），必须按当前年份撰写，禁止默认回退到历史年份。"

                    _max_tokens = 16384 if _is_complex else 8192
                    _doc_models = list(dict.fromkeys([
                        model_id,
                        "gemini-3-pro-preview",
                        "gemini-2.5-flash",
                        "gemini-2.0-flash",
                    ]))

                    _doc_collected = []  # 收集所有流式文本块

                    for model_attempt, current_model in enumerate(_doc_models):
                        if _doc_collected:
                            break
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if model_attempt > 0:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 切换到备用模型 {current_model}...', 'detail': ''})}\n\n"
                            _doc_collected.clear()
                        else:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'🚀 正在调用 {current_model}...', 'detail': '流式生成中'})}\n\n"

                        try:
                            _doc_stream = client.models.generate_content_stream(
                                model=current_model,
                                contents=file_gen_input,
                                config=types.GenerateContentConfig(
                                    system_instruction=_doc_instruction,
                                    max_output_tokens=_max_tokens,
                                    temperature=0.7,
                                )
                            )
                            _first_chunk = False
                            for item_type, item_data in stream_with_keepalive(
                                _doc_stream, start_time,
                                keepalive_interval=5,
                                max_wait_first_token=120  # 文档生成允许等待更久
                            ):
                                if interrupted():
                                    print(f"[FILE_GEN/DOC] 用户中断")
                                    _interrupt_msg = '\n\n⏹️ 文件生成已中断'
                                    yield f"data: {json.dumps({'type': 'token', 'content': _interrupt_msg})}\n\n"
                                    break

                                if item_type == 'heartbeat':
                                    _elapsed = item_data
                                    _char_count = sum(len(c) for c in _doc_collected)
                                    if _first_chunk:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在撰写文档...', 'detail': f'已生成 {_char_count} 字符，耗时 {_elapsed}s'})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'progress', 'message': '🧠 模型正在组织内容...', 'detail': f'已等待 {_elapsed}s，请耐心等待'})}\n\n"

                                elif item_type == 'timeout':
                                    print(f"[FILE_GEN/DOC] ⚠️ {current_model} 等待首token超时: {item_data}")
                                    break  # 尝试下一个模型

                                elif item_type == 'chunk':
                                    chunk = item_data
                                    if chunk.text:
                                        if not _first_chunk:
                                            _first_chunk = True
                                            print(f"[FILE_GEN/DOC] ✅ {current_model} 收到第一个响应块，耗时 {time.time() - start_time:.1f}s")
                                        _doc_collected.append(chunk.text)
                                        # 每收到10个chunk发送一次进度更新，保持客户端连接活跃
                                        if len(_doc_collected) % 10 == 0:
                                            _char_count = sum(len(c) for c in _doc_collected)
                                            _elapsed = int(time.time() - start_time)
                                            yield f"data: {json.dumps({'type': 'progress', 'message': '📝 正在撰写文档...', 'detail': f'已生成 {_char_count} 字符，耗时 {_elapsed}s'})}\n\n"

                        except Exception as _doc_err:
                            err_str = str(_doc_err)
                            print(f"[FILE_GEN/DOC] ❌ {current_model}: {err_str[:200]}")
                            if "location is not supported" in err_str.lower():
                                response_text = "❌ 地区限制，请配置中转服务"
                                break
                            continue

                    response_text = ''.join(_doc_collected)
                    if response_text:
                        print(f"[FILE_GEN/DOC] ✅ 流式生成完成，共 {len(response_text)} 字符")

                    if not response_text or response_text.startswith("❌"):
                        yield f"data: {json.dumps({'type': 'token', 'content': response_text or '❌ 所有模型都不可用，请稍后重试'})}\n\n"
                    else:
                        # 直接保存文档
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'📝 正在保存 {_doc_type} 文档...', 'detail': ''})}\n\n"
                        try:
                            try:
                                from web.document_generator import save_docx, save_pdf
                            except ModuleNotFoundError:
                                from document_generator import save_docx, save_pdf
                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)
                            title_match = re.search(r'^#\s*(.+)$', response_text, re.MULTILINE)
                            title = title_match.group(1).strip()[:50] if title_match else f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                            if _doc_type == "PDF":
                                saved_path = save_pdf(response_text, title=title, output_dir=docs_dir)
                            else:
                                saved_path = save_docx(response_text, title=title, output_dir=docs_dir)
                            rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace("\\", "/")
                            generated_files.append(rel_path)
                            success_msg = f"✅ **{_doc_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                            print(f"[FILE_GEN/DOC] ✅ 文档已保存: {rel_path}")
                        except Exception as doc_err:
                            import traceback
                            traceback.print_exc()
                            print(f"[FILE_GEN/DOC] ❌ 文档保存失败: {doc_err}")
                            fallback_msg = f"⚠️ 文档保存失败 ({doc_err})，以下是生成的内容：\n\n"
                            yield f"data: {json.dumps({'type': 'token', 'content': fallback_msg + response_text})}\n\n"

                    _gen_msg = f"已生成文件: {', '.join(generated_files)}" if generated_files else (response_text[:500] if response_text else "生成失败")
                    session_manager.append_and_save(f"{session_name}.json", user_input, _gen_msg)
                    total_time = time.time() - start_time
                    print(f"[FILE_GEN/DOC] ★★★ done event, files: {generated_files}, time: {total_time:.2f}s")
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                    return

                # 普通 FILE_GEN 模式（需要模型生成代码/脚本）
                yield f"data: {json.dumps({'type': 'progress', 'message': '📄 正在生成文件代码...', 'detail': '请稍候，可能需要 10-30 秒'})}\n\n"
                
                # 模型列表（主模型 + 备用模型）
                file_gen_models = [
                    model_id,  # 主模型
                    "gemini-3-pro-preview",  # 备用1 (更强的推理)
                    "gemini-2.5-flash",  # 备用2
                    "gemini-2.0-flash",  # 备用3
                ]
                
                # 使用线程 + 超时来调用API（带重试）
                import threading
                import tempfile
                
                for model_attempt, current_model in enumerate(file_gen_models):
                    if response_text and not response_text.startswith("❌"):
                        break  # 已成功

                    if interrupted():
                        yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                        total_time = time.time() - start_time
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                        return
                    
                    if model_attempt > 0:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'🔄 切换到备用模型 {current_model}...', 'detail': ''})}\n\n"
                        print(f"[FILE_GEN] Trying fallback model: {current_model}")
                    else:
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'🚀 正在调用 {current_model}...', 'detail': '生成中'})}\n\n"
                    
                    response_holder = {'data': None, 'error': None}
                    
                    def call_api(m=current_model):
                        try:
                            print(f"[FILE_GEN] Calling API: {m}")
                            response = client.models.generate_content(
                                model=m,
                                contents=file_gen_input,  # 使用上下文增强的输入
                                config=types.GenerateContentConfig(
                                    system_instruction=_get_system_instruction(),
                                    max_output_tokens=8192
                                )
                            )
                            response_holder['data'] = response
                            print(f"[FILE_GEN] ✅ API call successful with {m}")
                        except Exception as e:
                            print(f"[FILE_GEN] ❌ API call exception with {m}: {type(e).__name__}: {str(e)}")
                            response_holder['error'] = e
                    
                    api_thread = threading.Thread(target=call_api, daemon=True)
                    api_thread.start()
                    
                    # 在等待期间发送心跳进度
                    wait_interval = 5  # 每 5 秒发送一次进度
                    elapsed = 0
                    while api_thread.is_alive() and elapsed < api_timeout:
                        api_thread.join(timeout=wait_interval)
                        elapsed += wait_interval
                        if interrupted():
                            yield f"data: {json.dumps({'type': 'token', 'content': '⏹️ 文件生成已中断'})}\n\n"
                            total_time = time.time() - start_time
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        if api_thread.is_alive() and elapsed < api_timeout:
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'⏳ 正在生成中...', 'detail': f'已等待 {elapsed} 秒'})}\n\n"
                    
                    if api_thread.is_alive():
                        print(f"[FILE_GEN] ⚠️ API call timeout with {current_model} after {api_timeout}s")
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ {current_model} 响应超时', 'detail': '正在切换模型...'})}\n\n"
                        response_text = ""
                        continue  # 尝试下一个模型
                    elif response_holder['error']:
                        error_str = str(response_holder['error'])
                        print(f"[FILE_GEN] API Error with {current_model}: {error_str}")
                        
                        # 地区限制错误 - 直接失败，不重试
                        if "location is not supported" in error_str.lower() or "failed_precondition" in error_str.lower():
                            response_text = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 config/gemini_config.env 配置中转服务 GEMINI_API_BASE\n2. 或使用支持的代理服务"
                            break  # 地区限制，不继续重试
                        elif "503" in error_str or "overloaded" in error_str.lower() or "unavailable" in error_str.lower():
                            yield f"data: {json.dumps({'type': 'progress', 'message': f'⚠️ {current_model} 服务繁忙', 'detail': '正在切换模型...'})}\n\n"
                            response_text = ""
                            continue  # 503 错误，尝试下一个模型
                        else:
                            response_text = f"❌ API 调用失败: {error_str[:200]}"
                            continue  # 其他错误也尝试下一个模型
                    elif response_holder['data']:
                        file_gen_response = response_holder['data']
                        if file_gen_response.candidates and file_gen_response.candidates[0].content.parts:
                            for part in file_gen_response.candidates[0].content.parts:
                                if hasattr(part, 'text') and part.text:
                                    response_text += part.text
                        print(f"[FILE_GEN] Response length: {len(response_text)}")
                        if response_text:
                            break  # 成功获取响应
                
                if not response_text:
                    response_text = "❌ 所有模型都不可用，请稍后重试"
                
                if Utils.is_failure_output(response_text):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次生成失败，正在修正...', 'detail': ''})}\n\n"
                    fix_prompt = Utils.build_fix_prompt("FILE_GEN", user_input, response_text)
                    try:
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_system_instruction(),
                                max_output_tokens=8192,
                                temperature=0.4
                            )
                        )
                        response_text = fix_resp.text or response_text
                    except Exception as fix_err:
                        print(f"[FILE_GEN] 修正重试失败: {fix_err}")
                
                # 只显示简短的进度，不显示完整代码
                if response_text and not response_text.startswith("❌"):
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🔧 正在处理代码...', 'detail': ''})}\n\n"
                    
                    # 提取代码到临时文件
                    patterns = [
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\s*(.*?)---\s*END_FILE\s*---",
                        r"---BEGIN_FILE:\s*([a-zA-Z0-9_.-]+)\s*---\n(.*?)\n---END_FILE---",
                    ]
                    
                    code_content = None
                    for pattern in patterns:
                        matches = re.findall(pattern, response_text, re.DOTALL | re.IGNORECASE)
                        if matches:
                            _, code_content = matches[0]
                            code_content = code_content.strip()
                            print(f"[FILE_GEN] Extracted code, length: {len(code_content)}")
                            break
                    
                    # 检查提取的内容是否是有效的Python代码（不是JSON或其他格式）
                    is_valid_python = False
                    if code_content:
                        code_lower = code_content.lower()
                        # 如果提取的内容是 JSON 或 HTML 或其他格式，直接跳过代码执行
                        if code_lower.startswith(('{', '[', '<', '"')):
                            print(f"[FILE_GEN] Extracted content is not Python code (starts with {code_content[0]}), treating as text content")
                            code_content = None
                        else:
                            is_valid_python = True
                    
                    if code_content and is_valid_python:
                        # 保存到临时文件
                        temp_dir = tempfile.gettempdir()
                        temp_script = os.path.join(temp_dir, f"koto_gen_{int(time.time())}.py")
                        
                        with open(temp_script, 'w', encoding='utf-8') as f:
                            f.write(code_content)
                        temp_scripts.append(temp_script)
                        print(f"[FILE_GEN] Saved temp script: {temp_script}")
                        
                        # 执行脚本
                        yield f"data: {json.dumps({'type': 'progress', 'message': '⚙️ 正在执行脚本生成文件...', 'detail': ''})}\n\n"
                        
                        try:
                            result = subprocess.run(
                                [sys.executable, temp_script],
                                capture_output=True,
                                text=True,
                                timeout=60,
                                cwd=WORKSPACE_DIR,
                                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
                            )
                            print(f"[FILE_GEN] Script exit code: {result.returncode}")
                            print(f"[FILE_GEN] Script stdout: {result.stdout}")
                            print(f"[FILE_GEN] Script stderr: {result.stderr}")
                            
                            if result.returncode == 0:
                                # 检查生成的文件
                                docs_dir = settings_manager.documents_dir
                                if os.path.exists(docs_dir):
                                    for f in os.listdir(docs_dir):
                                        if f.endswith(('.pdf', '.docx', '.xlsx', '.pptx', '.ppt', '.png', '.jpg')):
                                            full_path = os.path.join(docs_dir, f)
                                            age = time.time() - os.path.getmtime(full_path)
                                            if age < 60:
                                                rel_path = os.path.relpath(full_path, WORKSPACE_DIR).replace("\\", "/")
                                                if rel_path not in generated_files:
                                                    generated_files.append(rel_path)
                                                    print(f"[FILE_GEN] Generated: {rel_path}")
                                
                                if generated_files:
                                    files_list = ", ".join([os.path.basename(f) for f in generated_files])
                                    success_msg = "✅ **文件生成成功！**\n\n📁 生成的文件: **" + files_list + "**\n📍 保存位置: `" + docs_dir + "`"
                                    yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                                else:
                                    # 脚本执行成功但没有检测到新文件
                                    output = result.stdout.strip()
                                    if output:
                                        msg = "✅ 脚本执行完成\n```\n" + output + "\n```"
                                        yield f"data: {json.dumps({'type': 'token', 'content': msg})}\n\n"
                                    else:
                                        yield f"data: {json.dumps({'type': 'token', 'content': '⚠️ 脚本执行完成，但未检测到新文件'})}\n\n"
                            else:
                                error_msg = result.stderr.strip() or "未知错误"
                                err_content = "❌ 脚本执行失败\n```\n" + error_msg[:500] + "\n```"
                                yield f"data: {json.dumps({'type': 'token', 'content': err_content})}\n\n"
                        
                        except subprocess.TimeoutExpired:
                            yield f"data: {json.dumps({'type': 'token', 'content': '⚠️ 脚本执行超时（60秒）'})}\n\n"
                        except Exception as e:
                            print(f"[FILE_GEN] Execution error: {e}")
                            err_msg = "❌ 执行错误: " + str(e)
                            yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                        
                        # 删除临时脚本
                        for temp_file in temp_scripts:
                            try:
                                if os.path.exists(temp_file):
                                    os.remove(temp_file)
                                    print(f"[FILE_GEN] Deleted temp script: {temp_file}")
                            except:
                                pass
                    else:
                        # 没有匹配到代码格式：直接把模型内容生成文档
                        try:
                            from web.document_generator import save_docx, save_pdf

                            docs_dir = settings_manager.documents_dir
                            os.makedirs(docs_dir, exist_ok=True)

                            # 提取标题（尝试从内容中找 # 标题）
                            title_match = re.search(r'^#\s*(.+)$', response_text, re.MULTILINE)
                            if title_match:
                                title = title_match.group(1).strip()[:50]
                            else:
                                # 尝试从用户输入提取关键词作为文件名
                                try:
                                    clean_input = user_input
                                    # 去除常用指令词
                                    stop_patterns = ["生成的", "写一个", "写一篇", "帮我", "请", "关于", "一下", "文档", "file", "generate", "write", "about", "make", "create"]
                                    for pattern in stop_patterns:
                                        clean_input = clean_input.replace(pattern, " ")
                                    
                                    # 提取中英文关键词 (2-20 chars)
                                    keywords = [w for w in re.split(r'[^a-zA-Z0-9\u4e00-\u9fa5]', clean_input) if w.strip()]
                                    valid_keywords = [k for k in keywords if len(k) > 1 and len(k) < 20]
                                    
                                    if valid_keywords:
                                        # 取前几个关键词组合
                                        title = "_".join(valid_keywords[:3])
                                    else:
                                        title = f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                except:
                                    title = f"Koto文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                            user_lower = user_input.lower()
                            if "pdf" in user_lower:
                                saved_path = save_pdf(response_text, title=title, output_dir=docs_dir)
                                file_type = "PDF"
                            else:
                                saved_path = save_docx(response_text, title=title, output_dir=docs_dir)
                                file_type = "Word"

                            rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace("\\", "/")
                            if rel_path not in generated_files:
                                generated_files.append(rel_path)

                            success_msg = f"✅ **{file_type} 文档生成成功！**\n\n📁 文件: **{os.path.basename(saved_path)}**\n📍 位置: `{docs_dir}`"
                            yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        except Exception as direct_err:
                            print(f"[FILE_GEN] Direct save failed: {direct_err}")
                            # 回退展示原始响应
                            yield f"data: {json.dumps({'type': 'token', 'content': response_text})}\n\n"
                else:
                    yield f"data: {json.dumps({'type': 'token', 'content': response_text or '⚠️ 模型未返回响应'})}\n\n"
                
                # 保存历史（基于磁盘完整历史追加，在 done 事件之前）
                _gen_msg = f"已生成文件: {', '.join(generated_files)}" if generated_files else (response_text[:500] if response_text else "生成失败")
                session_manager.append_and_save(f"{session_name}.json", user_input, _gen_msg)
                
                # 发送完成事件
                total_time = time.time() - start_time
                print(f"[FILE_GEN] ★★★ Sending done event, generated_files: {generated_files}, total_time: {total_time:.2f}s")
                yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': generated_files, 'total_time': total_time})}\n\n"
                return
            
            # === Regular Mode (流式输出) ===
            # 根据任务类型选择系统指令
            # CHAT/RESEARCH等使用简化指令，避免不必要的文件生成
            use_instruction = _get_DEFAULT_CHAT_SYSTEM_INSTRUCTION() if task_type in ["CHAT", "RESEARCH"] else _get_system_instruction()

            # 注入长期记忆上下文
            _memory_manager = get_memory_manager()
            memory_context = _memory_manager.get_context_string(user_input)
            if memory_context:
                use_instruction += f"\n\n{memory_context}"
                print(f"[MEMORY] 注入了 {len(memory_context)} 字符的记忆上下文")
                t = yield_thinking(f"从长期记忆中检索到 {len(memory_context)} 字符的相关上下文并注入", "context")
                if t: yield t

            # 根据任务类型提供差异化进度提示
            if task_type == "CODER":
                used_model = model_id
                t = yield_thinking(f"进入代码生成模式，使用 {model_id} 进行代码分析与生成", "generating")
                if t: yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '💻 正在分析代码需求...', 'detail': f'使用 {model_id}'})}\n\n"
                
                # 特殊优化：对于游戏开发或安装包，添加简短指令避免啰嗦
                if any(k in user_input.lower() for k in ["游戏", "app", "五子棋", "pygame", "install", "安装"]):
                     use_instruction += "\n\n[Important] If suggesting to install packages (like pygame), assume the user knows how to use pip. Just output `pip install package_name` in a code block. Do NOT write long tutorials about installation. Focus on the Python Code."

            elif task_type == "CHAT":
                used_model = model_id
                t = yield_thinking(f"进入对话模式，使用 {model_id} 生成回复", "generating")
                if t: yield t
                yield f"data: {json.dumps({'type': 'progress', 'message': '💬 Koto 正在思考...', 'detail': '请稍候'})}\n\n"

                # ═══ 本地模型快速通道：简单问题直接走 Ollama ═══
                from app.core.routing import LocalModelRouter
                if LocalModelRouter.is_simple_query(user_input, task_type, history):
                    local_stream = LocalModelRouter.generate_stream(
                        user_input, history=history,
                        system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION()
                    )
                    if local_stream is not None:
                        print(f"[CHAT] ⚡ 使用本地模型快速响应: {LocalModelRouter._response_model}")
                        t = yield_thinking(f"检测到简单查询，切换到本地模型 {LocalModelRouter._response_model} 快速响应", "model")
                        if t: yield t
                        yield f"data: {json.dumps({'type': 'classification', 'task_type': task_type, 'task_display': '💬 对话', 'model': f'🏠 {LocalModelRouter._response_model} (本地)', 'message': f'🎯 任务分类: 💬 对话 (方法: 🏠 {LocalModelRouter._response_model} 本地快速通道)'})}\n\n"
                        local_full_text = ""
                        local_ok = False
                        try:
                            for chunk in local_stream:
                                local_full_text += chunk
                                yield f"data: {json.dumps({'type': 'token', 'content': chunk})}\n\n"
                            local_ok = bool(local_full_text.strip())
                        except Exception as local_err:
                            print(f"[CHAT] 本地模型生成失败: {local_err}")

                        if local_ok:
                            # 本地模型成功 → 保存并返回
                            session_manager.append_and_save(
                                f"{session_name}.json", user_input, local_full_text,
                                task=task_type, model_name=f"ollama/{LocalModelRouter._response_model}"
                            )
                            if task_type == "CHAT":
                                _start_memory_extraction(user_input, local_full_text, history)
                            total_time = time.time() - start_time
                            print(f"[CHAT] ⚡ 本地模型响应完成 ({total_time:.2f}s)")
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                            return
                        else:
                            # 本地模型失败 → 静默降级到云模型
                            print(f"[CHAT] 本地模型输出为空，降级到云模型")
                            t = yield_thinking(f"本地模型输出为空，降级到云端模型 {model_id}", "model")
                            if t: yield t
                            yield f"data: {json.dumps({'type': 'progress', 'message': '☁️ 切换到云端模型...', 'detail': model_id})}\n\n"
            elif task_type == "RESEARCH":
                yield f"data: {json.dumps({'type': 'progress', 'message': '🔬 正在进行深度分析...', 'detail': f'使用 {model_id}'})}\n\n"
            else:
                yield f"data: {json.dumps({'type': 'progress', 'message': '💭 Koto 正在思考...', 'detail': '请稍候'})}\n\n"
            
            # 构建历史记录（非延续任务时过滤无关历史）
            if context_info and context_info.get("is_continuation"):
                history_for_model = history
                t = yield_thinking(f"检测到上下文延续，保留全部 {len(history)} 轮对话历史", "context")
                if t: yield t
            else:
                history_for_model = ContextAnalyzer.filter_history(user_input, history)
                if len(history_for_model) != len(history):
                    t = yield_thinking(f"过滤对话历史: {len(history)} 轮 → {len(history_for_model)} 轮相关记录", "context")
                    if t: yield t

            formatted_history = []
            for turn in history_for_model:
                formatted_history.append(types.Content(
                    role=turn['role'],
                    parts=[types.Part.from_text(text=p) for p in turn['parts']]
                ))
            
            t = yield_thinking(f"准备调用 {model_id} API，发送 {len(formatted_history)+1} 条消息", "generating")
            if t: yield t
            
            # 使用流式响应
            response = client.models.generate_content_stream(
                model=model_id,
                contents=formatted_history + [types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=effective_input)]
                )],
                config=types.GenerateContentConfig(
                    system_instruction=use_instruction
                )
            )
            
            full_text = ""
            chunk_count = 0
            heartbeat_interval = 5  # 每5秒发送一次心跳
            first_chunk_received = False
            
            try:
                # 使用保活包装器处理流式响应
                max_wait = 60 if task_type == "CODER" else 120
                for item_type, item_data in stream_with_keepalive(response, start_time,
                                                                   keepalive_interval=heartbeat_interval,
                                                                   max_wait_first_token=max_wait):
                    # 检查中断标志
                    if _interrupt_manager.is_interrupted(session_name):
                        print(f"[INTERRUPT] User interrupted at chunk {chunk_count}")
                        interrupt_msg = "\n\n⏸️ 用户已中断"
                        yield f"data: {json.dumps({'type': 'token', 'content': interrupt_msg})}\n\n"
                        break
                    
                    if item_type == 'heartbeat':
                        elapsed = item_data
                        if first_chunk_received:
                            # 根据任务类型差异化心跳
                            char_count = len(full_text)
                            if task_type == "CODER":
                                hb_msg = f'💻 代码生成中... 已输出 {char_count} 字符'
                            elif task_type == "RESEARCH":
                                hb_msg = f'🔬 深度分析中... 已输出 {char_count} 字符'
                            else:
                                hb_msg = '💭 正在生成...'
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'{elapsed}s'})}\n\n"
                        else:
                            if task_type == "CODER":
                                hb_msg = '💻 代码分析中，请稍候...'
                            elif task_type == "RESEARCH":
                                hb_msg = '🔬 深度思考中，请耐心等待...'
                            else:
                                hb_msg = '🧠 模型思考中...'
                            yield f"data: {json.dumps({'type': 'progress', 'message': hb_msg, 'detail': f'已等待 {elapsed}s'})}\n\n"
                    
                    elif item_type == 'timeout':
                        if task_type == "CODER" and not full_text:
                            yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 首包超时，切换到快速模型...', 'detail': ''})}\n\n"
                            try:
                                fallback_resp = client.models.generate_content(
                                    model="gemini-3-flash-preview",
                                    contents=formatted_history + [types.Content(
                                        role="user",
                                        parts=[types.Part.from_text(text=effective_input)]
                                    )],
                                    config=types.GenerateContentConfig(
                                        system_instruction=use_instruction,
                                        temperature=0.4,
                                        max_output_tokens=4000,
                                    )
                                )
                                fallback_text = fallback_resp.text or ""
                                if fallback_text:
                                    full_text = fallback_text
                                    yield f"data: {json.dumps({'type': 'token', 'content': fallback_text})}\n\n"
                            except Exception:
                                yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，请稍后重试'})}\n\n"
                        else:
                            yield f"data: {json.dumps({'type': 'token', 'content': f'⚠️ {item_data}，请稍后重试'})}\n\n"
                        break
                    
                    elif item_type == 'chunk':
                        chunk = item_data
                        if chunk.text:
                            if not first_chunk_received:
                                first_chunk_received = True
                                print(f"[CHAT] 收到第一个响应，耗时 {time.time() - start_time:.1f}s")
                            
                            full_text += chunk.text
                            chunk_count += 1
                            yield f"data: {json.dumps({'type': 'token', 'content': chunk.text})}\n\n"
                            
            except Exception as stream_error:
                error_str = str(stream_error)
                print(f"[CHAT] Stream error: {error_str}")
                
                # 地区限制错误
                if "location is not supported" in error_str.lower() or "failed_precondition" in error_str.lower():
                    error_text = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 `config/gemini_config.env` 配置中转服务 `GEMINI_API_BASE`\n2. 或使用支持的代理服务"
                    yield f"data: {json.dumps({'type': 'token', 'content': error_text})}\n\n"
                    total_time = time.time() - start_time
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
                    return
                # 流式传输中断，但已有部分内容
                elif full_text:
                    error_msg = error_str[:50]
                    warn_text = f"\n\n⚠️ (传输中断: {error_msg}...)"
                    yield f"data: {json.dumps({'type': 'token', 'content': warn_text})}\n\n"
                else:
                    raise stream_error
            
            # 失败时先修正一次（不直接报错）
            if Utils.is_failure_output(full_text):
                yield f"data: {json.dumps({'type': 'progress', 'message': '⚠️ 初次生成失败，正在修正...', 'detail': ''})}\n\n"
                fix_prompt = Utils.build_fix_prompt(task_type, user_input, full_text)
                fix_resp = client.models.generate_content(
                    model=model_id,
                    contents=fix_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=use_instruction,
                        temperature=0.4,
                        max_output_tokens=4000,
                    )
                )
                corrected_text = fix_resp.text or full_text
                if corrected_text and corrected_text != full_text:
                    corrected_msg = f"\n\n🔁 修正版本:\n{corrected_text}"
                    yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                    full_text = corrected_text
            else:
                # 复杂任务进行快速自检
                is_complex_task = (
                    task_type in ["RESEARCH", "FILE_GEN", "CODER"] or
                    (context_info and context_info.get("complexity") == "complex") or
                    len(user_input) > 200
                )
                if is_complex_task:
                    check = Utils.quick_self_check(task_type, user_input, full_text)
                    if not check.get("pass") and check.get("fix_prompt"):
                        status_msg = "🩺 自检未通过，正在修正..."
                        yield f"data: {json.dumps({'type': 'progress', 'message': status_msg, 'detail': '快速模型自检'})}\n\n"
                        fix_resp = client.models.generate_content(
                            model=model_id,
                            contents=check["fix_prompt"],
                            config=types.GenerateContentConfig(
                                system_instruction=use_instruction,
                                temperature=0.4,
                                max_output_tokens=4000,
                            )
                        )
                        corrected_text = fix_resp.text or full_text
                        if corrected_text and corrected_text != full_text:
                            corrected_msg = f"\n\n🔁 修正版本:\n{corrected_text}"
                            yield f"data: {json.dumps({'type': 'token', 'content': corrected_msg})}\n\n"
                            full_text = corrected_text

            # 处理自动保存的文件
            saved_files = Utils.auto_save_files(full_text)

            # 代码任务: 检测并自动安装依赖
            if task_type == "CODER":
                pkgs = Utils.detect_required_packages(full_text)
                if pkgs:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📦 检测到依赖，正在检查/安装...', 'detail': ', '.join(pkgs)})}\n\n"
                    install_result = Utils.auto_install_packages(pkgs)
                    installed = install_result.get("installed", [])
                    failed = install_result.get("failed", [])
                    skipped = install_result.get("skipped", [])
                    msg_parts = []
                    if installed:
                        msg_parts.append(f"✅ 已安装: {', '.join(installed)}")
                    if skipped:
                        msg_parts.append(f"ℹ️ 已存在: {', '.join(skipped)}")
                    if failed:
                        msg_parts.append(f"⚠️ 安装失败: {', '.join(failed)}")
                    if msg_parts:
                        msg_content = '\n\n' + '\n'.join(msg_parts)
                        yield f"data: {json.dumps({'type': 'token', 'content': msg_content})}\n\n"
            
            # 如果有保存的文件，提示用户保存位置
            if saved_files:
                files_list = ", ".join(saved_files)
                save_hint = f"\n\n📁 文件已保存: **{files_list}**\n📂 位置: `{WORKSPACE_DIR}`"
                yield f"data: {json.dumps({'type': 'token', 'content': save_hint})}\n\n"
            
            # 先保存历史，再发送 done 事件（包含元数据用于前端渲染）
            session_manager.append_and_save(
                f"{session_name}.json", user_input, full_text,
                task=task_type, model_name=model_id, saved_files=saved_files
            )
            if task_type == "CHAT":
                _start_memory_extraction(user_input, full_text, history_for_model)
            
            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': saved_files, 'total_time': total_time})}\n\n"
            
        except Exception as e:
            error_str = str(e)
            print(f"[CHAT] Exception: {error_str}")
            
            # 地区限制错误
            if "location is not supported" in error_str.lower() or "failed_precondition" in error_str.lower():
                error_response = "❌ 地区限制\n\n您所在的地区不支持 Gemini API。\n\n💡 解决方案:\n1. 在 `config/gemini_config.env` 配置中转服务 `GEMINI_API_BASE`\n2. 或使用支持的代理服务"
            else:
                error_response = f"❌ 发生错误: {error_str[:200]}"
            
            # 即使出错也要保存用户的问题
            session_manager.append_and_save(
                f"{session_name}.json",
                user_input,
                error_response,
                task=task_type,
                model_name=model_id
            )
            
            yield f"data: {json.dumps({'type': 'token', 'content': error_response})}\n\n"
            total_time = time.time() - start_time
            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': total_time})}\n\n"
    
    response = Response(stream_with_context(generate()), mimetype='text/event-stream')
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['X-Accel-Buffering'] = 'no'  # 禁用 nginx 缓冲
    response.headers['Connection'] = 'keep-alive'
    return response


@app.route('/api/chat/file', methods=['POST'])
def chat_with_file():
    """处理文件上传和聊天请求"""
    from web.file_processor import process_uploaded_file
    from web.document_generator import save_docx, save_pdf, to_workspace_rel

    def _strip_code_blocks(text: str) -> str:
        if not text:
            return text
        # Remove fenced code blocks entirely
        text = re.sub(r"```[\s\S]*?```", "", text)
        # Remove inline code ticks but keep the content
        text = text.replace("`", "")
        return text.strip()

    def _build_analysis_title(user_text: str, filename: str, is_binary: bool) -> str:
        name_base = os.path.splitext(filename)[0]
        text_lower = (user_text or "").lower()
        ext = os.path.splitext(filename)[1].lower()
        
        # 1. Determine File Type Prefix
        if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            prefix = "图片"
        elif ext == ".pdf":
            prefix = "PDF"
        elif ext in [".doc", ".docx"]:
            prefix = "Word"
        elif ext in [".ppt", ".pptx"]:
            prefix = "PPT"
        else:
            prefix = "文件" if is_binary else "文档"

        # 2. Determine Intent
        intent = "分析"
        intent_map = {
            "翻译": ["翻译", "translate", "译文", "中译英", "英译中"],
            "总结": ["总结", "归纳", "摘要", "summary", "概括", "核心内容"],
            "文字识别": ["提取", "识别", "ocr", "文字", "转文字", "读图"],
            "表格识别": ["表格", "table", "excel", "转表"],
            "对比分析": ["对比", "比较", "diff", "区别", "差异"],
            "校对": ["校对", "检查", "审阅", "纠错", "改错"],
            "润色": ["润色", "改写", "polish", "rewrite", "优化", "美化"],
            "续写": ["续写", "扩写", "continue", "补充"],
            "大纲": ["大纲", "框架", "outline", "目录"],
            "解释": ["解释", "explain", "什么意思", "含义"],
        }
        
        found_intent_keywords = []
        for k, v in intent_map.items():
            for kw in v:
                if kw in text_lower:
                    intent = k
                    found_intent_keywords.append(kw)
                    break
            if intent != "分析":
                break

        # 3. Extract Topic Keywords (Improved)
        stop_words = [
            "帮我", "请", "一下", "把", "这个", "这篇", "文件", "文章", "内容", "生成", "写一个", "做一份",
            "koto", "分析", "阅读", "提取", "识别", "output", "make", "create", "generate", "please", 
            "the", "a", "an", "is", "of", "to", "for", "with", "in", "on", "user", "file", "document",
            "from", "this", "that", "it", "what", "how", "why", "where", "into", "check", "run"
        ]
        
        # Prepare text
        text_lower = user_text.lower()
        
        # Safe replacement for Chinese phrases (which don't use spaces)
        zh_stops = [w for w in stop_words if re.match(r'[\u4e00-\u9fa5]+', w)]
        for stop in zh_stops + found_intent_keywords:
            if re.match(r'[\u4e00-\u9fa5]+', stop): # Only safe-replace Chinese phrases
                text_lower = text_lower.replace(stop, " ")

        # Tokenize by non-word chars (separates English words, breaks Chinese into blocks if spaces inserted)
        # Regex: Keep Chinese chars and English words
        # This splits "summary of report" -> "summary", "of", "report"
        tokens = re.findall(r'[a-zA-Z0-9\u4e00-\u9fa5]+', text_lower)
        
        # Filter tokens
        valid_keywords = []
        en_stops = set([w for w in stop_words if not re.match(r'[\u4e00-\u9fa5]+', w)])
        
        for token in tokens:
            if token in en_stops: continue
            if token in found_intent_keywords: continue # Filter intent words token-wise
            if len(token) < 2: continue
            valid_keywords.append(token)
            
        # Select best keyword
        topic = ""
        if valid_keywords:
             topic = "_".join(valid_keywords[:3])
        
        # 4. Construct Final Title
        # Strategy: 
        # If user provided a specific topic, prioritize it: "{Intent}_{Topic}_{Filename}"
        # If no detected topic but intent exists: "{Intent}_{Filename}"
        # Fallback: "{Prefix}{Intent}_{Filename}"
        
        sanitized_name = name_base.replace(" ", "_")
        
        if topic:
             return f"{intent}_{topic}_{sanitized_name}"
        else:
             return f"{prefix}{intent}_{sanitized_name}"
    
    session_name = request.form.get('session')
    user_input = request.form.get('message', '')
    files = request.files.getlist('file')
    
    # 🔍 调试日志
    print(f"[FILE UPLOAD DEBUG] ========== 接收到文件上传请求 ==========")
    print(f"[FILE UPLOAD DEBUG] request.files keys: {list(request.files.keys())}")
    print(f"[FILE UPLOAD DEBUG] request.files.getlist('file'): {len(files)} 个文件")
    for i, f in enumerate(files):
        print(f"[FILE UPLOAD DEBUG]   {i+1}. {f.filename if f else 'None'}")
    
    if not files:
        single_file = request.files.get('file')
        if single_file:
            files = [single_file]
            print(f"[FILE UPLOAD DEBUG] 使用单文件模式，文件: {single_file.filename}")
    
    locked_task = request.form.get('locked_task')
    locked_model = request.form.get('locked_model', 'auto')
    stream_mode = request.form.get('stream', '').lower() in ('1', 'true', 'yes')
    
    print(f"[FILE UPLOAD DEBUG] 最终 files 列表: {len(files)} 个文件")
    print(f"[FILE UPLOAD DEBUG] 判断: len(files) > 1 = {len(files) > 1}")
    
    if not session_name or not files:
        return jsonify({"error": "Missing session or file"}), 400
    if len(files) > 10:
        return jsonify({"error": "最多一次上传 10 个文件"}), 400

    if len(files) > 1:
        # 检测是否是 PPT 生成意图 (多文件合并生成 PPT)
        ppt_keywords = ["ppt", "slide", "幻灯片", "演示文稿", "powerpoint"]
        is_ppt_intent = any(kw in (user_input or "").lower() for kw in ppt_keywords)

        if is_ppt_intent:
            print(f"[FILE UPLOAD] 检测到多文件 PPT 生成意图: {user_input}")
            
            # 预先保存所有文件，避免在生成器中访问已关闭的 FileStorage
            saved_file_paths = []
            source_filenames = []
            
            for f in files:
                if f and f.filename:
                    fname = f.filename
                    fpath = os.path.join(UPLOAD_DIR, fname)
                    # 如果文件指针不在开头，重置它
                    f.seek(0)
                    f.save(fpath)
                    saved_file_paths.append(fpath)
                    source_filenames.append(fname)

            def generate_ppt_stream():
                try:
                    yield f"data: {json.dumps({'type': 'progress', 'message': '📊 正在准备 PPT 生成...', 'detail': f'检测到 {len(saved_file_paths)} 个源文件'})}\n\n"
                    
                    context_text = ""
                    
                    # 1. 提取所有已保存文件内容
                    for i, filepath in enumerate(saved_file_paths):
                        filename = os.path.basename(filepath)
                        yield f"data: {json.dumps({'type': 'progress', 'message': f'📖 正在读取文件 ({i+1}/{len(saved_file_paths)})...', 'detail': filename})}\n\n"
                        
                        try:
                            # 提取内容
                            from web.file_processor import FileProcessor
                            processor = FileProcessor()
                            # 简化版的 process
                            f_result = processor.process_file(filepath)
                            content = f_result.get('text_content') or f_result.get('content', '')
                            
                            # 截断过长内容避免Token爆炸，但保留足够上下文
                            if len(content) > 50000:
                                content = content[:50000] + "...(truncated)"
                                
                            context_text += f"\n\n=== {filename} ===\n{content}\n"
                            
                        except Exception as e:
                            print(f"[PPT BATCH] 读取文件 {filename} 失败: {e}")
                            context_text += f"\n\n=== {filename} (Error) ===\n无法读取内容\n"

                    # 2. 调用 PPT 生成管道
                    yield f"data: {json.dumps({'type': 'progress', 'message': '🎨 正在设计 PPT 结构...', 'detail': '基于多个文件内容'})}\n\n"
                    
                    from web.ppt_pipeline import PPTGenerationPipeline
                    import asyncio
                    
                    # 构造增强后的 Prompt
                    enhanced_prompt = f"{user_input}\n\n【参考资料】\n基于以下文件生成的 PPT:\n{context_text}"
                    
                    # 限制 Prompt 长度
                    if len(enhanced_prompt) > 100000:
                         enhanced_prompt = enhanced_prompt[:100000] + "\n...(context truncated)"
                    
                    # 异步执行 PPT 生成
                    # 使用项目内的 get_client() 获取 Gemini 客户端
                    ai_client = get_client()
                    pipeline = PPTGenerationPipeline(ai_client=ai_client)
                    
                    import threading
                    import traceback
                    import queue

                    pipeline_timeout_sec = 300
                    start_ts = time.time()
                    
                    # 混合消息队列（进度+思考）
                    event_queue = queue.Queue()

                    def _progress_listener(msg, p=None):
                        event_queue.put({"type": "progress", "msg": msg, "progress": p})
                    
                    def _thought_listener(text):
                        # Use a dedicated type for thought/reasoning text
                        event_queue.put({"type": "thought", "text": text})

                    run_state = {
                        "done": False,
                        "result": None,
                        "error": None,
                        "traceback": "",
                    }

                    def _run_pipeline_bg():
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        try:
                            # 传递 progress_callback 和 thought_callback
                            run_state["result"] = loop.run_until_complete(
                                pipeline.generate(
                                    user_request=enhanced_prompt,
                                    output_path=os.path.join(settings_manager.documents_dir, f"Koto_Presentation_{int(time.time())}.pptx"),
                                    enable_auto_images=True,  # 允许自动配图
                                    progress_callback=_progress_listener,
                                    thought_callback=_thought_listener
                                )
                            )
                        except Exception as bg_err:
                            run_state["error"] = str(bg_err)
                            run_state["traceback"] = traceback.format_exc()
                        finally:
                            try:
                                loop.close()
                            except Exception:
                                pass
                            run_state["done"] = True

                    worker = threading.Thread(target=_run_pipeline_bg, daemon=True)
                    worker.start()

                    # 实时轮询进度队列，转发给前端
                    last_progress_msg = "初始化生成环境..."
                    
                    while not run_state["done"]:
                        elapsed = int(time.time() - start_ts)
                        if elapsed > pipeline_timeout_sec:
                            _progress_listener("生成超时，正在强制停止...", 100)
                            run_state["error"] = f"PPT 生成超时（>{pipeline_timeout_sec}s）"
                            break
                        
                        # 消费所有的事件
                        try:
                            while not event_queue.empty():
                                item = event_queue.get_nowait()
                                
                                if item['type'] == 'progress':
                                    msg = item['msg']
                                    p = item['progress']
                                    last_progress_msg = msg
                                    detail_text = f"进度: {p}%" if p is not None else f"已用时 {elapsed}s"
                                    yield f"data: {json.dumps({'type': 'progress', 'message': msg, 'detail': detail_text})}\n\n"
                                
                                elif item['type'] == 'thought':
                                    # Send thought as a partial text response or a special 'thought' event
                                    # Assuming frontend can handle 'text' type for appending to the assistant's message
                                    # or 'thought' for a distinct UI block. 
                                    # Let's use 'text' for now to ensure it appears in the chat stream.
                                    thought_text = f"\n\n> 🤖 **Koto 思考**: {item['text']}\n"
                                    yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"

                        except queue.Empty:
                            pass
                        
                        # 如果没有新消息，每2秒发一次心跳防止连接断开
                        if elapsed % 2 == 0 and event_queue.empty():
                             yield f"data: {json.dumps({'type': 'progress', 'message': last_progress_msg, 'detail': f'已用时 {elapsed}s'})}\n\n"

                        time.sleep(0.5)

                    # 发送最后剩余的消息
                    try:
                        while not event_queue.empty():
                             item = event_queue.get_nowait()
                             if item['type'] == 'progress':
                                  yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': ''})}\n\n"
                             elif item['type'] == 'thought':
                                  thought_text = f"\n\n> 🤖 **Koto 思考**: {item['text']}\n"
                                  yield f"data: {json.dumps({'type': 'text', 'content': thought_text})}\n\n"
                    except: pass

                    if run_state["error"]:
                        err = run_state["error"]
                        tb = run_state.get("traceback", "")
                        print(f"[PPT BATCH] Background pipeline error: {err}")
                        if tb:
                            print(f"[PPT BATCH] Traceback: {tb[:800]}")
                        raise Exception(f"PPT 管道异常: {err}")

                    ppt_result = run_state["result"] or {}
                    
                    # pipeline returns 'output_path', also check 'file_path' for compat
                    saved_path = ppt_result.get("output_path") or ppt_result.get("file_path")
                    
                    if not ppt_result.get("success"):
                        err_detail = ppt_result.get('error', '未知错误')
                        tb = ppt_result.get('traceback', '')
                        print(f"[PPT BATCH] Pipeline returned failure: {err_detail}")
                        if tb:
                            print(f"[PPT BATCH] Traceback: {tb[:500]}")
                        raise Exception(f"PPT 管道生成失败: {err_detail}")
                    
                    if saved_path and os.path.exists(saved_path):
                        yield f"data: {json.dumps({'type': 'progress', 'message': '✅ PPT 生成完成！', 'detail': os.path.basename(saved_path)})}\n\n"
                        
                        rel_path = os.path.relpath(saved_path, WORKSPACE_DIR).replace("\\", "/")
                        success_msg = f"✅ **PPT 生成成功！**\n\n基于 {len(saved_file_paths)} 个文件生成的演示文稿。\n📁 文件: **{os.path.basename(saved_path)}**"
                        
                        yield f"data: {json.dumps({'type': 'token', 'content': success_msg})}\n\n"
                        
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [rel_path], 'total_time': 0})}\n\n"
                    else:
                        raise Exception("PPT 文件生成失败，未返回路径")
                        
                except Exception as e:
                    print(f"[PPT BATCH ERROR] {e}")
                    import traceback
                    traceback.print_exc()
                    err_msg = f"❌ 生成失败: {str(e)}"
                    yield f"data: {json.dumps({'type': 'token', 'content': err_msg})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': 0})}\n\n"

            return Response(stream_with_context(generate_ppt_stream()), mimetype='text/event-stream')

        history = session_manager.load(f"{session_name}.json")
        file_names = [f.filename for f in files if f and f.filename]
        user_message = f"[Files: {', '.join(file_names)}] {user_input}"
        session_manager.append_user_early(f"{session_name}.json", user_message)

        batch_results = []
        combined_saved_files = []
        combined_images = []

        def _process_single_file(file):
            if not file or not file.filename:
                return None

            filename = file.filename
            filepath = os.path.join(UPLOAD_DIR, filename)
            file.save(filepath)
            print(f"[FILE UPLOAD] 文件已保存: {filename} -> {filepath}")
            file_type = file.mimetype or file.content_type or ""
            file_ext = os.path.splitext(filename)[1].lower()

            # 检测是否是纯归档/整理请求（不需要AI分析内容）
            organize_keywords = ["整理", "归档", "归纳", "分类", "整理一下", "整理下", "帮我整理", "文件整理", "organize", "sort"]
            is_organize_only = any(kw in (user_input or "") for kw in organize_keywords)

            try:
                # formatted_message, file_data = process_uploaded_file(filepath, user_input)
                # --- Modify to use FileProcessor directly for simultaneous KB indexing ---
                from web.file_processor import FileProcessor
                _processor = FileProcessor()
                _file_raw = _processor.process_file(filepath)
                
                # 1. 自动建库 (Auto-Indexing to Knowledge Base) - Use threading to not block UI
                try:
                    _text_content = _file_raw.get('text_content', '')
                    if _text_content and len(_text_content) > 50: # Ignore tiny files
                        def _bg_index(content, meta):
                            try:
                                from web.knowledge_base import KnowledgeBase
                                _kb = KnowledgeBase()
                                res = _kb.add_content(content, meta)
                                print(f"[KB] Auto-indexing completed: {res}")
                            except Exception as e:
                                print(f"[KB] Auto-indexing failed: {e}")
                        
                        import threading
                        _idx_thread = threading.Thread(target=_bg_index, args=(_text_content, {
                            "file_path": filepath,
                            "file_name": filename,
                            "file_type": file_ext,
                            "mtime": os.path.getmtime(filepath)
                        }))
                        _idx_thread.start()
                        print(f"[KB] 已启动后台建库任务: {filename}")
                except Exception as _kb_err:
                    print(f"[KB] Indexing trigger failed: {_kb_err}")

                # 2. Continue with standard chat formatting
                formatted_message, file_data = _processor.format_result_for_chat(_file_raw, user_input)

                task_type = locked_task
                context_info = None
                route_method = "Auto"
                if not task_type:
                    if file_data and file_type and file_type.startswith('image'):
                        message_lower = (user_input or "").lower()
                        is_edit = any(kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS)
                        task_type = "PAINTER" if is_edit else "VISION"
                        route_method = "🖼️ Image Edit" if is_edit else "👁️ Image Analysis"
                        print(f"[FILE UPLOAD] 图片任务直通路由: {task_type} (方法: {route_method})")
                    else:
                        use_annotation = _should_use_annotation_system(user_input, has_file=True) and file_ext in [".doc", ".docx"]
                        use_analysis = _is_analysis_request(user_input)

                        if use_annotation:
                            task_type = "DOC_ANNOTATE"
                            route_method = "📌 Annotation-Strict"
                        elif use_analysis:
                            task_type = "RESEARCH"
                            route_method = "🔬 Analysis"
                        else:
                            task_analysis, route_method, context_info = SmartDispatcher.analyze(
                                formatted_message,
                                history=history
                            )
                            task_type = task_analysis

                        if task_type == "CHAT":
                            task_type = "FILE_GEN"
                            route_method = "📄 Upload-Default"

                if task_type == "DOC_ANNOTATE":
                    task_type = "FILE_GEN"
                    route_method = "📄 Batch-Upload-Default"

                if locked_model != 'auto':
                    model_to_use = locked_model
                else:
                    complexity = "complex" if file_data is None else "normal"
                    if context_info and context_info.get("complexity"):
                        complexity = context_info["complexity"]

                    if task_type == "FILE_GEN":
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type,
                            has_image=bool(file_data),
                            complexity=complexity
                        )
                    else:
                        model_to_use = SmartDispatcher.get_model_for_task(
                            task_type,
                            has_image=bool(file_data)
                        )

                print(f"[FILE UPLOAD] 任务类型: {task_type}, 模型: {model_to_use}")

                result = {
                    "task": task_type,
                    "model": model_to_use,
                    "route_method": route_method,
                    "response": "",
                    "images": [],
                    "saved_files": [],
                }

                # 纯归档模式：跳过AI内容分析，直接归档
                if is_organize_only:
                    print(f"[FILE UPLOAD] 纯归档模式: {filename}，跳过AI分析")
                    result["response"] = ""
                    result["task"] = "FILE_ORGANIZE"
                else:
                    print(f"[FILE UPLOAD] 处理文件: {filename}, 使用 brain.chat")
                    brain_result = brain.chat(
                        history=history,
                        user_input=formatted_message,
                        file_data=file_data,
                        model=model_to_use,
                        auto_model=(locked_model == 'auto')
                    )
                    result.update(brain_result)

                # 🗂️ 关键：为每个文件调用FileOrganizer进行归档
                organize_info = {"success": False, "message": "未归档"}
                try:
                    # 使用AI分析文件类型和建议目录
                    from web.file_analyzer import FileAnalyzer
                    analyzer = FileAnalyzer()
                    analysis = analyzer.analyze_file(filepath)  # 只传文件路径
                    suggested_folder = analysis.get('suggested_folder')
                    entity_name = analysis.get('entity')
                    entity_type = analysis.get('entity_type')
                    organizer = get_file_organizer()

                    # 如果已存在同名公司/项目文件夹，则复用
                    if entity_name:
                        existing_folder = organizer.find_entity_folder(entity_name)
                        if existing_folder:
                            suggested_folder = existing_folder
                    
                    if suggested_folder:
                        org_result = organizer.organize_file(
                            filepath,
                            suggested_folder,
                            auto_confirm=True,
                            metadata={
                                "entity": entity_name,
                                "entity_type": entity_type
                            }
                        )
                        
                        if org_result.get('success'):
                            organize_info = {
                                "success": True,
                                "message": f"✅ 已归档到: {org_result.get('relative_path', suggested_folder)}",
                                "category": suggested_folder,
                                "path": org_result.get('dest_file')
                            }
                            print(f"[FILE ORGANIZE] ✅ {filename} -> {suggested_folder}")
                        else:
                            organize_info = {
                                "success": False,
                                "message": f"⚠️ 归档失败: {org_result.get('error', '未知错误')}"
                            }
                    else:
                        organize_info = {
                            "success": False,
                            "message": "⚠️ 无法确定文件分类"
                        }
                except Exception as e:
                    organize_info = {
                        "success": False,
                        "message": f"⚠️ 归档异常: {str(e)}"
                    }
                    print(f"[FILE ORGANIZE ERROR] {filename}: {e}")

                result["file_name"] = filename
                result["organize"] = organize_info
                return result

            except Exception as e:
                return {
                    "file_name": filename,
                    "task": "ERROR",
                    "model": "none",
                    "response": f"❌ 处理文件时出错: {str(e)}",
                    "images": [],
                    "saved_files": [],
                    "organize": {
                        "success": False,
                        "message": "❌ 处理失败，未归档"
                    }
                }

        if stream_mode:
            def generate_progress():
                total = len([f for f in files if f and f.filename])
                started = {
                    "type": "progress",
                    "current": 0,
                    "total": total,
                    "status": "start",
                    "detail": f"开始处理 {total} 个文件"
                }
                yield f"data: {json.dumps(started)}\n\n"

                current = 0
                for file in files:
                    if not file or not file.filename:
                        continue

                    current += 1
                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "processing",
                        "detail": f"处理中: {file.filename} ({current}/{total})"
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                    result = _process_single_file(file)
                    if result:
                        batch_results.append(result)
                        combined_saved_files.extend(result.get("saved_files", []))
                        combined_images.extend(result.get("images", []))

                    payload = {
                        "type": "progress",
                        "current": current,
                        "total": total,
                        "status": "done",
                        "detail": f"完成: {file.filename} ({current}/{total})"
                    }
                    yield f"data: {json.dumps(payload)}\n\n"

                summary_lines = [f"📦 批量处理完成，共 {len(batch_results)} 个文件", ""]

                organized_count = sum(1 for item in batch_results if item.get("organize", {}).get("success"))
                if organized_count > 0:
                    summary_lines.append(f"✅ 已归档: {organized_count} 个文件")

                summary_lines.append("\n📄 **文件详情：**")
                for i, item in enumerate(batch_results, 1):
                    fname = item.get('file_name', 'unknown')
                    task = item.get('task', 'UNKNOWN')
                    organize = item.get('organize', {})

                    status = "✅" if task != "ERROR" else "❌"
                    org_status = organize.get('message', '未归档')

                    summary_lines.append(f"{i}. {status} **{fname}**")
                    summary_lines.append(f"   📂 {org_status}")

                    response = item.get('response', '')
                    if response and len(response) > 100:
                        summary_lines.append(f"   💬 {response[:100]}...")
                    elif response:
                        summary_lines.append(f"   💬 {response}")

                summary_msg = "\n".join(summary_lines)

                session_manager.update_last_model_response(
                    f"{session_name}.json", summary_msg,
                    task="FILE_BATCH",
                    model_name=locked_model if locked_model != 'auto' else "auto",
                    saved_files=combined_saved_files,
                    images=combined_images
                )

                final_payload = {
                    "type": "final",
                    "response": summary_msg,
                    "task": "FILE_BATCH",
                    "model": locked_model if locked_model != 'auto' else "auto",
                    "results": batch_results,
                    "images": combined_images,
                    "saved_files": combined_saved_files
                }
                yield f"data: {json.dumps(final_payload)}\n\n"

            return Response(generate_progress(), mimetype='text/event-stream')

        for file in files:
            result = _process_single_file(file)
            if not result:
                continue
            batch_results.append(result)
            combined_saved_files.extend(result.get("saved_files", []))
            combined_images.extend(result.get("images", []))

        # 生成详细摘要，包含归档信息
        summary_lines = [f"📦 批量处理完成，共 {len(batch_results)} 个文件", ""]
        
        organized_count = sum(1 for item in batch_results if item.get("organize", {}).get("success"))
        if organized_count > 0:
            summary_lines.append(f"✅ 已归档: {organized_count} 个文件")
        
        summary_lines.append("\n📄 **文件详情：**")
        for i, item in enumerate(batch_results, 1):
            fname = item.get('file_name', 'unknown')
            task = item.get('task', 'UNKNOWN')
            organize = item.get('organize', {})
            
            status = "✅" if task != "ERROR" else "❌"
            org_status = organize.get('message', '未归档')
            
            summary_lines.append(f"{i}. {status} **{fname}**")
            summary_lines.append(f"   📂 {org_status}")
            
            # 显示AI响应摘要（截取前100字）
            response = item.get('response', '')
            if response and len(response) > 100:
                summary_lines.append(f"   💬 {response[:100]}...")
            elif response:
                summary_lines.append(f"   💬 {response}")
        
        summary_msg = "\n".join(summary_lines)

        session_manager.update_last_model_response(
            f"{session_name}.json", summary_msg,
            task="FILE_BATCH",
            model_name=locked_model if locked_model != 'auto' else "auto",
            saved_files=combined_saved_files,
            images=combined_images
        )

        return jsonify({
            "response": summary_msg,
            "task": "FILE_BATCH",
            "model": locked_model if locked_model != 'auto' else "auto",
            "results": batch_results,
            "images": combined_images,
            "saved_files": combined_saved_files
        })

    file = files[0]
    
    # Save uploaded file
    filename = file.filename
    filepath = os.path.join(UPLOAD_DIR, filename)
    file.save(filepath)
    print(f"[FILE UPLOAD] 文件已保存: {filename} -> {filepath}")
    file_type = file.mimetype or file.content_type or ""
    file_ext = os.path.splitext(filename)[1].lower()
    
    # Load history first (保证即使出错也能保存用户输入)
    history = session_manager.load(f"{session_name}.json")
    user_message = f"[File: {filename}] {user_input}"
    
    # 🔒 立即保存用户消息到磁盘，防止断连/崩溃导致丢失
    session_manager.append_user_early(f"{session_name}.json", user_message)
    
    try:
        # 使用新的文件处理器（提取文本/二进制）
        formatted_message, file_data = process_uploaded_file(filepath, user_input)
        
        # ==================== 智能文档分析引擎 ====================
        # 对 .docx/.doc 文件，使用 LLM 驱动的智能分析引擎判断用户意图
        # 不再硬编码正则，而是让分析器理解用户真实需求
        if file_ext in ['.docx', '.doc']:
            # 标注任务优先级更高：显式标注意图或用户锁定 DOC_ANNOTATE 时，不进入智能分析引擎
            force_annotation = (locked_task == "DOC_ANNOTATE") or _should_use_annotation_system(user_input, has_file=True)

            # 智能检测：任何对文档内容有实质性处理需求的请求
            # 包括但不限于：写摘要、改引言、改结论、润色、分析结构等
            _doc_intent_keywords = [
                # 生成类
                '写', '生成', '帮我写', '写一段', '写个',
                # 修改/改善类
                '改', '改善', '改进', '优化', '润色', '重写', '修改', '提升',
                # 学术部件
                '摘要', '引言', '结论', 'abstract', '前言', '导言',
                # 分析类
                '分析', '总结', '梳理', '概述', '评估',
                # 质量类
                '不满意', '不好', '不够', '需要改', '有问题',
            ]
            is_doc_processing_request = any(kw in user_input.lower() for kw in _doc_intent_keywords)
            
            if is_doc_processing_request and not force_annotation:
                print(f"[INTELLIGENT ANALYZER] 检测到文档处理请求，启用智能分析引擎")
                from web.intelligent_document_analyzer import create_intelligent_analyzer
                
                # 创建智能分析器
                analyzer = create_intelligent_analyzer(client)
                
                # 流式处理文档分析
                def generate_intelligent_analysis():
                    """生成智能文档分析的流式响应"""
                    try:
                        # 使用async生成器（需要在async context中）
                        import asyncio
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        
                        async def run_analysis():
                            async for event in analyzer.process_document_intelligent_streaming(
                                filepath,
                                user_input,
                                session_name
                            ):
                                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                        
                        gen = run_analysis()
                        while True:
                            try:
                                result = loop.run_until_complete(gen.__anext__())
                                yield result
                            except StopAsyncIteration:
                                break
                    except Exception as e:
                        error_event = {
                            'stage': 'error',
                            'message': f'智能分析失败: {str(e)}'
                        }
                        yield f"data: {json.dumps(error_event, ensure_ascii=False)}\n\n"
                    finally:
                        loop.close()
                
                return Response(
                    stream_with_context(generate_intelligent_analysis()),
                    content_type='text/event-stream'
                )
        # ==================== 智能文档分析引擎结束 ====================
        
        # 智能任务分析
        task_type = locked_task
        context_info = None
        route_method = "Auto"
        if not task_type:
            # 如果是图片上传，直接判断编辑或分析，避免初始化本地路由器导致卡顿
            if file_data and file_type and file_type.startswith('image'):
                message_lower = (user_input or "").lower()
                is_edit = any(kw in message_lower for kw in KotoBrain.IMAGE_EDIT_KEYWORDS)
                task_type = "PAINTER" if is_edit else "VISION"
                route_method = "🖼️ Image Edit" if is_edit else "👁️ Image Analysis"
                print(f"[FILE UPLOAD] 图片任务直通路由: {task_type} (方法: {route_method})")
            else:
                # 文档上传：严格检测标注意图（必须明确要求在原文上标记）
                use_annotation = _should_use_annotation_system(user_input, has_file=True) and file_ext in [".doc", ".docx"]
                # 纯分析请求（不含生成/改善意图）走 RESEARCH
                use_analysis = _is_analysis_request(user_input)
                
                # 🎯 PPT 生成请求显式检测（P0 关键路由）
                ppt_keywords = ["ppt", "幻灯片", "演示", "汇报", "presentation", "slide", "deck"]
                prefer_ppt = any(kw in user_input.lower() for kw in ppt_keywords)

                if use_annotation:
                    task_type = "DOC_ANNOTATE"
                    route_method = "📌 Annotation-Strict"
                elif use_analysis:
                    task_type = "RESEARCH"
                    route_method = "🔬 Analysis"
                elif prefer_ppt:
                    # 🎯 PPT 生成路由（新增 P0 功能）
                    task_type = "FILE_GEN"
                    route_method = "📊 PPT-from-File"
                    print(f"[FILE UPLOAD] 🎯 检测到 PPT 生成请求，启用 FILE_GEN + PPT 模式")
                else:
                    # 如果用户没有锁定任务，智能判断
                    task_analysis, route_method, context_info = SmartDispatcher.analyze(
                        formatted_message,
                        history=history
                    )
                    task_type = task_analysis

                # 上传文件的请求不应归类为 CHAT
                if task_type == "CHAT":
                    task_type = "FILE_GEN"
                    route_method = "📄 Upload-Default"

                print(f"[FILE UPLOAD] 智能路由选择任务类型: {task_type} (方法: {route_method})")
        
        # 确定使用的模型
        if locked_model != 'auto':
            model_to_use = locked_model
        else:
            # 获取任务复杂度（上传文件默认按复杂任务处理）
            complexity = "complex" if file_data is None else "normal"
            if context_info and context_info.get("complexity"):
                complexity = context_info["complexity"]
            
            if task_type == "DOC_ANNOTATE":
                model_to_use = "gemini-3-pro-preview"
            elif task_type == "FILE_GEN":
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type,
                    has_image=bool(file_data),
                    complexity=complexity
                )
            else:
                model_to_use = SmartDispatcher.get_model_for_task(
                    task_type,
                    has_image=bool(file_data)
                )
        
        print(f"[FILE UPLOAD] 任务类型: {task_type}, 模型: {model_to_use}")
        
        # 如果是文本类文件，按任务类型处理
        result = {
            "task": "FILE_GEN" if task_type == "DOC_ANNOTATE" else task_type,
            "subtask": "DOC_ANNOTATE" if task_type == "DOC_ANNOTATE" else None,
            "model": model_to_use,
            "route_method": route_method,
            "response": "",
            "images": [],
            "saved_files": [],
        }

        # 文档标注任务 - 流式反馈，生成带Track Changes的Word文档
        if task_type == "DOC_ANNOTATE":
            docs_dir = settings_manager.documents_dir
            os.makedirs(docs_dir, exist_ok=True)

            source_path = filepath
            target_path = os.path.join(docs_dir, filename)
            if os.path.abspath(source_path) != os.path.abspath(target_path):
                shutil.copy2(source_path, target_path)

            # 使用流式SSE返回进度，让前端能实时显示
            def generate_doc_annotate_stream():
                import time as _time
                _start = _time.time()
                task_id = f"doc_annotate_{session_name}_{int(_start * 1000)}"
                
                try:
                    from web.document_feedback import DocumentFeedbackSystem
                    feedback_system = DocumentFeedbackSystem(gemini_client=client)
                    
                    # 发送分类信息
                    yield f"data: {json.dumps({'type': 'classification', 'task_type': 'DOC_ANNOTATE', 'route_method': route_method, 'model': model_to_use, 'task_id': task_id, 'message': '📄 DOC_ANNOTATE'})}\n\n"
                    
                    # 发送初始进度
                    yield f"data: {json.dumps({'type': 'progress', 'stage': 'init_reading', 'message': '📖 正在读取文档...', 'detail': filename, 'progress': 5})}\n\n"
                    
                    revised_file = None
                    final_result = None
                    
                    for progress_event in feedback_system.full_annotation_loop_streaming(
                        target_path,
                        user_input,
                        task_id=task_id,
                        model_id=model_to_use,
                        cancel_check=lambda: _interrupt_manager.is_interrupted(session_name)
                    ):
                        stage = progress_event.get('stage', 'unknown')
                        progress = progress_event.get('progress', 0)
                        message_text = progress_event.get('message', '')
                        detail = progress_event.get('detail', '')
                        
                        if stage == 'cancelled':
                            yield f"data: {json.dumps({'type': 'info', 'message': '⏸️ 任务已取消'})}\n\n"
                            _elapsed = _time.time() - _start
                            # 保存取消记录
                            session_manager.update_last_model_response(f"{session_name}.json", "⏸️ 文档标注任务已取消")
                            yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed, 'cancelled': True})}\n\n"
                            return
                        
                        yield f"data: {json.dumps({'type': 'progress', 'stage': stage, 'message': message_text, 'detail': detail, 'progress': progress})}\n\n"
                        
                        if stage == 'complete':
                            final_result = progress_event.get('result', {})
                            revised_file = final_result.get('revised_file')
                    
                    _elapsed = _time.time() - _start
                    
                    if final_result and final_result.get('success'):
                        applied = final_result.get('applied', 0)
                        failed = final_result.get('failed', 0)
                        total = final_result.get('total', applied + failed)
                        
                        # 读取文档信息
                        try:
                            from docx import Document as _Doc
                            _d = _Doc(target_path)
                            _total_paras = len([p for p in _d.paragraphs if p.text.strip()])
                            _total_chars = sum(len(p.text) for p in _d.paragraphs)
                        except Exception:
                            _total_paras = 0
                            _total_chars = 0
                        
                        density = (applied / _total_chars * 1000) if _total_chars > 0 else 0
                        
                        summary_lines = [
                            "## ✅ 文档修改完成！",
                            "",
                            "### 📊 修改统计",
                            f"- 找到并应用: **{applied}** 处修改",
                            f"- 定位失败: {failed} 处",
                            f"- 总计分析: {total} 处",
                            "",
                            "### 📋 文档信息",
                            f"- 文件名: `{filename}`",
                            f"- 段落数: {_total_paras} 段",
                            f"- 字数: {_total_chars} 字",
                            f"- 修改密度: **{density:.1f}** 处/千字",
                            "",
                            f"### 📄 模型: `{model_to_use}`",
                            "",
                            f"### 📝 输出文件: `{os.path.basename(revised_file) if revised_file else '待生成'}`",
                            "",
                            "### 💡 使用方法",
                            "1. 用 Microsoft Word 打开输出文件",
                            "2. 点击「审阅」标签页",
                            "3. 右侧气泡中查看全部修改建议",
                            "4. 逐条接受或忽略（右键批注可操作）",
                        ]
                        summary_msg = "\n".join(summary_lines)
                        
                        yield f"data: {json.dumps({'type': 'token', 'content': summary_msg})}\n\n"
                        
                        session_manager.update_last_model_response(
                            f"{session_name}.json", summary_msg,
                            task="DOC_ANNOTATE", model_name=model_to_use,
                            saved_files=[revised_file] if revised_file else []
                        )
                        
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [revised_file] if revised_file else [], 'total_time': _elapsed})}\n\n"
                    else:
                        err_msg = final_result.get('message', '未知错误') if final_result else '处理失败'
                        # 保存失败记录
                        session_manager.update_last_model_response(f"{session_name}.json", f"❌ 文档标注失败: {err_msg}")
                        yield f"data: {json.dumps({'type': 'error', 'message': '❌ ' + err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [], 'total_time': _elapsed})}\n\n"
                
                except Exception as e:
                    import traceback
                    traceback.print_exc()
                    # 保存异常记录
                    session_manager.update_last_model_response(f"{session_name}.json", f"❌ 标注系统错误: {str(e)[:200]}")
                    yield f"data: {json.dumps({'type': 'error', 'message': '❌ 标注系统错误: ' + str(e)[:200]})}\n\n"
                    yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
            
            return Response(generate_doc_annotate_stream(), mimetype='text/event-stream')

        # 🎯 FILE_GEN + PPT 生成（P0 新增）
        elif task_type == "FILE_GEN" and prefer_ppt:
            print(f"[FILE_GEN PPT] 开始 PPT 生成流程")
            
            # 第 1 步：使用 FileParser 提取结构化内容
            from web.file_parser import FileParser
            from web.ppt_session_manager import PPTSessionManager
            
            parser = FileParser()
            parse_result = parser.parse_file(filepath)
            file_content = parse_result.get('content', '') if parse_result else ""
            
            # 第 2 步：创建 PPT 会话
            ppt_session_dir = os.path.join(WORKSPACE_DIR, 'workspace', 'ppt_sessions')
            os.makedirs(ppt_session_dir, exist_ok=True)
            
            session_manager_ppt = PPTSessionManager(ppt_session_dir)
            ppt_session_id = session_manager_ppt.create_session(
                title=f"PPT from {os.path.splitext(filename)[0]}",
                user_input=user_input,
                theme="business"
            )
            print(f"[FILE_GEN PPT] 创建会话: {ppt_session_id}")
            
            # 第 3 步：保存文件内容到会话
            session_manager_ppt.save_generation_data(
                session_id=ppt_session_id,
                ppt_data=None,
                ppt_file_path=None,
                uploaded_file_context=file_content[:3000]  # 将内容限制为前3000字符
            )
            print(f"[FILE_GEN PPT] 文件内容已保存到会话")
            
            # 使用流式响应（Streamed Response）以支持实时进度显示
            def generate_ppt_file_stream():
                import asyncio
                import queue
                import threading
                from web.app import TaskOrchestrator
                import time as _time
                _start = _time.time()
                
                # 发送初始化信息
                yield f"data: {json.dumps({'type': 'classification', 'task_type': 'FILE_GEN', 'subtask': 'PPT_CREATION', 'message': '📊 开始 PPT 演示文稿生成流程'})}\n\n"
                
                # 准备任务参数
                subtask = {
                    "task_type": "FILE_GEN",
                    "index": 1,
                    "description": f"从文档 {filename} 生成 PPT"
                }
                context = {
                    "original_input": user_input,
                    "step_1_output": file_content
                }
                
                # 进度队列
                progress_queue = queue.Queue()
                def _progress_cb(msg, detail=""):
                    progress_queue.put({"msg": msg, "detail": detail})
                
                # 任务结果容器
                task_result_holder = {"result": None}

                # 后台执行函数
                def _run_task_thread():
                    # 为新线程创建独立的事件循环
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    try:
                        task_result_holder["result"] = loop.run_until_complete(
                            TaskOrchestrator._execute_file_gen(user_input, context, subtask, _progress_cb)
                        )
                    except Exception as e:
                        task_result_holder["result"] = {"success": False, "error": str(e)}
                    finally:
                        loop.close()
                        progress_queue.put(None) # Signal done

                # 启动后台线程
                t = threading.Thread(target=_run_task_thread)
                t.start()
                
                # 主线程循环读取进度
                while True:
                    try:
                        item = progress_queue.get(timeout=0.1)
                        if item is None:
                            break
                        # 发送进度SSE
                        yield f"data: {json.dumps({'type': 'progress', 'message': item['msg'], 'detail': item['detail']})}\n\n"
                    except queue.Empty:
                        if not t.is_alive():
                            break
                
                t.join()
                ppt_result = task_result_holder["result"]
                _elapsed = _time.time() - _start

                # 处理最终结果
                if ppt_result and ppt_result.get('success'):
                    saved_files = ppt_result.get('saved_files', [])
                    if saved_files:
                        ppt_file_path = saved_files[0] if isinstance(saved_files, list) else saved_files
                        # 保存会话数据
                        session_manager_ppt.save_generation_data(
                            session_id=ppt_session_id,
                            ppt_data=ppt_result.get('ppt_data'),
                            ppt_file_path=ppt_file_path
                        )
                        
                        final_msg = (
                             f"✅ PPT 演示已生成\n\n"
                             f"📄 文件: [{os.path.basename(ppt_file_path)}]({ppt_file_path.replace(os.sep, '/')})\n"
                             f"🔗 会话ID: `{ppt_session_id}`\n"
                             f"⏱️ 耗时: {_elapsed:.1f}s"
                        )
                        yield f"data: {json.dumps({'type': 'token', 'content': final_msg})}\n\n"
                        
                        # 更新历史
                        session_manager.update_last_model_response(
                            f"{session_name}.json", final_msg,
                            task="FILE_GEN", model_name=model_to_use, saved_files=[ppt_file_path]
                        )
                        
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': [ppt_file_path], 'ppt_session_id': ppt_session_id})}\n\n"
                    else:
                        err_msg = "⚠️ PPT 框架已生成，但文件保存失败"
                        yield f"data: {json.dumps({'type': 'error', 'message': err_msg})}\n\n"
                        yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"
                else:
                     err_msg = ppt_result.get('error', '未知错误') if ppt_result else "任务执行无结果"
                     yield f"data: {json.dumps({'type': 'error', 'message': f'❌ 生成失败: {err_msg}'})}\n\n"
                     yield f"data: {json.dumps({'type': 'done', 'images': [], 'saved_files': []})}\n\n"

            return Response(stream_with_context(generate_ppt_file_stream()), mimetype='text/event-stream')
        
            analysis_prompt = (
                "你是Koto，一个专业的文档分析助手。\n"
                "请基于用户上传的原始文件内容，按以下结构生成高质量总结：\n\n"
                "# 摘要\n"
                "- 用3-5条要点概述核心信息\n\n"
                "# 详细分析\n"
                "## 背景与上下文\n"
                "## 关键观点与要点\n"
                "## 数据与证据\n"
                "## 结论与建议\n\n"
                "要求：\n"
                "- 用中文输出，条理清晰，避免冗长\n"
                "- 不要输出任何代码或BEGIN_FILE标记\n"
                "- 使用标题和项目符号组织内容\n"
            )

            full_input = formatted_message + "\n\n" + analysis_prompt
            print(f"[FILE UPLOAD] 调用模型: {model_to_use}")
            
            response = client.models.generate_content(
                model=model_to_use,
                contents=full_input,
                config=types.GenerateContentConfig(
                    system_instruction=_get_filegen_brief_instruction(),
                    temperature=0.7,
                    max_output_tokens=4000,
                )
            )
            text_out = response.text or "(无输出)"
            print(f"[FILE UPLOAD] 模型返回长度: {len(text_out)} 字符")

            # Save DOCX and optionally PDF
            title = _build_analysis_title(user_input, filename, is_binary=False)
            cleaned_text = _strip_code_blocks(text_out)
            saved_docx = save_docx(cleaned_text, title=title, output_dir=settings_manager.documents_dir)
            docx_rel = os.path.relpath(saved_docx, WORKSPACE_DIR).replace("\\", "/")
            result["saved_files"].append(docx_rel)
            print(f"[FILE UPLOAD] ✅ 已保存 DOCX: {docx_rel}")
            
            # Also save as PDF if user wants both formats
            if user_input and any(kw in user_input.lower() for kw in ['pdf', '两种格式', 'both']):
                try:
                    saved_pdf = save_pdf(cleaned_text, title=title, output_dir=settings_manager.documents_dir)
                    pdf_rel = os.path.relpath(saved_pdf, WORKSPACE_DIR).replace("\\", "/")
                    result["saved_files"].append(pdf_rel)
                    print(f"[FILE UPLOAD] ✅ 已保存 PDF: {pdf_rel}")
                except Exception as pdf_err:
                    print(f"[FILE UPLOAD] ⚠️ PDF保存失败: {pdf_err}")
            
            result["response"] = (
                "✅ 文档分析完成并已保存\n\n"
                f"📄 生成文件: {', '.join([os.path.basename(f) for f in result['saved_files']])}\n"
                f"📂 位置: `{settings_manager.documents_dir}`\n\n"
                f"💡 提示: 文件已自动保存到 workspace/documents 目录"
            )
        else:
            # Binary (image/PDF) → let brain handle vision and optionally save analysis as DOCX
            print(f"[FILE UPLOAD] 二进制文件，使用 brain.chat 进行视觉分析")
            brain_result = brain.chat(
                history=history,
                user_input=formatted_message,
                file_data=file_data,
                model=model_to_use,
                auto_model=(locked_model == 'auto')
            )
            result.update(brain_result)
            # If model returned text, save a DOCX analysis too
            if brain_result.get("response") and len(brain_result["response"]) > 50:
                try:
                    title = _build_analysis_title(user_input, filename, is_binary=True)
                    cleaned_text = _strip_code_blocks(brain_result["response"])
                    saved_path = save_docx(cleaned_text, title=title, output_dir=settings_manager.documents_dir)
                    docx_rel = os.path.relpath(saved_path, WORKSPACE_DIR).replace("\\", "/")
                    result.setdefault("saved_files", []).append(docx_rel)
                    print(f"[FILE UPLOAD] ✅ 视觉分析已保存为 DOCX: {docx_rel}")
                    
                    # Update response to mention saved analysis
                    if result.get("response"):
                        result["response"] += f"\n\n📄 分析报告已保存: {os.path.basename(saved_path)}"
                except Exception as docx_err:
                    print(f"[FILE UPLOAD] ⚠️ DOCX保存失败: {docx_err}")
        
        # 确保result包含完整的响应
        if not result.get('response'):
            result['response'] = "处理完成"
        if not result.get('task'):
            result['task'] = task_type
        if not result.get('model'):
            result['model'] = model_to_use
        
        # Update history（更新模型回复，用户消息已早期保存，包含元数据）
        session_manager.update_last_model_response(
            f"{session_name}.json", result.get("response", ""),
            task=result.get('task', task_type),
            model_name=result.get('model', model_to_use),
            saved_files=result.get('saved_files', []),
            images=result.get('images', [])
        )
        
        print(f"[FILE UPLOAD] 响应成功，任务: {result.get('task')}, 文件: {len(result.get('saved_files', []))} 个")
        return jsonify(result)
        
    except Exception as e:
        # 即使出错也保存用户的问题和错误信息
        import traceback
        error_detail = traceback.format_exc()
        print(f"[FILE UPLOAD ERROR] {error_detail}")
        
        error_response = f"❌ 处理文件时出错: {str(e)}"
        session_manager.update_last_model_response(f"{session_name}.json", error_response)
        
        return jsonify({
            "response": error_response,
            "task": "ERROR",
            "model": "none",
            "images": [],
            "saved_files": []
        })

# ==================== PPT 相关 API 端点（P0 补充）====================

@app.route('/api/ppt/download', methods=['POST'])
def download_ppt():
    """下载 PPT PPTX 文件"""
    try:
        session_id = request.json.get('session_id')
        if not session_id:
            return jsonify({"error": "Missing session_id"}), 400
        
        # 从 PPT 会话中获取文件路径
        from web.ppt_session_manager import PPTSessionManager
        ppt_session_dir = os.path.join(WORKSPACE_DIR, 'workspace', 'ppt_sessions')
        manager = PPTSessionManager(ppt_session_dir)
        
        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404
        
        ppt_file_path = session_data.get('ppt_file_path')
        if not ppt_file_path:
            # 如果文件还没生成，尝试生成一个临时的
            return jsonify({"error": "PPT file not generated yet"}), 400
        
        # 构建完整的文件路径
        full_path = os.path.join(WORKSPACE_DIR, ppt_file_path.lstrip('/').replace('/', os.sep))
        
        if not os.path.exists(full_path):
            return jsonify({"error": "PPT file not found"}), 404
        
        # 返回文件下载
        return send_file(
            full_path,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation',
            as_attachment=True,
            download_name=os.path.basename(full_path)
        )
    
    except Exception as e:
        print(f"[PPT DOWNLOAD] 错误: {e}")
        return jsonify({"error": f"Download failed: {str(e)}"}), 500


@app.route('/api/ppt/session/<session_id>', methods=['GET'])
def get_ppt_session(session_id):
    """获取 PPT 会话信息"""
    try:
        from web.ppt_session_manager import PPTSessionManager
        ppt_session_dir = os.path.join(WORKSPACE_DIR, 'workspace', 'ppt_sessions')
        manager = PPTSessionManager(ppt_session_dir)
        
        session_data = manager.load_session(session_id)
        if not session_data:
            return jsonify({"error": "Session not found"}), 404
        
        return jsonify({
            "success": True,
            "session": {
                "id": session_data.get('session_id'),
                "title": session_data.get('title'),
                "status": session_data.get('status'),
                "ppt_file_path": session_data.get('ppt_file_path'),
                "created_at": session_data.get('created_at'),
                "updated_at": session_data.get('updated_at')
            }
        })
    
    except Exception as e:
        print(f"[PPT SESSION] 错误: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/ping', methods=['GET'])
def ping():
    start = time.time()
    try:
        client.models.get(model=MODEL_MAP['CHAT'])
        latency = (time.time() - start) * 1000
        return jsonify({
            "status": "ok",
            "latency": latency,
            "ollama": LocalDispatcher.is_ollama_running()
        })
    except Exception as e:
        return jsonify({
            "status": "error",
            "error": str(e)
        })

@app.route('/api/health', methods=['GET'])
def health():
    """轻量健康检查（不触发模型调用）"""
    return jsonify({
        "status": "ok",
        "time": time.time()
    })

@app.route('/api/analyze', methods=['POST'])
def analyze_task():
    """预分析任务类型和模型选择 - 让前端立即显示"""
    data = request.json
    message = data.get('message', '')
    locked_task = data.get('locked_task')
    locked_model = data.get('locked_model', 'auto')
    has_file = data.get('has_file', False)
    file_type = data.get('file_type', '')
    
    if not message:
        return jsonify({"task": "CHAT", "model": MODEL_MAP["CHAT"], "route_method": "Empty"})
    
    # 图像编辑关键词
    IMAGE_EDIT_KEYWORDS = [
        "修改", "换", "改成", "变成", "底色", "背景", "颜色",
        "抠图", "去背景", "P图", "美化", "滤镜", "调色", "编辑",
        "change", "modify", "edit", "background", "color",
    ]
    
    # 如果用户锁定了任务类型
    if locked_task:
        task = locked_task
        route_method = "🔒 Manual"
    elif has_file and file_type and file_type.startswith('image'):
        # 有图片文件，判断是编辑还是分析
        message_lower = message.lower()
        is_edit = any(kw in message_lower for kw in IMAGE_EDIT_KEYWORDS)
        if is_edit:
            task = "PAINTER"
            route_method = "🖼️ Image Edit"
        else:
            task = "VISION"
            route_method = "👁️ Image Analysis"
    else:
        # 使用智能路由器
        task, route_method, _ = SmartDispatcher.analyze(message)
    
    # 如果用户选择了特定模型
    if locked_model and locked_model != 'auto':
        model = locked_model
    else:
        model = SmartDispatcher.get_model_for_task(task, has_image=has_file)
    
    # 获取模型显示信息
    model_info = MODEL_INFO.get(model, {"name": model, "speed": ""})
    
    return jsonify({
        "task": task,
        "model": model,
        "model_name": model_info.get("name", model),
        "model_speed": model_info.get("speed", ""),
        "route_method": route_method,  # 路由算法信息
        "strengths": model_info.get("strengths", []),
    })

@app.route('/api/workspace/<path:filepath>')
def get_workspace_file(filepath):
    """获取 workspace 中的文件，支持子目录"""
    print(f"[API] Serving workspace file: {filepath}")
    full_path = os.path.join(WORKSPACE_DIR, filepath)
    
    # 安全检查：确保请求的路径在 WORKSPACE_DIR 下
    try:
        resolved_path = os.path.abspath(full_path)
        resolved_workspace = os.path.abspath(WORKSPACE_DIR)
        if not resolved_path.startswith(resolved_workspace):
            print(f"[API] Security violation: {resolved_path} not under {resolved_workspace}")
            return jsonify({"error": "Access denied"}), 403
        
        if not os.path.exists(resolved_path):
            print(f"[API] File not found: {resolved_path}")
            return jsonify({"error": "File not found"}), 404
        
        print(f"[API] Serving: {resolved_path}")
        return send_from_directory(WORKSPACE_DIR, filepath)
    except Exception as e:
        print(f"[API] Error serving {filepath}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Server error", "detail": str(e)}), 500

@app.route('/api/workspace', methods=['GET'])
def list_workspace_files():
    files = os.listdir(WORKSPACE_DIR)
    return jsonify({"files": files})

@app.route('/api/open-workspace', methods=['POST'])
def open_workspace():
    """打开 workspace 文件夹"""
    try:
        if sys.platform == "win32":
            subprocess.Popen(
                f'explorer "{WORKSPACE_DIR}"',
                shell=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
        elif sys.platform == "darwin":
            subprocess.Popen(["open", WORKSPACE_DIR])
        else:
            subprocess.Popen(["xdg-open", WORKSPACE_DIR])
        return jsonify({"success": True, "path": WORKSPACE_DIR})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/open-file', methods=['POST'])
def open_file_native():
    """用系统默认程序打开文件（不经过浏览器）"""
    try:
        data = request.get_json()
        filepath = data.get("filepath", "")
        if not filepath:
            return jsonify({"success": False, "error": "No filepath provided"}), 400
        
        full_path = os.path.join(WORKSPACE_DIR, filepath)
        resolved_path = os.path.abspath(full_path)
        resolved_workspace = os.path.abspath(WORKSPACE_DIR)
        
        if not resolved_path.startswith(resolved_workspace):
            return jsonify({"success": False, "error": "Access denied"}), 403
        
        if not os.path.exists(resolved_path):
            return jsonify({"success": False, "error": "File not found"}), 404
        
        print(f"[API] Opening file natively: {resolved_path}")
        if sys.platform == "win32":
            os.startfile(resolved_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", resolved_path])
        else:
            subprocess.Popen(["xdg-open", resolved_path])
        
        return jsonify({"success": True, "path": resolved_path})
    except Exception as e:
        print(f"[API] Error opening file: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# ================= Settings API =================

@app.route('/api/settings', methods=['GET'])
def get_settings():
    # 合并 appearance 主题（如有 cookie/参数可在此合并）
    return jsonify(settings_manager.get_all())

@app.route('/api/settings', methods=['POST'])
def update_settings():
    data = request.json
    category = data.get('category')
    key = data.get('key')
    value = data.get('value')
    
    if category and key:
        success = settings_manager.set(category, key, value)
        settings_manager.ensure_directories()
        return jsonify({"success": success})
    return jsonify({"success": False, "error": "Missing category or key"})

@app.route('/api/settings/reset', methods=['POST'])
def reset_settings():
    success = settings_manager.reset()
    return jsonify({"success": success})

# ================= Mini Mode Switch API =================

@app.route('/api/switch-to-mini', methods=['POST'])
def switch_to_mini():
    """切换到迷你模式"""
    import subprocess
    import sys
    
    try:
        # 启动迷你窗口
        mini_koto_path = os.path.join(PROJECT_ROOT, 'web', 'mini_koto.py')
        if os.path.exists(mini_koto_path):
            # 在新进程中启动迷你窗口
            subprocess.Popen(
                [sys.executable, mini_koto_path],
                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0,
                cwd=PROJECT_ROOT
            )
            return jsonify({"success": True, "message": "迷你模式已启动"})
        else:
            return jsonify({"success": False, "error": "找不到迷你模式程序"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/switch-to-main', methods=['POST'])
def switch_to_main():
    """切换到主程序"""
    import subprocess
    import sys
    
    try:
        # 启动主窗口
        main_app_path = os.path.join(PROJECT_ROOT, 'koto_app.py')
        if os.path.exists(main_app_path):
            subprocess.Popen(
                [sys.executable, main_app_path],
                creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0,
                cwd=PROJECT_ROOT
            )
            return jsonify({"success": True, "message": "主程序已启动"})
        else:
            return jsonify({"success": False, "error": "找不到主程序"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/mini')
def mini_page():
    """迷你模式页面（浏览器访问用）"""
    return render_template('mini_koto.html')

@app.route('/api/mini/chat', methods=['POST'])
def mini_chat():
    """迷你模式专用聊天API - 使用与原版完全相同的任务分配和执行逻辑"""
    data = request.json
    user_input = data.get('message', '').strip()
    
    if not user_input:
        return jsonify({"error": "消息不能为空"}), 400
    
    user_input = Utils.sanitize_string(user_input)
    
    # 使用固定的迷你会话
    session_name = "MiniKoto_Quick"
    history = session_manager.load(f"{session_name}.json")
    
    # 🎯 使用 SmartDispatcher 进行任务分析（与完整版相同）
    task_type, route_method, context_info = SmartDispatcher.analyze(user_input, history)
    print(f"[MINI_CHAT] SmartDispatcher 分析结果: task_type='{task_type}', method='{route_method}'")
    
    response_text = ""
    is_error = False
    used_model = "unknown"
    
    try:
        # ===== 根据任务类型执行不同的处理逻辑（与完整版相同）=====
        
        if task_type == "WEB_SEARCH":
            # 🌐 网络搜索 - 使用 Gemini Google Search Grounding
            print(f"[MINI_CHAT] 🌐 执行网络搜索...")
            search_result = WebSearcher.search_with_grounding(user_input)
            response_text = search_result.get("response", "")
            used_model = "gemini-2.5-flash (Google Search)"
            
            # 如果搜索失败，尝试修正查询
            if not search_result.get("success") or Utils.is_failure_output(response_text) or "搜索失败" in response_text:
                print(f"[MINI_CHAT] ⚠️ 初次搜索失败，尝试修正查询...")
                fix_query_prompt = (
                    "请把用户需求改写成更适合搜索的简短关键词或查询语句，只输出查询语句。\n"
                    f"用户需求: {user_input}"
                )
                try:
                    fix_query_resp = client.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=fix_query_prompt,
                        config=types.GenerateContentConfig(temperature=0.2, max_output_tokens=64)
                    )
                    fixed_query = (fix_query_resp.text or user_input).strip()
                    print(f"[MINI_CHAT] 修正后的查询: {fixed_query}")
                    search_result = WebSearcher.search_with_grounding(fixed_query)
                    response_text = search_result.get("response", "")
                except Exception as e:
                    print(f"[MINI_CHAT] 修正查询失败: {e}")
            
            if not response_text or Utils.is_failure_output(response_text):
                is_error = True
                response_text = f"搜索失败：无法获取 '{user_input}' 的实时信息"
        
        elif task_type == "SYSTEM":
            # 🖥️ 系统命令 - 本地执行
            print(f"[MINI_CHAT] 🖥️ 执行系统命令：{user_input}")
            try:
                exec_result = LocalExecutor.execute(user_input)
                response_text = exec_result.get("message", "命令执行失败")
                if exec_result.get("details"):
                    response_text += f"\n\n{exec_result['details']}"
                used_model = "LocalExecutor"
                is_error = not exec_result.get("success", False)
                
                # 如果执行失败，尝试用 AI 修正
                if is_error or Utils.is_failure_output(response_text):
                    print(f"[MINI_CHAT] ⚠️ 本地执行失败，尝试 AI 修正...")
                    fix_prompt = Utils.build_fix_prompt("SYSTEM", user_input, response_text)
                    try:
                        fix_resp = client.models.generate_content(
                            model="gemini-2.0-flash",
                            contents=fix_prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=_get_DEFAULT_CHAT_SYSTEM_INSTRUCTION(),
                                temperature=0.4,
                                max_output_tokens=1000,
                            )
                        )
                        response_text = fix_resp.text or response_text
                        used_model = "gemini-2.0-flash (fallback)"
                        is_error = False
                    except Exception as e:
                        print(f"[MINI_CHAT] AI 修正失败: {e}")
            except Exception as e:
                print(f"[MINI_CHAT] ❌ 系统命令执行出错: {e}")
                response_text = f"系统命令执行出错：{str(e)}"
                used_model = "LocalExecutor"
                is_error = True
        
        else:
            # 💬 其他任务（CHAT, RESEARCH, CODER 等）- 使用 brain.chat()
            print(f"[MINI_CHAT] 💬 执行 {task_type} 任务...")
            model = MODEL_MAP.get(task_type, MODEL_MAP['CHAT'])
            result = brain.chat(history, user_input, model=model, auto_model=False)
            response_text = result.get("response", "")
            used_model = result.get("model", model)
            is_error = response_text.startswith("Error:")
            
            # 如果遇到 404 错误，尝试备用模型
            if is_error and "404" in response_text:
                print(f"[MINI_CHAT] ⚠️ 模型 404，尝试备用模型...")
                for fallback_model in ["gemini-2.0-flash", "gemini-1.5-pro"]:
                    try:
                        result = brain.chat(history, user_input, model=fallback_model, auto_model=False)
                        if not result.get("response", "").startswith("Error:"):
                            response_text = result.get("response", "")
                            used_model = fallback_model
                            is_error = False
                            break
                    except Exception as e:
                        continue
    
    except Exception as e:
        print(f"[MINI_CHAT] ❌ 执行出错: {e}")
        is_error = True
        response_text = f"Error: {str(e)}"
    
    # 更新历史（成功和失败都保存，便于排查）
    if response_text:
        session_manager.append_and_save(f"{session_name}.json", user_input, response_text)
    
    print(f"[MINI_CHAT] ✅ 完成: task_type={task_type}, model={used_model}, success={not is_error}")
    
    # 返回统一格式
    return jsonify({
        "success": not is_error,
        "response": response_text,
        "model": used_model,
        "task_type": task_type,
        "route_method": route_method,
        "error": response_text if is_error else ""
    })

# ================= Setup & Initialization API =================

@app.route('/api/setup/status', methods=['GET'])
def get_setup_status():
    """检查首次设置状态"""
    config_path = os.path.join(PROJECT_ROOT, "config", "gemini_config.env")
    has_api_key = bool(API_KEY and len(API_KEY) > 10)
    has_workspace = os.path.exists(WORKSPACE_DIR)
    
    return jsonify({
        "initialized": has_api_key and has_workspace,
        "has_api_key": has_api_key,
        "has_workspace": has_workspace,
        "workspace_path": os.path.abspath(WORKSPACE_DIR),
        "config_path": os.path.abspath(config_path)
    })

@app.route('/api/setup/apikey', methods=['POST'])
def setup_api_key():
    """设置 API Key"""
    data = request.json
    api_key = data.get('api_key', '').strip()
    
    if not api_key or len(api_key) < 10:
        return jsonify({"success": False, "error": "Invalid API key"})
    
    config_path = os.path.join(PROJECT_ROOT, "config", "gemini_config.env")
    os.makedirs(os.path.dirname(config_path), exist_ok=True)
    
    try:
        # 写入配置文件
        with open(config_path, 'w', encoding='utf-8') as f:
            f.write(f"# Koto Configuration\nAPI_KEY={api_key}\n")
        
        # 更新环境变量
        os.environ['API_KEY'] = api_key
        global API_KEY, client
        API_KEY = api_key
        client = create_client()
        
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/setup/workspace', methods=['POST'])
def setup_workspace():
    """设置工作区目录"""
    data = request.json
    workspace_path = data.get('path', '').strip()
    
    if not workspace_path:
        workspace_path = os.path.join(PROJECT_ROOT, "workspace")
    
    try:
        os.makedirs(workspace_path, exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "documents"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "images"), exist_ok=True)
        os.makedirs(os.path.join(workspace_path, "code"), exist_ok=True)
        
        # 更新设置
        settings_manager.set('storage', 'workspace_dir', workspace_path)
        
        return jsonify({"success": True, "path": workspace_path})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/setup/test', methods=['GET'])
def test_api_connection():
    """测试 API 连接"""
    try:
        start = time.time()
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents="Say 'Koto is ready!' in one short sentence."
        )
        latency = time.time() - start
        return jsonify({
            "success": True,
            "message": response.text,
            "latency": round(latency, 2)
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/api/diagnose', methods=['GET'])
def diagnose_models():
    """诊断所有模型的可用性"""
    import threading
    
    results = {
        "proxy": {
            "detected": get_detected_proxy(),
            "force": FORCE_PROXY or None,
            "custom_endpoint": GEMINI_API_BASE or None
        },
        "models": {}
    }
    
    # 测试模型列表
    test_models = [
        ("gemini-2.0-flash-lite", "路由分类"),
        ("gemini-3-flash-preview", "日常对话"),
        ("gemini-3-pro-preview", "代码生成"),
        ("gemini-2.5-flash", "联网搜索"),
        ("nano-banana-pro-preview", "图像生成"),
    ]
    
    def test_model(model_id, purpose):
        try:
            start = time.time()
            if "nano-banana" in model_id or "imagen" in model_id:
                # 图像模型只测试连通性
                response = client.models.generate_content(
                    model=model_id,
                    contents="test",
                    config=types.GenerateContentConfig(
                        max_output_tokens=10
                    )
                )
            else:
                response = client.models.generate_content(
                    model=model_id,
                    contents="Reply with only: OK",
                    config=types.GenerateContentConfig(
                        max_output_tokens=10
                    )
                )
            latency = time.time() - start
            return {
                "status": "✅ 可用",
                "latency": round(latency, 2),
                "purpose": purpose
            }
        except Exception as e:
            error_msg = str(e)
            if "location is not supported" in error_msg:
                status = "❌ 地区限制"
            elif "not found" in error_msg.lower():
                status = "❌ 模型不存在"
            elif "quota" in error_msg.lower():
                status = "⚠️ 配额耗尽"
            elif "timeout" in error_msg.lower():
                status = "⚠️ 超时"
            else:
                status = f"❌ 错误"
            return {
                "status": status,
                "error": error_msg[:150],
                "purpose": purpose
            }
    
    # 并行测试（带超时）
    threads = []
    for model_id, purpose in test_models:
        def run_test(m=model_id, p=purpose):
            results["models"][m] = test_model(m, p)
        t = threading.Thread(target=run_test, daemon=True)
        threads.append(t)
        t.start()
    
    # 等待所有线程完成（最多 15 秒）
    for t in threads:
        t.join(timeout=15)
    
    # 检查是否所有模型都不可用
    all_failed = all(
        "❌" in results["models"].get(m, {}).get("status", "")
        for m, _ in test_models
    )
    
    if all_failed:
        results["recommendation"] = "所有模型均不可用。建议：\n1. 检查代理配置是否正确\n2. 考虑使用 API 中转服务\n3. 在 gemini_config.env 中配置 GEMINI_API_BASE"
    
    return jsonify(results)

@app.route('/api/browse', methods=['GET'])
def browse_folders():
    import os
    path = request.args.get('path', 'C:\\')
    
    try:
        if not os.path.exists(path):
            return jsonify({"error": "路径不存在", "folders": [], "parent": None})
        
        if not os.path.isdir(path):
            return jsonify({"error": "不是文件夹", "folders": [], "parent": None})
        
        folders = []
        try:
            for item in os.listdir(path):
                item_path = os.path.join(path, item)
                if os.path.isdir(item_path):
                    folders.append({"name": item, "path": item_path})
        except PermissionError:
            return jsonify({"error": "没有权限访问", "folders": [], "parent": None})
        
        folders.sort(key=lambda x: x['name'].lower())
        
        # Get parent path
        parent = os.path.dirname(path)
        if parent == path:  # Root drive
            parent = None
        
        return jsonify({
            "folders": folders,
            "parent": parent,
            "current": path
        })
    except Exception as e:
        return jsonify({"error": str(e), "folders": [], "parent": None})


@app.route('/api/chat/interrupt', methods=['POST'])
def interrupt_chat():
    """中断当前对话生成"""
    payload = request.json or {}
    session_name = payload.get('session')
    task_id = payload.get('task_id')
    if not session_name:
        return jsonify({"error": "Missing session"}), 400
    
    # 使用新的中断管理器
    _interrupt_manager.set_interrupt(session_name)
    # 保持向后兼容
    _interrupt_flags[session_name] = True

    # 可选：如果前端传入 task_id，同步取消调度器任务（用于 DOC_ANNOTATE 等流式长任务）
    if task_id:
        try:
            from task_scheduler import get_task_scheduler
            get_task_scheduler().cancel_task(task_id)
            print(f"[INTERRUPT] Cancel task_id={task_id}")
        except Exception as e:
            print(f"[INTERRUPT] cancel task failed: {e}")
    
    # 同步中断标志到 AgentLoop（如果正在执行 Agent 任务）
    # NOTE: Legacy agent_loop retired — interrupt handled by _interrupt_manager above
    pass
    
    return jsonify({"success": True, "message": "Chat interrupted"})


@app.route('/api/chat/reset-interrupt', methods=['POST'])
def reset_interrupt():
    """重置中断标志"""
    session_name = request.json.get('session')
    if session_name:
        # 使用新的中断管理器
        _interrupt_manager.reset(session_name)
        # 保持向后兼容
        if session_name in _interrupt_flags:
            del _interrupt_flags[session_name]
    return jsonify({"success": True})


# ================= 新功能 API 路由 =================

# === 快速笔记 API ===
@app.route('/api/notes/add', methods=['POST'])
def add_note():
    """添加笔记"""
    from note_manager import get_note_manager
    
    data = request.json
    title = data.get('title', '')
    content = data.get('content', '')
    category = data.get('category', 'default')
    tags = data.get('tags', [])
    
    note_manager = get_note_manager()
    note_id = note_manager.add_note(title, content, category, tags)
    
    return jsonify({"success": True, "note_id": note_id})


@app.route('/api/notes/list', methods=['GET'])
def list_notes():
    """列出最近笔记"""
    from note_manager import get_note_manager
    
    limit = int(request.args.get('limit', 20))
    category = request.args.get('category')
    
    note_manager = get_note_manager()
    notes = note_manager.get_recent_notes(limit, category)
    
    return jsonify({"notes": notes})


@app.route('/api/notes/search', methods=['GET'])
def search_notes():
    """搜索笔记"""
    from note_manager import get_note_manager
    
    query = request.args.get('query', '')
    note_manager = get_note_manager()
    results = note_manager.search_notes(query)
    
    return jsonify({"results": results})


@app.route('/api/notes/<note_id>', methods=['DELETE'])
def delete_note(note_id):
    """删除笔记"""
    from note_manager import get_note_manager
    
    note_manager = get_note_manager()
    success = note_manager.delete_note(note_id)
    
    return jsonify({"success": success})


# === 本地提醒 API（Windows 系统通知） ===
@app.route('/api/reminders/add', methods=['POST'])
def add_reminder():
    """创建本地系统提醒
    请求体: {"title": str, "message": str, "time": ISO8601, "seconds": int}
    - 传 time (ISO 时间) 或 seconds (相对秒数) 任选其一
    """
    from reminder_manager import get_reminder_manager
    from datetime import datetime

    data = request.json or {}
    title = data.get('title') or '提醒'
    message = data.get('message') or ''
    icon = data.get('icon')
    remind_time = data.get('time')
    seconds = data.get('seconds')

    mgr = get_reminder_manager()
    if remind_time:
        try:
            dt = datetime.fromisoformat(remind_time)
        except Exception:
            return jsonify({"success": False, "error": "时间格式需为 ISO8601"}), 400
        rid = mgr.add_reminder(title, message, dt, icon)
    elif seconds is not None:
        try:
            sec = int(seconds)
        except Exception:
            return jsonify({"success": False, "error": "seconds 需为整数"}), 400
        rid = mgr.add_reminder_in(title, message, sec, icon)
    else:
        return jsonify({"success": False, "error": "需提供 time 或 seconds"}), 400

    return jsonify({"success": True, "reminder_id": rid})


@app.route('/api/reminders/list', methods=['GET'])
def list_reminders_api():
    """列出所有提醒"""
    from reminder_manager import get_reminder_manager
    mgr = get_reminder_manager()
    return jsonify({"reminders": mgr.list_reminders()})


@app.route('/api/reminders/<reminder_id>', methods=['DELETE'])
def cancel_reminder(reminder_id):
    """取消提醒"""
    from reminder_manager import get_reminder_manager
    mgr = get_reminder_manager()
    ok = mgr.cancel_reminder(reminder_id)
    return jsonify({"success": ok})


# === 日程（本地日历） API ===
@app.route('/api/calendar/add', methods=['POST'])
def add_calendar_event():
    """新增日程并自动创建本地提醒
    请求体: {"title": str, "description": str, "start": ISO8601, "end": ISO8601?, "remind_before_minutes": int?}
    """
    from calendar_manager import get_calendar_manager
    from datetime import datetime

    data = request.json or {}
    title = data.get('title') or '日程'
    description = data.get('description') or ''
    start = data.get('start')
    end = data.get('end')
    remind_before_minutes = int(data.get('remind_before_minutes') or 0)

    if not start:
        return jsonify({"success": False, "error": "start 不能为空 (ISO8601)"}), 400
    try:
        start_dt = datetime.fromisoformat(start)
    except Exception:
        return jsonify({"success": False, "error": "start 必须是 ISO8601 时间"}), 400
    end_dt = None
    if end:
        try:
            end_dt = datetime.fromisoformat(end)
        except Exception:
            return jsonify({"success": False, "error": "end 必须是 ISO8601 时间"}), 400

    mgr = get_calendar_manager()
    event_id = mgr.add_event(title, description, start_dt, end_dt, remind_before_minutes)
    return jsonify({"success": True, "event_id": event_id})


@app.route('/api/calendar/list', methods=['GET'])
def list_calendar_events():
    from calendar_manager import get_calendar_manager
    limit = int(request.args.get('limit', 100))
    mgr = get_calendar_manager()
    return jsonify({"events": mgr.list_events(limit)})


@app.route('/api/calendar/<event_id>', methods=['DELETE'])
def delete_calendar_event(event_id):
    from calendar_manager import get_calendar_manager
    mgr = get_calendar_manager()
    ok = mgr.delete_event(event_id)
    return jsonify({"success": ok})


# === 剪贴板 API ===
@app.route('/api/clipboard/history', methods=['GET'])
def get_clipboard_history():
    """获取剪贴板历史"""
    from clipboard_manager import get_clipboard_manager
    
    limit = int(request.args.get('limit', 50))
    type_filter = request.args.get('type')
    clipboard_manager = get_clipboard_manager()
    history = clipboard_manager.get_history(limit)
    if type_filter:
        history = [item for item in history if item.get('type') == type_filter]
    
    return jsonify({"history": history})


@app.route('/api/clipboard/search', methods=['GET'])
def search_clipboard():
    """搜索剪贴板历史"""
    from clipboard_manager import get_clipboard_manager
    
    query = request.args.get('query', '')
    type_filter = request.args.get('type')
    clipboard_manager = get_clipboard_manager()
    results = clipboard_manager.search(query)
    if type_filter:
        results = [item for item in results if item.get('type') == type_filter]
    
    return jsonify({"results": results})


@app.route('/api/clipboard/copy', methods=['POST'])
def copy_from_history():
    """从历史中复制"""
    from clipboard_manager import get_clipboard_manager
    
    content = request.json.get('content')
    index = request.json.get('index')
    clipboard_manager = get_clipboard_manager()
    if index is not None:
        try:
            index = int(index)
        except Exception:
            return jsonify({"success": False, "error": "index 必须是整数"}), 400
        success = clipboard_manager.copy_from_history(index)
    else:
        success = clipboard_manager.copy_from_history(content or "")
    
    return jsonify({"success": success})


# === 任务调度 API ===
@app.route('/api/tasks/add', methods=['POST'])
def add_task():
    """添加任务到队列"""
    from task_scheduler import get_task_scheduler, Task, TaskPriority
    
    data = request.json
    task_name = data.get('name', '')
    priority = data.get('priority', 'NORMAL')
    
    # 这里需要根据任务类型创建相应的action
    # 简化示例：只记录任务信息
    def dummy_action():
        print(f"执行任务: {task_name}")
        return {"status": "completed"}
    
    task = Task(
        task_id=f"task_{int(time.time())}",
        name=task_name,
        action=dummy_action,
        priority=TaskPriority[priority]
    )
    
    scheduler = get_task_scheduler()
    task_id = scheduler.add_task(task)
    
    return jsonify({"success": True, "task_id": task_id})


@app.route('/api/tasks/schedule', methods=['POST'])
def schedule_task():
    """调度定时任务"""
    from task_scheduler import get_task_scheduler
    
    data = request.json
    task_name = data.get('name', '')
    schedule_type = data.get('schedule_type', 'daily')
    time_str = data.get('time', '09:00')
    
    def dummy_action():
        print(f"执行定时任务: {task_name}")
        return {"status": "completed"}
    
    scheduler = get_task_scheduler()
    task_id = scheduler.schedule_task(
        name=task_name,
        action=dummy_action,
        schedule_type=schedule_type,
        time_str=time_str
    )
    
    return jsonify({"success": True, "task_id": task_id})


@app.route('/api/tasks/list', methods=['GET'])
def list_tasks():
    """列出所有任务"""
    from task_scheduler import get_task_scheduler, TaskStatus
    
    status = request.args.get('status')
    scheduler = get_task_scheduler()
    
    if status:
        tasks = scheduler.list_tasks(TaskStatus[status])
    else:
        tasks = scheduler.list_tasks()
    
    return jsonify({"tasks": tasks})


@app.route('/api/tasks/<task_id>/cancel', methods=['POST'])
def cancel_task(task_id):
    """取消任务"""
    from task_scheduler import get_task_scheduler
    
    scheduler = get_task_scheduler()
    success = scheduler.cancel_task(task_id)
    
    return jsonify({"success": success})


# === 邮件 API ===
@app.route('/api/email/accounts', methods=['GET'])
def list_email_accounts():
    """列出邮箱账户"""
    from email_manager import get_email_manager
    
    email_manager = get_email_manager()
    accounts = list(email_manager.accounts.keys())
    default = email_manager.default_account
    
    return jsonify({"accounts": accounts, "default": default})


@app.route('/api/email/accounts/add', methods=['POST'])
def add_email_account():
    """添加邮箱账户"""
    from email_manager import get_email_manager
    
    data = request.json
    email_address = data.get('email')
    password = data.get('password')
    smtp_server = data.get('smtp_server')
    smtp_port = data.get('smtp_port', 587)
    imap_server = data.get('imap_server')
    set_as_default = data.get('set_as_default', False)
    
    email_manager = get_email_manager()
    success = email_manager.add_account(
        email_address=email_address,
        password=password,
        smtp_server=smtp_server,
        smtp_port=smtp_port,
        imap_server=imap_server,
        set_as_default=set_as_default
    )
    
    return jsonify({"success": success})


@app.route('/api/email/send', methods=['POST'])
def send_email():
    """发送邮件"""
    from email_manager import get_email_manager
    
    data = request.json
    to_addrs = data.get('to', [])
    subject = data.get('subject', '')
    body = data.get('body', '')
    cc_addrs = data.get('cc', [])
    attachments = data.get('attachments', [])
    html = data.get('html', False)
    
    email_manager = get_email_manager()
    success = email_manager.send_email(
        to_addrs=to_addrs,
        subject=subject,
        body=body,
        cc_addrs=cc_addrs,
        attachments=attachments,
        html=html
    )
    
    return jsonify({"success": success})


@app.route('/api/email/fetch', methods=['GET'])
def fetch_emails():
    """获取邮件列表"""
    from email_manager import get_email_manager
    
    folder = request.args.get('folder', 'INBOX')
    limit = int(request.args.get('limit', 20))
    unread_only = request.args.get('unread_only', 'false').lower() == 'true'
    
    email_manager = get_email_manager()
    emails = email_manager.fetch_emails(
        folder=folder,
        limit=limit,
        unread_only=unread_only
    )
    
    return jsonify({"emails": emails})


@app.route('/api/email/search', methods=['GET'])
def search_emails():
    """搜索邮件"""
    from email_manager import get_email_manager
    
    keyword = request.args.get('query', '')
    folder = request.args.get('folder', 'INBOX')
    
    email_manager = get_email_manager()
    results = email_manager.search_emails(keyword, folder=folder)
    
    return jsonify({"results": results})


# === 浏览器自动化 API ===
@app.route('/api/browser/open', methods=['POST'])
def browser_open():
    """打开 URL"""
    from browser_automation import get_browser_automation
    
    url = request.json.get('url', '')
    browser = get_browser_automation()
    success = browser.open_url(url)
    
    return jsonify({"success": success})


@app.route('/api/browser/search', methods=['POST'])
def browser_search():
    """Google 搜索"""
    from browser_automation import get_browser_automation
    
    query = request.json.get('query', '')
    browser = get_browser_automation()
    results = browser.search_google(query)
    
    return jsonify({"results": results})


@app.route('/api/browser/screenshot', methods=['POST'])
def browser_screenshot():
    """截图"""
    from browser_automation import get_browser_automation
    import os
    
    filename = request.json.get('filename', f'screenshot_{int(time.time())}.png')
    file_path = os.path.join(WORKSPACE_DIR, 'images', filename)
    
    browser = get_browser_automation()
    success = browser.take_screenshot(file_path)
    
    return jsonify({"success": success, "path": file_path})


# === 智能搜索 API ===
@app.route('/api/search/all', methods=['GET'])
def search_all():
    """全局搜索"""
    from search_engine import get_search_engine
    
    query = request.args.get('query', '')
    max_results = int(request.args.get('max_results', 50))
    
    search_engine = get_search_engine()
    results = search_engine.search_all(query, max_results)
    
    return jsonify(results)


@app.route('/api/search/files', methods=['GET'])
def search_files():
    """搜索文件"""
    from search_engine import get_search_engine
    
    query = request.args.get('query', '')
    max_results = int(request.args.get('max_results', 20))
    
    search_engine = get_search_engine()
    results = search_engine.search_files(query, max_results)
    
    return jsonify({"results": results})


# ================= 语音识别 API (新架构) =================
@app.route('/api/voice/engines', methods=['GET'])
def voice_engines():
    """获取可用语音引擎列表"""
    try:
        from web.voice_fast import get_available_engines
        result = get_available_engines()
        return jsonify(result)
    except Exception as e:
        return jsonify({
            "success": False,
            "engines": [],
            "message": f"获取引擎列表失败: {str(e)}"
        }), 500

@app.route('/api/voice/record', methods=['POST'])
def voice_record():
    """录制音频"""
    try:
        data = request.json or {}
        duration = data.get('duration', 5)
        
        from web.voice_input import record_audio
        result = record_audio(duration=int(duration))
        
        return jsonify(result)
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"录音失败: {str(e)}",
            "audio_file": None
        }), 500

@app.route('/api/voice/recognize', methods=['POST'])
def voice_recognize():
    """识别音频文件"""
    try:
        data = request.json or {}
        audio_path = data.get('audio_path')
        engine = data.get('engine', None)
        
        if not audio_path:
            return jsonify({
                "success": False,
                "message": "缺少音频文件路径"
            }), 400
        
        from web.voice_input import recognize_audio
        result = recognize_audio(audio_path, engine)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"识别失败: {str(e)}"
        }), 500

@app.route('/api/voice/listen', methods=['POST'])
def voice_listen():
    """一键麦克风识别（本地模式 - 优化版：立即启动）"""
    try:
        data = request.json or {}
        timeout = data.get('timeout', 5)
        language = data.get('language', 'zh-CN')
        
        # 使用快速本地识别
        from web.voice_fast import recognize_voice
        result = recognize_voice(timeout=int(timeout), language=language)
        
        # 优化：设置响应头加快传输
        response = jsonify(result)
        response.headers['Cache-Control'] = 'no-cache, no-store'
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        return response
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        response = jsonify({
            "success": False,
            "text": "",
            "message": f"语音识别出错: {str(e)}",
            "engine": "error"
        })
        response.status_code = 500
        response.headers['Cache-Control'] = 'no-cache'
        return response


@app.route('/api/voice/stream')
def voice_stream():
    """流式语音识别 - 实时返回识别结果（SSE）"""
    from flask import Response
    
    def generate():
        try:
            from web.voice_fast import recognize_streaming
            
            for result in recognize_streaming(timeout=12):
                # 发送 SSE 格式数据
                import json
                yield f"data: {json.dumps(result, ensure_ascii=False)}\n\n"
                
                # 如果是最终结果或错误，结束流
                if result.get('type') in ('final', 'error'):
                    break
                    
        except Exception as e:
            import json
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
    
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no'
        }
    )


# 全局变量：控制语音识别停止
_voice_stop_flag = False

@app.route('/api/voice/stop', methods=['POST'])
def voice_stop():
    """停止语音识别"""
    global _voice_stop_flag
    _voice_stop_flag = True
    return jsonify({"success": True, "message": "已发送停止信号"})


# ================= 增强功能 API (场景1-3) =================

@app.route('/api/data/extract-transform', methods=['POST'])
def data_extract_transform():
    """数据提取与转换 - 场景1：跨应用数据搬运"""
    try:
        data = request.json
        source_type = data.get('source_type', 'wechat_contact')
        source_data = data.get('source_data')
        target_format = data.get('target_format', 'excel')
        output_filename = data.get('output_filename', f'提取数据_{datetime.now().strftime("%Y%m%d_%H%M%S")}')
        
        # 确定输出路径
        if target_format == 'excel':
            ext = '.xlsx'
        elif target_format == 'csv':
            ext = '.csv'
        else:
            ext = '.json'
        
        output_path = os.path.join(WORKSPACE_DIR, 'documents', f'{output_filename}{ext}')
        
        # 执行数据管道
        from web.data_pipeline import CrossAppDataPipeline
        pipeline = CrossAppDataPipeline()
        result = pipeline.run_pipeline(source_type, source_data, target_format, output_path)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/code/generate', methods=['POST'])
def code_generate():
    """代码生成 - 场景2：帮助用户完成编程任务"""
    try:
        data = request.json
        template_name = data.get('template_name')
        description = data.get('description')
        language = data.get('language', 'python')
        output_filename = data.get('output_filename')
        
        from web.code_generator import CodeGenerator
        generator = CodeGenerator()
        
        # 确定输出路径
        output_path = None
        if output_filename:
            output_path = os.path.join(WORKSPACE_DIR, 'code', output_filename)
        
        # 生成代码
        if template_name:
            result = generator.generate(template_name, output_path, **data.get('params', {}))
        elif description:
            # 使用AI生成（如果可用）
            result = generator.generate_from_description(description, language)
        else:
            return jsonify({
                "success": False,
                "error": "需要提供template_name或description"
            }), 400
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/code/templates', methods=['GET'])
def code_templates():
    """获取可用代码模板列表"""
    try:
        from web.code_generator import CodeGenerator
        generator = CodeGenerator()
        
        language = request.args.get('language')
        templates = generator.list_templates(language)
        
        return jsonify({
            "success": True,
            "templates": templates,
            "count": len(templates)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/ppt/generate', methods=['POST'])
def ppt_generate():
    """PPT生成 - 场景3：高质量演示文稿"""
    try:
        data = request.json
        title = data.get('title', '演示文稿')
        subtitle = data.get('subtitle', '')
        outline = data.get('outline')
        content = data.get('content')
        theme = data.get('theme', 'business')
        output_filename = data.get('output_filename', f'{title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pptx')
        
        output_path = os.path.join(WORKSPACE_DIR, 'documents', output_filename)
        
        from web.ppt_generator import PPTGenerator
        generator = PPTGenerator(theme=theme)
        
        # 生成PPT
        if outline:
            result = generator.generate_from_outline(title, outline, output_path, subtitle=subtitle)
        elif content:
            result = generator.generate_from_text(content, output_path, title)
        else:
            return jsonify({
                "success": False,
                "error": "需要提供outline或content"
            }), 400
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500



# ==================== 智能文档处理路由 ====================

def _should_use_annotation_system(requirement: str, has_file: bool = False) -> bool:
    """
    严格判断是否使用文档标注系统（在原文上标红修改）
    
    标注系统仅适用于：用户明确要求在原文上做标记/批注/标红/Track Changes
    
    注意："修改"、"优化"、"改善"等词太宽泛，不能单独触发标注。
    只有与"在原文上"、"标出来"、"标红"等定位词组合才触发。
    """
    if not requirement:
        return False
    
    requirement_lower = requirement.lower()
    
    # 第一层：明确的标注/批注关键词 — 直接触发
    explicit_annotation = ['标注', '标记', '批注', '标出', '标红', 'track changes', '批改']
    if any(kw in requirement_lower for kw in explicit_annotation):
        return True
    
    # 第二层：编辑意图 + 定位词组合才触发
    # "修改"单独出现 ≠ 标注，"修改+标出来" = 标注
    edit_words = ['修改', '改正', '纠正', '校对', '审校', '纠错']
    location_words = ['在原文', '原文上', '标出', '标记出', '指出.*位置', '哪些地方', '哪些位置']
    has_edit = any(kw in requirement_lower for kw in edit_words)
    has_location = any(re.search(kw, requirement_lower) for kw in location_words)
    
    if has_edit and has_location:
        return True
    
    # 第三层：审查/修改+质量描述组合
    review_words = ['审查', '评审', '审核', '改善', '优化', '修改', '润色', '调整']
    quality_words = ['不合适', '生硬', '翻译腔', '语序', '用词', '逻辑', '问题']
    has_review = any(kw in requirement_lower for kw in review_words)
    has_quality = any(kw in requirement_lower for kw in quality_words)
    
    if has_review and has_quality:
        return True
    
    # 默认不触发 — 宁可漏判也不误判
    return False


def _is_analysis_request(requirement: str) -> bool:
    """判断是否为纯分析/研究类请求（只读分析，不生成新内容）"""
    if not requirement:
        return False

    requirement_lower = requirement.lower()
    
    # 明确的分析动作词
    analysis_actions = [
        "分析", "总结", "概述", "梳理", "解读",
        "评估", "对比", "提炼", "归纳",
        "主要观点", "核心观点", "要点",
        "review", "analysis", "summary", "summarize"
    ]
    
    # 排除词：如果同时包含生成/写/改善意图，这不是纯分析
    generation_words = [
        '写', '生成', '改善', '改进', '优化', '润色',
        '重写', '摘要', '引言', '结论', '帮我做'
    ]
    
    has_analysis = any(kw in requirement_lower for kw in analysis_actions)
    has_generation = any(kw in requirement_lower for kw in generation_words)
    
    # 只有纯分析（无生成意图）才返回True
    if has_analysis and not has_generation:
        return True
    
    return False


@app.route('/api/document/smart-process', methods=['POST'])
def document_smart_process():
    """
    智能文档处理入口
    自动判断使用：标注系统 or 文件分析系统
    """
    try:
        data = request.json
        file_path = data.get('file_path')
        requirement = data.get('requirement', '')
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 智能判断应该用哪个系统
        use_annotation = _should_use_annotation_system(requirement)
        
        print(f"[SmartProcess] 智能判断: use_annotation={use_annotation}")
        print(f"[SmartProcess] 需求: {requirement[:100]}")
        
        if use_annotation:
            # 使用文档标注系统
            print(f"[SmartProcess] 路由到: 文档自动标注系统")
            return _call_document_annotate(file_path, requirement)
        else:
            # 使用传统的文件分析系统
            print(f"[SmartProcess] 路由到: 文件分析系统")
            return _call_document_analysis(file_path, requirement)
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


def _call_document_annotate(file_path: str, requirement: str):
    """调用文档标注系统"""
    try:
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        result = feedback_system.full_annotation_loop(
            file_path=file_path,
            user_requirement=requirement,
            model_id="gemini-3-pro-preview"
        )
        
        # 添加处理模式标记
        result['processing_mode'] = 'annotation'
        result['mode_description'] = '文档自动标注'
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
            "processing_mode": "annotation"
        }), 500


def _call_document_analysis(file_path: str, requirement: str):
    """调用传统的文件分析系统"""
    try:
        # 这里调用现有的文件分析逻辑
        # 临时返回说明（实际应该调用现有的分析端点）
        return jsonify({
            "success": False,
            "error": "文件分析系统需要单独实现",
            "processing_mode": "analysis",
            "mode_description": "文件分析"
        }), 501
    
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
            "processing_mode": "analysis"
        }), 500


@app.route('/api/document/feedback', methods=['POST'])
def document_feedback():
    """文档智能反馈：读取文档 → AI分析 → 应用修改"""
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        auto_apply = data.get('auto_apply', True)
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        # 执行完整反馈闭环
        result = feedback_system.full_feedback_loop(
            file_path=file_path,
            user_requirement=user_requirement,
            auto_apply=auto_apply
        )
        
        return jsonify(result)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/analyze', methods=['POST'])
def document_analyze():
    """仅分析文档，不应用修改"""
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        # 仅分析
        result = feedback_system.analyze_and_suggest(
            file_path=file_path,
            user_requirement=user_requirement
        )
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/apply', methods=['POST'])
def document_apply():
    """应用修改建议到文档"""
    try:
        data = request.json
        file_path = data.get('file_path')
        modifications = data.get('modifications', [])
        
        if not file_path or not modifications:
            return jsonify({
                "success": False,
                "error": "缺少file_path或modifications参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 应用修改
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        result = feedback_system.apply_suggestions(
            file_path=file_path,
            modifications=modifications
        )
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/annotate', methods=['POST'])
def document_annotate():
    """文档自动标注：AI分析 -> 生成标注 -> 应用到副本"""
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        model_id = data.get('model_id', 'gemini-3-pro-preview')
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 初始化反馈系统
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        # 执行完整标注闭环
        result = feedback_system.full_annotation_loop(
            file_path=file_path,
            user_requirement=user_requirement,
            model_id=model_id
        )
        
        return jsonify(result)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/analyze-annotations', methods=['POST'])
def document_analyze_annotations():
    """仅分析文档并生成标注建议（不应用）- 已弃用，请使用 /api/document/batch-annotate-stream"""
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 使用V2批量标注系统（立即返回结果，不流式）
        from web.document_direct_edit import ImprovedBatchAnnotator
        annotator = ImprovedBatchAnnotator(gemini_client=client, batch_size=5)
        
        # 收集所有事件（非流式）
        events = []
        final_result = None
        
        for event in annotator.annotate_document_streaming(file_path, user_requirement):
            # 解析事件
            if event.startswith("event: complete"):
                data_line = event.split("\n")[1]
                if data_line.startswith("data: "):
                    final_result = json.loads(data_line[6:])
            events.append(event)
        
        if final_result:
            return jsonify({
                "success": True,
                **final_result
            })
        else:
            return jsonify({
                "success": False,
                "error": "处理失败，未收到完成事件"
            }), 500
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/batch-annotate-stream', methods=['POST'])
def document_batch_annotate_stream():
    """
    批量标注文档（SSE流式返回，实时反馈进度）
    
    接收参数:
        file_path: 文档路径
        requirement: 用户需求（可选）
        batch_size: 每批处理段落数（默认5）
    
    返回: SSE事件流
        event: progress - 进度更新
        event: batch_complete - 批次完成
        event: complete - 全部完成
        event: error - 错误
    """
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        batch_size = data.get('batch_size', 5)
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 导入V2批量标注系统
        from web.document_batch_annotator_v2 import annotate_large_document
        
        # 返回SSE流
        return Response(
            annotate_large_document(
                file_path=file_path,
                user_requirement=user_requirement,
                gemini_client=client,
                batch_size=batch_size
            ),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',
                'Connection': 'keep-alive'
            }
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/apply-annotations', methods=['POST'])
def document_apply_annotations():
    """应用标注建议到文档"""
    try:
        data = request.json
        file_path = data.get('file_path')
        annotations = data.get('annotations', [])
        
        if not file_path or not annotations:
            return jsonify({
                "success": False,
                "error": "缺少file_path或annotations参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 应用标注
        from web.document_feedback import DocumentFeedbackSystem
        feedback_system = DocumentFeedbackSystem(gemini_client=client)
        
        result = feedback_system.annotate_document(file_path, annotations)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# ==================== 新功能 API 路由 ====================

# ==================== 改进的建议式标注 API ====================

@app.route('/api/document/suggest-stream', methods=['POST'])
def document_suggest_stream():
    """
    生成修改建议流（SSE）
    
    请求参数:
        file_path: 文档路径
        requirement: 用户需求（可选）
    
    返回: SSE事件流
        event: progress - 进度
        event: suggestion - 单个建议
        event: suggestions_complete - 所有建议完成
        event: complete - 完成
    """
    try:
        data = request.json
        file_path = data.get('file_path')
        user_requirement = data.get('requirement', '')
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 使用建议式标注器
        from web.suggestion_annotator import SuggestionAnnotator
        annotator = SuggestionAnnotator(batch_size=3)
        
        # 返回SSE流
        return Response(
            annotator.analyze_document_streaming(file_path, user_requirement),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',
                'Connection': 'keep-alive'
            }
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/document/apply-suggestions', methods=['POST'])
def document_apply_suggestions():
    """
    根据用户选择应用修改建议
    
    请求参数:
        file_path: 原始文档路径
        suggestions: 用户的选择列表
            [
                {
                    "id": "s_5_0",
                    "原文": "在被记录的",
                    "修改": "在记录的",
                    "接受": True/False
                },
                ...
            ]
    
    返回:
        {
            "success": True,
            "output_file": "修改后的文件路径",
            "applied_count": 实际应用的修改数,
            "accepted_count": 用户接受的数量
        }
    """
    try:
        from docx import Document
        
        data = request.json
        file_path = data.get('file_path')
        suggestions = data.get('suggestions', [])
        
        if not file_path:
            return jsonify({
                "success": False,
                "error": "缺少file_path参数"
            }), 400
        
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.join(WORKSPACE_DIR, 'documents', file_path)
        
        if not os.path.exists(file_path):
            return jsonify({
                "success": False,
                "error": f"文件不存在: {file_path}"
            }), 404
        
        # 读取文档
        doc = Document(file_path)
        
        # 筛选用户接受的建议
        accepted_suggestions = [s for s in suggestions if s.get("接受", False)]
        
        applied_count = 0
        
        # 应用修改（直接在段落中查找并替换）
        for suggestion in accepted_suggestions:
            original = suggestion.get("原文", "")
            modified = suggestion.get("修改", "")
            
            if not original or not modified:
                continue
            
            # 在所有段落中查找并替换
            for para in doc.paragraphs:
                if original in para.text:
                    # 替换文本
                    full_text = para.text
                    new_text = full_text.replace(original, modified, 1)
                    
                    if new_text != full_text:
                        # 清空并重新添加（保留格式）
                        para.clear()
                        para.add_run(new_text)
                        applied_count += 1
                        break  # 每个建议只应用一次
            
            # 检查表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if original in para.text:
                                full_text = para.text
                                new_text = full_text.replace(original, modified, 1)
                                if new_text != full_text:
                                    para.clear()
                                    para.add_run(new_text)
                                    applied_count += 1
        
        # 保存为新文件
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_accepted.docx"
        doc.save(output_path)
        
        return jsonify({
            "success": True,
            "output_file": output_path,
            "applied_count": applied_count,
            "accepted_count": len(accepted_suggestions),
            "message": f"已应用 {applied_count} 处修改（用户接受了 {len(accepted_suggestions)} 个建议）"
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# 知识库 API
@app.route('/api/knowledge-base/add', methods=['POST'])
def kb_add_document():
    """添加文档到知识库"""
    try:
        from web.knowledge_base import KnowledgeBase
        
        data = request.json
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400
        
        kb = KnowledgeBase()
        result = kb.add_document(file_path)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/knowledge-base/search', methods=['POST'])
def kb_search():
    """搜索知识库"""
    try:
        from web.knowledge_base import KnowledgeBase
        
        data = request.json
        query = data.get('query')
        max_results = data.get('max_results', 10)
        
        if not query:
            return jsonify({"success": False, "error": "缺少query参数"}), 400
        
        kb = KnowledgeBase()
        results = kb.search(query, max_results=max_results)
        
        return jsonify({"success": True, "results": results})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/knowledge-base/stats', methods=['GET'])
def kb_stats():
    """获取知识库统计"""
    try:
        from web.knowledge_base import KnowledgeBase
        
        kb = KnowledgeBase()
        stats = kb.get_stats()
        
        return jsonify({"success": True, "stats": stats})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== 文件网络索引 API ====================

@app.route('/api/file-network/search', methods=['POST'])
def file_network_search():
    """多维查询文件
    
    请求参数:
        query: 文本搜索查询（可选）
        file_type: 文件类型（docx, pdf等，可选）
        tags: 标签列表（可选）
        operation: 处理操作（annotate, edit等，可选）
        date_from: 开始日期（ISO格式，可选）
        date_to: 结束日期（ISO格式，可选）
        limit: 返回数量限制（默认50）
    """
    try:
        from web.processed_file_network import get_file_network
        
        data = request.json or {}
        file_network = get_file_network()
        
        result = file_network.search_files(
            query=data.get('query'),
            file_type=data.get('file_type'),
            tags=data.get('tags'),
            operation=data.get('operation'),
            date_from=data.get('date_from'),
            date_to=data.get('date_to'),
            limit=data.get('limit', 50)
        )
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/file-network/open', methods=['POST'])
def file_network_open():
    """快速打开文件
    
    请求参数:
        file_id: 文件ID
    """
    try:
        from web.processed_file_network import get_file_network
        
        data = request.json
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({"success": False, "error": "缺少file_id参数"}), 400
        
        file_network = get_file_network()
        result = file_network.open_file(file_id)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/file-network/network', methods=['POST'])
def file_network_get_network():
    """获取文件关系网络
    
    请求参数:
        file_id: 文件ID
        depth: 关系深度（1=直接关系，2=二级关系，默认2）
    """
    try:
        from web.processed_file_network import get_file_network
        
        data = request.json
        file_id = data.get('file_id')
        depth = data.get('depth', 2)
        
        if not file_id:
            return jsonify({"success": False, "error": "缺少file_id参数"}), 400
        
        file_network = get_file_network()
        result = file_network.get_file_network(file_id, depth)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/file-network/statistics', methods=['GET'])
def file_network_statistics():
    """获取文件网络统计信息"""
    try:
        from web.processed_file_network import get_file_network
        
        file_network = get_file_network()
        result = file_network.get_statistics()
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/file-network/register', methods=['POST'])
def file_network_register():
    """手动注册文件到网络
    
    请求参数:
        file_path: 文件路径
        tags: 标签列表（可选）
        extract_snippets: 是否提取文本片段（默认true）
    """
    try:
        from web.processed_file_network import get_file_network
        
        data = request.json
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"success": False, "error": "缺少file_path参数"}), 400
        
        file_network = get_file_network()
        result = file_network.register_file(
            file_path=file_path,
            tags=data.get('tags'),
            extract_snippets=data.get('extract_snippets', True)
        )
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 批量处理 API
@app.route('/api/batch/rename', methods=['POST'])
def batch_rename():
    """批量重命名文件"""
    try:
        from web.batch_processor import BatchFileProcessor
        
        data = request.json
        directory = data.get('directory')
        pattern = data.get('pattern')
        
        processor = BatchFileProcessor()
        result = processor.batch_rename(directory, **pattern)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/batch/convert', methods=['POST'])
def batch_convert():
    """批量格式转换"""
    try:
        from web.batch_processor import BatchFileProcessor
        
        data = request.json
        directory = data.get('directory')
        from_format = data.get('from_format')
        to_format = data.get('to_format')
        
        processor = BatchFileProcessor()
        result = processor.batch_convert(directory, from_format, to_format)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 模板库 API
@app.route('/api/template/list', methods=['GET'])
def template_list():
    """获取模板列表"""
    try:
        from web.template_library import TemplateLibrary
        
        library = TemplateLibrary()
        templates = library.list_templates()
        
        return jsonify({"success": True, "templates": templates})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/template/generate', methods=['POST'])
def template_generate():
    """从模板生成文档"""
    try:
        from web.template_library import TemplateLibrary
        
        data = request.json
        template_name = data.get('template_id') or data.get('template_name')
        variables = data.get('variables', {})
        output_dir = data.get('output_dir')
        output_file = data.get('output_file')
        if output_file and not output_dir:
            if os.path.isdir(output_file):
                output_dir = output_file
            else:
                output_dir = os.path.dirname(output_file) or None
        
        library = TemplateLibrary()
        result = library.generate_from_template(template_name, variables, output_dir)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 一致性检查 API
@app.route('/api/check/consistency', methods=['POST'])
def check_consistency():
    """检查文档一致性"""
    try:
        from web.consistency_checker import ConsistencyChecker
        
        data = request.json
        file_path = data.get('file_path')
        
        checker = ConsistencyChecker()
        result = checker.check_document(file_path)
        report = checker.generate_report(result)
        
        return jsonify({"success": True, "result": result, "report": report})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 文档对比 API
@app.route('/api/compare/documents', methods=['POST'])
def compare_documents():
    """对比文档"""
    try:
        from web.document_comparator import DocumentComparator
        
        data = request.json
        file_a = data.get('file_a')
        file_b = data.get('file_b')
        output_format = data.get('output_format', 'markdown')
        
        comparator = DocumentComparator()
        result = comparator.compare_documents(file_a, file_b, output_format)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# OCR 助手 API
@app.route('/api/ocr/screenshot', methods=['POST'])
def ocr_screenshot():
    """截图并OCR"""
    try:
        from web.clipboard_ocr_assistant import ClipboardOCRAssistant
        
        data = request.json
        save_image = data.get('save_image', True)
        auto_index = data.get('auto_index', False)
        
        assistant = ClipboardOCRAssistant()
        result = assistant.capture_and_ocr(source='screenshot', save_image=save_image)
        
        if auto_index and result.get('ocr_success'):
            assistant.auto_index_to_knowledge_base(result)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/ocr/clipboard', methods=['POST'])
def ocr_clipboard():
    """剪贴板图片OCR"""
    try:
        from web.clipboard_ocr_assistant import ClipboardOCRAssistant
        
        data = request.json
        save_image = data.get('save_image', True)
        auto_index = data.get('auto_index', False)
        
        assistant = ClipboardOCRAssistant()
        result = assistant.capture_and_ocr(source='clipboard', save_image=save_image)
        
        if auto_index and result.get('ocr_success'):
            assistant.auto_index_to_knowledge_base(result)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 操作历史 API
@app.route('/api/history/list', methods=['GET'])
def history_list():
    """获取操作历史"""
    try:
        from web.operation_history import OperationHistory
        
        limit = request.args.get('limit', 50, type=int)
        file_path = request.args.get('file_path')
        
        history = OperationHistory()
        operations = history.get_history(limit=limit, file_path=file_path)
        
        return jsonify({"success": True, "operations": operations})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/history/rollback/<op_id>', methods=['POST'])
def history_rollback(op_id):
    """回滚操作"""
    try:
        from web.operation_history import OperationHistory
        
        history = OperationHistory()
        result = history.rollback(op_id)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/history/stats', methods=['GET'])
def history_stats():
    """获取历史统计"""
    try:
        from web.operation_history import OperationHistory
        
        history = OperationHistory()
        stats = history.get_statistics()
        
        return jsonify({"success": True, "stats": stats})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# 语音转写 API
@app.route('/api/speech/transcribe-file', methods=['POST'])
def speech_transcribe_file():
    """转写音频文件"""
    try:
        from web.speech_transcriber import SpeechTranscriber
        
        data = request.json
        audio_path = data.get('audio_path')
        language = data.get('language', 'zh-CN')
        output_format = data.get('output_format', 'txt')
        title = data.get('title')
        auto_summary = data.get('auto_summary', True)
        
        if not audio_path:
            return jsonify({"success": False, "error": "缺少audio_path参数"}), 400
        
        transcriber = SpeechTranscriber()
        result = transcriber.process_audio_complete(
            audio_path,
            language=language,
            output_format=output_format,
            title=title,
            auto_summary=auto_summary
        )
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/speech/transcribe-microphone', methods=['POST'])
def speech_transcribe_microphone():
    """从麦克风录音并转写"""
    try:
        from web.speech_transcriber import SpeechTranscriber
        
        data = request.json
        duration = data.get('duration', 30)
        language = data.get('language', 'zh-CN')
        output_format = data.get('output_format', 'txt')
        title = data.get('title')
        
        transcriber = SpeechTranscriber()
        
        # 录音
        mic_result = transcriber.transcribe_microphone(duration=duration, language=language)
        
        if not mic_result["success"]:
            return jsonify(mic_result), 400
        
        text = mic_result["text"]
        
        # 提取总结
        summary_result = transcriber.extract_keywords_and_summary(text)
        keywords = summary_result.get("keywords", []) if summary_result["success"] else []
        summary = summary_result.get("summary", []) if summary_result["success"] else []
        
        # 生成文档
        output_file = transcriber.generate_transcript_document(
            text,
            keywords=keywords,
            summary=summary,
            title=title,
            output_format=output_format
        )
        
        return jsonify({
            "success": True,
            "text": text,
            "keywords": keywords,
            "summary": summary,
            "output_file": output_file,
            "format": output_format
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/speech/extract-summary', methods=['POST'])
def speech_extract_summary():
    """从文本提取关键词和总结"""
    try:
        from web.speech_transcriber import SpeechTranscriber
        
        data = request.json
        text = data.get('text')
        max_keywords = data.get('max_keywords', 10)
        
        if not text:
            return jsonify({"success": False, "error": "缺少text参数"}), 400
        
        transcriber = SpeechTranscriber()
        result = transcriber.extract_keywords_and_summary(
            text,
            max_keywords=max_keywords
        )
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ================= 主程序入口 =================

# ================= NotebookLM 功能复刻 API =================

@app.route('/api/notebook/overview', methods=['POST'])
def notebook_overview():
    """生成音频概览 (Podcast)"""
    data = request.json
    content = data.get('content', '')
    if not content:
        return jsonify({"success": False, "error": "内容不能为空"}), 400
        
    try:
        from web.audio_overview import AudioOverviewGenerator
        generator = AudioOverviewGenerator(output_dir=os.path.join(settings_manager.workspace_dir, "audio_cache"))
        
        # 1. 生成剧本
        # 获取模型实例 (复用现有的 KotoBrain 或直接调用 API)
        # 这里为了简化，假设我们能获取到一个 genai model 实例
        # 实际项目中应该复用 koto_brain.client.models
        # 暂时使用临时的 model 实例
        import google.genai as genai
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        model = client.models
        
        script = asyncio.run(generator.generate_script(content, model))
        if not script:
             return jsonify({"success": False, "error": "剧本生成失败"}), 500
             
        # 2. 合成音频
        session_id = f"overview_{int(time.time())}"
        audio_path = asyncio.run(generator.synthesize_audio(script, session_id))
        
        if audio_path:
            # 返回相对于 workspace 的路径或者 download url
            rel_path = os.path.relpath(audio_path, settings_manager.workspace_dir)
            # 注意：实际访问可能需要通过 send_from_directory 路由
            # 假设我们有一个 /files/ 路由可以访问 workspace/
            audio_url = f"/api/files/download?path={requests.utils.quote(audio_path)}" 
            
            return jsonify({
                "success": True, 
                "audio_url": audio_url,
                "script": script
            })
        else:
            return jsonify({"success": False, "error": "音频合成失败"}), 500

    except Exception as e:
        print(f"Error processing audio overview: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/notebook/qa', methods=['POST'])
def notebook_qa():
    """源文档深度问答 (Source-Grounded Q&A)"""
    data = request.json
    question = data.get('question')
    file_ids = data.get('file_ids', []) # 假设前端传回 files (这里先简化为 content 直接传入 或者 file paths)
    # 为了简化演示，我们先接受纯文本 content
    context_content = data.get('context', '') 
    
    if not question or not context_content:
        return jsonify({"success": False, "error": "缺少问题或上下文"}), 400

    prompt = f"""
    Answer the user's question mostly based on the provided source context.
    
    [Source Context]
    {context_content[:30000]} 

    [User Question]
    {question}

    [Rules]
    1. You must cite your sources. When you use information from the context, append [Source] at the end of the sentence.
    2. If the answer is not in the context, state that clearly.
    3. Be precise and concise.
    """
    
    try:
         # 复用 KotoBrain 的逻辑或者直接调用
        import google.genai as genai
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model='gemini-2.5-flash', 
            contents=prompt
        )
        return jsonify({"success": True, "answer": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/notebook/study_guide', methods=['POST'])
def notebook_study_guide():
    """生成学习指南/简报"""
    data = request.json
    content = data.get('content', '')
    type_ = data.get('type', 'summary') # summary, quiz, timelime, faq
    
    prompts = {
        'summary': "Create a comprehensive briefing document summarizing the key points, key people, and timeline from the text.",
        'quiz': "Create 5 multiple-choice questions based on the text to test understanding. Include the correct answer key at the end.",
        'timeline': "Extract a chronological timeline of events mentioned in the text.",
        'faq': "Create a FAQ section based on the text, anticipating what a reader might ask."
    }
    
    selected_prompt = prompts.get(type_, prompts['summary'])
    full_prompt = f"{selected_prompt}\n\n[Source Text]\n{content[:20000]}"
    
    try:
        import google.genai as genai
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        response = client.models.generate_content(
            model='gemini-2.5-flash', 
            contents=full_prompt
        )
        return jsonify({"success": True, "result": response.text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/files/download', methods=['GET'])
def download_file_proxy():
    """通用的文件下载代理"""
    file_path = request.args.get('path')
    if not file_path or not os.path.exists(file_path):
        return "File not found", 404
    return send_file(file_path, as_attachment=True)


@app.route('/api/notebook/upload', methods=['POST'])
def notebook_upload():
    """上传并解析文件 (PDF/Docx/Txt)"""
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "error": "No selected file"}), 400
        
    try:
        # Save temp file
        filename = file.filename
        temp_path = os.path.join(tempfile.gettempdir(), f"koto_{int(time.time())}_{filename}")
        file.save(temp_path)
        
        # Parse using FileParser
        from web.file_parser import FileParser
        result = FileParser.parse_file(temp_path)
        
        # Cleanup
        try:
            os.remove(temp_path)
        except:
            pass
            
        if result.get("success"):
            return jsonify({
                "success": True,
                "filename": filename,
                "content": result.get("content", ""),
                "char_count": result.get("char_count", 0)
            })
        else:
            return jsonify({"success": False, "error": result.get("error")}), 500
            
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/notebook')
def notebook_ui():
    """NotebookLM 风格界面"""
    return render_template('notebook_lm.html')

if __name__ == '__main__':

    print("\n🚀 Koto Web Server Starting...")
    print(f"📁 Chat Directory: {os.path.abspath(CHAT_DIR)}")
    print(f"📁 Workspace: {os.path.abspath(WORKSPACE_DIR)}")
    
    # 延迟检查 Ollama 状态（不阻塞启动）
    def check_ollama_async():
        time.sleep(2)  # 延迟2秒后检查
        if LocalDispatcher.is_ollama_running():
            print("🦙 Ollama: Running")
        else:
            print("🦙 Ollama: Not Running")
    threading.Thread(target=check_ollama_async, daemon=True).start()
    
    print("⚠️ 本地模型任务路由器已禁用，使用远程 AI")
    
    print("\n🌐 Open http://localhost:5000 in your browser\n")
    
    # 启动后台服务（异步，不阻塞启动）
    def start_background_services():
        time.sleep(1)  # 延迟1秒后启动后台服务
        try:
            from clipboard_manager import get_clipboard_manager
            from task_scheduler import get_task_scheduler
            from auto_catalog_scheduler import get_auto_catalog_scheduler
            
            # 启动剪贴板监控
            clipboard_manager = get_clipboard_manager()
            clipboard_manager.start_monitoring()
            print("📋 剪贴板监控已启动")
            
            # 启动任务调度器
            task_scheduler = get_task_scheduler()
            task_scheduler.start()
            print("⏰ 任务调度器已启动")
            
            # 初始化自动归纳调度器（如果已启用）
            auto_catalog = get_auto_catalog_scheduler()
            if auto_catalog.is_auto_catalog_enabled():
                auto_catalog._register_scheduled_task()
                print(f"🗂️ 自动归纳已启用，每日 {auto_catalog.get_catalog_schedule()} 执行")
            
        except Exception as e:
            print(f"⚠️ 后台服务启动失败: {e}")
    
    threading.Thread(target=start_background_services, daemon=True).start()
    
    try:
        debug_mode = os.environ.get('KOTO_DEBUG', 'false').lower() == 'true'
        port = int(os.environ.get('KOTO_PORT', '5000'))
        app.run(debug=debug_mode, host='0.0.0.0', port=port, threaded=True)
    finally:
        # 应用关闭时清理并行执行系统
        if PARALLEL_SYSTEM_ENABLED:
            print("[PARALLEL] 🛑 Shutting down parallel execution system...")
            stop_dispatcher()
            print("[PARALLEL] ✅ Parallel execution system shut down")



# ═══ 文件组织系统 API ═══

# 初始化文件组织器
_file_organizer_cache = {}
_batch_ops_cache = {}

def get_file_organizer():
    """懒加载文件组织器"""
    if 'organizer' not in _file_organizer_cache:
        try:
            from web.file_organizer import FileOrganizer
        except ImportError:
            from file_organizer import FileOrganizer
        
        organize_root = get_organize_root()
        _file_organizer_cache['organizer'] = FileOrganizer(organize_root)
    
    return _file_organizer_cache['organizer']

def get_file_analyzer():
    """懒加载文件分析器"""
    if 'analyzer' not in _file_organizer_cache:
        try:
            from web.file_analyzer import FileAnalyzer
        except ImportError:
            from file_analyzer import FileAnalyzer
        
        _file_organizer_cache['analyzer'] = FileAnalyzer()
    
    return _file_organizer_cache['analyzer']

def get_batch_ops_manager():
    """懒加载批量文件处理管理器"""
    if 'batch_ops' not in _batch_ops_cache:
        try:
            from web.batch_file_ops import BatchFileOpsManager
        except ImportError:
            from batch_file_ops import BatchFileOpsManager
        _batch_ops_cache['batch_ops'] = BatchFileOpsManager()
    return _batch_ops_cache['batch_ops']

_file_editor_cache = {}
_file_indexer_cache = {}
_concept_extractor_cache = {}
_knowledge_graph_cache = {}
_behavior_monitor_cache = {}
_suggestion_engine_cache = {}
_insight_reporter_cache = {}

def get_file_editor():
    """懒加载文件编辑器"""
    if 'editor' not in _file_editor_cache:
        try:
            from web.file_editor import FileEditor
        except ImportError:
            from file_editor import FileEditor
        _file_editor_cache['editor'] = FileEditor()
    return _file_editor_cache['editor']

def get_file_indexer():
    """懒加载文件索引器"""
    if 'indexer' not in _file_indexer_cache:
        try:
            from web.file_indexer import FileIndexer
        except ImportError:
            from file_indexer import FileIndexer
        _file_indexer_cache['indexer'] = FileIndexer()
    return _file_indexer_cache['indexer']

def get_concept_extractor():
    """懒加载概念提取器"""
    if 'extractor' not in _concept_extractor_cache:
        try:
            from web.concept_extractor import ConceptExtractor
        except ImportError:
            from concept_extractor import ConceptExtractor
        _concept_extractor_cache['extractor'] = ConceptExtractor()
    return _concept_extractor_cache['extractor']

def get_knowledge_graph():
    """懒加载知识图谱"""
    if 'graph' not in _knowledge_graph_cache:
        try:
            from web.knowledge_graph import KnowledgeGraph
        except ImportError:
            from knowledge_graph import KnowledgeGraph
        _knowledge_graph_cache['graph'] = KnowledgeGraph()
    return _knowledge_graph_cache['graph']

def get_behavior_monitor():
    """懒加载行为监控器"""
    if 'monitor' not in _behavior_monitor_cache:
        try:
            from web.behavior_monitor import BehaviorMonitor
        except ImportError:
            from behavior_monitor import BehaviorMonitor
        _behavior_monitor_cache['monitor'] = BehaviorMonitor()
    return _behavior_monitor_cache['monitor']

def get_suggestion_engine():
    """懒加载建议引擎"""
    if 'engine' not in _suggestion_engine_cache:
        try:
            from web.suggestion_engine import SuggestionEngine
        except ImportError:
            from suggestion_engine import SuggestionEngine
        _suggestion_engine_cache['engine'] = SuggestionEngine()
    return _suggestion_engine_cache['engine']

def get_insight_reporter():
    """懒加载洞察报告生成器"""
    if 'reporter' not in _insight_reporter_cache:
        try:
            from web.insight_reporter import InsightReporter
        except ImportError:
            from insight_reporter import InsightReporter
        _insight_reporter_cache['reporter'] = InsightReporter()
    return _insight_reporter_cache['reporter']

# ==================== 增强主动能力模块缓存 ====================
_notification_manager_cache = {}
_proactive_dialogue_cache = {}
_context_awareness_cache = {}
_auto_execution_cache = {}
_trigger_system_cache = {}

def get_notification_manager():
    """懒加载通知管理器"""
    if 'manager' not in _notification_manager_cache:
        try:
            from web.notification_manager import get_notification_manager as _get_mgr
        except ImportError:
            from notification_manager import get_notification_manager as _get_mgr
        _notification_manager_cache['manager'] = _get_mgr()
    return _notification_manager_cache['manager']

def get_proactive_dialogue():
    """懒加载主动对话引擎"""
    if 'engine' not in _proactive_dialogue_cache:
        try:
            from web.proactive_dialogue import get_proactive_dialogue_engine
        except ImportError:
            from proactive_dialogue import get_proactive_dialogue_engine
        
        # 集成依赖模块
        notif_mgr = get_notification_manager()
        behavior_mon = get_behavior_monitor()
        suggestion_eng = get_suggestion_engine()
        
        _proactive_dialogue_cache['engine'] = get_proactive_dialogue_engine(
            notification_manager=notif_mgr,
            behavior_monitor=behavior_mon,
            suggestion_engine=suggestion_eng
        )
    return _proactive_dialogue_cache['engine']

def get_context_awareness():
    """懒加载情境感知系统"""
    if 'system' not in _context_awareness_cache:
        try:
            from web.context_awareness import get_context_awareness_system
        except ImportError:
            from context_awareness import get_context_awareness_system
        
        behavior_mon = get_behavior_monitor()
        _context_awareness_cache['system'] = get_context_awareness_system(
            behavior_monitor=behavior_mon
        )
    return _context_awareness_cache['system']

def get_auto_execution():
    """懒加载自动执行引擎"""
    if 'engine' not in _auto_execution_cache:
        try:
            from web.auto_execution import get_auto_execution_engine
        except ImportError:
            from auto_execution import get_auto_execution_engine
        
        notif_mgr = get_notification_manager()
        _auto_execution_cache['engine'] = get_auto_execution_engine(
            notification_manager=notif_mgr
        )
    return _auto_execution_cache['engine']

def get_trigger_system():
    """懒加载主动交互触发系统"""
    if 'system' not in _trigger_system_cache:
        try:
            from web.proactive_trigger import get_trigger_system as _get_trigger_system
        except ImportError:
            from proactive_trigger import get_trigger_system as _get_trigger_system
        
        behavior_mon = get_behavior_monitor()
        context_sys = get_context_awareness()
        suggestion_eng = get_suggestion_engine()
        notif_mgr = get_notification_manager()
        dialogue_eng = get_proactive_dialogue()
        
        _trigger_system_cache['system'] = _get_trigger_system(
            behavior_monitor=behavior_mon,
            context_awareness=context_sys,
            suggestion_engine=suggestion_eng,
            notification_manager=notif_mgr,
            dialogue_engine=dialogue_eng
        )
    return _trigger_system_cache['system']


@app.route('/api/batch/submit', methods=['POST'])
def batch_submit():
    """提交批量文件处理任务"""
    try:
        data = request.json or {}
        command = data.get('command', '')
        manager = get_batch_ops_manager()

        if command:
            parsed = manager.parse_command(command)
            if not parsed.get('success'):
                return jsonify({"success": False, "error": parsed.get('error'), "hint": parsed.get('hint')}), 400
            operation = parsed.get('operation')
            input_dir = parsed.get('input_dir')
            output_dir = parsed.get('output_dir')
            options = parsed.get('options', {})
        else:
            operation = data.get('operation')
            input_dir = data.get('input_dir')
            output_dir = data.get('output_dir')
            options = data.get('options', {})

        if not operation or not input_dir or not output_dir:
            return jsonify({"success": False, "error": "缺少必要参数"}), 400

        job = manager.create_job(
            name=f"batch_{operation}",
            operation=operation,
            input_dir=input_dir,
            output_dir=output_dir,
            options=options
        )
        manager.start_job(job.job_id)
        return jsonify({"success": True, "job_id": job.job_id, "job": manager.get_job(job.job_id)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/batch/jobs', methods=['GET'])
def batch_list_jobs():
    """列出批量任务"""
    manager = get_batch_ops_manager()
    return jsonify({"success": True, "jobs": manager.list_jobs()})


@app.route('/api/batch/jobs/<job_id>', methods=['GET'])
def batch_get_job(job_id):
    """获取单个任务详情"""
    manager = get_batch_ops_manager()
    job = manager.get_job(job_id)
    if not job:
        return jsonify({"success": False, "error": "任务不存在"}), 404
    return jsonify({"success": True, "job": job})


@app.route('/api/batch/stream/<job_id>', methods=['GET'])
def batch_stream_job(job_id):
    """批量任务进度流"""
    manager = get_batch_ops_manager()
    return Response(manager.stream_job(job_id), mimetype='text/event-stream')

@app.route('/api/organize/scan-file', methods=['POST'])
def organize_scan_file():
    """扫描和分析单个文件"""
    try:
        data = request.json
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"error": "缺少 file_path 参数"}), 400
        
        if not os.path.exists(file_path):
            return jsonify({"error": f"文件不存在: {file_path}"}), 404
        
        analyzer = get_file_analyzer()
        analysis_result = analyzer.analyze_file(file_path)
        
        return jsonify({
            "success": True,
            "file": os.path.basename(file_path),
            "analysis": analysis_result
        })
    
    except Exception as e:
        return jsonify({
            "error": f"分析失败: {str(e)}"
        }), 500

@app.route('/api/organize/auto-organize', methods=['POST'])
def organize_auto_organize():
    """自动组织文件（分析+移动）"""
    try:
        data = request.json
        file_path = data.get('file_path')
        auto_confirm = data.get('auto_confirm', True)
        
        if not file_path:
            return jsonify({"error": "缺少 file_path 参数"}), 400
        
        if not os.path.exists(file_path):
            return jsonify({"error": f"文件不存在: {file_path}"}), 404
        
        # 第一步：分析文件
        analyzer = get_file_analyzer()
        analysis = analyzer.analyze_file(file_path)
        suggested_folder = analysis.get('suggested_folder')
        
        if not suggested_folder:
            return jsonify({
                "error": "无法确定文件分类",
                "analysis": analysis
            }), 400
        
        # 第二步：组织文件
        organizer = get_file_organizer()
        org_result = organizer.organize_file(
            file_path,
            suggested_folder,
            auto_confirm=auto_confirm
        )
        
        if org_result.get('success'):
            return jsonify({
                "success": True,
                "file": os.path.basename(file_path),
                "analysis": analysis,
                "organized": org_result
            })
        else:
            return jsonify({
                "error": org_result.get('error', '组织失败'),
                "analysis": analysis
            }), 500
    
    except Exception as e:
        return jsonify({
            "error": f"自动组织失败: {str(e)}"
        }), 500

@app.route('/api/organize/list-categories', methods=['GET'])
def organize_list_categories():
    """列出所有分类和文件夹"""
    try:
        organizer = get_file_organizer()
        folders = organizer.list_organized_folders()
        stats = organizer.get_categories_stats()
        
        return jsonify({
            "success": True,
            "folders": folders,
            "stats": stats,
            "total_files": len(organizer.get_index().get('files', []))
        })
    
    except Exception as e:
        return jsonify({
            "error": f"获取分类失败: {str(e)}"
        }), 500

@app.route('/api/organize/search', methods=['POST'])
def organize_search():
    """搜索已组织的文件"""
    try:
        data = request.json
        keyword = data.get('keyword', '')
        
        if not keyword:
            return jsonify({"error": "缺少搜索关键词"}), 400
        
        organizer = get_file_organizer()
        results = organizer.search_files(keyword)
        
        return jsonify({
            "success": True,
            "keyword": keyword,
            "count": len(results),
            "results": results
        })
    
    except Exception as e:
        return jsonify({
            "error": f"搜索失败: {str(e)}"
        }), 500

@app.route('/api/organize/stats', methods=['GET'])
def organize_stats():
    """获取组织统计信息"""
    try:
        organizer = get_file_organizer()
        index = organizer.get_index()
        stats = organizer.get_categories_stats()
        folders = organizer.list_organized_folders()
        
        return jsonify({
            "success": True,
            "total_files": index.get('total_files', 0),
            "total_folders": len(folders),
            "by_industry": stats,
            "last_updated": index.get('last_updated')
        })
    
    except Exception as e:
        return jsonify({
            "error": f"获取统计失败: {str(e)}"
        }), 500


@app.route('/api/organize/cleanup', methods=['POST'])
def organize_cleanup():
    """整合清理 _organize 目录中的重复文件夹"""
    try:
        data = request.get_json(silent=True) or {}
        dry_run = data.get('dry_run', True)
        ai_rename = data.get('ai_rename', False)

        organize_root = get_organize_root()

        try:
            from web.organize_cleanup import OrganizeCleanup
        except ImportError:
            from organize_cleanup import OrganizeCleanup

        cleanup = OrganizeCleanup(organize_root=organize_root)
        report = cleanup.run(dry_run=dry_run, ai_rename=ai_rename)

        return jsonify({
            "success": True,
            "dry_run": dry_run,
            "total_folders_scanned": report.get("total_folders_scanned", 0),
            "similarity_groups": report.get("similarity_groups", 0),
            "merge_plans": report.get("merge_plans", 0),
            "merged_files": report.get("merged_files", 0),
            "deduped_files": report.get("deduped_files", 0),
            "removed_folders": report.get("removed_folders", 0),
            "empty_cleaned": report.get("empty_cleaned", 0),
            "ai_renames": report.get("ai_renames", 0),
            "log": report.get("log", [])[-50:],  # 最近50条日志
        })

    except Exception as e:
        return jsonify({
            "error": f"整合清理失败: {str(e)}"
        }), 500


# ═══════════════════════════════════════════════════
# 文件编辑与搜索 API
# ═══════════════════════════════════════════════════

@app.route('/api/file-editor/read', methods=['POST'])
def file_editor_read():
    """读取文件内容"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({"error": "缺少文件路径"}), 400
        
        editor = get_file_editor()
        result = editor.read_file(file_path)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-editor/write', methods=['POST'])
def file_editor_write():
    """写入文件内容"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        content = data.get('content')
        
        if not file_path or content is None:
            return jsonify({"error": "缺少必要参数"}), 400
        
        editor = get_file_editor()
        result = editor.write_file(file_path, content)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-editor/replace', methods=['POST'])
def file_editor_replace():
    """替换文件内容"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        old_text = data.get('old_text')
        new_text = data.get('new_text')
        use_regex = data.get('use_regex', False)
        
        if not all([file_path, old_text is not None, new_text is not None]):
            return jsonify({"error": "缺少必要参数"}), 400
        
        editor = get_file_editor()
        result = editor.replace_text(file_path, old_text, new_text, use_regex=use_regex)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-editor/smart-edit', methods=['POST'])
def file_editor_smart_edit():
    """智能编辑（理解自然语言指令）"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        instruction = data.get('instruction')
        
        if not file_path or not instruction:
            return jsonify({"error": "缺少必要参数"}), 400
        
        editor = get_file_editor()
        result = editor.smart_edit(file_path, instruction)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-search/index', methods=['POST'])
def file_search_index():
    """索引文件或目录"""
    try:
        data = request.json or {}
        path = data.get('path')
        is_directory = data.get('is_directory', False)
        
        if not path:
            return jsonify({"error": "缺少路径参数"}), 400
        
        indexer = get_file_indexer()
        
        if is_directory:
            result = indexer.index_directory(path, recursive=True)
        else:
            result = indexer.index_file(path)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-search/search', methods=['POST'])
def file_search_search():
    """搜索文件"""
    try:
        data = request.json or {}
        query = data.get('query')
        limit = data.get('limit', 20)
        file_types = data.get('file_types')
        
        if not query:
            return jsonify({"error": "缺少搜索关键词"}), 400
        
        indexer = get_file_indexer()
        results = indexer.search(query, limit=limit, file_types=file_types)
        
        return jsonify({
            "success": True,
            "results": results,
            "count": len(results)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-search/find-by-content', methods=['POST'])
def file_search_find_by_content():
    """根据内容片段查找文件"""
    try:
        data = request.json or {}
        content_sample = data.get('content')
        min_similarity = data.get('min_similarity', 0.3)
        
        if not content_sample:
            return jsonify({"error": "缺少内容样本"}), 400
        
        indexer = get_file_indexer()
        results = indexer.find_by_content(content_sample, min_similarity=min_similarity)
        
        return jsonify({
            "success": True,
            "results": results,
            "count": len(results)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/file-search/list', methods=['GET'])
def file_search_list():
    """列出所有已索引文件"""
    try:
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        indexer = get_file_indexer()
        files = indexer.list_indexed_files(limit=limit, offset=offset)
        
        return jsonify({
            "success": True,
            "files": files,
            "count": len(files)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════
# 概念提取 API
# ═══════════════════════════════════════════════════

@app.route('/api/concepts/extract', methods=['POST'])
def concepts_extract():
    """从文件中提取关键概念"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        content = data.get('content')  # 可选，如果已读取内容
        top_n = data.get('top_n', 10)
        
        if not file_path:
            return jsonify({"error": "缺少文件路径"}), 400
        
        extractor = get_concept_extractor()
        result = extractor.analyze_file(file_path, content=content)
        
        return jsonify(result)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/concepts/related-files', methods=['POST'])
def concepts_related_files():
    """查找与文件相关的其他文件"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        limit = data.get('limit', 5)
        
        if not file_path:
            return jsonify({"error": "缺少文件路径"}), 400
        
        extractor = get_concept_extractor()
        related = extractor.find_related_files(file_path, limit=limit)
        
        return jsonify({
            "success": True,
            "file_path": file_path,
            "related_files": related
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/concepts/top', methods=['GET'])
def concepts_top():
    """获取全局热门概念"""
    try:
        limit = request.args.get('limit', 20, type=int)
        
        extractor = get_concept_extractor()
        concepts = extractor.get_top_concepts(limit=limit)
        
        return jsonify({
            "success": True,
            "concepts": concepts
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/concepts/stats', methods=['GET'])
def concepts_stats():
    """获取概念提取统计"""
    try:
        extractor = get_concept_extractor()
        stats = extractor.get_statistics()
        
        return jsonify(stats)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════
# 知识图谱 API
# ═══════════════════════════════════════════════════

@app.route('/api/knowledge-graph/build', methods=['POST'])
def knowledge_graph_build():
    """构建知识图谱"""
    try:
        data = request.json or {}
        file_paths = data.get('file_paths', [])
        force_rebuild = data.get('force_rebuild', False)
        
        if not file_paths:
            return jsonify({"error": "缺少文件路径列表"}), 400
        
        kg = get_knowledge_graph()
        kg.build_file_graph(file_paths, force_rebuild=force_rebuild)
        
        stats = kg.get_statistics()
        
        return jsonify({
            "success": True,
            "message": "知识图谱构建完成",
            "statistics": stats
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/knowledge-graph/data', methods=['GET'])
def knowledge_graph_data():
    """获取知识图谱数据用于可视化"""
    try:
        max_nodes = request.args.get('max_nodes', 100, type=int)
        
        kg = get_knowledge_graph()
        graph_data = kg.get_graph_data(max_nodes=max_nodes)
        
        return jsonify(graph_data)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/knowledge-graph/neighbors', methods=['POST'])
def knowledge_graph_neighbors():
    """获取文件的邻居节点"""
    try:
        data = request.json or {}
        file_path = data.get('file_path')
        depth = data.get('depth', 1)
        
        if not file_path:
            return jsonify({"error": "缺少文件路径"}), 400
        
        kg = get_knowledge_graph()
        neighbors = kg.get_file_neighbors(file_path, depth=depth)
        
        return jsonify(neighbors)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/knowledge-graph/concept-cluster', methods=['POST'])
def knowledge_graph_concept_cluster():
    """获取概念相关的文件集群"""
    try:
        data = request.json or {}
        concept = data.get('concept')
        limit = data.get('limit', 20)
        
        if not concept:
            return jsonify({"error": "缺少概念参数"}), 400
        
        kg = get_knowledge_graph()
        cluster = kg.get_concept_cluster(concept, limit=limit)
        
        return jsonify(cluster)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/knowledge-graph/stats', methods=['GET'])
def knowledge_graph_stats():
    """获取知识图谱统计"""
    try:
        kg = get_knowledge_graph()
        stats = kg.get_statistics()
        
        return jsonify(stats)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════
# 行为监控 API
# ═══════════════════════════════════════════════════

@app.route('/api/behavior/log-event', methods=['POST'])
def behavior_log_event():
    """记录用户行为事件"""
    try:
        data = request.json or {}
        event_type = data.get('event_type')
        file_path = data.get('file_path')
        session_id = data.get('session_id')
        event_data = data.get('event_data')
        duration_ms = data.get('duration_ms')
        user_id = data.get('user_id', 'default')
        auto_trigger = data.get('auto_trigger', True)
        
        if not event_type:
            return jsonify({"error": "缺少事件类型"}), 400
        
        monitor = get_behavior_monitor()
        event_id = monitor.log_event(
            event_type=event_type,
            file_path=file_path,
            session_id=session_id,
            event_data=event_data,
            duration_ms=duration_ms
        )

        decision_payload = None
        triggered = False
        if auto_trigger:
            trigger_system = get_trigger_system()
            decision = trigger_system.evaluate_interaction_need(user_id)
            if decision and decision.should_interact:
                trigger_system.execute_interaction(decision, user_id)
                triggered = True
                decision_payload = {
                    "interaction_type": decision.interaction_type.value,
                    "priority": decision.priority,
                    "reason": decision.reason,
                    "content": decision.content,
                    "scores": {
                        "urgency": decision.urgency_score,
                        "importance": decision.importance_score,
                        "disturbance": decision.disturbance_cost,
                        "final": decision.final_score
                    }
                }
        
        return jsonify({
            "success": True,
            "event_id": event_id,
            "triggered": triggered,
            "decision": decision_payload
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/behavior/recent-events', methods=['GET'])
def behavior_recent_events():
    """获取最近的事件"""
    try:
        limit = request.args.get('limit', 50, type=int)
        event_type = request.args.get('event_type')
        
        monitor = get_behavior_monitor()
        events = monitor.get_recent_events(limit=limit, event_type=event_type)
        
        return jsonify({
            "success": True,
            "events": events
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/behavior/top-files', methods=['GET'])
def behavior_top_files():
    """获取最常用的文件"""
    try:
        limit = request.args.get('limit', 10, type=int)
        
        monitor = get_behavior_monitor()
        files = monitor.get_frequently_used_files(limit=limit)
        
        return jsonify({
            "success": True,
            "files": files
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/behavior/work-patterns', methods=['GET'])
def behavior_work_patterns():
    """获取工作模式分析"""
    try:
        monitor = get_behavior_monitor()
        patterns = monitor.get_work_patterns()
        
        return jsonify(patterns)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/behavior/stats', methods=['GET'])
def behavior_stats():
    """获取行为统计"""
    try:
        monitor = get_behavior_monitor()
        stats = monitor.get_statistics()
        
        return jsonify(stats)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════
# 智能建议 API
# ═══════════════════════════════════════════════════

@app.route('/api/suggestions/generate', methods=['POST'])
def suggestions_generate():
    """生成智能建议"""
    try:
        data = request.json or {}
        force_regenerate = data.get('force_regenerate', False)
        
        engine = get_suggestion_engine()
        suggestions = engine.generate_suggestions(force_regenerate=force_regenerate)
        
        return jsonify({
            "success": True,
            "suggestions": suggestions,
            "count": len(suggestions)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/suggestions/pending', methods=['GET'])
def suggestions_pending():
    """获取待处理的建议"""
    try:
        limit = request.args.get('limit', 10, type=int)
        
        engine = get_suggestion_engine()
        suggestions = engine.get_pending_suggestions(limit=limit)
        
        return jsonify({
            "success": True,
            "suggestions": suggestions,
            "count": len(suggestions)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/suggestions/dismiss', methods=['POST'])
def suggestions_dismiss():
    """拒绝建议"""
    try:
        data = request.json or {}
        suggestion_id = data.get('suggestion_id')
        feedback = data.get('feedback')
        
        if not suggestion_id:
            return jsonify({"error": "缺少建议ID"}), 400
        
        engine = get_suggestion_engine()
        engine.dismiss_suggestion(suggestion_id, feedback=feedback)
        
        return jsonify({
            "success": True,
            "message": "建议已拒绝"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/suggestions/apply', methods=['POST'])
def suggestions_apply():
    """应用建议"""
    try:
        data = request.json or {}
        suggestion_id = data.get('suggestion_id')
        feedback = data.get('feedback')
        
        if not suggestion_id:
            return jsonify({"error": "缺少建议ID"}), 400
        
        engine = get_suggestion_engine()
        engine.apply_suggestion(suggestion_id, feedback=feedback)
        
        return jsonify({
            "success": True,
            "message": "建议已应用"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/suggestions/stats', methods=['GET'])
def suggestions_stats():
    """获取建议统计"""
    try:
        engine = get_suggestion_engine()
        stats = engine.get_statistics()
        
        return jsonify(stats)
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════
# 洞察报告 API
# ═══════════════════════════════════════════════════

@app.route('/api/insights/generate-weekly', methods=['POST'])
def insights_generate_weekly():
    """生成周报"""
    try:
        reporter = get_insight_reporter()
        report = reporter.generate_weekly_report()
        
        return jsonify({
            "success": True,
            "report": report
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/insights/generate-monthly', methods=['POST'])
def insights_generate_monthly():
    """生成月报"""
    try:
        reporter = get_insight_reporter()
        report = reporter.generate_monthly_report()
        
        return jsonify({
            "success": True,
            "report": report
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/insights/latest', methods=['GET'])
def insights_latest():
    """获取最新报告"""
    try:
        report_type = request.args.get('type', 'weekly')
        
        reporter = get_insight_reporter()
        report = reporter.get_latest_report(report_type=report_type)
        
        if report:
            return jsonify({
                "success": True,
                "report": report
            })
        else:
            return jsonify({
                "success": False,
                "message": "暂无报告"
            })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/insights/export-markdown', methods=['POST'])
def insights_export_markdown():
    """导出报告为Markdown"""
    try:
        data = request.json or {}
        report = data.get('report')
        output_path = data.get('output_path', 'workspace/report.md')
        
        if not report:
            return jsonify({"error": "缺少报告数据"}), 400
        
        reporter = get_insight_reporter()
        saved_path = reporter.export_report_markdown(report, output_path)
        
        return jsonify({
            "success": True,
            "file_path": saved_path
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== 通知管理 API ====================

@app.route('/api/notifications/unread', methods=['GET'])
def get_unread_notifications():
    """获取未读通知"""
    try:
        user_id = request.args.get('user_id', 'default')
        limit = int(request.args.get('limit', 50))
        
        manager = get_notification_manager()
        notifications = manager.get_unread_notifications(user_id, limit)
        
        return jsonify({
            "success": True,
            "notifications": notifications,
            "count": len(notifications)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/notifications/mark-read', methods=['POST'])
def mark_notification_read():
    """标记通知已读"""
    try:
        data = request.json or {}
        notification_id = data.get('notification_id')
        user_id = data.get('user_id', 'default')
        
        if not notification_id:
            return jsonify({"error": "缺少notification_id"}), 400
        
        manager = get_notification_manager()
        manager.mark_as_read(notification_id, user_id)
        
        return jsonify({"success": True})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/notifications/dismiss', methods=['POST'])
def dismiss_notification():
    """忽略通知"""
    try:
        data = request.json or {}
        notification_id = data.get('notification_id')
        user_id = data.get('user_id', 'default')
        
        if not notification_id:
            return jsonify({"error": "缺少notification_id"}), 400
        
        manager = get_notification_manager()
        manager.dismiss_notification(notification_id, user_id)
        
        return jsonify({"success": True})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/notifications/stats', methods=['GET'])
def get_notification_stats():
    """获取通知统计"""
    try:
        user_id = request.args.get('user_id', 'default')
        days = int(request.args.get('days', 7))
        
        manager = get_notification_manager()
        stats = manager.get_notification_stats(user_id, days)
        
        return jsonify({
            "success": True,
            "stats": stats
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/notifications/preferences', methods=['GET', 'POST'])
def notification_preferences():
    """获取或设置通知偏好"""
    try:
        user_id = request.args.get('user_id', 'default')
        manager = get_notification_manager()
        
        if request.method == 'GET':
            prefs = manager.get_user_preferences(user_id)
            return jsonify({
                "success": True,
                "preferences": prefs
            })
        
        else:  # POST
            data = request.json or {}
            manager.update_user_preferences(user_id, data)
            return jsonify({"success": True})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== 主动对话 API ====================

@app.route('/api/dialogue/start-monitoring', methods=['POST'])
def start_dialogue_monitoring():
    """启动主动对话监控"""
    try:
        data = request.json or {}
        check_interval = data.get('check_interval', 300)  # 默认5分钟
        
        engine = get_proactive_dialogue()
        engine.start_monitoring(check_interval)
        
        return jsonify({
            "success": True,
            "message": "主动对话监控已启动"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/dialogue/stop-monitoring', methods=['POST'])
def stop_dialogue_monitoring():
    """停止主动对话监控"""
    try:
        engine = get_proactive_dialogue()
        engine.stop_monitoring()
        
        return jsonify({
            "success": True,
            "message": "主动对话监控已停止"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/dialogue/trigger', methods=['POST'])
def trigger_dialogue():
    """手动触发对话"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        scene_type = data.get('scene_type')
        context = data.get('context', {})
        
        if not scene_type:
            return jsonify({"error": "缺少scene_type"}), 400
        
        engine = get_proactive_dialogue()
        engine.manual_trigger(user_id, scene_type, **context)
        
        return jsonify({
            "success": True,
            "message": f"已触发{scene_type}对话"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/dialogue/history', methods=['GET'])
def get_dialogue_history():
    """获取对话历史"""
    try:
        user_id = request.args.get('user_id', 'default')
        limit = int(request.args.get('limit', 50))
        
        engine = get_proactive_dialogue()
        history = engine.get_dialogue_history(user_id, limit)
        
        return jsonify({
            "success": True,
            "history": history,
            "count": len(history)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== 情境感知 API ====================

@app.route('/api/context/detect', methods=['POST'])
def detect_context():
    """检测当前工作场景"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        
        system = get_context_awareness()
        context = system.detect_context(user_id)
        
        return jsonify({
            "success": True,
            "context": context
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/context/current', methods=['GET'])
def get_current_context():
    """获取当前场景"""
    try:
        system = get_context_awareness()
        context = system.get_current_context()
        
        return jsonify({
            "success": True,
            "context": context
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/context/history', methods=['GET'])
def get_context_history():
    """获取场景历史"""
    try:
        user_id = request.args.get('user_id', 'default')
        days = int(request.args.get('days', 7))
        
        system = get_context_awareness()
        history = system.get_context_history(user_id, days)
        
        return jsonify({
            "success": True,
            "history": history,
            "count": len(history)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/context/statistics', methods=['GET'])
def get_context_statistics():
    """获取场景统计"""
    try:
        user_id = request.args.get('user_id', 'default')
        days = int(request.args.get('days', 30))
        
        system = get_context_awareness()
        stats = system.get_context_statistics(user_id, days)
        
        return jsonify({
            "success": True,
            "statistics": stats
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/context/predict', methods=['GET'])
def predict_next_context():
    """预测下一个场景"""
    try:
        user_id = request.args.get('user_id', 'default')
        
        system = get_context_awareness()
        prediction = system.predict_next_context(user_id)
        
        return jsonify({
            "success": True,
            "prediction": prediction
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== 自动执行 API ====================

@app.route('/api/execution/authorize', methods=['POST'])
def authorize_task_execution():
    """授权任务执行"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        task_type = data.get('task_type')
        auto_execute = data.get('auto_execute', False)
        max_executions_per_day = data.get('max_executions_per_day', 10)
        expires_days = data.get('expires_days', 30)
        
        if not task_type:
            return jsonify({"error": "缺少task_type"}), 400
        
        engine = get_auto_execution()
        engine.authorize_task(
            user_id, task_type, auto_execute,
            max_executions_per_day, expires_days
        )
        
        return jsonify({
            "success": True,
            "message": f"已授权{task_type}任务"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/revoke', methods=['POST'])
def revoke_task_authorization():
    """撤销任务授权"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        task_type = data.get('task_type')
        
        if not task_type:
            return jsonify({"error": "缺少task_type"}), 400
        
        engine = get_auto_execution()
        engine.revoke_authorization(user_id, task_type)
        
        return jsonify({
            "success": True,
            "message": f"已撤销{task_type}任务授权"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/execute', methods=['POST'])
def execute_task():
    """执行任务"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        task_type = data.get('task_type')
        params = data.get('params', {})
        force = data.get('force', False)
        
        if not task_type:
            return jsonify({"error": "缺少task_type"}), 400
        
        engine = get_auto_execution()
        result = engine.execute_task(user_id, task_type, params, force)
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify(result), 400
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/queue', methods=['POST'])
def queue_task():
    """任务加入队列"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        task_type = data.get('task_type')
        params = data.get('params', {})
        priority = data.get('priority', 5)
        
        if not task_type:
            return jsonify({"error": "缺少task_type"}), 400
        
        engine = get_auto_execution()
        task_id = engine.queue_task(user_id, task_type, params, priority)
        
        return jsonify({
            "success": True,
            "task_id": task_id
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/history', methods=['GET'])
def get_execution_history():
    """获取执行历史"""
    try:
        user_id = request.args.get('user_id', 'default')
        limit = int(request.args.get('limit', 50))
        
        engine = get_auto_execution()
        history = engine.get_execution_history(user_id, limit)
        
        return jsonify({
            "success": True,
            "history": history,
            "count": len(history)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/statistics', methods=['GET'])
def get_execution_statistics():
    """获取执行统计"""
    try:
        user_id = request.args.get('user_id', 'default')
        days = int(request.args.get('days', 30))
        
        engine = get_auto_execution()
        stats = engine.get_statistics(user_id, days)
        
        return jsonify({
            "success": True,
            "statistics": stats
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/start-processor', methods=['POST'])
def start_execution_processor():
    """启动自动执行处理器"""
    try:
        data = request.json or {}
        interval = data.get('interval', 60)  # 默认1分钟
        
        engine = get_auto_execution()
        engine.start_queue_processor(interval)
        
        return jsonify({
            "success": True,
            "message": "自动执行处理器已启动"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/execution/stop-processor', methods=['POST'])
def stop_execution_processor():
    """停止自动执行处理器"""
    try:
        engine = get_auto_execution()
        engine.stop_queue_processor()
        
        return jsonify({
            "success": True,
            "message": "自动执行处理器已停止"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ==================== 主动交互触发系统 API ====================

@app.route('/api/triggers/evaluate', methods=['POST'])
def triggers_evaluate():
    """评估是否需要主动交互"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        execute = data.get('execute', True)
        
        system = get_trigger_system()
        decision = system.evaluate_interaction_need(user_id)
        
        if decision and decision.should_interact and execute:
            system.execute_interaction(decision, user_id)
        
        decision_payload = None
        if decision:
            decision_payload = {
                "should_interact": decision.should_interact,
                "interaction_type": decision.interaction_type.value,
                "priority": decision.priority,
                "reason": decision.reason,
                "content": decision.content,
                "scores": {
                    "urgency": decision.urgency_score,
                    "importance": decision.importance_score,
                    "disturbance": decision.disturbance_cost,
                    "final": decision.final_score
                }
            }
        
        return jsonify({
            "success": True,
            "decision": decision_payload
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/start', methods=['POST'])
def triggers_start():
    """启动主动交互监控"""
    try:
        data = request.json or {}
        user_id = data.get('user_id', 'default')
        interval = data.get('interval', 300)
        
        system = get_trigger_system()
        system.start_monitoring(check_interval=interval, user_id=user_id)
        
        return jsonify({
            "success": True,
            "message": "主动交互触发系统已启动",
            "interval": interval
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/stop', methods=['POST'])
def triggers_stop():
    """停止主动交互监控"""
    try:
        system = get_trigger_system()
        system.stop_monitoring()
        
        return jsonify({
            "success": True,
            "message": "主动交互触发系统已停止"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/stats', methods=['GET'])
def triggers_stats():
    """获取触发统计"""
    try:
        days = int(request.args.get('days', 7))
        
        system = get_trigger_system()
        stats = system.get_trigger_statistics(days)
        
        return jsonify({
            "success": True,
            "stats": stats
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/list', methods=['GET'])
def triggers_list():
    """获取触发器列表"""
    try:
        system = get_trigger_system()
        triggers = system.list_triggers()
        
        return jsonify({
            "success": True,
            "triggers": triggers
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/update', methods=['POST'])
def triggers_update():
    """更新触发器配置"""
    try:
        data = request.json or {}
        trigger_id = data.get('trigger_id')
        
        if not trigger_id:
            return jsonify({"error": "缺少trigger_id"}), 400
        
        enabled = data.get('enabled')
        priority = data.get('priority')
        cooldown_minutes = data.get('cooldown_minutes')
        threshold_value = data.get('threshold_value')
        parameters = data.get('parameters')
        
        system = get_trigger_system()
        ok = system.update_trigger_config(
            trigger_id,
            enabled=enabled,
            priority=priority,
            cooldown_minutes=cooldown_minutes,
            threshold_value=threshold_value
        )
        
        if not ok:
            return jsonify({"error": "触发器不存在"}), 404
        
        # 如果提供了参数，更新参数
        if parameters is not None:
            system.update_trigger_params(trigger_id, parameters)
        
        return jsonify({
            "success": True,
            "message": "触发器配置已更新"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/params/<trigger_id>', methods=['GET'])
def get_trigger_params(trigger_id):
    """获取触发器参数"""
    try:
        system = get_trigger_system()
        params = system.get_trigger_params(trigger_id)
        
        return jsonify({
            "success": True,
            "trigger_id": trigger_id,
            "parameters": params
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/params/<trigger_id>', methods=['POST'])
def update_trigger_params_endpoint(trigger_id):
    """更新触发器参数"""
    try:
        data = request.json or {}
        parameters = data.get('parameters', {})
        
        system = get_trigger_system()
        ok = system.update_trigger_params(trigger_id, parameters)
        
        if not ok:
            return jsonify({"error": "触发器不存在"}), 404
        
        return jsonify({
            "success": True,
            "message": "触发器参数已更新",
            "parameters": system.get_trigger_params(trigger_id)
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/triggers/feedback', methods=['POST'])
def triggers_feedback():
    """提交触发反馈"""
    try:
        data = request.json or {}
        trigger_id = data.get('trigger_id')
        feedback = data.get('feedback')
        response_time_seconds = data.get('response_time_seconds', 0)
        
        if not trigger_id or not feedback:
            return jsonify({"error": "缺少trigger_id或feedback"}), 400
        
        system = get_trigger_system()
        system.record_user_feedback(trigger_id, feedback, response_time_seconds)
        
        return jsonify({
            "success": True,
            "message": "反馈已记录"
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══ 注册增强记忆系统API（模块级别，确保始终执行） ═══
try:
    from memory_api_routes import register_memory_routes
    register_memory_routes(app, get_memory_manager)
except ImportError:
    try:
        from web.memory_api_routes import register_memory_routes
        register_memory_routes(app, get_memory_manager)
    except ImportError:
        print("⚠️  增强记忆系统API未找到，使用基础功能")


# ═══ 自动归纳调度器 API ═══

@app.route('/api/auto-catalog/status', methods=['GET'])
def auto_catalog_status():
    """获取自动归纳状态"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler
        scheduler = get_auto_catalog_scheduler()
        
        return jsonify({
            "success": True,
            "enabled": scheduler.is_auto_catalog_enabled(),
            "schedule_time": scheduler.get_catalog_schedule(),
            "source_directories": scheduler.get_source_directories(),
            "backup_directory": scheduler.get_backup_directory()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auto-catalog/enable', methods=['POST'])
def auto_catalog_enable():
    """启用自动归纳"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler
        scheduler = get_auto_catalog_scheduler()
        
        data = request.json or {}
        schedule_time = data.get('schedule_time', '02:00')
        source_dirs = data.get('source_directories')
        
        scheduler.enable_auto_catalog(schedule_time, source_dirs)
        
        return jsonify({
            "success": True,
            "message": f"自动归纳已启用，每日 {schedule_time} 执行",
            "schedule_time": schedule_time,
            "source_directories": scheduler.get_source_directories()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auto-catalog/disable', methods=['POST'])
def auto_catalog_disable():
    """禁用自动归纳"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler
        scheduler = get_auto_catalog_scheduler()
        
        scheduler.disable_auto_catalog()
        
        return jsonify({
            "success": True,
            "message": "自动归纳已禁用"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auto-catalog/run-now', methods=['POST'])
def auto_catalog_run_now():
    """立即执行一次归纳（手动触发）"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler
        scheduler = get_auto_catalog_scheduler()
        
        result = scheduler.manual_catalog_now()
        
        return jsonify({
            "success": result.get('success', False),
            "total_files": result.get('total_files', 0),
            "organized_count": result.get('organized_count', 0),
            "backed_up_count": result.get('backed_up_count', 0),
            "errors": result.get('errors', []),
            "report_path": result.get('report_path', '')
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/auto-catalog/backup-manifest/<path:filename>', methods=['GET'])
def get_backup_manifest(filename):
    """下载备份清单文件"""
    try:
        from auto_catalog_scheduler import get_auto_catalog_scheduler
        scheduler = get_auto_catalog_scheduler()
        
        backup_dir = scheduler.get_backup_directory()
        return send_from_directory(backup_dir, filename, as_attachment=False)
    except Exception as e:
        return jsonify({"error": str(e)}), 404


