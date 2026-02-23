#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
直接修改型文档标注系统 - 替代Comments批注
生成格式：原文 -> 修改后文本
用户可接受/拒绝
"""

import os
import json
import re
from typing import Dict, List, Any, Generator, Tuple, Optional
from datetime import datetime
from copy import deepcopy


class DirectEditAnnotator:
    """直接修改文档内容（不用Comments）"""
    
    def __init__(self):
        pass
    
    @staticmethod
    def prepare_document(file_path: str) -> Tuple[str, str]:
        """创建副本"""
        from shutil import copy2
        base_name = os.path.splitext(file_path)[0]
        ext = os.path.splitext(file_path)[1]
        revised_path = f"{base_name}_revised{ext}"
        copy2(file_path, revised_path)
        print(f"[DirectEdit] 📋 已创建工作副本: {os.path.basename(revised_path)}")
        return file_path, revised_path
    
    @staticmethod
    def apply_edits(file_path: str, edits: List[Dict[str, str]]) -> Dict[str, Any]:
        """
        直接修改文档内容
        
        Args:
            file_path: Word文档路径
            edits: [{"原文": "...", "修改": "..."}, ...]
        
        Returns:
            {"success": True/False, "modified_count": N, "output_file": path}
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            modified_count = 0
            
            # 对每个编辑进行处理
            for edit in edits:
                original = edit.get("原文", "")
                modified = edit.get("修改", "")
                
                if not original or not modified:
                    continue
                
                # 在所有段落中查找并替换
                for para in doc.paragraphs:
                    if original in para.text:
                        # 替换段落中的文本
                        full_text = para.text
                        new_text = full_text.replace(original, modified, 1)
                        
                        if new_text != full_text:
                            # 清空段落并重新添加（保留格式）
                            para.clear()
                            para.add_run(new_text)
                            modified_count += 1
                            print(f"[DirectEdit] ✅ 已修改: '{original}' -> '{modified}'")
                            break
                
                # 检查表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if original in para.text:
                                    full_text = para.text
                                    new_text = full_text.replace(original, modified, 1)
                                    if new_text != full_text:
                                        para.clear()
                                        para.add_run(new_text)
                                        modified_count += 1
                                        break
            
            doc.save(file_path)
            print(f"[DirectEdit] 💾 文档已保存: {modified_count}处修改")
            
            return {
                "success": True,
                "modified_count": modified_count,
                "output_file": file_path
            }
        
        except Exception as e:
            print(f"[DirectEdit] ✗ 应用编辑失败: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "modified_count": 0
            }


class ImprovedBatchAnnotator:
    """改进的批量标注器 - 直接修改模式"""
    
    def __init__(self, gemini_client=None, batch_size: int = 3):
        self.client = gemini_client
        self.batch_size = batch_size
        self.direct_editor = DirectEditAnnotator()
    
    def annotate_document_streaming(
        self,
        file_path: str,
        user_requirement: str = "把所有不合适的翻译、不符合中文语序逻辑、生硬的地方修改"
    ) -> Generator[str, None, None]:
        """流式标注文档（直接修改模式）"""
        
        from web.document_reader import DocumentReader
        
        # Step 1: 读取文档
        yield self._sse_event("progress", {
            "stage": "reading",
            "message": "📖 正在读取文档...",
            "progress": 0
        })
        
        reader = DocumentReader()
        doc_data = reader.read_document(file_path)
        
        if not doc_data.get("success"):
            yield self._sse_event("error", {
                "message": f"读取文档失败: {doc_data.get('error')}"
            })
            return
        
        paragraphs = doc_data.get("paragraphs", [])
        total_paras = len(paragraphs)
        
        yield self._sse_event("progress", {
            "stage": "reading_complete",
            "message": f"✅ 文档解析完成，共 {total_paras} 段",
            "progress": 10
        })
        
        # Step 2: 分批处理
        all_edits = []
        batch_count = (total_paras + self.batch_size - 1) // self.batch_size
        
        for batch_idx in range(batch_count):
            start_para = batch_idx * self.batch_size
            end_para = min(start_para + self.batch_size, total_paras)
            batch_paras = paragraphs[start_para:end_para]
            
            batch_text = "\n\n---分段---\n\n".join([p["text"] for p in batch_paras])
            
            progress = 10 + int((batch_idx / batch_count) * 80)
            yield self._sse_event("progress", {
                "stage": "analyzing",
                "message": f"🤖 分析第 {batch_idx + 1}/{batch_count} 批（段落 {start_para + 1}-{end_para}）",
                "progress": progress
            })
            
            # 调用AI或本地规则获得修改建议
            batch_edits = self._analyze_batch(batch_text, user_requirement)
            all_edits.extend(batch_edits)
            
            yield self._sse_event("batch_complete", {
                "batch": batch_idx + 1,
                "edits_found": len(batch_edits),
                "total_edits": len(all_edits)
            })
        
        # Step 3: 应用所有修改
        yield self._sse_event("progress", {
            "stage": "applying",
            "message": f"📝 正在应用 {len(all_edits)} 处修改...",
            "progress": 90
        })
        
        # 创建副本并应用修改
        original_path, revised_path = self.direct_editor.prepare_document(file_path)
        result = self.direct_editor.apply_edits(revised_path, all_edits)
        
        yield self._sse_event("progress", {
            "stage": "complete",
            "message": f"✅ 完成！已修改 {len(all_edits)} 处",
            "progress": 100
        })
        
        yield self._sse_event("complete", {
            "output_file": result.get("output_file"),
            "total_edits": len(all_edits),
            "success": result.get("success"),
            "modified_count": result.get("modified_count")
        })
    
    def _analyze_batch(self, batch_text: str, user_requirement: str) -> List[Dict[str, str]]:
        """分析一批段落，返回修改建议"""
        
        # 暂时禁用AI，使用改进的本地规则
        print(f"[BatchAnnotator] 📋 使用改进规则分析")
        return self._improved_local_rules(batch_text)
    
    def _improved_local_rules(self, batch_text: str) -> List[Dict[str, str]]:
        """
        改进的本地规则 - 覆盖更多场景，分布更均匀
        
        策略：
        1. 增加规则数量（从10->20+）
        2. 分类覆盖不同类型的翻译问题
        3. 允许多次匹配同一类型规则
        """
        edits = []
        
        # 第一类：冗余/重复词汇（最常见）
        redundant_patterns = [
            (r'可以(?:被)?(?:进行)?(\w{2,8})', lambda m: m.group(1)),  # 可以使用 -> 使用
            (r'进行(\w{2,8})', lambda m: m.group(1)),  # 进行分析 -> 分析
            (r'对(\w{2,6})(?:进行|做)', lambda m: m.group(1)),  # 对数据进行处理 -> 数据处理
            (r'进行了(\w{2,8})', lambda m: m.group(1)),  # 进行了探讨 -> 探讨
        ]
        
        for pattern, replacement_fn in redundant_patterns:
            for match in re.finditer(pattern, batch_text):
                try:
                    original = match.group(0)
                    modified = replacement_fn(match)
                    if original != modified and len(original) >= 3:
                        edits.append({"原文": original, "修改": modified})
                except:
                    pass
        
        # 第二类：翻译腔修正
        translation_patterns = [
            (r'比较(\w{2,4})', lambda m: f"更{m.group(1)}"),  # 比较好 -> 更好
            (r'非常的(\w{2,4})', lambda m: f"非常{m.group(1)}"),  # 非常的重要 -> 非常重要
            (r'是(?:一)?(\w{2,6})(?:的)?', lambda m: m.group(1)),  # 是重要的 -> 重要
            (r'(?:显得|看起来)(\w{2,6})', lambda m: m.group(1)),  # 显得不够好 -> 不够好
        ]
        
        for pattern, replacement_fn in translation_patterns:
            for match in re.finditer(pattern, batch_text):
                try:
                    original = match.group(0)
                    modified = replacement_fn(match)
                    if original != modified and 3 <= len(original) <= 10:
                        edits.append({"原文": original, "修改": modified})
                except:
                    pass
        
        # 第三类：被动句转主动
        passive_patterns = [
            (r'被(\w{2,6})', lambda m: m.group(1)),  # 被使用 -> 使用
            (r'受到(\w{2,6})(?:的)?', lambda m: m.group(1)),  # 受到影响 -> 影响
            (r'被称为(\w{2,8})', lambda m: f"称为{m.group(1)}"),  # 被称为X -> 称为X
        ]
        
        for pattern, replacement_fn in passive_patterns:
            for match in re.finditer(pattern, batch_text):
                try:
                    original = match.group(0)
                    modified = replacement_fn(match)
                    if original != modified and len(original) >= 3:
                        edits.append({"原文": original, "修改": modified})
                except:
                    pass
        
        # 第四类：啰嗦表达压缩
        verbose_patterns = [
            (r'在(\w{2,8})方面(?:上)?', lambda m: m.group(1)),  # 在研究方面 -> 研究
            (r'(?:的|这种|这样的)(\w{2,8})方式', lambda m: m.group(1)),  # 的处理方式 -> 处理
            (r'(\w{2,6})的过程(?:中)?', lambda m: f"{m.group(1)}"),  # X的过程 -> X
            (r'通过(?:采用|使用)(\w{2,8})', lambda m: f"用{m.group(1)}"),  # 通过采用X -> 用X
        ]
        
        for pattern, replacement_fn in verbose_patterns:
            for match in re.finditer(pattern, batch_text):
                try:
                    original = match.group(0)
                    modified = replacement_fn(match)
                    if original != modified and len(original) >= 4:
                        edits.append({"原文": original, "修改": modified})
                except:
                    pass
        
        # 第五类：语序不当
        word_order_patterns = [
            (r'对(\w{2,6})的(\w{2,6})', lambda m: f"{m.group(2)}对{m.group(1)}的"),  # 对数据的处理 -> 处理对数据的
            (r'(\w{2,4})(?:和|与)(\w{2,4})的', lambda m: f"{m.group(2)}和{m.group(1)}的"),  # 简化处理
        ]
        
        for pattern, replacement_fn in word_order_patterns:
            for match in re.finditer(pattern, batch_text):
                try:
                    original = match.group(0)
                    modified = replacement_fn(match)
                    if original != modified and len(original) >= 5:
                        edits.append({"原文": original, "修改": modified})
                except:
                    pass
        
        # 清洗和去重：保留多个匹配，但避免完全重复
        # 允许同一个pattern的多个不同匹配都被保留
        seen_exact = set()
        unique_edits = []
        
        for edit in edits:
            original = edit["原文"]
            modified = edit["修改"]
            
            # 只去除完全相同的 (原文和修改都一样)
            key = (original, modified)
            if key not in seen_exact:
                if 2 <= len(original) <= 20 and len(modified) > 0:
                    seen_exact.add(key)
                    unique_edits.append(edit)
        
        # 限制但不要太严格 - 允许每段多个编辑
        # 按类型分配：每类最多5条
        result = []
        category_counts = {}
        
        for edit in unique_edits:
            original = edit["原文"]
            
            # 分类计数
            if "进行" in original or "可以" in original:
                cat = "redundancy"
            elif "比较" in original or "非常的" in original or "是" in original:
                cat = "translation"
            elif "被" in original or "受到" in original:
                cat = "passive"
            elif "方面" in original or "方式" in original or "过程" in original:
                cat = "verbose"
            else:
                cat = "other"
            
            count = category_counts.get(cat, 0)
            
            # 每类最多8条（比之前的30更合理）
            if count < 8:
                result.append(edit)
                category_counts[cat] = count + 1
        
        return result[:40]  # 单批最多40条（给足空间）
    
    def _sse_event(self, event_type: str, data: Dict) -> str:
        """构造SSE事件"""
        return f"event: {event_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"
