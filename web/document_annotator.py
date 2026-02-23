#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
æ–‡æ¡£è‡ªåŠ¨æ ‡æ³¨ç³»ç»Ÿ - æ ¸å¿ƒæ¨¡å—
å·¥ä½œæµï¼šè¯»å– -> åˆ†æ -> å®šä½ -> æ³¨å…¥

é»˜è®¤é‡‡ç”¨æ–¹æ¡ˆAï¼ˆWord åŸç”Ÿæ‰¹æ³¨æ°”æ³¡ï¼‰ï¼š
- åœ¨åŸæ–‡ä½ç½®æ·»åŠ å³ä¾§æ‰¹æ³¨
- æ ¼å¼æ— æŸï¼ŒWord åŸç”Ÿä½“éªŒ
- å¤±è´¥æ—¶è‡ªåŠ¨å›é€€ä¸ºçº¢è‰²æ ‡æ³¨
"""

import os
import shutil
import json
from typing import Dict, List, Any, Tuple, Optional
from copy import deepcopy
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path


class DocumentAnnotator:
    """æ–‡æ¡£è‡ªåŠ¨æ ‡æ³¨ç³»ç»Ÿ"""
    
    def __init__(self, min_similarity: float = 0.8, annotation_mode: str = "comment"):
        """
        Args:
            min_similarity: æ¨¡ç³ŠåŒ¹é…çš„ç›¸ä¼¼åº¦é˜ˆå€¼ (0-1)
            annotation_mode: æ ‡æ³¨æ–¹å¼ ("comment" | "highlight")
        """
        self.min_similarity = min_similarity
        self.annotation_mode = annotation_mode
    
    # ==================== Step 1: æ–‡ä»¶é¢„å¤„ç† ====================
    
    @staticmethod
    def prepare_document(file_path: str) -> Tuple[str, str]:
        """
        Step 1: ç”¨æˆ·è¾“å…¥ä¸é¢„å¤„ç†
        
        åŠ¨ä½œï¼š
        1. åœ¨åå°å¤åˆ¶ä¸€ä»½å‰¯æœ¬ï¼ˆå‘½åä¸º _revised.docxï¼‰
        2. æ‰€æœ‰æ“ä½œéƒ½åœ¨å‰¯æœ¬ä¸Šè¿›è¡Œï¼Œç¡®ä¿åŸä»¶å®‰å…¨
        
        Args:
            file_path: åŸWordæ–‡ä»¶è·¯å¾„
        
        Returns:
            (åŸæ–‡ä»¶è·¯å¾„, å‰¯æœ¬è·¯å¾„)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")
        
        # ç”Ÿæˆå‰¯æœ¬è·¯å¾„
        base_name = os.path.splitext(file_path)[0]
        ext = os.path.splitext(file_path)[1]
        revised_path = f"{base_name}_revised{ext}"
        
        # å¤åˆ¶æ–‡ä»¶
        shutil.copy2(file_path, revised_path)
        print(f"[Annotator] ğŸ“‹ å·²åˆ›å»ºå·¥ä½œå‰¯æœ¬: {os.path.basename(revised_path)}")
        
        return file_path, revised_path
    
    # ==================== Step 1.5: æ–‡æœ¬æå– ====================
    
    @staticmethod
    def extract_text_from_word(file_path: str) -> Dict[str, Any]:
        """
        æå–Wordæ–‡æ¡£çš„çº¯æ–‡æœ¬å†…å®¹ï¼ˆåŒ…æ‹¬è¡¨æ ¼ï¼‰+ æ ¼å¼ä¿¡æ¯
        
        Returns:
            {
                "full_text": "å®Œæ•´æ–‡æœ¬",
                "full_text_with_format": "å¸¦æ ¼å¼æ ‡è®°çš„æ–‡æœ¬",  # æ–°å¢
                "paragraphs": [
                    {
                        "index": 0,
                        "text": "æ®µè½æ–‡æœ¬",
                        "text_with_format": "<bold>æ ‡é¢˜</bold>æ­£æ–‡",  # æ–°å¢
                        "runs": [{"text": "x", "bold": True, "size": 14}],  # æ–°å¢
                        "start_pos": 0,
                        "end_pos": 15,
                        "source": "body" | "table",
                        "table_idx": 0,
                        "row_idx": 0,
                        "cell_idx": 0
                    },
                    ...
                ]
            }
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs_info = []
            full_text = ""
            full_text_with_format = ""  # æ–°å¢ï¼šå¸¦æ ¼å¼çš„æ–‡æœ¬
            para_global_idx = 0
            
            # æå–æ­£æ–‡æ®µè½
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text
                
                # è·³è¿‡ç©ºæ®µè½
                if not text.strip():
                    continue
                
                start_pos = len(full_text)
                full_text += text + "\n"
                end_pos = len(full_text)
                
                # æ–°å¢ï¼šæå–runsçš„æ ¼å¼ä¿¡æ¯
                runs_info = []
                text_with_format = ""
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue
                    
                    run_data = {
                        "text": run_text,
                        "bold": run.bold if run.bold is not None else False,
                        "italic": run.italic if run.italic is not None else False,
                        "underline": run.underline if run.underline is not None else False,
                    }
                    
                    # æå–å­—ä½“å¤§å°
                    if run.font.size:
                        run_data["size_pt"] = run.font.size.pt
                    
                    # æå–å­—ä½“é¢œè‰²
                    if run.font.color and run.font.color.rgb:
                        run_data["color"] = str(run.font.color.rgb)
                    
                    # æå–å­—ä½“åç§°
                    if run.font.name:
                        run_data["font_name"] = run.font.name
                    
                    runs_info.append(run_data)
                    
                    # æ„å»ºå¸¦æ ¼å¼æ ‡è®°çš„æ–‡æœ¬
                    formatted = run_text
                    if run.bold:
                        formatted = f"<b>{formatted}</b>"
                    if run.italic:
                        formatted = f"<i>{formatted}</i>"
                    if run.underline:
                        formatted = f"<u>{formatted}</u>"
                    text_with_format += formatted
                
                full_text_with_format += text_with_format + "\n"
                
                paragraphs_info.append({
                    "index": para_global_idx,
                    "text": text,
                    "text_with_format": text_with_format,  # æ–°å¢
                    "runs": runs_info,  # æ–°å¢
                    "start_pos": start_pos,
                    "end_pos": end_pos,
                    "para_obj": para,  # ä¿ç•™åŸå¯¹è±¡ï¼Œç”¨äºåç»­ä¿®æ”¹
                    "source": "body",
                    "body_idx": para_idx
                })
                para_global_idx += 1
            
            # æå–è¡¨æ ¼æ®µè½ï¼ˆåŒæ ·æå–æ ¼å¼ä¿¡æ¯ï¼‰
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for cell_para_idx, cell_para in enumerate(cell.paragraphs):
                            text = cell_para.text
                            if not text.strip():
                                continue
                            
                            start_pos = len(full_text)
                            full_text += text + "\n"
                            end_pos = len(full_text)
                            
                            # æå–è¡¨æ ¼å†…çš„æ ¼å¼ä¿¡æ¯
                            runs_info = []
                            text_with_format = ""
                            for run in cell_para.runs:
                                run_text = run.text
                                if not run_text:
                                    continue
                                
                                run_data = {
                                    "text": run_text,
                                    "bold": run.bold if run.bold is not None else False,
                                    "italic": run.italic if run.italic is not None else False,
                                }
                                
                                if run.font.size:
                                    run_data["size_pt"] = run.font.size.pt
                                if run.font.color and run.font.color.rgb:
                                    run_data["color"] = str(run.font.color.rgb)
                                
                                runs_info.append(run_data)
                                
                                formatted = run_text
                                if run.bold:
                                    formatted = f"<b>{formatted}</b>"
                                if run.italic:
                                    formatted = f"<i>{formatted}</i>"
                                text_with_format += formatted
                            
                            full_text_with_format += text_with_format + "\n"
                            
                            paragraphs_info.append({
                                "index": para_global_idx,
                                "text": text,
                                "text_with_format": text_with_format,  # æ–°å¢
                                "runs": runs_info,  # æ–°å¢
                                "start_pos": start_pos,
                                "end_pos": end_pos,
                                "para_obj": cell_para,  # ä¿ç•™åŸå¯¹è±¡
                                "source": "table",
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "cell_idx": cell_idx,
                                "cell_para_idx": cell_para_idx
                            })
                            para_global_idx += 1
            
            return {
                "success": True,
                "file_path": file_path,
                "full_text": full_text,
                "full_text_with_format": full_text_with_format,  # æ–°å¢ï¼šå¸¦æ ¼å¼çš„æ–‡æœ¬
                "paragraphs": paragraphs_info,
                "total_chars": len(full_text)
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"æå–æ–‡æœ¬å¤±è´¥: {str(e)}"
            }
    
    # ==================== Step 3: é”šç‚¹å®šä½ ====================
    
    @staticmethod
    def locate_text_in_runs(para_obj, target_text: str) -> Optional[Dict[str, Any]]:
        """
        åœ¨æ®µè½çš„Runçº§åˆ«ç²¾ç¡®å®šä½æ–‡æœ¬
        
        Returns:
            {
                "found": True,
                "start_run_index": 0,  # èµ·å§‹Runç´¢å¼•
                "end_run_index": 2,    # ç»“æŸRunç´¢å¼•
                "start_char_offset": 5,  # åœ¨èµ·å§‹Runä¸­çš„å­—ç¬¦åç§»
                "end_char_offset": 10,   # åœ¨ç»“æŸRunä¸­çš„å­—ç¬¦åç§»
                "matched_text": "..."
            }
        """
        # æ„å»ºæ®µè½çš„å®Œæ•´æ–‡æœ¬å’ŒRunæ˜ å°„
        runs = para_obj.runs
        if not runs:
            return None
        
        # æ„å»ºæ–‡æœ¬åˆ°Runçš„æ˜ å°„
        full_text = ""
        run_map = []  # [(run_index, start_pos, end_pos, run_obj)]
        
        for i, run in enumerate(runs):
            start = len(full_text)
            full_text += run.text
            end = len(full_text)
            run_map.append((i, start, end, run))
        
        # æŸ¥æ‰¾ç›®æ ‡æ–‡æœ¬
        pos = full_text.find(target_text)
        if pos == -1:
            return None
        
        target_end = pos + len(target_text)
        
        # æ‰¾åˆ°èµ·å§‹å’Œç»“æŸRun
        start_run = None
        end_run = None
        start_offset = 0
        end_offset = 0
        
        for run_idx, start_pos, end_pos, run_obj in run_map:
            # ç›®æ ‡æ–‡æœ¬å¼€å§‹ä½ç½®
            if start_run is None and start_pos <= pos < end_pos:
                start_run = run_idx
                start_offset = pos - start_pos
            
            # ç›®æ ‡æ–‡æœ¬ç»“æŸä½ç½®
            if start_pos < target_end <= end_pos:
                end_run = run_idx
                end_offset = target_end - start_pos
                break
        
        if start_run is not None and end_run is not None:
            return {
                "found": True,
                "start_run_index": start_run,
                "end_run_index": end_run,
                "start_char_offset": start_offset,
                "end_char_offset": end_offset,
                "matched_text": target_text,
                "match_type": "precise"
            }
        
        return None

    @staticmethod
    def _set_run_text(run_element, text: str) -> None:
        """è®¾ç½®runå…ƒç´ æ–‡æœ¬ï¼Œä¿ç•™åŸæœ‰æ ¼å¼"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for t in run_element.findall(qn('w:t')):
            run_element.remove(t)
        t = OxmlElement('w:t')
        if text and (text[0].isspace() or text[-1].isspace()):
            t.set(qn('xml:space'), 'preserve')
        t.text = text
        run_element.append(t)

    @staticmethod
    def _split_run_at(para_obj, run_index: int, offset: int) -> int:
        """åœ¨æŒ‡å®šrunå†…æŒ‰å­—ç¬¦åç§»æ‹†åˆ†ï¼Œè¿”å›å³ä¾§æ–°runçš„ç´¢å¼•"""
        runs = para_obj.runs
        if run_index >= len(runs):
            return run_index

        run = runs[run_index]
        text = run.text or ""
        if offset <= 0 or offset >= len(text):
            return run_index

        left_text = text[:offset]
        right_text = text[offset:]

        # æ›´æ–°å½“å‰runä¸ºå·¦ä¾§æ–‡æœ¬
        run.text = left_text

        # å…‹éš†runå¹¶è®¾ç½®å³ä¾§æ–‡æœ¬
        new_run = deepcopy(run._element)
        DocumentAnnotator._set_run_text(new_run, right_text)
        run._element.addnext(new_run)

        return run_index + 1

    @staticmethod
    def _isolate_target_runs(para_obj, precise_location: Dict[str, Any], target_text: str) -> Tuple[int, int]:
        """æ‹†åˆ†runï¼Œç¡®ä¿ç›®æ ‡æ–‡æœ¬ç‹¬å runèŒƒå›´"""
        start_idx = precise_location["start_run_index"]
        end_idx = precise_location["end_run_index"]
        start_offset = precise_location["start_char_offset"]
        end_offset = precise_location["end_char_offset"]

        if start_idx == end_idx:
            # å•runå†…ï¼šå…ˆæ‹†èµ·ç‚¹ï¼Œå†æ‹†ç»ˆç‚¹
            right_idx = DocumentAnnotator._split_run_at(para_obj, start_idx, start_offset)
            DocumentAnnotator._split_run_at(para_obj, right_idx, len(target_text))
            return right_idx, right_idx

        # å¤šrunï¼šæ‹†èµ·ç‚¹ä¸ç»ˆç‚¹
        right_idx = DocumentAnnotator._split_run_at(para_obj, start_idx, start_offset)
        if right_idx != start_idx:
            end_idx += 1  # èµ·ç‚¹æ‹†åˆ†åï¼Œç»ˆç‚¹ç´¢å¼•å³ç§»
        DocumentAnnotator._split_run_at(para_obj, end_idx, end_offset)
        return right_idx, end_idx
    
    def locate_text_in_paragraphs(
        self,
        paragraphs_info: List[Dict],
        target_text: str
    ) -> Optional[Dict[str, Any]]:
        """
        Step 3: é”šç‚¹å®šä½ï¼ˆThe Locatorï¼‰
        
        åœ¨Wordæ®µè½ä¸­æŸ¥æ‰¾ç›®æ ‡æ–‡æœ¬
        
        ç­–ç•¥ï¼š
        1. å…ˆå°è¯•ç²¾ç¡®åŒ¹é…ï¼ˆå¿«é€Ÿï¼‰
        2. å¦‚æœå¤±è´¥ï¼Œå¯ç”¨æ¨¡ç³ŠåŒ¹é…ï¼ˆå®¹é”™ï¼‰

        Returns:
            {
                "found": True,
                "para_index": 0,
                "position": 5,  # åœ¨æ®µè½ä¸­çš„ä½ç½®
                "matched_text": "å®é™…åŒ¹é…çš„æ–‡æœ¬",
                "para_obj": æ®µè½å¯¹è±¡
            }
        """
        if not target_text or len(target_text.strip()) == 0:
            return None
        
        # ç¬¬1æ­¥ï¼šç²¾ç¡®åŒ¹é…
        for para_info in paragraphs_info:
            para_text = para_info["text"]
            
            if target_text in para_text:
                position = para_text.find(target_text)
                return {
                    "found": True,
                    "para_index": para_info["index"],
                    "position": position,
                    "matched_text": target_text,
                    "para_obj": para_info.get("para_obj"),
                    "full_para_text": para_text,
                    "match_type": "exact"
                }
        
        # ç¬¬2æ­¥ï¼šæ¨¡ç³ŠåŒ¹é…ï¼ˆå®¹é”™å¤„ç†ï¼‰
        print(f"[Annotator] âš ï¸ ç²¾ç¡®åŒ¹é…å¤±è´¥ï¼Œå¯ç”¨æ¨¡ç³ŠåŒ¹é…: {target_text[:20]}...")
        
        best_match = None
        best_ratio = 0
        
        for para_info in paragraphs_info:
            para_text = para_info["text"]
            
            # è®¡ç®—ç›¸ä¼¼åº¦
            ratio = SequenceMatcher(None, target_text, para_text).ratio()
            
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para_info
        
        # å¦‚æœç›¸ä¼¼åº¦è¶…è¿‡é˜ˆå€¼ï¼Œä½¿ç”¨æ¨¡ç³ŠåŒ¹é…
        if best_ratio >= self.min_similarity and best_match:
            print(f"[Annotator] âœ“ æ¨¡ç³ŠåŒ¹é…æˆåŠŸ (ç›¸ä¼¼åº¦: {best_ratio:.2%})")
            
            return {
                "found": True,
                "para_index": best_match["index"],
                "position": 0,
                "matched_text": target_text,
                "para_obj": best_match.get("para_obj"),
                "full_para_text": best_match["text"],
                "match_type": "fuzzy",
                "similarity": best_ratio
            }
        
        return None
    
    # ==================== Step 4: æ ¼å¼æ— æŸæ³¨å…¥ ====================
    
    def inject_annotation_to_paragraph(
        self,
        doc,
        para_obj,
        position: int,
        original_text: str,
        suggestion: str,
        color: str = "FF0000"  # çº¢è‰²
    ) -> bool:
        """
        Step 4: æ ¼å¼æ— æŸæ³¨å…¥ï¼ˆThe Surgeryï¼‰
        
        åœ¨æ‰¾åˆ°çš„ä½ç½®æ’å…¥æ ‡æ³¨
        
        é»˜è®¤æ–¹æ¡ˆAï¼šæ’å…¥ Word åŸç”Ÿæ‰¹æ³¨ï¼ˆç²¾ç¡®æ ‡æ³¨original_textç‰‡æ®µï¼‰
        å¤±è´¥æ—¶å›é€€æ–¹æ¡ˆBï¼šåœ¨åŸæ–‡åç›´æ¥æ’å…¥é«˜äº®æ–‡å­—
        æ ¼å¼ï¼šåŸæ–‡... [å»ºè®®ä¿®æ”¹: suggestion]
        
        Args:
            para_obj: python-docxçš„æ®µè½å¯¹è±¡
            position: åœ¨æ®µè½ä¸­çš„ä½ç½®ï¼ˆä»0å¼€å§‹ï¼‰
            original_text: åŸæ–‡ç‰‡æ®µï¼ˆç”¨äºç²¾ç¡®å®šä½ï¼‰
            suggestion: ä¿®æ”¹å»ºè®®
            color: é«˜äº®é¢œè‰² (RGB hexï¼Œä¸å«#)
        
        Returns:
            æ˜¯å¦æˆåŠŸ
        """
        try:
            if self.annotation_mode == "comment":
                # ä¼ é€’original_textä»¥å®ç°ç²¾ç¡®æ ‡æ³¨
                return self._inject_comment(doc, para_obj, suggestion, original_text)
        except Exception as e:
            print(f"[Annotator] âš ï¸ æ‰¹æ³¨æ³¨å…¥å¤±è´¥ï¼Œå›é€€é«˜äº®: {str(e)}")

        # å›é€€ä¸ºé«˜äº®æ–‡å­—
        try:
            from docx.shared import RGBColor

            new_run = para_obj.add_run(f" [å»ºè®®ä¿®æ”¹: {suggestion}]")
            new_run.font.color.rgb = RGBColor(
                int(color[0:2], 16),
                int(color[2:4], 16),
                int(color[4:6], 16)
            )
            new_run.font.bold = True
            return True
        except Exception as e:
            print(f"[Annotator] âœ— æ³¨å…¥æ ‡æ³¨å¤±è´¥: {str(e)}")
            return False

    @staticmethod
    def _inject_comment(doc, para_obj, suggestion: str, target_text: str = None) -> bool:
        """æ’å…¥çœŸæ­£çš„WordåŸç”Ÿæ‰¹æ³¨ï¼ˆå³ä¾§Commentï¼Œä¸å½±å“æ–‡æ¡£æ ¼å¼ï¼‰
        
        Args:
            doc: Documentå¯¹è±¡
            para_obj: æ®µè½å¯¹è±¡
            suggestion: æ‰¹æ³¨å»ºè®®å†…å®¹
            target_text: éœ€è¦æ ‡æ³¨çš„å…·ä½“æ–‡æœ¬ç‰‡æ®µï¼ˆå¦‚æœæä¾›ï¼Œä¼šç²¾ç¡®æ ‡æ³¨è¯¥ç‰‡æ®µï¼‰
        """
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import qn
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            import datetime
            
            # 1. è·å–æˆ–åˆ›å»ºcommentséƒ¨åˆ†
            comments_part = None
            for rel in doc.part.rels.values():
                if "comments" in rel.target_ref:
                    comments_part = rel.target_part
                    break
            
            if comments_part is None:
                # åˆ›å»ºæ–°çš„comments.xmléƒ¨åˆ†
                from docx.opc.part import XmlPart
                from docx.opc.packuri import PackURI
                
                # åˆ›å»ºcomments part
                partname = PackURI('/word/comments.xml')
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
                
                # åˆ›å»ºç©ºçš„comments XMLç»“æ„
                comments_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            mc:Ignorable="w14 w15">
</w:comments>'''
                
                comments_part = XmlPart.load(partname, content_type, comments_xml.encode('utf-8'), doc.part.package)
                
                # æ·»åŠ å…³ç³»
                doc.part.relate_to(comments_part, RT.COMMENTS)
            
            # 2. ç”Ÿæˆå”¯ä¸€çš„æ‰¹æ³¨ID
            comments_root = comments_part.element
            existing_ids = [int(c.get(qn('w:id'))) for c in comments_root.findall(qn('w:comment')) if c.get(qn('w:id'))]
            comment_id = str(max(existing_ids) + 1 if existing_ids else 0)
            
            # 3. åœ¨æ®µè½ä¸­æ’å…¥æ‰¹æ³¨æ ‡è®°
            p = para_obj._element
            
            # å°è¯•ç²¾ç¡®å®šä½ç›®æ ‡æ–‡æœ¬
            precise_location = None
            if target_text:
                precise_location = DocumentAnnotator.locate_text_in_runs(para_obj, target_text)
            
            if precise_location and precise_location["found"]:
                # ç²¾ç¡®æ ‡æ³¨ï¼šåªæ ‡æ³¨å…·ä½“çš„æ–‡æœ¬ç‰‡æ®µ
                start_idx, end_idx = DocumentAnnotator._isolate_target_runs(
                    para_obj,
                    precise_location,
                    target_text
                )
                runs = list(p.findall(qn('w:r')))
                
                # åœ¨èµ·å§‹Runä¹‹å‰æ’å…¥commentRangeStart
                if start_idx < len(runs):
                    commentRangeStart = parse_xml(
                        f'<w:commentRangeStart w:id="{comment_id}" '
                        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                    )
                    start_run = runs[start_idx]
                    p.insert(list(p).index(start_run), commentRangeStart)
                    
                    # åœ¨ç»“æŸRunä¹‹åæ’å…¥commentRangeEnd
                    if end_idx < len(runs):
                        commentRangeEnd = parse_xml(
                            f'<w:commentRangeEnd w:id="{comment_id}" '
                            f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                        )
                        end_run = runs[end_idx]
                        p.insert(list(p).index(end_run) + 1, commentRangeEnd)
                        
                        # åœ¨ç»“æŸæ ‡è®°åæ’å…¥commentReference
                        commentRef_run = parse_xml(
                            f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                                <w:rPr>
                                    <w:rStyle w:val="CommentReference"/>
                                </w:rPr>
                                <w:commentReference w:id="{comment_id}"/>
                            </w:r>'''
                        )
                        p.insert(list(p).index(end_run) + 2, commentRef_run)
                        print(f"[Annotator] âœ… ç²¾ç¡®æ ‡æ³¨: {target_text[:30]}...")
                    else:
                        raise Exception("End run not found")
                else:
                    raise Exception("Start run not found")
            else:
                # å›é€€ï¼šæ ‡æ³¨æ•´æ®µ
                commentRangeStart = parse_xml(
                    f'<w:commentRangeStart w:id="{comment_id}" '
                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
                p.insert(0, commentRangeStart)
                
                commentRangeEnd = parse_xml(
                    f'<w:commentRangeEnd w:id="{comment_id}" '
                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
                )
                p.append(commentRangeEnd)
                
                commentRef_run = parse_xml(
                    f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:rPr>
                            <w:rStyle w:val="CommentReference"/>
                        </w:rPr>
                        <w:commentReference w:id="{comment_id}"/>
                    </w:r>'''
                )
                p.append(commentRef_run)
                print(f"[Annotator] âš ï¸ å›é€€åˆ°æ•´æ®µæ ‡æ³¨")
            
            # 4. åœ¨comments.xmlä¸­æ·»åŠ æ‰¹æ³¨å†…å®¹
            date_str = datetime.datetime.now().isoformat()
            
            # è½¬ä¹‰XMLç‰¹æ®Šå­—ç¬¦
            suggestion_escaped = suggestion.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
            
            comment_xml = parse_xml(
                f'''<w:comment w:id="{comment_id}" 
                              w:author="Koto AI" 
                              w:initials="K" 
                              w:date="{date_str}"
                              xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:p>
                        <w:pPr>
                            <w:pStyle w:val="CommentText"/>
                        </w:pPr>
                        <w:r>
                            <w:rPr>
                                <w:rStyle w:val="CommentReference"/>
                            </w:rPr>
                            <w:annotationRef/>
                        </w:r>
                        <w:r>
                            <w:t xml:space="preserve">{suggestion_escaped}</w:t>
                        </w:r>
                    </w:p>
                </w:comment>'''
            )
            
            comments_root.append(comment_xml)
            
            print(f"[Annotator] âœ… Wordæ‰¹æ³¨å·²æ·»åŠ ï¼ˆID: {comment_id}ï¼‰")
            return True
            
        except Exception as e:
            print(f"[Annotator] âš ï¸ Wordæ‰¹æ³¨æ’å…¥å¤±è´¥: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    # ==================== å®Œæ•´å·¥ä½œæµ ====================
    
    def process_document(
        self,
        file_path: str,
        annotations: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """
        å¤„ç†æ–‡æ¡£ï¼šå®šä½ -> æ³¨å…¥
        
        Args:
            file_path: å·¥ä½œå‰¯æœ¬æ–‡ä»¶è·¯å¾„ï¼ˆå·²å¤åˆ¶ï¼‰
            annotations: æ ‡æ³¨åˆ—è¡¨
                [
                    {
                        "åŸæ–‡ç‰‡æ®µ": "éœ€è¦ä¿®æ”¹çš„æ–‡æœ¬",
                        "ä¿®æ”¹å»ºè®®": "å»ºè®®ä¿®æ”¹ä¸º..."
                    },
                    ...
                ]
        
        Returns:
            {
                "success": True,
                "applied": æˆåŠŸåº”ç”¨çš„æ ‡æ³¨æ•°,
                "failed": å¤±è´¥çš„æ ‡æ³¨æ•°,
                "details": [...]
            }
        """
        print(f"[Annotator] ğŸ”„ å¼€å§‹å¤„ç†æ–‡æ¡£...")
        
        # ç¬¬1æ­¥ï¼šæå–æ–‡æœ¬
        text_data = self.extract_text_from_word(file_path)
        if not text_data.get("success"):
            return {
                "success": False,
                "error": text_data.get("error")
            }
        
        paragraphs_info = text_data["paragraphs"]
        applied_count = 0
        failed_count = 0
        details = []
        
        # ç¬¬2æ­¥ï¼šé€ä¸ªå¤„ç†æ ‡æ³¨
        from docx import Document
        doc = Document(file_path)
        
        for anno in annotations:
            original_text = anno.get("åŸæ–‡ç‰‡æ®µ", "").strip()
            suggestion = anno.get("ä¿®æ”¹å»ºè®®", "").strip()
            
            if not original_text or not suggestion:
                failed_count += 1
                details.append({
                    "status": "skip",
                    "reason": "ç¼ºå°‘åŸæ–‡ç‰‡æ®µæˆ–ä¿®æ”¹å»ºè®®"
                })
                continue
            
            # ç¬¬3æ­¥ï¼šå®šä½æ–‡æœ¬
            location = self.locate_text_in_paragraphs(paragraphs_info, original_text)
            
            if not location or not location.get("found"):
                failed_count += 1
                details.append({
                    "original": original_text[:30],
                    "suggestion": suggestion[:30],
                    "status": "failed",
                    "reason": "æ–‡æœ¬æœªæ‰¾åˆ°"
                })
                continue
            
            # ç¬¬4æ­¥ï¼šè·å–æ®µè½å¯¹è±¡ï¼ˆå¯èƒ½åœ¨è¡¨æ ¼ä¸­ï¼‰
            para_info = paragraphs_info[location["para_index"]]
            
            if para_info["source"] == "body":
                # æ­£æ–‡æ®µè½
                para_obj = doc.paragraphs[para_info["body_idx"]]
            elif para_info["source"] == "table":
                # è¡¨æ ¼æ®µè½
                table = doc.tables[para_info["table_idx"]]
                cell = table.rows[para_info["row_idx"]].cells[para_info["cell_idx"]]
                para_obj = cell.paragraphs[para_info["cell_para_idx"]]
            else:
                failed_count += 1
                details.append({
                    "original": original_text[:30],
                    "suggestion": suggestion[:30],
                    "status": "failed",
                    "reason": "æœªçŸ¥æ®µè½æ¥æº"
                })
                continue
            
            # ç¬¬5æ­¥ï¼šæ³¨å…¥æ ‡æ³¨
            success = self.inject_annotation_to_paragraph(
                doc=doc,
                para_obj=para_obj,
                position=location["position"],
                original_text=original_text,
                suggestion=suggestion
            )
            
            if success:
                applied_count += 1
                details.append({
                    "original": original_text[:30],
                    "suggestion": suggestion[:30],
                    "status": "success",
                    "match_type": location.get("match_type")
                })
            else:
                failed_count += 1
                details.append({
                    "original": original_text[:30],
                    "suggestion": suggestion[:30],
                    "status": "failed",
                    "reason": "æ³¨å…¥å¤±è´¥"
                })
        
        # ç¬¬5æ­¥ï¼šä¿å­˜ä¿®æ”¹
        try:
            doc.save(file_path)
            print(f"[Annotator] ğŸ’¾ æ–‡æ¡£å·²ä¿å­˜")
        except Exception as e:
            return {
                "success": False,
                "error": f"ä¿å­˜æ–‡æ¡£å¤±è´¥: {str(e)}"
            }
        
        return {
            "success": True,
            "applied": applied_count,
            "failed": failed_count,
            "total": len(annotations),
            "success_rate": f"{applied_count / len(annotations) * 100:.1f}%" if annotations else "0%",
            "details": details
        }
    
    # ==================== å®Œæ•´é—­ç¯ ====================
    
    def annotate_document(
        self,
        file_path: str,
        annotations: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """
        å®Œæ•´é—­ç¯ï¼šé¢„å¤„ç† -> å®šä½ -> æ³¨å…¥ -> ä¿å­˜
        
        Args:
            file_path: åŸå§‹Wordæ–‡ä»¶è·¯å¾„
            annotations: æ ‡æ³¨åˆ—è¡¨
        
        Returns:
            {
                "success": True,
                "original_file": "...",
                "revised_file": "...",
                "applied": 5,
                "failed": 1
            }
        """
        print("=" * 60)
        print("ğŸ“‘ Koto æ–‡æ¡£è‡ªåŠ¨æ ‡æ³¨ç³»ç»Ÿ")
        print("=" * 60)
        
        # Step 1: é¢„å¤„ç† - åˆ›å»ºå‰¯æœ¬
        try:
            original_file, revised_file = self.prepare_document(file_path)
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
        
        # Step 2: å¤„ç†æ–‡æ¡£ - å®šä½ + æ³¨å…¥
        result = self.process_document(revised_file, annotations)
        
        if result.get("success"):
            result["original_file"] = original_file
            result["revised_file"] = revised_file
            print(f"[Annotator] âœ… å®Œæˆï¼å·²åº”ç”¨ {result['applied']} ä¸ªæ ‡æ³¨")
            return result
        else:
            return result


if __name__ == "__main__":
    print("æ–‡æ¡£è‡ªåŠ¨æ ‡æ³¨ç³»ç»Ÿå·²å‡†å¤‡å°±ç»ª")
